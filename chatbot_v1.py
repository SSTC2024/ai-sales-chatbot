import os
import json
import sqlite3
import pandas as pd
import numpy as np
from datetime import datetime
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox, font
import threading
import queue
import time
import logging
from pathlib import Path
import re
import yaml

# AI/ML imports
import torch
from transformers import (
    AutoTokenizer, AutoModelForCausalLM, 
    BitsAndBytesConfig, pipeline
)
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Web search imports
import requests
from bs4 import BeautifulSoup
from googlesearch import search
import urllib.parse

# Document processing
import docx
from PyPDF2 import PdfReader
from openpyxl import load_workbook
import easyocr
import socket

# Language detection and translation
try:
    from googletrans import Translator, LANGUAGES
    import langdetect
    TRANSLATION_AVAILABLE = True
    print("✅ Translation libraries loaded successfully")
except ImportError:
    TRANSLATION_AVAILABLE = False
    print("⚠️ Translation libraries not available. Install with: pip install googletrans==4.0.0rc1 langdetect")
    
# Web interface imports
from flask import Flask, render_template, request, jsonify
import argparse
import socket

class VietnameseAISalesBot:
    """
    Enhanced AI Sales ChatBot with Vietnamese language support (Text-only version)
    Features robust error handling for response generation and document processing
    """
    
    def __init__(self):
        self.load_config()
        self.setup_logging()
        self.setup_language_support()
        self.setup_database()
        self.initialize_ai_models()
        self.conversation_context = []
        self.conversation_summary = ""
        self.setup_gui()
    
    def load_config(self):
        """Load configuration with Vietnamese language settings"""
        try:
            with open('chatbot_config.yaml', 'r', encoding='utf-8') as file:
                self.config = yaml.safe_load(file)
            print("✅ Configuration loaded from chatbot_config.yaml")
            
            self.config = self.update_config_for_vietnamese()
            
        except FileNotFoundError:
            print("⚠️ chatbot_config.yaml not found, using default Vietnamese settings")
            self.config = self.get_default_vietnamese_config()
        except Exception as e:
            print(f"❌ Error loading config: {e}")
            self.config = self.get_default_vietnamese_config()
    
    def update_config_for_vietnamese(self):
        """Update existing config to support Vietnamese features"""
        config = self.config.copy()
        
        if 'language_config' not in config:
            config['language_config'] = {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True,
                'translate_responses': True
            }
        
        if 'conversation_flows' not in config:
            config['conversation_flows'] = {}
            
        config['conversation_flows'].update({
            'greeting_vietnamese': {
                'triggers': ['xin chào', 'chào', 'chào bạn', 'chào anh', 'chào chị'],
                'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng của bạn. Tôi có thể giúp bạn tìm kiếm sản phẩm phù hợp với nhu cầu. Bạn đang tìm kiếm gì hôm nay?',
                'next_stage': 'needs_assessment'
            }
        })
        
        config['vietnamese_templates'] = {
            'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp với tiêu chí của bạn.',
            'out_of_stock': 'Rất tiếc, {product_name} hiện đang hết hàng.',
            'processing': 'Đang xử lý yêu cầu của bạn...',
            'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn. Vui lòng thử lại.',
            'model_error': 'Mô hình AI không khả dụng. Đang sử dụng phản hồi dự phòng.',
            'generation_error': 'Có lỗi khi tạo phản hồi. Vui lòng thử lại.',
            'empty_response': 'Tôi hiểu câu hỏi của bạn nhưng cần thêm thông tin để trả lời tốt hơn.',
            'connection_error': 'Lỗi kết nối. Vui lòng kiểm tra mạng và thử lại.'
        }
        
        return config
    
    def get_default_vietnamese_config(self):
        """Get default configuration with Vietnamese support and error handling"""
        return {
            'version': '2.0',
            'environment': 'production',
            'ai_models': {
                'primary_llm': 'microsoft/DialoGPT-medium',
                'fallback_llm': 'microsoft/DialoGPT-small',
                'embedding_model': 'sentence-transformers/all-MiniLM-L6-v2'
            },
            'language_config': {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True,
                'translate_responses': True
            },
            'gpu_config': {
                'primary_device': 'cuda:0' if torch.cuda.is_available() else 'cpu',
                'use_quantization': True,
                'mixed_precision': True,
                'max_memory_per_gpu': 0.85,
                'batch_size': 2
            },
            'search_config': {
                'local_similarity_threshold': 0.1,
                'enable_google_search': True,
                'max_google_results': 3,
                'search_timeout': 10
            },
            'performance': {
                'max_response_length': 150,
                'temperature': 0.7,
                'top_p': 0.9,
                'repetition_penalty': 1.1,
                'do_sample': True,
                'num_return_sequences': 1
            },
            'conversation_flows': {
                'greeting_vietnamese': {
                    'triggers': ['xin chào', 'chào', 'chào bạn'],
                    'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng của bạn.'
                },
                'greeting': {
                    'triggers': ['hello', 'hi', 'hey'],
                    'response_template': 'Hello! I am your AI sales assistant.'
                }
            },
            'vietnamese_templates': {
                'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp.',
                'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn.',
                'model_error': 'Mô hình AI không khả dụng.',
                'generation_error': 'Có lỗi khi tạo phản hồi.',
                'empty_response': 'Tôi cần thêm thông tin để trả lời tốt hơn.'
            },
            'analytics': {
                'track_conversations': True,
                'track_language_usage': True,
                'track_errors': True
            }
        }
    
    def setup_language_support(self):
        """Initialize language detection and translation"""
        self.current_language = self.config['language_config']['default_language']
        
        if TRANSLATION_AVAILABLE:
            try:
                self.translator = Translator()
                print("✅ Translation service initialized")
            except Exception as e:
                print(f"⚠️ Translation service initialization failed: {e}")
                self.translator = None
        else:
            self.translator = None
    
    def detect_language(self, text):
        """Detect language of input text with error handling"""
        if not TRANSLATION_AVAILABLE or not text or not text.strip():
            return 'vi'
        
        try:
            detected = langdetect.detect(text)
            if detected in self.config['language_config']['supported_languages']:
                return detected
            return 'vi'
        except Exception as e:
            print(f"Language detection error: {e}")
            return 'vi'
    
    def translate_text(self, text, target_language):
        """Translate text to target language with error handling"""
        if not self.translator or not text or not text.strip():
            return text
        
        try:
            translated = self.translator.translate(text, dest=target_language)
            return translated.text if translated and translated.text else text
        except Exception as e:
            print(f"Translation error: {e}")
            return text
        
    def setup_logging(self):
        """Setup logging configuration with Vietnamese support"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('chatbot.log', encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
        
    def setup_database(self):
        """Initialize SQLite database with Vietnamese support"""
        self.conn = sqlite3.connect('chatbot_data.db', check_same_thread=False)
        self.conn.execute("PRAGMA encoding = 'UTF-8'")
        
        cursor = self.conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS products (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                name_vietnamese TEXT,
                description TEXT,
                description_vietnamese TEXT,
                category TEXT,
                category_vietnamese TEXT,
                price REAL,
                features TEXT,
                features_vietnamese TEXT,
                specifications TEXT,
                specifications_vietnamese TEXT,
                availability TEXT,
                availability_vietnamese TEXT,
                source_file TEXT,
                embedding BLOB,
                created_at TEXT,
                updated_at TEXT
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT,
                user_input TEXT,
                user_language TEXT,
                bot_response TEXT,
                bot_language TEXT,
                intent TEXT,
                confidence REAL,
                data_source TEXT,
                response_time REAL
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS knowledge_base (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                topic TEXT,
                content TEXT,
                source TEXT,
                embedding BLOB,
                created_at TEXT
            )
        ''')
        
        try:
            cursor.execute('ALTER TABLE conversations ADD COLUMN error_type TEXT')
            print("✅ Added error_type column to conversations table")
        except sqlite3.OperationalError:
            pass
        
        try:
            cursor.execute('ALTER TABLE conversations ADD COLUMN error_message TEXT')
            print("✅ Added error_message column to conversations table")
        except sqlite3.OperationalError:
            pass
        
        self.conn.commit()
        print("✅ Database schema updated successfully")
        
    def initialize_ai_models(self):
        """Initialize AI models with enhanced error handling"""
        self.logger.info("Initializing AI models from configuration...")
        
        try:
            self.setup_devices()
            self.load_language_model()
            self.load_embedding_model()
            self.load_ocr_model()
            
            self.logger.info("✅ All AI models loaded successfully!")
            
        except Exception as e:
            self.logger.error(f"❌ Error initializing AI models: {e}")
            self.text_generator = None
            self.embedding_model = None
            self.ocr_reader = None
    
    def setup_devices(self):
        """Device setup using configuration with fallback"""
        try:
            if torch.cuda.is_available():
                gpu_count = torch.cuda.device_count()
                print(f"Found {gpu_count} GPU(s)")
                
                best_gpu = 0
                max_memory = 0
                
                for i in range(gpu_count):
                    props = torch.cuda.get_device_properties(i)
                    memory_gb = props.total_memory / (1024**3)
                    print(f"GPU {i}: {props.name} - {memory_gb:.1f}GB")
                    
                    if memory_gb > max_memory:
                        max_memory = memory_gb
                        best_gpu = i
                
                primary_device = self.config['gpu_config'].get('primary_device', f'cuda:{best_gpu}')
                
                self.primary_device = primary_device
                self.secondary_device = f'cuda:{(best_gpu + 1) % gpu_count}' if gpu_count > 1 else self.primary_device
                self.primary_memory_gb = max_memory
                
            else:
                self.primary_device = 'cpu'
                self.secondary_device = 'cpu'
                self.primary_memory_gb = 8
                print("CUDA not available. Using CPU mode.")
        except Exception as e:
            print(f"Device setup error: {e}")
            self.primary_device = 'cpu'
            self.secondary_device = 'cpu'
            self.primary_memory_gb = 8
    
    def load_language_model(self):
        """Load language model with comprehensive error handling"""
        try:
            primary_llm = self.config['ai_models']['primary_llm']
            fallback_llm = self.config['ai_models']['fallback_llm']
            use_quantization_config = self.config['gpu_config']['use_quantization']
            
            primary_memory = getattr(self, 'primary_memory_gb', 8)
            
            if primary_memory >= 12:
                use_quantization = use_quantization_config
                model_name = primary_llm
            else:
                use_quantization = True
                model_name = fallback_llm
            
            print(f"Loading language model: {model_name}")
            
            try:
                self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                if self.tokenizer.pad_token is None:
                    self.tokenizer.pad_token = self.tokenizer.eos_token
                print("✅ Tokenizer loaded successfully")
            except Exception as e:
                print(f"❌ Tokenizer loading failed: {e}")
                self.tokenizer = None
                self.text_generator = None
                return
            
            quantization_config = None
            if use_quantization and torch.cuda.is_available():
                try:
                    quantization_config = BitsAndBytesConfig(
                        load_in_8bit=True,
                        llm_int8_threshold=6.0,
                    )
                except Exception as e:
                    print(f"❌ Quantization config failed: {e}")
                    quantization_config = None
            
            try:
                self.llm_model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    quantization_config=quantization_config,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                    trust_remote_code=True,
                    low_cpu_mem_usage=True
                )
                print("✅ Language model loaded successfully")
            except Exception as e:
                print(f"❌ Model loading failed: {e}")
                self.llm_model = None
                self.text_generator = None
                return
            
            try:
                try:
                    self.text_generator = pipeline(
                        "text-generation",
                        model=self.llm_model,
                        tokenizer=self.tokenizer,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
                    )
                    print("✅ Text generation pipeline created successfully (accelerate mode)")
                except Exception as accelerate_error:
                    print(f"⚠️ Accelerate mode failed, trying with device specification: {accelerate_error}")
                    self.text_generator = pipeline(
                        "text-generation",
                        model=self.llm_model,
                        tokenizer=self.tokenizer,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                        device=0 if torch.cuda.is_available() else -1
                    )
                    print("✅ Text generation pipeline created successfully (device mode)")
            except Exception as e:
                print(f"❌ Pipeline creation failed: {e}")
                self.text_generator = None
            
        except Exception as e:
            print(f"❌ Error loading language model: {e}")
            self.text_generator = None
            
    def load_embedding_model(self):
        """Load embedding model with error handling"""
        try:
            embedding_model_name = self.config['ai_models']['embedding_model']
            print(f"Loading embedding model: {embedding_model_name}")
            
            self.embedding_model = SentenceTransformer(embedding_model_name)
            if torch.cuda.is_available():
                try:
                    self.embedding_model = self.embedding_model.to(self.secondary_device)
                except Exception as e:
                    print(f"⚠️ Could not move embedding model to GPU: {e}")
            print("✅ Embedding model loaded successfully")
        except Exception as e:
            print(f"❌ Error loading embedding model: {e}")
            self.embedding_model = None
            
    def load_ocr_model(self):
        """Load OCR model with error handling"""
        try:
            self.ocr_reader = easyocr.Reader(['en', 'vi'])
            print("✅ OCR model loaded successfully with Vietnamese support")
        except Exception as e:
            print(f"❌ Error loading OCR model: {e}")
            self.ocr_reader = None
    
    def safe_generate_response(self, prompt, user_language='vi'):
        """Safely generate response with comprehensive error handling"""
        try:
            if not self.text_generator or not self.tokenizer:
                self.logger.warning("Text generator not available")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['model_error']
                else:
                    return "I apologize, but the AI text generation system is not available right now."
            
            if not prompt or not isinstance(prompt, str) or not prompt.strip():
                self.logger.warning("Invalid prompt provided")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need more information to provide a helpful response."
            
            performance_config = self.config.get('performance', {})
            
            try:
                with torch.no_grad():
                    if self.tokenizer.pad_token_id is None:
                        self.tokenizer.pad_token_id = self.tokenizer.eos_token_id
                    
                    generated = self.text_generator(
                        prompt,
                        max_new_tokens=performance_config.get('max_response_length', 150),
                        min_new_tokens=10,
                        temperature=performance_config.get('temperature', 0.7),
                        top_p=performance_config.get('top_p', 0.9),
                        repetition_penalty=performance_config.get('repetition_penalty', 1.1),
                        do_sample=performance_config.get('do_sample', True),
                        pad_token_id=self.tokenizer.pad_token_id,
                        num_return_sequences=performance_config.get('num_return_sequences', 1),
                        return_full_text=False,
                        batch_size=1
                    )
                
                if not generated:
                    self.logger.warning("Empty generation result")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['generation_error']
                    else:
                        return "I'm having trouble generating a response. Please try again."
                
                if isinstance(generated, list) and len(generated) > 0:
                    first_result = generated[0]
                    if isinstance(first_result, dict) and 'generated_text' in first_result:
                        response = first_result['generated_text']
                    else:
                        response = str(first_result)
                else:
                    response = str(generated)
                
                if not response or not response.strip():
                    self.logger.warning("Empty response generated")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question but need more details to respond properly."
                
                cleaned_response = self.clean_generated_response(response)
                
                if not cleaned_response or len(cleaned_response.strip()) < 5:
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question. Could you provide more details?"
                
                return cleaned_response
                
            except torch.cuda.OutOfMemoryError:
                self.logger.error("CUDA out of memory")
                if user_language == 'vi':
                    return "Hệ thống đang quá tải. Vui lòng thử lại sau."
                else:
                    return "System is overloaded. Please try again later."
                    
            except Exception as generation_error:
                self.logger.error(f"Text generation error: {generation_error}")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['generation_error']
                else:
                    return "I encountered an error while generating a response. Please try again."
            
        except Exception as e:
            self.logger.error(f"Safe generation error: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having technical difficulties. Please try again."
    
    def clean_generated_response(self, response):
        """Clean up generated response with error handling"""
        try:
            if not response:
                return ""
            
            response = str(response)
            
            response = re.sub(r'<[^>]+>', '', response)
            response = re.sub(r'\[.*?\]', '', response)
            response = re.sub(r'\n+', ' ', response)
            response = re.sub(r'\s+', ' ', response)
            
            response = response.replace('<|endoftext|>', '')
            response = response.replace('</s>', '')
            response = response.replace('<s>', '')
            
            try:
                sentences = response.split('.')
                if len(sentences) > 1 and len(sentences[-1].strip()) < 10:
                    response = '.'.join(sentences[:-1]) + '.'
            except Exception:
                pass
            
            response = response.strip()
            
            if len(response) < 3:
                return ""
                
            return response
            
        except Exception as e:
            self.logger.error(f"Response cleaning error: {e}")
            return str(response) if response else ""
    
    def process_message(self, user_input):
        """Main processing pipeline with Vietnamese support and comprehensive error handling"""
        start_time = time.time()
        
        try:
            if not user_input or not isinstance(user_input, str):
                error_msg = "Vui lòng nhập tin nhắn hợp lệ." if self.current_language == 'vi' else "Please enter a valid message."
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'invalid_input'
                }
            
            user_input = user_input.strip()
            if not user_input:
                error_msg = "Vui lòng nhập tin nhắn." if self.current_language == 'vi' else "Please enter a message."
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'empty_input'
                }
            
            try:
                user_language = self.detect_language(user_input)
            except Exception as e:
                self.logger.warning(f"Language detection error: {e}")
                user_language = self.current_language or 'vi'
            
            context_references = self.check_context_references(user_input, user_language)
            
            local_results = []
            knowledge_results = []
            database_searched = False
            
            try:
                if context_references:
                    search_query = self.build_contextual_search_query(user_input, context_references)
                    self.logger.info(f"Contextual search for: {search_query}")
                else:
                    search_query = user_input
                    
                self.logger.info(f"Searching database for: {search_query}")
                local_results = self.search_local_database(search_query)
                database_searched = True
                
                if not local_results:
                    self.logger.info("No products found, searching knowledge base...")
                    knowledge_results = self.search_knowledge_base(search_query)
                    
            except Exception as e:
                self.logger.error(f"Database search error: {e}")
                local_results = []
                knowledge_results = []
            
            try:
                response = ""
                data_source = "unknown"
                
                if local_results:
                    self.logger.info(f"Found {len(local_results)} products, generating AI response with product context")
                    response = self.generate_natural_response(
                        user_input, 
                        context_data=local_results,
                        data_source="database",
                        user_language=user_language
                    )
                    data_source = "local_database"
                    
                elif knowledge_results:
                    self.logger.info(f"Found {len(knowledge_results)} knowledge entries, generating AI response with knowledge context")
                    response = self.generate_natural_response(
                        user_input,
                        context_data=knowledge_results,
                        data_source="knowledge_base",
                        user_language=user_language
                    )
                    data_source = "knowledge_base"
                    
                else:
                    self.logger.info("No database results found, using AI general knowledge")
                    
                    if database_searched:
                        if user_language == 'vi':
                            no_results_msg = "Tôi đã tìm kiếm trong cơ sở dữ liệu nhưng không tìm thấy sản phẩm phù hợp. Để tôi trả lời dựa trên kiến thức chung:\n\n"
                        else:
                            no_results_msg = "I searched our database but couldn't find matching products. Let me answer based on general knowledge:\n\n"
                    else:
                        no_results_msg = ""
                    
                    ai_response = self.generate_natural_response(
                        user_input, 
                        context_data=None,
                        data_source="general",
                        user_language=user_language
                    )
                    
                    response = no_results_msg + ai_response
                    data_source = "ai_knowledge"
                    
            except Exception as e:
                self.logger.error(f"Response generation error: {e}")
                if user_language == 'vi':
                    response = self.config['vietnamese_templates']['generation_error']
                else:
                    response = "I encountered an error while generating a response. Please try again."
                data_source = "error"
            
            conversation_entry = {
                'user': user_input,
                'bot': response,
                'user_language': user_language,
                'timestamp': datetime.now().isoformat(),
                'data_source': data_source,
                'results_found': len(local_results) + len(knowledge_results),
                'products_mentioned': [p.get('name', '') for p in local_results[:3]] if local_results else [],
                'context_references': context_references
            }
            
            try:
                self.conversation_context.append(conversation_entry)
                
                if len(self.conversation_context) > 20:
                    self.conversation_context = self.conversation_context[-20:]
                    
                if len(self.conversation_context) % 5 == 0:
                    self.update_conversation_summary()
                    
            except Exception as e:
                self.logger.warning(f"Context update error: {e}")
                
            try:
                processing_time = time.time() - start_time
                self.store_conversation(user_input, response, data_source, processing_time, user_language)
            except Exception as e:
                self.logger.warning(f"Conversation storage error: {e}")
            
            self.logger.info(f"Response generated using: {data_source}")
            
            return {
                'response': response,
                'data_source': data_source,
                'local_results_count': len(local_results) + len(knowledge_results),
                'processing_time': time.time() - start_time,
                'user_language': user_language,
                'response_language': user_language,
                'database_searched': database_searched,
                'context_used': bool(context_references)
            }
            
        except Exception as e:
            self.logger.error(f"Message processing error: {e}")
            error_language = getattr(self, 'current_language', 'vi')
            if error_language == 'vi':
                error_msg = self.config['vietnamese_templates']['error']
            else:
                error_msg = "I apologize, but I encountered an error processing your request."
            
            return {
                'response': error_msg,
                'data_source': 'error',
                'local_results_count': 0,
                'processing_time': time.time() - start_time,
                'user_language': error_language,
                'response_language': error_language,
                'error': str(e)
            }
    
    def generate_natural_response(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Generate natural language response with Vietnamese support and error handling"""
        try:
            if not user_input or not isinstance(user_input, str):
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need a question or message to respond to."
            
            user_input = user_input.strip()
            if not user_input:
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "Please provide a question or message."
            
            try:
                if user_language == 'vi':
                    greeting_flow = self.config['conversation_flows'].get('greeting_vietnamese', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Xin chào!')
                else:
                    greeting_flow = self.config['conversation_flows'].get('greeting', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Hello!')
            except Exception as e:
                self.logger.warning(f"Greeting pattern matching error: {e}")
            
            if context_data and len(context_data) > 0:
                if data_source == "database":
                    prompt = self.build_product_aware_prompt(user_input, context_data, user_language)
                    ai_response = self.safe_generate_response(prompt, user_language)
                    
                    if not ai_response or ai_response == self.config['vietnamese_templates']['model_error']:
                        return self.format_product_response(context_data, user_input, user_language)
                    
                    return ai_response
                    
                elif data_source == "knowledge_base":
                    prompt = self.build_knowledge_aware_prompt(user_input, context_data, user_language)
                    ai_response = self.safe_generate_response(prompt, user_language)
                    
                    if not ai_response or ai_response == self.config['vietnamese_templates']['model_error']:
                        return self.format_knowledge_response(context_data, user_input, user_language)
                    
                    return ai_response
            
            prompt = self.build_sales_prompt(user_input, None, "general", user_language)
            response = self.safe_generate_response(prompt, user_language)
            
            if not response or response in [self.config['vietnamese_templates']['model_error'], 
                                           self.config['vietnamese_templates']['generation_error']]:
                if user_language == 'vi':
                    return "Xin lỗi, tôi đang gặp khó khăn trong việc tạo phản hồi. Bạn có thể mô tả chi tiết hơn về sản phẩm bạn đang tìm kiếm không?"
                else:
                    return "I apologize, I'm having trouble generating a response. Could you describe in more detail what product you're looking for?"
            
            return response
            
        except Exception as e:
            self.logger.error(f"Error generating natural response: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having trouble generating a response right now."
    
    def build_product_aware_prompt(self, user_input, products, user_language='vi'):
        """Build a prompt that incorporates product data for natural AI response"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý bán hàng AI thân thiện. Dựa trên câu hỏi của khách hàng và thông tin sản phẩm có sẵn, hãy tạo một phản hồi tự nhiên, hữu ích.

"""
            else:
                prompt = """You are a friendly AI sales assistant. Based on the customer's question and available product information, create a natural, helpful response.

"""
            
            summary_prompt = self.get_conversation_context_prompt(user_language)
            if summary_prompt:
                prompt += summary_prompt
            
            if self.conversation_context and len(self.conversation_context) > 0:
                if user_language == 'vi':
                    prompt += "Cuộc trò chuyện gần đây:\n"
                else:
                    prompt += "Recent conversation:\n"
                
                for turn in self.conversation_context[-3:]:
                    user_msg = turn.get('user', '')[:150]
                    bot_msg = turn.get('bot', '')[:150]
                    prompt += f"User: {user_msg}\nAssistant: {bot_msg}...\n"
                prompt += "\n"
            
            if user_language == 'vi':
                prompt += "Sản phẩm có sẵn:\n"
            else:
                prompt += "Available products:\n"
                
            for product in products[:3]:
                name = product.get('name_vietnamese', product.get('name', '')) if user_language == 'vi' else product.get('name', '')
                desc = product.get('description_vietnamese', product.get('description', '')) if user_language == 'vi' else product.get('description', '')
                price = product.get('price', 0)
                features = product.get('features_vietnamese', product.get('features', '')) if user_language == 'vi' else product.get('features', '')
                
                prompt += f"\n- {name}"
                if desc:
                    prompt += f": {desc[:100]}"
                if price > 0:
                    if user_language == 'vi':
                        prompt += f" (Giá: ${price:,.0f})"
                    else:
                        prompt += f" (Price: ${price:,.2f})"
                if features:
                    prompt += f" - {features[:50]}"
            
            if user_language == 'vi':
                prompt += f"\n\nCâu hỏi khách hàng: {user_input}\n\n"
                prompt += "Hãy trả lời một cách tự nhiên, nhớ ngữ cảnh cuộc trò chuyện, giới thiệu sản phẩm phù hợp và giải thích tại sao chúng đáp ứng nhu cầu của khách hàng. "
                prompt += "Nếu khách hàng đề cập đến sản phẩm đã nói trước đó, hãy nhắc lại thông tin đó:"
            else:
                prompt += f"\n\nCustomer question: {user_input}\n\n"
                prompt += "Provide a natural response that remembers the conversation context, introduces suitable products and explains why they meet the customer's needs. "
                prompt += "If the customer refers to previously mentioned products, recall that information:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Error building product prompt: {e}")
            return self.build_sales_prompt(user_input, products, "database", user_language)
    
    def build_knowledge_aware_prompt(self, user_input, knowledge_entries, user_language='vi'):
        """Build a prompt that incorporates knowledge base data for natural AI response"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý AI thông minh. Dựa trên thông tin từ cơ sở tri thức và câu hỏi của người dùng, hãy tạo một câu trả lời tự nhiên và hữu ích.

"""
            else:
                prompt = """You are an intelligent AI assistant. Based on information from the knowledge base and the user's question, create a natural and helpful response.

"""
            
            summary_prompt = self.get_conversation_context_prompt(user_language)
            if summary_prompt:
                prompt += summary_prompt
            
            if self.conversation_context and len(self.conversation_context) > 0:
                if user_language == 'vi':
                    prompt += "Ngữ cảnh cuộc trò chuyện:\n"
                else:
                    prompt += "Conversation context:\n"
                
                for turn in self.conversation_context[-2:]:
                    user_msg = turn.get('user', '')[:100]
                    prompt += f"User: {user_msg}\n"
                prompt += "\n"
            
            if user_language == 'vi':
                prompt += "Thông tin liên quan:\n"
            else:
                prompt += "Related information:\n"
                
            for entry in knowledge_entries[:2]:
                topic = entry.get('topic', '')
                content = entry.get('content', '')[:200]
                
                if topic:
                    prompt += f"\n{topic}: "
                if content:
                    prompt += f"{content}..."
            
            if user_language == 'vi':
                prompt += f"\n\nCâu hỏi: {user_input}\n\n"
                prompt += "Trả lời dựa trên ngữ cảnh cuộc trò chuyện và thông tin có sẵn:"
            else:
                prompt += f"\n\nQuestion: {user_input}\n\n"
                prompt += "Response based on conversation context and available information:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Error building knowledge prompt: {e}")
            return self.build_sales_prompt(user_input, knowledge_entries, "knowledge_base", user_language)
    
    def build_sales_prompt(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Build an optimized prompt for sales conversations with Vietnamese support"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý AI bán hàng chuyên nghiệp. Hãy trả lời bằng tiếng Việt một cách thân thiện và hữu ích.

"""
            else:
                prompt = """You are a professional AI sales assistant. Respond in English in a friendly and helpful manner.

"""
            
            try:
                if self.conversation_context and len(self.conversation_context) > 0:
                    if user_language == 'vi':
                        prompt += "Lịch sử cuộc trò chuyện:\n"
                    else:
                        prompt += "Conversation history:\n"
                    
                    recent_context = self.conversation_context[-5:]
                    for i, turn in enumerate(recent_context):
                        if isinstance(turn, dict) and 'user' in turn and 'bot' in turn:
                            user_msg = str(turn['user'])[:200]
                            bot_msg = str(turn['bot'])[:200]
                            
                            if user_language == 'vi':
                                prompt += f"\nLượt {i+1}:\n"
                            else:
                                prompt += f"\nTurn {i+1}:\n"
                            
                            prompt += f"User: {user_msg}\nAssistant: {bot_msg}\n"
                    
                    if len(self.conversation_context) > 5:
                        if user_language == 'vi':
                            prompt += f"\n(Đã có {len(self.conversation_context)} lượt trò chuyện trước đó)\n"
                        else:
                            prompt += f"\n(There were {len(self.conversation_context)} previous conversation turns)\n"
                    
                    prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Context building error: {e}")
                
            try:
                if context_data and isinstance(context_data, list) and len(context_data) > 0:
                    if data_source == "database":
                        if user_language == 'vi':
                            prompt += "Sản phẩm liên quan:\n"
                        else:
                            prompt += "Relevant products:\n"
                        
                        for item in context_data[:3]:
                            if isinstance(item, dict):
                                name = item.get('name_vietnamese', item.get('name', '')) if user_language == 'vi' else item.get('name', '')
                                desc = item.get('description_vietnamese', item.get('description', '')) if user_language == 'vi' else item.get('description', '')
                                
                                if name:
                                    prompt += f"- {str(name)[:100]}"
                                    if desc:
                                        prompt += f": {str(desc)[:150]}...\n"
                                    else:
                                        prompt += "\n"
                                        
                                    if item.get('price'):
                                        try:
                                            price_label = "Giá" if user_language == 'vi' else "Price"
                                            prompt += f"  {price_label}: ${float(item['price']):.2f}\n"
                                        except (ValueError, TypeError):
                                            pass
                        prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Product context building error: {e}")
                
            try:
                truncated_input = str(user_input)[:300]
                if user_language == 'vi':
                    prompt += f"Câu hỏi hiện tại của khách hàng: {truncated_input}\n\n"
                    prompt += "Hãy trả lời dựa trên lịch sử cuộc trò chuyện và duy trì ngữ cảnh. "
                    prompt += "Nếu khách hàng đề cập đến điều gì đó đã nói trước đó, hãy nhớ và tham chiếu đến nó.\n"
                    prompt += "Trả lời bằng tiếng Việt:"
                else:
                    prompt += f"Current customer question: {truncated_input}\n\n"
                    prompt += "Please respond based on the conversation history and maintain context. "
                    prompt += "If the customer refers to something mentioned earlier, remember and reference it.\n"
                    prompt += "Response:"
            except Exception as e:
                self.logger.error(f"Question formatting error: {e}")
                if user_language == 'vi':
                    prompt += f"Câu hỏi: {user_input}\nTrả lời:"
                else:
                    prompt += f"Question: {user_input}\nResponse:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Prompt building error: {e}")
            if user_language == 'vi':
                return f"Câu hỏi: {user_input}\nTrả lời bằng tiếng Việt:"
            else:
                return f"Question: {user_input}\nResponse:"
    
    def format_product_response(self, products, user_input, user_language='vi'):
        """Format product search results into a natural response"""
        try:
            if user_language == 'vi':
                response = "Tôi tìm thấy các sản phẩm sau phù hợp với yêu cầu của bạn:\n\n"
                
                for i, product in enumerate(products[:3], 1):
                    name = product.get('name_vietnamese') or product.get('name', '')
                    desc = product.get('description_vietnamese') or product.get('description', '')
                    price = product.get('price', 0)
                    
                    response += f"{i}. **{name}**\n"
                    if desc:
                        response += f"   {desc[:150]}...\n"
                    if price > 0:
                        response += f"   💰 Giá: ${price:,.2f}\n"
                    response += "\n"
                
                response += "Bạn muốn biết thêm thông tin chi tiết về sản phẩm nào?"
            else:
                response = "I found the following products that match your request:\n\n"
                
                for i, product in enumerate(products[:3], 1):
                    name = product.get('name', '')
                    desc = product.get('description', '')
                    price = product.get('price', 0)
                    
                    response += f"{i}. **{name}**\n"
                    if desc:
                        response += f"   {desc[:150]}...\n"
                    if price > 0:
                        response += f"   💰 Price: ${price:,.2f}\n"
                    response += "\n"
                
                response += "Which product would you like to know more about?"
                
            return response
        except Exception as e:
            self.logger.error(f"Error formatting product response: {e}")
            if user_language == 'vi':
                return "Tôi tìm thấy một số sản phẩm phù hợp. Bạn có thể cho tôi biết thêm về nhu cầu của bạn không?"
            else:
                return "I found some matching products. Could you tell me more about your needs?"
    
    def format_knowledge_response(self, knowledge_entries, user_input, user_language='vi'):
        """Format knowledge base search results into a natural response"""
        try:
            if user_language == 'vi':
                response = "Dựa trên thông tin trong cơ sở dữ liệu, tôi có thể chia sẻ:\n\n"
                
                for entry in knowledge_entries[:2]:
                    topic = entry.get('topic', '')
                    content = entry.get('content', '')
                    
                    if topic:
                        response += f"📚 **{topic}**\n"
                    if content:
                        response += f"{content[:300]}...\n\n"
                
                response += "Bạn cần thông tin chi tiết hơn về vấn đề nào?"
            else:
                response = "Based on the information in our knowledge base:\n\n"
                
                for entry in knowledge_entries[:2]:
                    topic = entry.get('topic', '')
                    content = entry.get('content', '')
                    
                    if topic:
                        response += f"📚 **{topic}**\n"
                    if content:
                        response += f"{content[:300]}...\n\n"
                
                response += "Would you like more detailed information on any topic?"
                
            return response
        except Exception as e:
            self.logger.error(f"Error formatting knowledge response: {e}")
            if user_language == 'vi':
                return "Tôi tìm thấy một số thông tin liên quan. Bạn có thể hỏi cụ thể hơn không?"
            else:
                return "I found some related information. Could you be more specific?"
    
    def search_local_database(self, user_input, similarity_threshold=None):
        """Search local database with Vietnamese support and error handling"""
        try:
            if not user_input or not user_input.strip():
                return []
            
            if similarity_threshold is None:
                similarity_threshold = self.config['search_config']['local_similarity_threshold']
            
            if self.embedding_model:
                try:
                    query_embedding = self.embedding_model.encode([user_input.strip()])
                    if query_embedding is None or len(query_embedding) == 0:
                        self.logger.warning("Failed to generate query embedding")
                        return self.keyword_search_fallback(user_input)
                except Exception as e:
                    self.logger.error(f"Query embedding error: {e}")
                    return self.keyword_search_fallback(user_input)
                
                try:
                    cursor = self.conn.cursor()
                    cursor.execute("""
                        SELECT name, name_vietnamese, description, description_vietnamese, 
                               category, price, embedding 
                        FROM products 
                        WHERE embedding IS NOT NULL
                    """)
                    products = cursor.fetchall()
                except Exception as e:
                    self.logger.error(f"Database query error: {e}")
                    return []
                
                if not products:
                    self.logger.info("No products with embeddings found, trying keyword search")
                    return self.keyword_search_fallback(user_input)
                    
                best_matches = []
                
                for product in products:
                    try:
                        if len(product) < 7:
                            continue
                            
                        stored_embedding_blob = product[6]
                        
                        if stored_embedding_blob:
                            try:
                                stored_embedding = np.frombuffer(stored_embedding_blob, dtype=np.float32)
                                stored_embedding = stored_embedding.reshape(1, -1)
                            except Exception as e:
                                self.logger.warning(f"Embedding decode error: {e}")
                                continue
                            
                            try:
                                similarity = cosine_similarity(query_embedding, stored_embedding)[0][0]
                            except Exception as e:
                                self.logger.warning(f"Similarity calculation error: {e}")
                                continue
                            
                            if similarity > similarity_threshold:
                                best_matches.append({
                                    'name': product[0] or '',
                                    'name_vietnamese': product[1] or '',
                                    'description': product[2] or '',
                                    'description_vietnamese': product[3] or '',
                                    'category': product[4] or '',
                                    'price': product[5],
                                    'similarity': float(similarity)
                                })
                        
                    except Exception as e:
                        self.logger.warning(f"Product processing error: {e}")
                        continue
                            
                try:
                    best_matches.sort(key=lambda x: x.get('similarity', 0), reverse=True)
                    
                    if not best_matches:
                        self.logger.info("No embedding matches found, trying keyword search")
                        return self.keyword_search_fallback(user_input)
                        
                    return best_matches[:3]
                except Exception as e:
                    self.logger.error(f"Sorting error: {e}")
                    return best_matches[:3] if best_matches else []
            else:
                self.logger.info("Embedding model not available, using keyword search")
                return self.keyword_search_fallback(user_input)
                
        except Exception as e:
            self.logger.error(f"Database search error: {e}")
            return []
    
    def keyword_search_fallback(self, user_input):
        """Fallback keyword-based search when embedding search fails or is unavailable"""
        try:
            if not user_input or not user_input.strip():
                return []
            
            keywords = user_input.lower().split()
            
            cursor = self.conn.cursor()
            
            query = """
                SELECT name, name_vietnamese, description, description_vietnamese, 
                       category, price, NULL as embedding
                FROM products 
                WHERE 1=0
            """
            
            conditions = []
            params = []
            
            for keyword in keywords:
                if len(keyword) > 2:
                    keyword_pattern = f'%{keyword}%'
                    conditions.append("""
                        (LOWER(name) LIKE ? OR 
                         LOWER(name_vietnamese) LIKE ? OR 
                         LOWER(description) LIKE ? OR 
                         LOWER(description_vietnamese) LIKE ? OR
                         LOWER(category) LIKE ? OR
                         LOWER(category_vietnamese) LIKE ?)
                    """)
                    params.extend([keyword_pattern] * 6)
            
            if conditions:
                query = f"""
                    SELECT name, name_vietnamese, description, description_vietnamese, 
                           category, price, NULL as embedding
                    FROM products 
                    WHERE {' OR '.join(conditions)}
                    LIMIT 10
                """
                
                cursor.execute(query, params)
                products = cursor.fetchall()
                
                results = []
                for product in products:
                    if len(product) >= 6:
                        results.append({
                            'name': product[0] or '',
                            'name_vietnamese': product[1] or '',
                            'description': product[2] or '',
                            'description_vietnamese': product[3] or '',
                            'category': product[4] or '',
                            'price': product[5],
                            'similarity': 0.5
                        })
                
                self.logger.info(f"Keyword search found {len(results)} results")
                return results[:3]
            
            return []
            
        except Exception as e:
            self.logger.error(f"Keyword search error: {e}")
            return []

    def search_knowledge_base(self, user_input, similarity_threshold=None):
        """Search knowledge base for relevant information"""
        try:
            if not self.embedding_model:
                self.logger.warning("Embedding model not available for knowledge base search")
                return []
            
            if not user_input or not user_input.strip():
                return []
            
            if similarity_threshold is None:
                similarity_threshold = self.config['search_config']['local_similarity_threshold']
            
            try:
                query_embedding = self.embedding_model.encode([user_input.strip()])
                if query_embedding is None or len(query_embedding) == 0:
                    return []
            except Exception as e:
                self.logger.error(f"Knowledge base query embedding error: {e}")
                return []
            
            try:
                cursor = self.conn.cursor()
                cursor.execute("""
                    SELECT topic, content, source, embedding
                    FROM knowledge_base
                    WHERE embedding IS NOT NULL
                """)
                knowledge_entries = cursor.fetchall()
            except Exception as e:
                self.logger.error(f"Knowledge base query error: {e}")
                return []
            
            if not knowledge_entries:
                self.logger.info("No knowledge base entries with embeddings found")
                return []
            
            best_matches = []
            
            for entry in knowledge_entries:
                try:
                    if len(entry) < 4:
                        continue
                    
                    topic, content, source, embedding_blob = entry
                    
                    if embedding_blob:
                        try:
                            stored_embedding = np.frombuffer(embedding_blob, dtype=np.float32)
                            stored_embedding = stored_embedding.reshape(1, -1)
                        except Exception as e:
                            self.logger.warning(f"Knowledge embedding decode error: {e}")
                            continue
                        
                        try:
                            similarity = cosine_similarity(query_embedding, stored_embedding)[0][0]
                        except Exception as e:
                            self.logger.warning(f"Knowledge similarity calculation error: {e}")
                            continue
                        
                        if similarity > similarity_threshold:
                            best_matches.append({
                                'topic': topic or '',
                                'content': content[:500] if content else '',
                                'source': source or '',
                                'similarity': float(similarity),
                                'type': 'knowledge'
                            })
                
                except Exception as e:
                    self.logger.warning(f"Knowledge entry processing error: {e}")
                    continue
            
            try:
                best_matches.sort(key=lambda x: x.get('similarity', 0), reverse=True)
                return best_matches[:3]
            except Exception as e:
                self.logger.error(f"Knowledge sorting error: {e}")
                return best_matches[:3] if best_matches else []
                
        except Exception as e:
            self.logger.error(f"Knowledge base search error: {e}")
            return []
    
    def check_context_references(self, user_input, user_language='vi'):
        """Check if user is referring to previous conversation items"""
        try:
            if user_language == 'vi':
                context_keywords = [
                    'cái đó', 'cái này', 'sản phẩm đó', 'sản phẩm này',
                    'như trên', 'đã nói', 'vừa nói', 'trước đó',
                    'cái thứ', 'cái đầu', 'cái cuối', 'nó',
                    'chúng', 'những cái', 'mấy cái'
                ]
            else:
                context_keywords = [
                    'that', 'this', 'those', 'these',
                    'the one', 'previous', 'mentioned',
                    'above', 'it', 'them', 'the first',
                    'the last', 'the second'
                ]
            
            user_input_lower = user_input.lower()
            
            references = []
            for keyword in context_keywords:
                if keyword in user_input_lower:
                    references.append(keyword)
            
            if self.conversation_context:
                for turn in self.conversation_context[-3:]:
                    if 'products_mentioned' in turn and turn['products_mentioned']:
                        for product in turn['products_mentioned']:
                            if product.lower() in user_input_lower:
                                references.append(f"product:{product}")
            
            return references
            
        except Exception as e:
            self.logger.error(f"Context reference check error: {e}")
            return []
    
    def build_contextual_search_query(self, user_input, context_references):
        """Build search query that includes context from previous conversations"""
        try:
            search_terms = [user_input]
            
            for turn in self.conversation_context[-3:]:
                if 'products_mentioned' in turn and turn['products_mentioned']:
                    for product in turn['products_mentioned']:
                        if product and product not in search_terms:
                            search_terms.append(product)
            
            enhanced_query = ' '.join(search_terms[:3])
            
            self.logger.info(f"Enhanced search query: {enhanced_query}")
            return enhanced_query
            
        except Exception as e:
            self.logger.error(f"Contextual search query error: {e}")
            return user_input
    
    def update_conversation_summary(self):
        """Update conversation summary for long-term context"""
        try:
            if not self.conversation_context:
                return
            
            topics = []
            products_discussed = set()
            
            for turn in self.conversation_context:
                if 'products_mentioned' in turn and turn['products_mentioned']:
                    products_discussed.update(turn['products_mentioned'])
                
                user_msg = turn.get('user', '').lower()
                if 'laptop' in user_msg or 'máy tính' in user_msg:
                    topics.append('laptops')
                if 'gaming' in user_msg or 'game' in user_msg:
                    topics.append('gaming')
                if 'mouse' in user_msg or 'chuột' in user_msg:
                    topics.append('mouse')
                if 'keyboard' in user_msg or 'bàn phím' in user_msg:
                    topics.append('keyboard')
            
            self.conversation_summary = {
                'topics': list(set(topics)),
                'products_discussed': list(products_discussed),
                'turn_count': len(self.conversation_context),
                'last_update': datetime.now().isoformat()
            }
            
            self.logger.info(f"Updated conversation summary: {self.conversation_summary}")
            
        except Exception as e:
            self.logger.error(f"Conversation summary update error: {e}")
    
    def get_conversation_context_prompt(self, user_language='vi'):
        """Get enhanced conversation context for prompts"""
        try:
            context_prompt = ""
            
            if hasattr(self, 'conversation_summary') and self.conversation_summary:
                if user_language == 'vi':
                    context_prompt += "Tóm tắt cuộc trò chuyện:\n"
                    if self.conversation_summary.get('topics'):
                        context_prompt += f"- Chủ đề đã thảo luận: {', '.join(self.conversation_summary['topics'])}\n"
                    if self.conversation_summary.get('products_discussed'):
                        context_prompt += f"- Sản phẩm đã xem: {', '.join(self.conversation_summary['products_discussed'][:5])}\n"
                else:
                    context_prompt += "Conversation summary:\n"
                    if self.conversation_summary.get('topics'):
                        context_prompt += f"- Topics discussed: {', '.join(self.conversation_summary['topics'])}\n"
                    if self.conversation_summary.get('products_discussed'):
                        context_prompt += f"- Products viewed: {', '.join(self.conversation_summary['products_discussed'][:5])}\n"
                context_prompt += "\n"
            
            return context_prompt
            
        except Exception as e:
            self.logger.error(f"Context prompt error: {e}")
            return ""
    
    def store_conversation(self, user_input, response, data_source, processing_time, 
                          user_language, error_type=None, error_message=None):
        """Store conversation in database with language tracking and error logging"""
        try:
            cursor = self.conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT INTO conversations 
                    (timestamp, user_input, user_language, bot_response, bot_language, 
                     data_source, response_time, error_type, error_message)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    datetime.now().isoformat(),
                    user_input,
                    user_language,
                    response,
                    user_language,
                    data_source,
                    processing_time,
                    error_type,
                    error_message
                ))
            except sqlite3.OperationalError as e:
                if "no column named error_type" in str(e):
                    cursor.execute('''
                        INSERT INTO conversations 
                        (timestamp, user_input, user_language, bot_response, bot_language, 
                         data_source, response_time)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        datetime.now().isoformat(),
                        user_input,
                        user_language,
                        response,
                        user_language,
                        data_source,
                        processing_time
                    ))
                else:
                    raise e
            
            self.conn.commit()
        except Exception as e:
            self.logger.error(f"Error storing conversation: {e}")
    
    def process_excel_file(self, file_path):
        """Process Excel files for product data with Vietnamese support"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, encoding='utf-8')
            else:
                df = pd.read_excel(file_path)
                
            self.update_training_status(f"📊 Found {len(df)} rows in Excel file")
                
            column_mapping = {
                'name': ['name', 'tên', 'ten', 'product_name', 'sản phẩm'],
                'name_vietnamese': ['name_vietnamese', 'tên_tiếng_việt', 'ten_tieng_viet'],
                'description': ['description', 'mô tả', 'mo_ta', 'desc'],
                'description_vietnamese': ['description_vietnamese', 'mô_tả_tiếng_việt'],
                'category': ['category', 'danh mục', 'danh_muc', 'loại'],
                'category_vietnamese': ['category_vietnamese', 'danh_mục_tiếng_việt'],
                'price': ['price', 'giá', 'gia', 'cost'],
                'features': ['features', 'tính năng', 'tinh_nang'],
                'features_vietnamese': ['features_vietnamese', 'tính_năng_tiếng_việt'],
                'specifications': ['specifications', 'thông số', 'thong_so', 'specs'],
                'specifications_vietnamese': ['specifications_vietnamese', 'thông_số_tiếng_việt']
            }
            
            mapped_columns = {}
            for standard_col, possible_cols in column_mapping.items():
                for col in possible_cols:
                    if col in df.columns:
                        mapped_columns[standard_col] = col
                        break
            
            if 'name' not in mapped_columns and 'name_vietnamese' not in mapped_columns:
                raise ValueError("Excel file must contain at least a name column (English or Vietnamese)")
                
            cursor = self.conn.cursor()
            added_count = 0
            
            for _, row in df.iterrows():
                try:
                    name = str(row.get(mapped_columns.get('name', ''), '')) if 'name' in mapped_columns else ''
                    name_vietnamese = str(row.get(mapped_columns.get('name_vietnamese', ''), '')) if 'name_vietnamese' in mapped_columns else ''
                    description = str(row.get(mapped_columns.get('description', ''), '')) if 'description' in mapped_columns else ''
                    description_vietnamese = str(row.get(mapped_columns.get('description_vietnamese', ''), '')) if 'description_vietnamese' in mapped_columns else ''
                    category = str(row.get(mapped_columns.get('category', ''), 'General')) if 'category' in mapped_columns else 'General'
                    category_vietnamese = str(row.get(mapped_columns.get('category_vietnamese', ''), '')) if 'category_vietnamese' in mapped_columns else ''
                    features = str(row.get(mapped_columns.get('features', ''), '')) if 'features' in mapped_columns else ''
                    features_vietnamese = str(row.get(mapped_columns.get('features_vietnamese', ''), '')) if 'features_vietnamese' in mapped_columns else ''
                    specifications = str(row.get(mapped_columns.get('specifications', ''), '')) if 'specifications' in mapped_columns else ''
                    specifications_vietnamese = str(row.get(mapped_columns.get('specifications_vietnamese', ''), '')) if 'specifications_vietnamese' in mapped_columns else ''

                    price = 0
                    if 'price' in mapped_columns:
                        try:
                            price_val = row.get(mapped_columns['price'], 0)
                            if pd.notna(price_val):
                                price = float(str(price_val).replace(',', '').replace('$', ''))
                        except (ValueError, TypeError):
                            price = 0
                    
                    if not name and not name_vietnamese:
                        continue
                        
                    product_text_en = f"{name} {description} {features} {specifications}"
                    product_text_vi = f"{name_vietnamese} {description_vietnamese} {features_vietnamese} {specifications_vietnamese}"
                    combined_text = f"{product_text_en} {product_text_vi}".strip()
                    
                    embedding_blob = None
                    if self.embedding_model and combined_text:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.update_training_status(f"⚠️ Embedding generation failed for {name or name_vietnamese}: {e}")
                    
                    cursor.execute('''
                        INSERT INTO products 
                        (name, name_vietnamese, description, description_vietnamese, 
                         category, category_vietnamese, price, features, features_vietnamese,
                         specifications, specifications_vietnamese, source_file, embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        name or name_vietnamese,
                        name_vietnamese,
                        description or description_vietnamese,
                        description_vietnamese,
                        category,
                        category_vietnamese,
                        price,
                        features,
                        features_vietnamese,
                        specifications,
                        specifications_vietnamese,
                        file_path,
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                    added_count += 1
                    
                except Exception as row_error:
                    self.update_training_status(f"⚠️ Error processing row: {row_error}")
                    continue
                
            self.conn.commit()
            self.update_training_status(f"✅ Added {added_count} products from Excel file with embeddings")
            
        except Exception as e:
            raise Exception(f"Error processing Excel file: {e}")
            
    def process_pdf_file(self, file_path):
        """Process PDF files for knowledge base with Vietnamese support"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    self.update_training_status(f"📄 Processed page {page_num + 1}/{len(reader.pages)}")
                except Exception as e:
                    self.update_training_status(f"⚠️ Error extracting page {page_num + 1}: {e}")
                    continue
                
            if text.strip():
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from PDF: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from PDF")
                
        except Exception as e:
            raise Exception(f"Error processing PDF file: {e}")
            
    def process_word_file(self, file_path):
        """Process Word documents for knowledge base with Vietnamese support"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                if para_num % 10 == 0:
                    self.update_training_status(f"📝 Processed {para_num} paragraphs...")
                    
            for table_num, table in enumerate(doc.tables):
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                self.update_training_status(f"📊 Processed table {table_num + 1}/{len(doc.tables)}")
                
            if text.strip():
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from Word doc: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from Word document")
                
        except Exception as e:
            raise Exception(f"Error processing Word file: {e}")
            
    def process_image_file(self, file_path):
        """Process image files using OCR with Vietnamese support"""
        try:
            if not self.ocr_reader:
                raise Exception("OCR model not available")
                
            self.update_training_status(f"🖼️ Running OCR on image...")
            
            results = self.ocr_reader.readtext(file_path)
            
            extracted_texts = []
            for (bbox, text, confidence) in results:
                if confidence > 0.5:
                    extracted_texts.append(text.strip())
                    
            if extracted_texts:
                combined_text = " ".join(extracted_texts)
                
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding = self.embedding_model.encode([combined_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    f"OCR: {os.path.basename(file_path)}",
                    combined_text,
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Extracted {len(extracted_texts)} text segments from image")
                self.update_training_status(f"📝 Text preview: {combined_text[:100]}...")
            else:
                self.update_training_status(f"⚠️ No text detected in image")
                
        except Exception as e:
            raise Exception(f"Error processing image file: {e}")
    
    def setup_gui(self):
        """Setup the GUI interface with Vietnamese font support"""
        self.root = tk.Tk()
        self.root.title("Trợ lý AI Bán hàng - Vietnamese AI Sales ChatBot")
        self.root.geometry("1400x900")
        
        self.setup_vietnamese_fonts()
        
        self.setup_main_interface()
        
    def setup_vietnamese_fonts(self):
        """Setup fonts that support Vietnamese characters with enhanced input support"""
        vietnamese_fonts = [
            ('Segoe UI', 12),
            ('Times New Roman', 12),
            ('Arial Unicode MS', 12),
            ('Calibri', 12),
            ('Tahoma', 12),
            ('Microsoft Sans Serif', 12),
            ('DejaVu Sans', 12),
            ('Liberation Sans', 12),
            ('Verdana', 11)
        ]
        
        self.vietnamese_font = None
        self.input_font = None
        available_fonts = font.families()
        
        for font_name, size in vietnamese_fonts:
            if font_name in available_fonts:
                try:
                    test_font = font.Font(family=font_name, size=size)
                    
                    self.vietnamese_font = font.Font(family=font_name, size=size-1)
                    self.input_font = font.Font(
                        family=font_name, 
                        size=size, 
                        weight='normal'
                    )
                    
                    print(f"✅ Using Vietnamese font: {font_name}")
                    break
                except Exception as e:
                    print(f"⚠️ Font {font_name} failed: {e}")
                    continue
        
        if not self.vietnamese_font:
            self.vietnamese_font = font.Font(family="TkDefaultFont", size=11)
            self.input_font = font.Font(family="TkDefaultFont", size=12)
            print("⚠️ Using default font (Vietnamese characters may not display correctly)")
        
        self.root.option_add('*Font', self.vietnamese_font)
        
        try:
            self.root.tk.call('encoding', 'system', 'utf-8')
            print("✅ System encoding set to UTF-8 for Vietnamese support")
        except Exception as e:
            print(f"⚠️ Could not set UTF-8 encoding: {e}")
    
    def setup_main_interface(self):
        """Setup the main GUI interface with notebook tabs"""
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.setup_chat_tab()
        self.setup_database_tab()
        self.setup_training_tab()
        self.setup_analytics_tab()
    
    def setup_chat_tab(self):
        """Setup the main chat interface"""
        chat_frame = ttk.Frame(self.notebook)
        self.notebook.add(chat_frame, text="💬 Chat")
        
        top_frame = ttk.Frame(chat_frame)
        top_frame.pack(fill=tk.X, padx=10, pady=5)
        
        lang_frame = ttk.Frame(top_frame)
        lang_frame.pack(side=tk.LEFT)
        
        ttk.Label(lang_frame, text="Ngôn ngữ / Language:", font=self.vietnamese_font).pack(side=tk.LEFT, padx=5)
        
        self.language_var = tk.StringVar(value=self.config['language_config']['default_language'])
        language_combo = ttk.Combobox(lang_frame, textvariable=self.language_var, 
                                     values=['vi', 'en'], state='readonly', width=8)
        language_combo.pack(side=tk.LEFT, padx=5)
        language_combo.bind('<<ComboboxSelected>>', self.on_language_change)
        
        self.chat_display = scrolledtext.ScrolledText(
            chat_frame, 
            height=30, 
            state=tk.DISABLED,
            font=self.vietnamese_font,
            wrap=tk.WORD
        )
        self.chat_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.chat_display.tag_configure("user", foreground="blue", font=(self.vietnamese_font['family'], 10, 'bold'))
        self.chat_display.tag_configure("bot", foreground="green", font=(self.vietnamese_font['family'], 10))
        self.chat_display.tag_configure("system", foreground="gray", font=(self.vietnamese_font['family'], 9, 'italic'))
        self.chat_display.tag_configure("error", foreground="red", font=(self.vietnamese_font['family'], 9, 'italic'))
        
        input_frame = ttk.Frame(chat_frame)
        input_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.user_input = tk.Text(
            input_frame, 
            height=3, 
            font=self.input_font,
            wrap=tk.WORD
        )
        self.user_input.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        self.user_input.bind('<Control-Return>', self.send_message)
        self.user_input.bind('<Return>', self.on_enter_key)
        
        button_frame = ttk.Frame(input_frame)
        button_frame.pack(side=tk.RIGHT, fill=tk.Y)
        
        send_btn = ttk.Button(button_frame, text="Gửi / Send", command=self.send_message)
        send_btn.pack(fill=tk.X, pady=(0, 5))
        
        clear_btn = ttk.Button(button_frame, text="Xóa / Clear", command=self.clear_chat)
        clear_btn.pack(fill=tk.X)
        
        status_frame = ttk.Frame(chat_frame)
        status_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.status_label = ttk.Label(status_frame, text="Sẵn sàng / Ready", font=self.vietnamese_font)
        self.status_label.pack(side=tk.LEFT)
        
        self.processing_label = ttk.Label(status_frame, text="", font=self.vietnamese_font)
        self.processing_label.pack(side=tk.RIGHT)
    
    def setup_database_tab(self):
        """Setup database management interface"""
        db_frame = ttk.Frame(self.notebook)
        self.notebook.add(db_frame, text="🗄️ Database")
        
        product_frame = ttk.LabelFrame(db_frame, text="Product Database Management")
        product_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        btn_frame = ttk.Frame(product_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(btn_frame, text="Add Sample Products", 
                  command=self.add_sample_products).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="View Products", 
                  command=self.view_products).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Search Test", 
                  command=self.test_search).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Regenerate Embeddings", 
                  command=self.regenerate_all_embeddings).pack(side=tk.LEFT, padx=5)
        
        self.product_display = scrolledtext.ScrolledText(
            product_frame, height=15, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.product_display.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
    
    def setup_training_tab(self):
        """Setup training interface with file upload capabilities"""
        training_frame = ttk.Frame(self.notebook)
        self.notebook.add(training_frame, text="📚 Training Data")
        
        upload_frame = ttk.LabelFrame(training_frame, text="Upload Training Data")
        upload_frame.pack(fill=tk.X, padx=10, pady=5)
        
        upload_btn_frame = ttk.Frame(upload_frame)
        upload_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(upload_btn_frame, text="📊 Upload Excel Files", 
                  command=lambda: self.process_files('excel')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📄 Upload PDF Files", 
                  command=lambda: self.process_files('pdf')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📝 Upload Word Files", 
                  command=lambda: self.process_files('word')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="🖼️ Upload Images (OCR)", 
                  command=lambda: self.process_files('image')).pack(side=tk.LEFT, padx=5, pady=5)
        
        options_frame = ttk.LabelFrame(training_frame, text="Training Options")
        options_frame.pack(fill=tk.X, padx=10, pady=5)
        
        options_btn_frame = ttk.Frame(options_frame)
        options_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(options_btn_frame, text="🔄 Regenerate All Embeddings", 
                  command=self.regenerate_all_embeddings).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="🧹 Clear Training Data", 
                  command=self.clear_training_data).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="📈 Training Statistics", 
                  command=self.show_training_stats).pack(side=tk.LEFT, padx=5)
        
        self.training_status = scrolledtext.ScrolledText(
            training_frame, height=20, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.training_status.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def setup_analytics_tab(self):
        """Setup analytics interface"""
        analytics_frame = ttk.Frame(self.notebook)
        self.notebook.add(analytics_frame, text="📊 Analytics")
        
        metrics_frame = ttk.LabelFrame(analytics_frame, text="Performance Metrics")
        metrics_frame.pack(fill=tk.X, padx=10, pady=5)
        
        metrics_btn_frame = ttk.Frame(metrics_frame)
        metrics_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(metrics_btn_frame, text="🔄 Refresh Analytics", 
                  command=self.refresh_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📤 Export Data", 
                  command=self.export_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📋 Conversation History", 
                  command=self.view_conversation_history).pack(side=tk.LEFT, padx=5, pady=5)
        
        self.analytics_display = scrolledtext.ScrolledText(
            analytics_frame, height=25, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.analytics_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def on_enter_key(self, event):
        """Handle Enter key press"""
        if event.state & 0x4:
            self.send_message()
        else:
            return None
    
    def on_language_change(self, event=None):
        """Handle language change"""
        try:
            self.current_language = self.language_var.get()
            print(f"Language changed to: {self.current_language}")
        except Exception as e:
            print(f"Language change error: {e}")
    
    def send_message(self, event=None):
        """Handle sending user message with error handling"""
        try:
            user_text = self.user_input.get(1.0, tk.END).strip()
            if not user_text:
                return
                
            self.user_input.delete(1.0, tk.END)
            
            self.display_message("Bạn / You", user_text, "user")
            
            if hasattr(self, 'processing_label'):
                status_text = "Đang xử lý..." if self.current_language == 'vi' else "Processing..."
                self.processing_label.config(text=status_text)
                self.root.update()
            
            threading.Thread(target=self.process_message_thread, args=(user_text,), daemon=True).start()
        except Exception as e:
            print(f"Send message error: {e}")
            error_msg = f"Lỗi gửi tin nhắn / Send error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
    
    def process_message_thread(self, user_input):
        """Process message in separate thread with comprehensive error handling"""
        try:
            result = self.process_message(user_input)
            
            response_tag = "error" if result.get('data_source') == 'error' else "bot"
            self.display_message("Trợ lý AI / AI Assistant", result.get('response', ''), response_tag)
            
            user_language = result.get('user_language', 'vi')
            if user_language == 'vi':
                if result.get('data_source') == 'local_database':
                    status_text = f"✅ Tìm thấy {result.get('local_results_count', 0)} sản phẩm ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'knowledge_base':
                    status_text = f"📚 Tìm thấy trong cơ sở tri thức ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'error':
                    status_text = f"❌ Lỗi ({result.get('processing_time', 0):.1f}s)"
                else:
                    status_text = f"💬 Hoàn thành ({result.get('processing_time', 0):.1f}s)"
            else:
                if result.get('data_source') == 'local_database':
                    status_text = f"✅ Found {result.get('local_results_count', 0)} products ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'knowledge_base':
                    status_text = f"📚 Found in knowledge base ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'error':
                    status_text = f"❌ Error ({result.get('processing_time', 0):.1f}s)"
                else:
                    status_text = f"💬 Complete ({result.get('processing_time', 0):.1f}s)"
            
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text=status_text))
            
        except AttributeError as ae:
            self.logger.error(f"AttributeError in message thread: {ae}")
            error_msg = f"Method not found error: {str(ae)}"
            self.display_message("Hệ thống / System", error_msg, "error")
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text="Method Error"))
        except Exception as e:
            self.logger.error(f"Message thread error: {e}")
            error_msg = f"Lỗi xử lý / Processing error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text="Lỗi / Error"))
    
    def display_message(self, sender, message, tag):
        """Display message in chat window with error handling"""
        def update_display():
            try:
                self.chat_display.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime("%H:%M:%S")
                
                self.chat_display.insert(tk.END, f"[{timestamp}] {sender}:\n", tag)
                self.chat_display.insert(tk.END, f"{message}\n\n")
                
                self.chat_display.config(state=tk.DISABLED)
                self.chat_display.see(tk.END)
            except Exception as e:
                print(f"Display update error: {e}")
            
        try:
            if threading.current_thread() != threading.main_thread():
                self.root.after(0, update_display)
            else:
                update_display()
        except Exception as e:
            print(f"Display message error: {e}")
    
    def clear_chat(self):
        """Clear chat display with error handling"""
        try:
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)
            self.chat_display.config(state=tk.DISABLED)
            
            self.conversation_context.clear()
            self.conversation_summary = ""
            
            self.logger.info("Chat history and context cleared")
            
        except Exception as e:
            print(f"Clear chat error: {e}")
    
    def update_training_status(self, message):
        """Update training status display with Vietnamese support"""
        def update():
            if hasattr(self, 'training_status'):
                self.training_status.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime('%H:%M:%S')
                self.training_status.insert(tk.END, f"[{timestamp}] {message}\n")
                self.training_status.config(state=tk.DISABLED)
                self.training_status.see(tk.END)
            
        if threading.current_thread() != threading.main_thread():
            self.root.after(0, update)
        else:
            update()
    
    def process_files(self, file_type):
        """Process different types of files for training data"""
        file_types = {
            'excel': [('Excel files', '*.xlsx *.xls *.csv')],
            'pdf': [('PDF files', '*.pdf')],
            'word': [('Word files', '*.docx *.doc')],
            'image': [('Image files', '*.png *.jpg *.jpeg *.bmp *.tiff')]
        }
        
        files = filedialog.askopenfilenames(
            title=f"Select {file_type} files",
            filetypes=file_types[file_type]
        )
        
        if files:
            self.update_training_status(f"📂 Starting to process {len(files)} {file_type} file(s)...")
            threading.Thread(
                target=self.process_files_worker,
                args=(files, file_type),
                daemon=True
            ).start()
            
    def process_files_worker(self, files, file_type):
        """Worker thread for processing files with Vietnamese support"""
        for file_path in files:
            try:
                self.update_training_status(f"🔄 Processing: {os.path.basename(file_path)}")
                
                if file_type == 'excel':
                    self.process_excel_file(file_path)
                elif file_type == 'pdf':
                    self.process_pdf_file(file_path)
                elif file_type == 'word':
                    self.process_word_file(file_path)
                elif file_type == 'image':
                    self.process_image_file(file_path)
                    
                self.update_training_status(f"✅ Completed: {os.path.basename(file_path)}")
                
            except Exception as e:
                self.update_training_status(f"❌ Error processing {os.path.basename(file_path)}: {e}")
    
    def add_sample_products(self):
        """Add Vietnamese sample products with error handling"""
        sample_products = [
            {
                'name': 'Gaming Laptop Pro X1',
                'name_vietnamese': 'Laptop Gaming Pro X1',
                'description': 'High-performance gaming laptop with RTX 4070, 16GB RAM, 1TB SSD, perfect for gaming and professional work',
                'description_vietnamese': 'Laptop gaming hiệu suất cao với RTX 4070, RAM 16GB, SSD 1TB, hoàn hảo cho gaming và công việc chuyên nghiệp',
                'category': 'Laptops',
                'category_vietnamese': 'Laptop',
                'price': 1899.99,
                'features': 'RTX 4070, Intel i7, 16GB RAM, 1TB SSD, 15.6" 144Hz display',
                'features_vietnamese': 'RTX 4070, Intel i7, RAM 16GB, SSD 1TB, màn hình 15.6" 144Hz'
            },
            {
                'name': 'Wireless Gaming Mouse RGB Pro',
                'name_vietnamese': 'Chuột Gaming Không dây RGB Pro',
                'description': 'High-precision wireless gaming mouse with RGB lighting, 12000 DPI, programmable buttons',
                'description_vietnamese': 'Chuột gaming không dây độ chính xác cao với đèn RGB, 12000 DPI, nút bấm có thể lập trình',
                'category': 'Gaming Accessories',
                'category_vietnamese': 'Phụ kiện Gaming',
                'price': 79.99,
                'features': '12000 DPI, RGB lighting, 7 programmable buttons, 70-hour battery',
                'features_vietnamese': '12000 DPI, đèn RGB, 7 nút lập trình, pin 70 giờ'
            },
            {
                'name': 'Mechanical Gaming Keyboard',
                'name_vietnamese': 'Bàn phím Gaming Cơ',
                'description': 'Premium mechanical gaming keyboard with RGB backlighting and blue switches',
                'description_vietnamese': 'Bàn phím gaming cơ cao cấp với đèn nền RGB và switch xanh',
                'category': 'Gaming Accessories',
                'category_vietnamese': 'Phụ kiện Gaming',
                'price': 129.99,
                'features': 'Blue mechanical switches, RGB per-key lighting, aluminum frame',
                'features_vietnamese': 'Switch cơ xanh, đèn RGB từng phím, khung nhôm'
            },
            {
                'name': 'Gaming Headset 7.1 Surround',
                'name_vietnamese': 'Tai nghe Gaming 7.1 Surround',
                'description': '7.1 surround sound gaming headset with noise-canceling microphone',
                'description_vietnamese': 'Tai nghe gaming âm thanh vòm 7.1 với microphone chống ồn',
                'category': 'Audio',
                'category_vietnamese': 'Âm thanh',
                'price': 89.99,
                'features': '7.1 surround sound, noise-canceling mic, comfortable padding',
                'features_vietnamese': 'Âm thanh vòm 7.1, mic chống ồn, đệm êm ái'
            },
            {
                'name': 'Business Laptop UltraSlim',
                'name_vietnamese': 'Laptop Văn phòng UltraSlim',
                'description': 'Lightweight business laptop with long battery life, perfect for professionals',
                'description_vietnamese': 'Laptop văn phòng siêu mỏng với thời lượng pin dài, hoàn hảo cho dân văn phòng',
                'category': 'Laptops',
                'category_vietnamese': 'Laptop',
                'price': 899.99,
                'features': 'Intel i5, 8GB RAM, 512GB SSD, 14" display, 12-hour battery',
                'features_vietnamese': 'Intel i5, RAM 8GB, SSD 512GB, màn hình 14", pin 12 giờ'
            }
        ]
        
        try:
            cursor = self.conn.cursor()
            for product in sample_products:
                try:
                    english_text = f"{product['name']} {product['description']} {product.get('features', '')}"
                    vietnamese_text = f"{product['name_vietnamese']} {product['description_vietnamese']} {product.get('features_vietnamese', '')}"
                    combined_text = f"{english_text} {vietnamese_text}"
                    
                    embedding_blob = None
                    if self.embedding_model:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.logger.warning(f"Embedding generation error for {product['name']}: {e}")
                    
                    cursor.execute('''
                        INSERT OR IGNORE INTO products 
                        (name, name_vietnamese, description, description_vietnamese, 
                         category, category_vietnamese, price, features, features_vietnamese,
                         embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        product['name'], product['name_vietnamese'],
                        product['description'], product['description_vietnamese'],
                        product['category'], product['category_vietnamese'],
                        product['price'], product.get('features', ''), product.get('features_vietnamese', ''),
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                except Exception as e:
                    self.logger.error(f"Error inserting product {product.get('name', 'unknown')}: {e}")
                    continue
            
            self.conn.commit()
            print("✅ Sample products added successfully")
            
        except Exception as e:
            self.logger.error(f"Error adding sample products: {e}")
    
    def view_products(self):
        """View products in database with Vietnamese support"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT name, name_vietnamese, description, description_vietnamese, 
                       category, price, source_file, embedding 
                FROM products 
                ORDER BY created_at DESC
            """)
            products = cursor.fetchall()
            
            self.product_display.config(state=tk.NORMAL)
            self.product_display.delete(1.0, tk.END)
            
            if products:
                self.product_display.insert(tk.END, f"📦 Found {len(products)} products in database:\n\n")
                
                for i, product in enumerate(products, 1):
                    name = product[0] or product[1] or "Unnamed Product"
                    name_vi = product[1] or ""
                    desc = product[2] or product[3] or "No description"
                    category = product[4] or "General"
                    price = product[5] or 0
                    source = os.path.basename(product[6]) if product[6] else "Manual"
                    embedding_status = "✅" if product[7] else "❌"
                    
                    self.product_display.insert(tk.END, f"{i}. {name} {embedding_status}\n")
                    if name_vi:
                        self.product_display.insert(tk.END, f"   Vietnamese: {name_vi}\n")
                    self.product_display.insert(tk.END, f"   Description: {desc[:100]}...\n")
                    self.product_display.insert(tk.END, f"   Category: {category} | Price: ${price} | Source: {source}\n\n")
            else:
                self.product_display.insert(tk.END, "📦 No products found in database.\n\n")
                self.product_display.insert(tk.END, "💡 To add products:\n")
                self.product_display.insert(tk.END, "• Click 'Add Sample Products' for demo data\n")
                self.product_display.insert(tk.END, "• Use the Training tab to upload Excel/CSV files\n")
                self.product_display.insert(tk.END, "• Use PDFs/Word docs for knowledge base\n")
                
            self.product_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing products: {e}")
    
    def test_search(self):
        """Test search functionality with Vietnamese support"""
        test_queries = [
            "gaming laptop", "laptop gaming",
            "business computer", "máy tính văn phòng", 
            "wireless mouse", "chuột không dây",
            "cheap laptop under $1500", "laptop rẻ dưới 1500 đô",
            "RGB accessories", "phụ kiện RGB",
            "4K monitor", "màn hình 4K",
            "mechanical keyboard", "bàn phím cơ",
            "laptop for programming", "laptop lập trình"
        ]
        
        self.product_display.config(state=tk.NORMAL)
        self.product_display.delete(1.0, tk.END)
        self.product_display.insert(tk.END, "🔍 Testing search functionality...\n")
        self.product_display.insert(tk.END, f"Threshold: {self.config['search_config']['local_similarity_threshold']}\n\n")
        
        for query in test_queries:
            self.product_display.insert(tk.END, f"Query: '{query}'\n")
            results = self.search_local_database(query)
            
            if results:
                for result in results[:3]:
                    self.product_display.insert(tk.END, f"  ✓ {result['name']} (similarity: {result['similarity']:.3f})\n")
            else:
                self.product_display.insert(tk.END, "  ✗ No results found\n")
            
            self.product_display.insert(tk.END, "\n")
            
        self.product_display.config(state=tk.DISABLED)
    
    def regenerate_all_embeddings(self):
        """Regenerate all embeddings for products and knowledge base"""
        result = messagebox.askyesno(
            "Regenerate Embeddings", 
            "This will regenerate all embeddings for products and knowledge base entries.\n\n"
            "This may take several minutes depending on the amount of data.\n\n"
            "Continue?"
        )
        
        if result:
            threading.Thread(target=self.regenerate_embeddings_worker, daemon=True).start()
            
    def regenerate_embeddings_worker(self):
        """Worker thread for regenerating embeddings"""
        try:
            if not self.embedding_model:
                self.update_training_status("❌ Embedding model not available")
                return
                
            cursor = self.conn.cursor()
            
            self.update_training_status("🔄 Regenerating product embeddings...")
            cursor.execute("""
                SELECT id, name, name_vietnamese, description, description_vietnamese, 
                       features, features_vietnamese, specifications, specifications_vietnamese 
                FROM products
            """)
            products = cursor.fetchall()
            
            for i, (product_id, name, name_vi, desc, desc_vi, feat, feat_vi, spec, spec_vi) in enumerate(products):
                try:
                    text_parts = [name or '', name_vi or '', desc or '', desc_vi or '', 
                                feat or '', feat_vi or '', spec or '', spec_vi or '']
                    combined_text = ' '.join(part for part in text_parts if part.strip())
                    
                    if combined_text.strip():
                        embedding = self.embedding_model.encode([combined_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                        
                        cursor.execute("""
                            UPDATE products 
                            SET embedding = ?, updated_at = ?
                            WHERE id = ?
                        """, (embedding_blob, datetime.now().isoformat(), product_id))
                        
                        if (i + 1) % 10 == 0:
                            self.update_training_status(f"📦 Processed {i + 1}/{len(products)} products...")
                            
                except Exception as e:
                    self.update_training_status(f"⚠️ Error processing product {name}: {e}")
                    continue
                    
            self.update_training_status("🔄 Regenerating knowledge base embeddings...")
            cursor.execute("SELECT id, content FROM knowledge_base")
            knowledge_entries = cursor.fetchall()
            
            for i, (kb_id, content) in enumerate(knowledge_entries):
                try:
                    if content and content.strip():
                        embedding = self.embedding_model.encode([content[:5000]])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                        
                        cursor.execute("""
                            UPDATE knowledge_base 
                            SET embedding = ?
                            WHERE id = ?
                        """, (embedding_blob, kb_id))
                        
                        if (i + 1) % 5 == 0:
                            self.update_training_status(f"📚 Processed {i + 1}/{len(knowledge_entries)} knowledge entries...")
                            
                except Exception as e:
                    self.update_training_status(f"⚠️ Error processing knowledge entry {kb_id}: {e}")
                    continue
                    
            self.conn.commit()
            self.update_training_status(f"✅ Successfully regenerated embeddings for {len(products)} products and {len(knowledge_entries)} knowledge entries!")
            
        except Exception as e:
            self.update_training_status(f"❌ Error regenerating embeddings: {e}")
    
    def clear_training_data(self):
        """Clear all training data with confirmation"""
        result = messagebox.askyesno(
            "Clear Training Data", 
            "Are you sure you want to clear all training data?\n\n"
            "This will delete:\n"
            "• All products\n"
            "• All knowledge base entries\n"
            "• All embeddings\n\n"
            "This action cannot be undone!"
        )
        
        if result:
            try:
                cursor = self.conn.cursor()
                cursor.execute("DELETE FROM products")
                cursor.execute("DELETE FROM knowledge_base")
                self.conn.commit()
                
                self.update_training_status("🧹 All training data cleared successfully")
                messagebox.showinfo("Success", "All training data has been cleared.")
                
            except Exception as e:
                self.update_training_status(f"❌ Error clearing data: {e}")
                messagebox.showerror("Error", f"Error clearing training data: {e}")
    
    def show_training_stats(self):
        """Show training data statistics"""
        try:
            cursor = self.conn.cursor()
            
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(DISTINCT category) FROM products")
            category_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base WHERE embedding IS NOT NULL")
            knowledge_with_embeddings = cursor.fetchone()[0]
            
            self.training_status.config(state=tk.NORMAL)
            self.training_status.delete(1.0, tk.END)
            
            self.training_status.insert(tk.END, "=== TRAINING DATA STATISTICS ===\n\n")
            self.training_status.insert(tk.END, f"📦 Products: {product_count}\n")
            self.training_status.insert(tk.END, f"🔗 Products with embeddings: {products_with_embeddings}/{product_count}\n")
            self.training_status.insert(tk.END, f"📂 Categories: {category_count}\n\n")
            self.training_status.insert(tk.END, f"📚 Knowledge base entries: {knowledge_count}\n")
            self.training_status.insert(tk.END, f"🔗 Knowledge with embeddings: {knowledge_with_embeddings}/{knowledge_count}\n\n")
            
            cursor.execute("""
                SELECT name, created_at, source_file 
                FROM products 
                ORDER BY created_at DESC 
                LIMIT 5
            """)
            recent_products = cursor.fetchall()
            
            if recent_products:
                self.training_status.insert(tk.END, "📈 Recent Products:\n")
                for name, created_at, source_file in recent_products:
                    source = os.path.basename(source_file) if source_file else "Manual"
                    self.training_status.insert(tk.END, f"  • {name} ({source})\n")
                    
            self.training_status.config(state=tk.DISABLED)
            
        except Exception as e:
            self.update_training_status(f"❌ Error showing statistics: {e}")
    
    def refresh_analytics(self):
        """Refresh analytics display with Vietnamese language support"""
        try:
            cursor = self.conn.cursor()
            
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_conversations,
                    AVG(response_time) as avg_response_time,
                    data_source,
                    COUNT(*) as count_by_source
                FROM conversations 
                GROUP BY data_source
            """)
            
            source_stats = cursor.fetchall()
            
            cursor.execute("""
                SELECT timestamp, user_input, user_language, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 10
            """)
            
            recent_convs = cursor.fetchall()
            
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            cursor.execute("""
                SELECT user_language, COUNT(*) as count
                FROM conversations 
                WHERE user_language IS NOT NULL
                GROUP BY user_language
            """)
            language_stats = cursor.fetchall()
            
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            self.analytics_display.insert(tk.END, "=== VIETNAMESE AI CHATBOT ANALYTICS ===\n\n")
            self.analytics_display.insert(tk.END, f"📦 Products in Database: {product_count}\n")
            self.analytics_display.insert(tk.END, f"🔗 Products with Embeddings: {products_with_embeddings}/{product_count}\n")
            self.analytics_display.insert(tk.END, f"📚 Knowledge Base Entries: {knowledge_count}\n")
            
            self.analytics_display.insert(tk.END, f"\n=== LANGUAGE USAGE ===\n")
            for lang, count in language_stats:
                lang_name = "Vietnamese" if lang == 'vi' else "English" if lang == 'en' else lang
                self.analytics_display.insert(tk.END, f"{lang_name}: {count} conversations\n")
            
            self.analytics_display.insert(tk.END, "\n=== RESPONSE SOURCES ===\n")
            for stat in source_stats:
                if len(stat) >= 4:
                    self.analytics_display.insert(tk.END, f"{stat[2]}: {stat[3]} responses\n")
                
            if source_stats:
                total_conversations = sum(stat[3] for stat in source_stats if len(stat) >= 4)
                if total_conversations > 0:
                    avg_response_time = sum(stat[1] * stat[3] for stat in source_stats if stat[1] and len(stat) >= 4) / total_conversations
                    self.analytics_display.insert(tk.END, f"\nTotal Conversations: {total_conversations}\n")
                    self.analytics_display.insert(tk.END, f"Average Response Time: {avg_response_time:.2f} seconds\n")
            
            self.analytics_display.insert(tk.END, "\n=== RECENT CONVERSATIONS ===\n")
            for conv in recent_convs:
                if len(conv) >= 5:
                    timestamp, user_input, user_lang, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    self.analytics_display.insert(tk.END, f"{timestamp} {lang_flag} - {data_source} ({response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input[:100]}...\n\n")
                
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error refreshing analytics: {e}")
    
    def export_analytics(self):
        """Export analytics data to CSV with Vietnamese support"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT * FROM conversations")
            
            column_names = [description[0] for description in cursor.description]
            
            df = pd.DataFrame(cursor.fetchall(), columns=column_names)
            
            export_path = f"vietnamese_chatbot_analytics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            df.to_csv(export_path, index=False, encoding='utf-8')
            
            messagebox.showinfo("Success", f"Analytics exported to: {export_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error exporting analytics: {e}")
    
    def view_conversation_history(self):
        """View detailed conversation history"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT timestamp, user_input, user_language, bot_response, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 20
            """)
            conversations = cursor.fetchall()
            
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            self.analytics_display.insert(tk.END, "=== CONVERSATION HISTORY (Last 20) ===\n\n")
            
            for i, conv in enumerate(conversations, 1):
                if len(conv) >= 6:
                    timestamp, user_input, user_lang, bot_response, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    
                    self.analytics_display.insert(tk.END, f"{i}. [{timestamp}] {lang_flag} ({data_source}, {response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input}\n")
                    self.analytics_display.insert(tk.END, f"Bot: {bot_response[:200]}...\n\n")
            
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing conversation history: {e}")
    
    def run(self):
        """Start the application with comprehensive error handling"""
        try:
            print("📦 Adding sample products...")
            self.add_sample_products()
            
            self.display_welcome_message()
            
            print("🚀 Starting GUI...")
            self.root.mainloop()
            
        except Exception as e:
            print(f"❌ Error starting application: {e}")
            self.logger.error(f"Application start error: {e}")
    
    def display_welcome_message(self):
        """Display welcome message with system status"""
        try:
            if self.current_language == 'vi':
                welcome_msg = f"""🤖 Chào mừng đến với Trợ lý AI Bán hàng!

Trạng thái hệ thống:
✅ Hỗ trợ tiếng Việt và tiếng Anh
{'✅' if self.text_generator else '❌'} Mô hình AI: {'Hoạt động' if self.text_generator else 'Không khả dụng'}
{'✅' if self.embedding_model else '❌'} Tìm kiếm thông minh: {'Hoạt động' if self.embedding_model else 'Không khả dụng'}

Tính năng có sẵn:
📦 Tìm kiếm sản phẩm thông minh
💬 Ngữ cảnh cuộc trò chuyện
🌐 Hỗ trợ đa ngôn ngữ
📚 Xử lý file Excel, PDF, Word
🖼️ OCR từ hình ảnh

Hãy thử đặt câu hỏi như:
• "Tôi muốn mua laptop gaming"
• "Bạn có chuột không dây nào không?"
• "Giá của laptop là bao nhiêu?"
• "So sánh các sản phẩm gaming"

📝 Sử dụng bàn phím để nhập tin nhắn
{f'🔧 Một số tính năng AI có thể bị hạn chế do lỗi tải mô hình' if not self.text_generator else ''}"""
            else:
                welcome_msg = f"""🤖 Welcome to the AI Sales Assistant!

System Status:
✅ Vietnamese and English support
{'✅' if self.text_generator else '❌'} AI Model: {'Available' if self.text_generator else 'Not available'}
{'✅' if self.embedding_model else '❌'} Smart Search: {'Available' if self.embedding_model else 'Not available'}

Available features:
📦 Smart product search
💬 Conversation context
🌐 Multi-language support
📚 Excel, PDF, Word processing
🖼️ Image OCR

Try asking questions like:
• "I want to buy a gaming laptop"
• "What wireless mice do you have?"
• "How much does the laptop cost?"
• "Compare gaming products"

📝 Use keyboard to type messages
{f'🔧 Some AI features may be limited due to model loading errors' if not self.text_generator else ''}"""

            self.display_message("Hệ thống / System", welcome_msg, "system")
        except Exception as e:
            print(f"Welcome message error: {e}")
        
    def __del__(self):
        """Cleanup with error handling"""
        try:
            if hasattr(self, 'conn'):
                self.conn.close()
        except:
            pass

# COMPLETE ENHANCED WebChatInterface CLASS - WITH BEAUTIFUL FEMALE SALESPERSON

class WebChatInterface:
    """Web interface for the chatbot with personalized addressing and beautiful salesperson"""
    
    def __init__(self, chatbot_instance, host='0.0.0.0', port=5000):
        self.chatbot = chatbot_instance
        self.host = host
        self.port = port
        self.app = Flask(__name__)
        # Store user sessions for personalized addressing
        self.user_sessions = {}
        self.setup_routes()
    
    def get_addressing_terms(self, gender, age):
        """Get appropriate addressing terms based on gender and age"""
        if age < 25:
            return {
                'customer_address': 'Bạn',  # Call customer "You"
                'bot_address': 'tôi',       # Bot uses "I"
                'customer_address_en': 'You',
                'bot_address_en': 'I'
            }
        elif 25 <= age <= 45:
            if gender.lower() == 'male':
                return {
                    'customer_address': 'Anh',  # Call customer "Anh"
                    'bot_address': 'em',        # Bot uses "Em"
                    'customer_address_en': 'Sir',
                    'bot_address_en': 'I'
                }
            else:  # female
                return {
                    'customer_address': 'Chị',  # Call customer "Chi"
                    'bot_address': 'em',        # Bot uses "Em"
                    'customer_address_en': 'Madam',
                    'bot_address_en': 'I'
                }
        else:  # 45+
            if gender.lower() == 'male':
                return {
                    'customer_address': 'Chú',  # Call customer "Chu"
                    'bot_address': 'con',       # Bot uses "Con"
                    'customer_address_en': 'Uncle',
                    'bot_address_en': 'I'
                }
            else:  # female
                return {
                    'customer_address': 'Cô',   # Call customer "Co"
                    'bot_address': 'con',       # Bot uses "Con"
                    'customer_address_en': 'Aunt',
                    'bot_address_en': 'I'
                }
    
    def setup_routes(self):
        """Setup Flask routes"""
        
        @self.app.route('/')
        def index():
            """Main chat page with beautiful female salesperson"""
            return '''
<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI Sales Assistant - Tech Store</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            height: 100vh;
            display: flex;
            align-items: center;
            padding: 20px;
            overflow: hidden;
        }
        
        .main-container {
            display: flex;
            width: 100%;
            max-width: 1400px;
            height: 90vh;
            margin: 0 auto;
            gap: 30px;
            align-items: stretch;
        }
        
        .chat-section {
            flex: 1;
            max-width: 750px;
            display: flex;
            flex-direction: column;
        }
        
        .salesperson-section {
            width: 420px;
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
            background: rgba(255, 255, 255, 0.95);
            border-radius: 25px;
            padding: 30px;
            box-shadow: 0 25px 50px rgba(0, 0, 0, 0.2);
            backdrop-filter: blur(10px);
            border: 1px solid rgba(255, 255, 255, 0.3);
        }
        
        .salesperson-image-container {
            width: 320px;
            height: 400px;
            border-radius: 20px;
            margin-bottom: 25px;
            background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
            display: flex;
            align-items: center;
            justify-content: center;
            border: 3px solid #667eea;
            position: relative;
            overflow: hidden;
            box-shadow: 0 15px 35px rgba(102, 126, 234, 0.3);
        }
        
        .salesperson-image {
            width: 100%;
            height: 100%;
            background: url('data:image/svg+xml,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 320 400"><defs><linearGradient id="skin" x1="0%" y1="0%" x2="0%" y2="100%"><stop offset="0%" style="stop-color:%23fef3e2"/><stop offset="100%" style="stop-color:%23fed7aa"/></linearGradient><linearGradient id="hair" x1="0%" y1="0%" x2="0%" y2="100%"><stop offset="0%" style="stop-color:%23374151"/><stop offset="100%" style="stop-color:%23111827"/></linearGradient><linearGradient id="suit" x1="0%" y1="0%" x2="0%" y2="100%"><stop offset="0%" style="stop-color:%232563eb"/><stop offset="100%" style="stop-color:%231d4ed8"/></linearGradient><linearGradient id="shirt" x1="0%" y1="0%" x2="0%" y2="100%"><stop offset="0%" style="stop-color:%23ffffff"/><stop offset="100%" style="stop-color:%23f8fafc"/></linearGradient></defs><rect width="320" height="400" fill="%23f1f5f9"/><ellipse cx="160" cy="100" rx="55" ry="65" fill="url(%23skin)"/><path d="M105,70 Q160,30 215,70 Q215,90 160,95 Q105,90 105,70" fill="url(%23hair)"/><circle cx="140" cy="85" r="4" fill="%23374151"/><circle cx="180" cy="85" r="4" fill="%23374151"/><ellipse cx="160" cy="95" rx="8" ry="4" fill="%23f59e0b"/><path d="M150,105 Q160,110 170,105" stroke="%23ef4444" stroke-width="3" fill="none"/><rect x="130" y="80" width="25" height="15" rx="12" fill="none" stroke="%23374151" stroke-width="2"/><rect x="165" y="80" width="25" height="15" rx="12" fill="none" stroke="%23374151" stroke-width="2"/><line x1="155" y1="87" x2="165" y2="87" stroke="%23374151" stroke-width="2"/><ellipse cx="160" cy="120" rx="45" ry="35" fill="url(%23skin)"/><rect x="120" y="160" width="80" height="100" rx="10" fill="url(%23shirt)"/><rect x="110" y="170" width="100" height="120" rx="15" fill="url(%23suit)"/><rect x="155" y="170" width="10" height="80" fill="%23dc2626"/><ellipse cx="160" cy="140" rx="25" ry="15" fill="url(%23skin)"/><ellipse cx="160" cy="155" rx="35" ry="20" fill="url(%23skin)"/><rect x="140" y="290" width="40" height="60" rx="8" fill="%23374151"/><rect x="135" y="350" width="50" height="50" rx="25" fill="%2392400e"/><ellipse cx="160" cy="180" rx="20" ry="25" fill="url(%23skin)"/><circle cx="160" cy="120" r="35" fill="none" stroke="%23fbbf24" stroke-width="1" opacity="0.3"/><text x="160" y="380" text-anchor="middle" font-family="Arial" font-size="14" fill="%23374151" font-weight="bold">AI Sales Assistant</text><circle cx="125" cy="75" r="15" fill="none" stroke="%23667eea" stroke-width="2" opacity="0.8"/><circle cx="195" cy="75" r="15" fill="none" stroke="%23667eea" stroke-width="2" opacity="0.8"/></svg>') center center;
            background-size: cover;
            position: relative;
        }
        
        .salesperson-info {
            text-align: center;
            color: #1e293b;
        }
        
        .salesperson-info h2 {
            font-size: 26px;
            margin-bottom: 15px;
            color: #667eea;
            font-weight: bold;
        }
        
        .salesperson-info p {
            font-size: 16px;
            line-height: 1.6;
            color: #64748b;
            margin-bottom: 15px;
        }
        
        .salesperson-features {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            justify-content: center;
            margin-top: 20px;
        }
        
        .feature-badge {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 8px 16px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: bold;
        }
        
        .chat-container {
            background: rgba(255, 255, 255, 0.95);
            border-radius: 25px;
            box-shadow: 0 25px 50px rgba(0, 0, 0, 0.2);
            display: flex;
            flex-direction: column;
            overflow: hidden;
            backdrop-filter: blur(10px);
            border: 1px solid rgba(255, 255, 255, 0.3);
            height: 100%;
        }
        
        .chat-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 25px;
            text-align: center;
            font-size: 22px;
            font-weight: bold;
            position: relative;
        }
        
        .chat-header::after {
            content: '';
            position: absolute;
            bottom: -10px;
            left: 50%;
            transform: translateX(-50%);
            width: 0;
            height: 0;
            border-left: 15px solid transparent;
            border-right: 15px solid transparent;
            border-top: 10px solid #764ba2;
        }
        
        .personal-info-modal {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0, 0, 0, 0.7);
            display: flex;
            justify-content: center;
            align-items: center;
            z-index: 1000;
            backdrop-filter: blur(5px);
        }
        
        .modal-content {
            background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
            padding: 40px;
            border-radius: 20px;
            box-shadow: 0 25px 50px rgba(0, 0, 0, 0.3);
            max-width: 450px;
            width: 90%;
            text-align: center;
            border: 1px solid rgba(102, 126, 234, 0.2);
        }
        
        .modal-content h2 {
            color: #667eea;
            margin-bottom: 25px;
            font-size: 28px;
        }
        
        .form-group {
            margin-bottom: 25px;
            text-align: left;
        }
        
        .form-group label {
            display: block;
            margin-bottom: 10px;
            font-weight: bold;
            color: #374151;
            font-size: 16px;
        }
        
        .form-group select,
        .form-group input {
            width: 100%;
            padding: 15px;
            border: 2px solid #e5e7eb;
            border-radius: 12px;
            font-size: 16px;
            outline: none;
            transition: all 0.3s;
            background: white;
        }
        
        .form-group select:focus,
        .form-group input:focus {
            border-color: #667eea;
            box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
        }
        
        .submit-btn {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            padding: 15px 35px;
            border-radius: 12px;
            font-size: 18px;
            font-weight: bold;
            cursor: pointer;
            transition: all 0.3s;
            width: 100%;
        }
        
        .submit-btn:hover:not(:disabled) {
            transform: translateY(-2px);
            box-shadow: 0 10px 25px rgba(102, 126, 234, 0.3);
        }
        
        .submit-btn:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none;
        }
        
        .chat-messages {
            flex: 1;
            overflow-y: auto;
            padding: 25px;
            background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
            min-height: 400px;
        }
        
        .message {
            margin-bottom: 20px;
            display: flex;
            align-items: flex-start;
            animation: slideIn 0.3s ease-out;
        }
        
        @keyframes slideIn {
            from {
                opacity: 0;
                transform: translateY(10px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        .message.user {
            justify-content: flex-end;
        }
        
        .message.bot {
            justify-content: flex-start;
        }
        
        .message-content {
            max-width: 75%;
            padding: 18px 25px;
            border-radius: 20px;
            word-wrap: break-word;
            line-height: 1.5;
            position: relative;
            font-size: 16px;
        }
        
        .message.user .message-content {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border-bottom-right-radius: 5px;
            margin-left: 60px;
        }
        
        .message.user .message-content::after {
            content: '';
            position: absolute;
            bottom: 0;
            right: -10px;
            width: 0;
            height: 0;
            border-left: 10px solid #764ba2;
            border-top: 10px solid transparent;
            border-bottom: 10px solid transparent;
        }
        
        .message.bot .message-content {
            background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
            border: 2px solid #e5e7eb;
            color: #374151;
            border-bottom-left-radius: 5px;
            margin-right: 60px;
        }
        
        .message.bot .message-content::after {
            content: '';
            position: absolute;
            bottom: 0;
            left: -12px;
            width: 0;
            height: 0;
            border-right: 10px solid #ffffff;
            border-top: 10px solid transparent;
            border-bottom: 10px solid transparent;
        }
        
        .chat-input-container {
            padding: 25px;
            background: rgba(255, 255, 255, 0.95);
            border-top: 1px solid #e5e7eb;
            display: flex;
            gap: 15px;
            align-items: center;
        }
        
        .language-selector {
            padding: 12px 15px;
            border: 2px solid #667eea;
            border-radius: 25px;
            background: white;
            color: #667eea;
            font-weight: bold;
            cursor: pointer;
            outline: none;
            transition: all 0.3s;
        }
        
        .language-selector:hover {
            background: #667eea;
            color: white;
        }
        
        .chat-input {
            flex: 1;
            padding: 15px 25px;
            border: 2px solid #e5e7eb;
            border-radius: 25px;
            font-size: 16px;
            outline: none;
            transition: all 0.3s;
            background: white;
        }
        
        .chat-input:focus {
            border-color: #667eea;
            box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
        }
        
        .send-button {
            padding: 15px 30px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 25px;
            cursor: pointer;
            font-size: 16px;
            font-weight: bold;
            transition: all 0.3s;
        }
        
        .send-button:hover:not(:disabled) {
            transform: translateY(-2px);
            box-shadow: 0 10px 25px rgba(102, 126, 234, 0.3);
        }
        
        .send-button:disabled {
            opacity: 0.6;
            cursor: not-allowed;
            transform: none;
        }
        
        .typing-indicator {
            display: none;
            padding: 18px 25px;
            background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
            border: 2px solid #e5e7eb;
            border-radius: 20px;
            border-bottom-left-radius: 5px;
            max-width: 75%;
            margin-bottom: 20px;
            margin-right: 60px;
        }
        
        .typing-dots {
            display: flex;
            gap: 6px;
        }
        
        .typing-dots span {
            width: 10px;
            height: 10px;
            background: #667eea;
            border-radius: 50%;
            animation: typing 1.4s infinite ease-in-out;
        }
        
        .typing-dots span:nth-child(2) {
            animation-delay: 0.2s;
        }
        
        .typing-dots span:nth-child(3) {
            animation-delay: 0.4s;
        }
        
        @keyframes typing {
            0%, 60%, 100% {
                transform: translateY(0);
                opacity: 0.4;
            }
            30% {
                transform: translateY(-15px);
                opacity: 1;
            }
        }
        
        .status-info {
            font-size: 14px;
            color: #64748b;
            margin-top: 10px;
            text-align: center;
            padding: 0 25px;
        }
        
        .user-info-display {
            background: linear-gradient(135deg, #dbeafe 0%, #e0e7ff 100%);
            padding: 15px 25px;
            margin: 0 25px 20px 25px;
            border-radius: 15px;
            font-size: 14px;
            color: #1e40af;
            text-align: center;
            border: 1px solid #3b82f6;
        }
        
        .error-message {
            color: #dc2626;
            background: #fef2f2;
            padding: 10px;
            border-radius: 8px;
            margin: 10px 0;
            border: 1px solid #fecaca;
        }
        
        @media (max-width: 1200px) {
            .main-container {
                flex-direction: column;
                height: auto;
                gap: 20px;
            }
            
            .salesperson-section {
                width: 100%;
                order: -1;
                max-width: 500px;
                margin: 0 auto;
            }
            
            .chat-section {
                max-width: none;
            }
            
            .chat-container {
                height: 70vh;
            }
        }
        
        @media (max-width: 768px) {
            body {
                padding: 10px;
            }
            
            .salesperson-image-container {
                width: 250px;
                height: 300px;
            }
            
            .message-content {
                max-width: 85%;
                font-size: 14px;
                padding: 15px 20px;
            }
            
            .chat-header {
                font-size: 18px;
                padding: 20px;
            }
            
            .chat-input-container {
                flex-wrap: wrap;
                gap: 10px;
            }
            
            .language-selector {
                order: 3;
                flex: 1;
                text-align: center;
            }
        }
    </style>
</head>
<body>
    <!-- Personal Information Modal -->
    <div class="personal-info-modal" id="personalInfoModal">
        <div class="modal-content">
            <h2>🤖 Xin chào! Chúc bạn một ngày tốt lành!</h2>
            <p style="margin-bottom: 25px; color: #64748b; font-size: 16px;">
                Vui lòng cung cấp một số thông tin cá nhân để tiện xưng hô nhé!<br>
                <small>Hello! Wish you a good day. Please provide some personal information for convenience in addressing.</small>
            </p>
            
            <div id="errorContainer"></div>
            
            <form id="personalInfoForm">
                <div class="form-group">
                    <label for="gender">Giới tính / Gender:</label>
                    <select id="gender" required>
                        <option value="">Chọn giới tính / Select gender</option>
                        <option value="male">Nam / Male</option>
                        <option value="female">Nữ / Female</option>
                    </select>
                </div>
                
                <div class="form-group">
                    <label for="age">Tuổi / Age:</label>
                    <input type="number" id="age" min="15" max="100" required 
                           placeholder="Nhập tuổi của bạn / Enter your age">
                </div>
                
                <button type="submit" class="submit-btn" id="submitBtn">
                    Bắt đầu trò chuyện / Start Chat
                </button>
            </form>
        </div>
    </div>

    <div class="main-container">
        <!-- Chat Section (Left Side) -->
        <div class="chat-section">
            <div class="chat-container">
                <div class="chat-header">
                    💬 AI Sales Assistant / Trợ lý AI Bán hàng
                </div>
                
                <div class="user-info-display" id="userInfoDisplay" style="display: none;">
                    <span id="userInfoText"></span>
                </div>
                
                <div class="chat-messages" id="chatMessages">
                    <div class="typing-indicator" id="typingIndicator">
                        <div class="typing-dots">
                            <span></span>
                            <span></span>
                            <span></span>
                        </div>
                    </div>
                </div>
                
                <div class="chat-input-container">
                    <select class="language-selector" id="languageSelect">
                        <option value="vi">🇻🇳 Tiếng Việt</option>
                        <option value="en">🇺🇸 English</option>
                    </select>
                    <input type="text" class="chat-input" id="messageInput" 
                           placeholder="Nhập tin nhắn của bạn / Type your message..." 
                           autocomplete="off" disabled>
                    <button class="send-button" id="sendButton" onclick="sendMessage()" disabled>
                        Gửi / Send
                    </button>
                </div>
                
                <div class="status-info" id="statusInfo">
                    Vui lòng nhập thông tin cá nhân để bắt đầu / Please enter personal information to start
                </div>
            </div>
        </div>

        <!-- Beautiful Female Salesperson Section (Right Side) -->
        <div class="salesperson-section">
            <div class="salesperson-image-container">
                <div class="salesperson-image"></div>
            </div>
            
            <div class="salesperson-info">
                <h2>👩‍💼 Trợ lý AI Bán hàng</h2>
                <p><strong>Xin chào!</strong> Tôi là trợ lý AI xinh đẹp của cửa hàng. Với kiến thức chuyên sâu về công nghệ, tôi sẵn sàng giúp bạn tìm kiếm những sản phẩm phù hợp nhất.</p>
                
                <p><strong>Hello!</strong> I'm the beautiful AI assistant of the store. With deep knowledge of technology, I'm ready to help you find the most suitable products.</p>
                
                <div class="salesperson-features">
                    <div class="feature-badge">Chuyên gia IT</div>
                    <div class="feature-badge">Tư vấn 24/7</div>
                    <div class="feature-badge">Đa ngôn ngữ</div>
                    <div class="feature-badge">Thân thiện</div>
                </div>
            </div>
        </div>
    </div>

    <script>
        let isProcessing = false;
        let userInfo = null;
        let addressingTerms = null;
        
        function showError(message) {
            const errorContainer = document.getElementById('errorContainer');
            errorContainer.innerHTML = `<div class="error-message">${message}</div>`;
        }
        
        function clearError() {
            const errorContainer = document.getElementById('errorContainer');
            errorContainer.innerHTML = '';
        }
        
        // Handle personal information form submission
        document.getElementById('personalInfoForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            clearError();
            
            console.log('Form submission started');
            
            const gender = document.getElementById('gender').value;
            const age = parseInt(document.getElementById('age').value);
            
            console.log('Form values - Gender:', gender, 'Age:', age);
            
            if (!gender) {
                showError('Vui lòng chọn giới tính / Please select gender');
                return;
            }
            
            if (!age || age < 15 || age > 100) {
                showError('Vui lòng nhập tuổi hợp lệ (15-100) / Please enter valid age (15-100)');
                return;
            }
            
            userInfo = { gender, age };
            console.log('UserInfo prepared:', userInfo);
            
            // Update submit button
            const submitBtn = document.getElementById('submitBtn');
            submitBtn.disabled = true;
            submitBtn.textContent = 'Đang xử lý... / Processing...';
            
            try {
                console.log('Sending request to /set_user_info');
                
                const response = await fetch('/set_user_info', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify(userInfo)
                });
                
                console.log('Response received, status:', response.status);
                
                if (!response.ok) {
                    throw new Error(`Server responded with status: ${response.status}`);
                }
                
                const data = await response.json();
                console.log('Response data:', data);
                
                if (!data.success) {
                    throw new Error(data.error || 'Server returned error');
                }
                
                addressingTerms = data.addressing_terms;
                console.log('Addressing terms received:', addressingTerms);
                
                // Hide modal and enable chat
                document.getElementById('personalInfoModal').style.display = 'none';
                document.getElementById('messageInput').disabled = false;
                document.getElementById('sendButton').disabled = false;
                
                // Show user info
                const userInfoDisplay = document.getElementById('userInfoDisplay');
                const userInfoText = document.getElementById('userInfoText');
                const language = document.getElementById('languageSelect').value;
                
                if (language === 'vi') {
                    userInfoText.textContent = `Chế độ xưng hô: ${addressingTerms.customer_address} - ${addressingTerms.bot_address} | Tuổi: ${age} | Giới tính: ${gender === 'male' ? 'Nam' : 'Nữ'}`;
                } else {
                    userInfoText.textContent = `Addressing mode: ${addressingTerms.customer_address_en} - ${addressingTerms.bot_address_en} | Age: ${age} | Gender: ${gender}`;
                }
                userInfoDisplay.style.display = 'block';
                
                // Show welcome message
                if (data.welcome_message) {
                    addMessage(data.welcome_message, false, language);
                }
                
                // Update status
                updateStatus(language === 'vi' ? 'Sẵn sàng trò chuyện' : 'Ready to chat');
                
                // Focus on input
                document.getElementById('messageInput').focus();
                
                console.log('Setup completed successfully');
                
            } catch (error) {
                console.error('Error in form submission:', error);
                showError(`Lỗi: ${error.message} / Error: ${error.message}`);
                
                // Re-enable submit button
                submitBtn.disabled = false;
                submitBtn.textContent = 'Bắt đầu trò chuyện / Start Chat';
            }
        });
        
        function addMessage(content, isUser = false, language = 'vi') {
            const messagesContainer = document.getElementById('chatMessages');
            const typingIndicator = document.getElementById('typingIndicator');
            
            const messageDiv = document.createElement('div');
            messageDiv.className = `message ${isUser ? 'user' : 'bot'}`;
            
            const messageContent = document.createElement('div');
            messageContent.className = 'message-content';
            messageContent.innerHTML = content.replace(/\\n/g, '<br>');
            
            messageDiv.appendChild(messageContent);
            messagesContainer.insertBefore(messageDiv, typingIndicator);
            
            // Scroll to bottom
            messagesContainer.scrollTop = messagesContainer.scrollHeight;
        }
        
        function showTyping() {
            document.getElementById('typingIndicator').style.display = 'block';
            const messagesContainer = document.getElementById('chatMessages');
            messagesContainer.scrollTop = messagesContainer.scrollHeight;
        }
        
        function hideTyping() {
            document.getElementById('typingIndicator').style.display = 'none';
        }
        
        function updateStatus(message) {
            document.getElementById('statusInfo').textContent = message;
        }
        
        async function sendMessage() {
            if (isProcessing || !userInfo) return;
            
            const messageInput = document.getElementById('messageInput');
            const sendButton = document.getElementById('sendButton');
            const languageSelect = document.getElementById('languageSelect');
            
            const message = messageInput.value.trim();
            if (!message) return;
            
            const language = languageSelect.value;
            
            // Disable input and show processing state
            isProcessing = true;
            messageInput.disabled = true;
            sendButton.disabled = true;
            updateStatus(language === 'vi' ? 'Đang xử lý...' : 'Processing...');
            
            // Add user message
            addMessage(message, true, language);
            messageInput.value = '';
            
            // Show typing indicator
            showTyping();
            
            try {
                const response = await fetch('/chat', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({
                        message: message,
                        language: language,
                        user_info: userInfo
                    })
                });
                
                const data = await response.json();
                
                // Hide typing indicator
                hideTyping();
                
                if (data.success) {
                    addMessage(data.response, false, language);
                    
                    // Update status with processing info
                    const processingTime = data.processing_time ? data.processing_time.toFixed(1) : '0.0';
                    const dataSource = data.data_source || 'unknown';
                    const statusMsg = language === 'vi' ? 
                        `✅ Hoàn thành (${processingTime}s, ${dataSource})` :
                        `✅ Complete (${processingTime}s, ${dataSource})`;
                    updateStatus(statusMsg);
                } else {
                    const errorMsg = language === 'vi' ? 
                        'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn.' :
                        'Sorry, I encountered an error processing your request.';
                    addMessage(data.error || errorMsg, false, language);
                    updateStatus(language === 'vi' ? '❌ Lỗi' : '❌ Error');
                }
            } catch (error) {
                console.error('Error:', error);
                hideTyping();
                
                const errorMsg = language === 'vi' ? 
                    'Lỗi kết nối. Vui lòng thử lại.' :
                    'Connection error. Please try again.';
                addMessage(errorMsg, false, language);
                updateStatus(language === 'vi' ? '❌ Lỗi kết nối' : '❌ Connection Error');
            } finally {
                // Re-enable input
                isProcessing = false;
                messageInput.disabled = false;
                sendButton.disabled = false;
                messageInput.focus();
            }
        }
        
        // Enter key support
        document.getElementById('messageInput').addEventListener('keypress', function(e) {
            if (e.key === 'Enter' && !e.shiftKey) {
                e.preventDefault();
                sendMessage();
            }
        });
        
        // Language change handler
        document.getElementById('languageSelect').addEventListener('change', function(e) {
            const language = e.target.value;
            const messageInput = document.getElementById('messageInput');
            
            if (language === 'vi') {
                messageInput.placeholder = 'Nhập tin nhắn của bạn...';
            } else {
                messageInput.placeholder = 'Type your message...';
            }
            
            // Update user info display
            if (userInfo && addressingTerms) {
                const userInfoText = document.getElementById('userInfoText');
                if (language === 'vi') {
                    userInfoText.textContent = `Chế độ xưng hô: ${addressingTerms.customer_address} - ${addressingTerms.bot_address} | Tuổi: ${userInfo.age} | Giới tính: ${userInfo.gender === 'male' ? 'Nam' : 'Nữ'}`;
                } else {
                    userInfoText.textContent = `Addressing mode: ${addressingTerms.customer_address_en} - ${addressingTerms.bot_address_en} | Age: ${userInfo.age} | Gender: ${userInfo.gender}`;
                }
            }
        });
    </script>
</body>
</html>
            '''
        
        @self.app.route('/set_user_info', methods=['POST'])
        def set_user_info():
            """Set user personal information and return addressing terms"""
            try:
                data = request.get_json()
                if not data or 'gender' not in data or 'age' not in data:
                    return jsonify({'success': False, 'error': 'Missing user information'})
                
                gender = data['gender']
                age = int(data['age'])
                
                # Store user info in session (simple implementation)
                session_id = request.remote_addr  # Use IP as simple session ID
                addressing_terms = self.get_addressing_terms(gender, age)
                
                self.user_sessions[session_id] = {
                    'gender': gender,
                    'age': age,
                    'addressing_terms': addressing_terms
                }
                
                # Generate personalized welcome message
                customer_addr = addressing_terms['customer_address']
                bot_addr = addressing_terms['bot_address']
                
                welcome_message = f"""Xin chào {customer_addr}! {bot_addr.title()} là trợ lý AI bán hàng. {bot_addr.title()} rất vui được hỗ trợ {customer_addr} hôm nay.

{customer_addr} đang tìm kiếm sản phẩm gì? {bot_addr.title()} có thể giúp {customer_addr} tìm laptop, phụ kiện gaming, hoặc các sản phẩm công nghệ khác phù hợp với nhu cầu của {customer_addr}.

---

Hello! I am your AI sales assistant. I'm very happy to help you today.

What products are you looking for? I can help you find laptops, gaming accessories, or other technology products that suit your needs."""
                
                return jsonify({
                    'success': True,
                    'addressing_terms': addressing_terms,
                    'welcome_message': welcome_message
                })
                
            except Exception as e:
                return jsonify({'success': False, 'error': str(e)})
        
        @self.app.route('/chat', methods=['POST'])
        def chat():
            """Handle chat messages with personalized addressing"""
            try:
                data = request.get_json()
                if not data or 'message' not in data:
                    return jsonify({'success': False, 'error': 'No message provided'})
                
                user_message = data['message']
                language = data.get('language', 'vi')
                user_info = data.get('user_info', {})
                
                # Get user session info
                session_id = request.remote_addr
                session_data = self.user_sessions.get(session_id, {})
                addressing_terms = session_data.get('addressing_terms', {})
                
                # Set the chatbot's current language
                self.chatbot.current_language = language
                
                # Process the message using existing chatbot logic
                result = self.chatbot.process_message(user_message)
                
                # Personalize the response with appropriate addressing
                response = result.get('response', '')
                if addressing_terms and language == 'vi':
                    response = self.personalize_response(response, addressing_terms)
                
                return jsonify({
                    'success': True,
                    'response': response,
                    'data_source': result.get('data_source', 'unknown'),
                    'processing_time': result.get('processing_time', 0),
                    'user_language': result.get('user_language', language),
                    'local_results_count': result.get('local_results_count', 0)
                })
                
            except Exception as e:
                return jsonify({
                    'success': False,
                    'error': str(e),
                    'processing_time': 0
                })
    
    def personalize_response(self, response, addressing_terms):
        """Personalize response with appropriate Vietnamese addressing terms"""
        try:
            customer_addr = addressing_terms.get('customer_address', 'Bạn')
            bot_addr = addressing_terms.get('bot_address', 'tôi')
            
            # Replace common addressing terms in the response
            replacements = {
                'bạn': customer_addr.lower(),
                'Bạn': customer_addr,
                'tôi': bot_addr,
                'Tôi': bot_addr.title(),
                'mình': bot_addr,
                'Mình': bot_addr.title()
            }
            
            personalized_response = response
            for old_term, new_term in replacements.items():
                # Use word boundaries to avoid partial replacements
                import re
                pattern = r'\b' + re.escape(old_term) + r'\b'
                personalized_response = re.sub(pattern, new_term, personalized_response)
            
            return personalized_response
            
        except Exception as e:
            # If personalization fails, return original response
            return response
    
    def run(self):
        """Run the web server"""
        try:
            print(f"🌐 Starting enhanced AI sales interface on http://{self.host}:{self.port}")
            print(f"📱 Local access: http://localhost:{self.port}")
            
            # Get local IP for network access
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
                s.connect(("8.8.8.8", 80))
                local_ip = s.getsockname()[0]
                s.close()
                print(f"🌐 Network access: http://{local_ip}:{self.port}")
            except:
                pass
            
            print("✨ Features:")
            print("   • Beautiful female AI salesperson with glasses")
            print("   • Personalized greeting and addressing system")
            print("   • Professional dialogue-style chat interface")
            print("   • Vietnamese/English language support")
            print("   • Responsive design for all devices")
            
            self.app.run(host=self.host, port=self.port, debug=False, threaded=True)
        except Exception as e:
            print(f"❌ Error starting web server: {e}")

def create_config_file():
    """Create a sample configuration file"""
    config_content = """# Vietnamese AI Sales ChatBot Configuration (Text-only version)
version: '2.0'
environment: 'production'

# AI Models Configuration
ai_models:
  primary_llm: 'microsoft/DialoGPT-medium'
  fallback_llm: 'microsoft/DialoGPT-small'
  embedding_model: 'sentence-transformers/all-MiniLM-L6-v2'

# Language Support
language_config:
  default_language: 'vi'
  supported_languages: ['vi', 'en']
  auto_detect_language: true
  translate_responses: true

# GPU Configuration
gpu_config:
  primary_device: 'cuda:0'
  use_quantization: true
  mixed_precision: true
  max_memory_per_gpu: 0.85
  batch_size: 2

# Search Configuration
search_config:
  local_similarity_threshold: 0.1
  enable_google_search: true
  max_google_results: 3
  search_timeout: 10

# Performance Settings
performance:
  max_response_length: 150
  temperature: 0.7
  top_p: 0.9
  repetition_penalty: 1.1
  do_sample: true
  num_return_sequences: 1

# Analytics
analytics:
  track_conversations: true
  track_language_usage: true
  track_errors: true
"""
    
    try:
        with open('chatbot_config.yaml', 'w', encoding='utf-8') as f:
            f.write(config_content)
        print("✅ Configuration file created: chatbot_config.yaml")
    except Exception as e:
        print(f"❌ Error creating config file: {e}")


def main():
    """Main function with enhanced error handling and setup"""
    print("""
    ===============================================
    🚀 Vietnamese AI Sales ChatBot (Text-only)
    ===============================================
    
    🔧 Features:
    ✅ Vietnamese + English language support
    ✅ Smart product search with embeddings
    ✅ Excel, PDF, Word document processing
    ✅ Image OCR with Vietnamese support
    ✅ Web interface for local network access
    ✅ Admin GUI for testing and database management
    
    Loading chatbot...
    """)
    
    # Parse command line arguments
    parser = argparse.ArgumentParser(description='Vietnamese AI Sales ChatBot')
    parser.add_argument('--web', action='store_true', 
                       help='Start web interface instead of GUI')
    parser.add_argument('--host', default='0.0.0.0', 
                       help='Web interface host (default: 0.0.0.0)')
    parser.add_argument('--port', type=int, default=5000, 
                       help='Web interface port (default: 5000)')
    args = parser.parse_args()
    
    try:
        os.makedirs('data', exist_ok=True)
        os.makedirs('logs', exist_ok=True)
        
        if not os.path.exists('chatbot_config.yaml'):
            print("📝 Creating configuration file...")
            create_config_file()
        
        print("🤖 Initializing chatbot...")
        chatbot = VietnameseAISalesBot()
        
        # Add sample products
        print("📦 Adding sample products...")
        chatbot.add_sample_products()
        
        if args.web:
            # Start web interface
            print("🌐 Starting web interface...")
            web_interface = WebChatInterface(chatbot, host=args.host, port=args.port)
            web_interface.run()
        else:
            # Start GUI interface (existing functionality)
            print("🖥️ Starting GUI interface...")
            chatbot.display_welcome_message()
            chatbot.root.mainloop()
        
    except KeyboardInterrupt:
        print("\n👋 ChatBot stopped by user")
    except Exception as e:
        print(f"❌ Error starting chatbot: {e}")
        print("\n🔧 Troubleshooting tips:")
        print("1. Ensure all required packages are installed: pip install flask")
        print("2. Check if the port is available")
        print("3. For GUI mode, ensure tkinter is available")
        
        import traceback
        traceback.print_exc()
        
        input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()
