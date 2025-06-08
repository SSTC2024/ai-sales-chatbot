import os
import json
import sqlite3
import pandas as pd
import numpy as np
from datetime import datetime
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox, font
import threading
import queue
import time
import logging
from pathlib import Path
import re
import yaml

# AI/ML imports
import torch
from transformers import (
    AutoTokenizer, AutoModelForCausalLM, 
    BitsAndBytesConfig, pipeline
)
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Web search imports
import requests
from bs4 import BeautifulSoup
from googlesearch import search
import urllib.parse

# Document processing
import docx
from PyPDF2 import PdfReader
from openpyxl import load_workbook
import easyocr

# Voice support imports with better error handling
VOICE_AVAILABLE = False
MICROPHONE_AVAILABLE = False
TTS_AVAILABLE = False

try:
    import speech_recognition as sr
    import pyttsx3
    import pyaudio
    
    # Test if microphone is available
    try:
        # Test microphone availability
        recognizer = sr.Recognizer()
        mic_list = sr.Microphone.list_microphone_names()
        if mic_list:
            # Try to initialize a microphone
            test_mic = sr.Microphone()
            with test_mic as source:
                recognizer.adjust_for_ambient_noise(source, duration=0.1)
            MICROPHONE_AVAILABLE = True
            print(f"✅ Found {len(mic_list)} microphone(s): {mic_list[0] if mic_list else 'None'}")
        else:
            print("⚠️ No microphones detected")
    except Exception as mic_error:
        print(f"⚠️ Microphone test failed: {mic_error}")
        MICROPHONE_AVAILABLE = False
    
    # Test TTS availability
    try:
        test_engine = pyttsx3.init()
        test_engine.stop()
        TTS_AVAILABLE = True
        print("✅ Text-to-speech engine available")
    except Exception as tts_error:
        print(f"⚠️ TTS test failed: {tts_error}")
        TTS_AVAILABLE = False
    
    VOICE_AVAILABLE = True
    print("✅ Voice support libraries loaded successfully")
    
except ImportError as import_error:
    VOICE_AVAILABLE = False
    print(f"⚠️ Voice support libraries not available: {import_error}")
    print("Install with: pip install SpeechRecognition pyttsx3 pyaudio")

# Language detection and translation
try:
    from googletrans import Translator, LANGUAGES
    import langdetect
    TRANSLATION_AVAILABLE = True
    print("✅ Translation libraries loaded successfully")
except ImportError:
    TRANSLATION_AVAILABLE = False
    print("⚠️ Translation libraries not available. Install with: pip install googletrans==4.0.0rc1 langdetect")

class VietnameseVoiceChatBot:
    """
    Enhanced AI Sales ChatBot with Vietnamese language and voice support
    Features robust error handling for audio devices and response generation
    """
    
    def __init__(self):
        self.load_config()
        self.setup_logging()
        self.setup_language_support()
        self.setup_voice_support()
        self.setup_database()
        self.initialize_ai_models()
        self.conversation_context = []
        self.setup_gui()
    
    def load_config(self):
        """Load configuration with Vietnamese language settings"""
        try:
            with open('chatbot_config.yaml', 'r', encoding='utf-8') as file:
                self.config = yaml.safe_load(file)
            print("✅ Configuration loaded from chatbot_config.yaml")
            
            # Update configuration to enable Vietnamese and voice features
            self.config = self.update_config_for_vietnamese()
            
        except FileNotFoundError:
            print("⚠️ chatbot_config.yaml not found, using default Vietnamese settings")
            self.config = self.get_default_vietnamese_config()
        except Exception as e:
            print(f"❌ Error loading config: {e}")
            self.config = self.get_default_vietnamese_config()
    
    def update_config_for_vietnamese(self):
        """Update existing config to support Vietnamese and voice features"""
        config = self.config.copy()
        
        # Add Vietnamese language settings if not present
        if 'language_config' not in config:
            config['language_config'] = {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True,
                'translate_responses': True,
                'voice_language': 'vi-VN'
            }
        
        # Add voice settings if not present
        if 'voice_config' not in config:
            config['voice_config'] = {
                'enable_voice_input': MICROPHONE_AVAILABLE,
                'enable_voice_output': TTS_AVAILABLE,
                'voice_rate': 150,
                'voice_volume': 0.8,
                'vietnamese_voice_id': 0,
                'english_voice_id': 1,
                'auto_send_voice': False,
                'graceful_fallback': True  # Fallback to text if voice fails
            }
        
        # Add Vietnamese conversation flows
        if 'conversation_flows' not in config:
            config['conversation_flows'] = {}
            
        config['conversation_flows'].update({
            'greeting_vietnamese': {
                'triggers': ['xin chào', 'chào', 'chào bạn', 'chào anh', 'chào chị'],
                'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng của bạn. Tôi có thể giúp bạn tìm kiếm sản phẩm phù hợp với nhu cầu. Bạn đang tìm kiếm gì hôm nay?',
                'next_stage': 'needs_assessment'
            }
        })
        
        # Add Vietnamese templates with error handling messages
        config['vietnamese_templates'] = {
            'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp với tiêu chí của bạn. Hãy để tôi tìm kiếm các lựa chọn tương tự hoặc kiểm tra xem chúng tôi có sản phẩm tương đương nào sắp ra mắt không.',
            'out_of_stock': 'Rất tiếc, {product_name} hiện đang hết hàng. Tuy nhiên, tôi có thể đề xuất các sản phẩm thay thế tương tự hoặc kiểm tra khi nào sản phẩm sẽ có hàng trở lại.',
            'processing': 'Đang xử lý yêu cầu của bạn...',
            'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn. Vui lòng thử lại.',
            'voice_error': 'Không thể nhận diện giọng nói. Vui lòng thử lại hoặc gõ tin nhắn.',
            'voice_unavailable': 'Tính năng giọng nói không khả dụng. Vui lòng sử dụng gõ văn bản.',
            'microphone_error': 'Không tìm thấy microphone. Kiểm tra kết nối và thử lại.',
            'model_error': 'Mô hình AI không khả dụng. Đang sử dụng phản hồi dự phòng.',
            'generation_error': 'Có lỗi khi tạo phản hồi. Vui lòng thử lại.',
            'empty_response': 'Tôi hiểu câu hỏi của bạn nhưng cần thêm thông tin để trả lời tốt hơn.',
            'connection_error': 'Lỗi kết nối. Vui lòng kiểm tra mạng và thử lại.'
        }
        
        return config
    
    def get_default_vietnamese_config(self):
        """Get default configuration with Vietnamese support and error handling"""
        return {
            'version': '2.0',
            'environment': 'production',
            'ai_models': {
                'primary_llm': 'microsoft/DialoGPT-medium',  # More stable fallback
                'fallback_llm': 'microsoft/DialoGPT-small',
                'embedding_model': 'sentence-transformers/all-MiniLM-L6-v2'
            },
            'language_config': {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True,
                'translate_responses': True,
                'voice_language': 'vi-VN'
            },
            'voice_config': {
                'enable_voice_input': MICROPHONE_AVAILABLE,
                'enable_voice_output': TTS_AVAILABLE,
                'voice_rate': 150,
                'voice_volume': 0.8,
                'graceful_fallback': True
            },
            'gpu_config': {
                'primary_device': 'cuda:0' if torch.cuda.is_available() else 'cpu',
                'use_quantization': True,
                'mixed_precision': True,
                'max_memory_per_gpu': 0.85,
                'batch_size': 2  # Reduced for stability
            },
            'search_config': {
                'local_similarity_threshold': 0.1,
                'enable_google_search': True,
                'max_google_results': 3,
                'search_timeout': 10
            },
            'performance': {
                'max_response_length': 150,  # Reduced for reliability
                'temperature': 0.7,
                'top_p': 0.9,
                'repetition_penalty': 1.1,
                'do_sample': True,
                'num_return_sequences': 1
            },
            'conversation_flows': {
                'greeting_vietnamese': {
                    'triggers': ['xin chào', 'chào', 'chào bạn'],
                    'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng của bạn. Tôi có thể giúp bạn tìm kiếm sản phẩm phù hợp với nhu cầu. Bạn đang tìm kiếm gì hôm nay?'
                },
                'greeting': {
                    'triggers': ['hello', 'hi', 'hey'],
                    'response_template': 'Hello! I\'m your AI sales assistant. I can help you find the perfect products for your needs. What are you looking for today?'
                }
            },
            'vietnamese_templates': {
                'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp với tiêu chí của bạn.',
                'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn. Vui lòng thử lại.',
                'voice_unavailable': 'Tính năng giọng nói không khả dụng. Vui lòng sử dụng gõ văn bản.',
                'model_error': 'Mô hình AI không khả dụng. Đang sử dụng phản hồi dự phòng.',
                'generation_error': 'Có lỗi khi tạo phản hồi. Vui lòng thử lại.',
                'empty_response': 'Tôi hiểu câu hỏi của bạn nhưng cần thêm thông tin để trả lời tốt hơn.'
            },
            'analytics': {
                'track_conversations': True,
                'track_language_usage': True,
                'track_voice_usage': True,
                'track_errors': True
            }
        }
    
    def setup_language_support(self):
        """Initialize language detection and translation"""
        self.current_language = self.config['language_config']['default_language']
        
        if TRANSLATION_AVAILABLE:
            try:
                self.translator = Translator()
                print("✅ Translation service initialized")
            except Exception as e:
                print(f"⚠️ Translation service initialization failed: {e}")
                self.translator = None
        else:
            self.translator = None
    
    def setup_voice_support(self):
        """Initialize voice recognition and text-to-speech with error handling"""
        self.voice_enabled = False
        self.tts_enabled = False
        self.recognizer = None
        self.microphone = None
        self.tts_engine = None
        
        if not VOICE_AVAILABLE:
            print("⚠️ Voice libraries not available - running in text-only mode")
            return
        
        # Initialize speech recognition
        if MICROPHONE_AVAILABLE and self.config['voice_config']['enable_voice_input']:
            try:
                self.recognizer = sr.Recognizer()
                
                # Try to find an available microphone
                mic_list = sr.Microphone.list_microphone_names()
                if mic_list:
                    self.microphone = sr.Microphone()
                    
                    # Adjust for ambient noise with shorter duration to avoid hanging
                    with self.microphone as source:
                        print("🎤 Adjusting microphone for ambient noise...")
                        self.recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    
                    self.voice_enabled = True
                    print(f"✅ Speech recognition initialized with microphone: {mic_list[0]}")
                else:
                    print("❌ No microphones found")
                    
            except Exception as e:
                print(f"❌ Speech recognition initialization failed: {e}")
                self.voice_enabled = False
        else:
            print("⚠️ Microphone not available or disabled in config")
        
        # Initialize text-to-speech
        if TTS_AVAILABLE and self.config['voice_config']['enable_voice_output']:
            try:
                self.tts_engine = pyttsx3.init()
                self.setup_tts_voices()
                self.tts_enabled = True
                print("✅ Text-to-speech initialized")
            except Exception as e:
                print(f"❌ TTS initialization failed: {e}")
                self.tts_enabled = False
        else:
            print("⚠️ TTS not available or disabled in config")
        
        # Report final voice status
        if self.voice_enabled or self.tts_enabled:
            print(f"🎙️ Voice features: Input={'✅' if self.voice_enabled else '❌'} | Output={'✅' if self.tts_enabled else '❌'}")
        else:
            print("📝 Running in text-only mode - all voice features disabled")
    
    def setup_tts_voices(self):
        """Setup TTS voices for Vietnamese and English"""
        if not self.tts_engine:
            return
        
        try:
            voices = self.tts_engine.getProperty('voices')
            self.vietnamese_voice = None
            self.english_voice = None
            
            print("🔍 Available TTS voices:")
            for i, voice in enumerate(voices):
                print(f"  {i}: {voice.name}")
                
                # Try to identify Vietnamese voice
                voice_str = str(voice.name).lower()
                if any(lang in voice_str for lang in ['vi', 'vietnamese', 'viet']):
                    self.vietnamese_voice = voice.id
                    print(f"    ✅ Vietnamese voice found: {voice.name}")
                
                # Try to identify English voice
                if any(lang in voice_str for lang in ['en', 'english', 'us', 'gb']) and not self.english_voice:
                    self.english_voice = voice.id
                    print(f"    ✅ English voice found: {voice.name}")
            
            # Fallback to system default if specific voices not found
            if not self.vietnamese_voice and voices:
                self.vietnamese_voice = voices[0].id
                print("⚠️ No specific Vietnamese voice found, using default")
            
            if not self.english_voice and voices:
                self.english_voice = voices[1].id if len(voices) > 1 else voices[0].id
                print("⚠️ No specific English voice found, using default")
            
            # Set TTS properties
            self.tts_engine.setProperty('rate', self.config['voice_config']['voice_rate'])
            self.tts_engine.setProperty('volume', self.config['voice_config']['voice_volume'])
            
        except Exception as e:
            print(f"❌ TTS voice setup failed: {e}")
            self.tts_enabled = False
    
    def detect_language(self, text):
        """Detect language of input text with error handling"""
        if not TRANSLATION_AVAILABLE or not text or not text.strip():
            return 'vi'  # Default to Vietnamese
        
        try:
            detected = langdetect.detect(text)
            if detected in self.config['language_config']['supported_languages']:
                return detected
            return 'vi'  # Default to Vietnamese
        except Exception as e:
            print(f"Language detection error: {e}")
            return 'vi'  # Default to Vietnamese on error
    
    def translate_text(self, text, target_language):
        """Translate text to target language with error handling"""
        if not self.translator or not text or not text.strip():
            return text
        
        try:
            translated = self.translator.translate(text, dest=target_language)
            return translated.text if translated and translated.text else text
        except Exception as e:
            print(f"Translation error: {e}")
            return text
    
    def speak_text(self, text, language='vi'):
        """Convert text to speech with robust error handling"""
        if not self.tts_enabled or not self.tts_engine or not text or not text.strip():
            return
        
        try:
            # Select appropriate voice
            if language == 'vi' and self.vietnamese_voice:
                self.tts_engine.setProperty('voice', self.vietnamese_voice)
            elif language == 'en' and self.english_voice:
                self.tts_engine.setProperty('voice', self.english_voice)
            
            # Speak the text
            self.tts_engine.say(text)
            self.tts_engine.runAndWait()
            
        except Exception as e:
            print(f"TTS error: {e}")
            # Graceful fallback - don't crash the application
    
    def listen_for_speech(self, language='vi-VN', timeout=5):
        """Listen for speech input with robust error handling"""
        if not self.voice_enabled or not self.recognizer or not self.microphone:
            return None, None
        
        try:
            with self.microphone as source:
                print("🎤 Listening...")
                # Listen for audio with timeout
                audio = self.recognizer.listen(source, timeout=timeout, phrase_time_limit=10)
            
            print("🔄 Processing speech...")
            
            # Try Vietnamese first, then English
            try:
                text = self.recognizer.recognize_google(audio, language='vi-VN')
                return text, 'vi'
            except sr.UnknownValueError:
                try:
                    text = self.recognizer.recognize_google(audio, language='en-US')
                    return text, 'en'
                except sr.UnknownValueError:
                    print("Could not understand audio")
                    return None, None
            except sr.RequestError as e:
                print(f"Speech recognition service error: {e}")
                return None, None
                    
        except sr.WaitTimeoutError:
            print("⏰ Speech timeout - no audio detected")
            return None, None
        except Exception as e:
            print(f"Speech recognition error: {e}")
            return None, None
    
    def setup_logging(self):
        """Setup logging configuration with Vietnamese support"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('chatbot.log', encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
        
    def setup_database(self):
        """Initialize SQLite database with Vietnamese support"""
        self.conn = sqlite3.connect('chatbot_data.db', check_same_thread=False)
        self.conn.execute("PRAGMA encoding = 'UTF-8'")
        
        cursor = self.conn.cursor()
        
        # Enhanced database schema with language support
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS products (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                name_vietnamese TEXT,
                description TEXT,
                description_vietnamese TEXT,
                category TEXT,
                category_vietnamese TEXT,
                price REAL,
                features TEXT,
                features_vietnamese TEXT,
                specifications TEXT,
                specifications_vietnamese TEXT,
                availability TEXT,
                availability_vietnamese TEXT,
                source_file TEXT,
                embedding BLOB,
                created_at TEXT,
                updated_at TEXT
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT,
                user_input TEXT,
                user_language TEXT,
                bot_response TEXT,
                bot_language TEXT,
                intent TEXT,
                confidence REAL,
                data_source TEXT,
                response_time REAL,
                voice_input BOOLEAN DEFAULT 0,
                voice_output BOOLEAN DEFAULT 0
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS knowledge_base (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                topic TEXT,
                content TEXT,
                source TEXT,
                embedding BLOB,
                created_at TEXT
            )
        ''')
        
        # Add new columns if they don't exist (for existing databases)
        try:
            cursor.execute('ALTER TABLE conversations ADD COLUMN error_type TEXT')
            print("✅ Added error_type column to conversations table")
        except sqlite3.OperationalError:
            pass  # Column already exists
        
        try:
            cursor.execute('ALTER TABLE conversations ADD COLUMN error_message TEXT')
            print("✅ Added error_message column to conversations table")
        except sqlite3.OperationalError:
            pass  # Column already exists
        
        self.conn.commit()
        print("✅ Database schema updated successfully")
        
    def initialize_ai_models(self):
        """Initialize AI models with enhanced error handling"""
        self.logger.info("Initializing AI models from configuration...")
        
        try:
            self.setup_devices()
            self.load_language_model()
            self.load_embedding_model()
            self.load_ocr_model()
            
            self.logger.info("✅ All AI models loaded successfully!")
            
        except Exception as e:
            self.logger.error(f"❌ Error initializing AI models: {e}")
            # Set fallback states
            self.text_generator = None
            self.embedding_model = None
            self.ocr_reader = None
    
    def setup_devices(self):
        """Device setup using configuration with fallback"""
        try:
            if torch.cuda.is_available():
                gpu_count = torch.cuda.device_count()
                print(f"Found {gpu_count} GPU(s)")
                
                best_gpu = 0
                max_memory = 0
                
                for i in range(gpu_count):
                    props = torch.cuda.get_device_properties(i)
                    memory_gb = props.total_memory / (1024**3)
                    print(f"GPU {i}: {props.name} - {memory_gb:.1f}GB")
                    
                    if memory_gb > max_memory:
                        max_memory = memory_gb
                        best_gpu = i
                
                primary_device = self.config['gpu_config'].get('primary_device', f'cuda:{best_gpu}')
                
                self.primary_device = primary_device
                self.secondary_device = f'cuda:{(best_gpu + 1) % gpu_count}' if gpu_count > 1 else self.primary_device
                self.primary_memory_gb = max_memory
                
            else:
                self.primary_device = 'cpu'
                self.secondary_device = 'cpu'
                self.primary_memory_gb = 8
                print("CUDA not available. Using CPU mode.")
        except Exception as e:
            print(f"Device setup error: {e}")
            self.primary_device = 'cpu'
            self.secondary_device = 'cpu'
            self.primary_memory_gb = 8
    
    def load_language_model(self):
        """Load language model with comprehensive error handling"""
        try:
            primary_llm = self.config['ai_models']['primary_llm']
            fallback_llm = self.config['ai_models']['fallback_llm']
            use_quantization_config = self.config['gpu_config']['use_quantization']
            
            primary_memory = getattr(self, 'primary_memory_gb', 8)
            
            # Choose model based on available memory
            if primary_memory >= 12:
                use_quantization = use_quantization_config
                model_name = primary_llm
            else:
                use_quantization = True
                model_name = fallback_llm
            
            print(f"Loading language model: {model_name}")
            
            # Load tokenizer with error handling
            try:
                self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                if self.tokenizer.pad_token is None:
                    self.tokenizer.pad_token = self.tokenizer.eos_token
                print("✅ Tokenizer loaded successfully")
            except Exception as e:
                print(f"❌ Tokenizer loading failed: {e}")
                self.tokenizer = None
                self.text_generator = None
                return
            
            # Configure quantization
            quantization_config = None
            if use_quantization and torch.cuda.is_available():
                try:
                    quantization_config = BitsAndBytesConfig(
                        load_in_8bit=True,
                        llm_int8_threshold=6.0,
                    )
                except Exception as e:
                    print(f"❌ Quantization config failed: {e}")
                    quantization_config = None
            
            # Load model
            try:
                self.llm_model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    quantization_config=quantization_config,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                    trust_remote_code=True,
                    low_cpu_mem_usage=True
                )
                print("✅ Language model loaded successfully")
            except Exception as e:
                print(f"❌ Model loading failed: {e}")
                self.llm_model = None
                self.text_generator = None
                return
            
            # Create text generation pipeline
            try:
                # Try without device first (for accelerate compatibility)
                try:
                    self.text_generator = pipeline(
                        "text-generation",
                        model=self.llm_model,
                        tokenizer=self.tokenizer,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
                    )
                    print("✅ Text generation pipeline created successfully (accelerate mode)")
                except Exception as accelerate_error:
                    # Fallback: try with device specification for non-accelerate models
                    print(f"⚠️ Accelerate mode failed, trying with device specification: {accelerate_error}")
                    self.text_generator = pipeline(
                        "text-generation",
                        model=self.llm_model,
                        tokenizer=self.tokenizer,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                        device=0 if torch.cuda.is_available() else -1
                    )
                    print("✅ Text generation pipeline created successfully (device mode)")
            except Exception as e:
                print(f"❌ Pipeline creation failed: {e}")
                self.text_generator = None
            
        except Exception as e:
            print(f"❌ Error loading language model: {e}")
            self.text_generator = None
            
    def load_embedding_model(self):
        """Load embedding model with error handling"""
        try:
            embedding_model_name = self.config['ai_models']['embedding_model']
            print(f"Loading embedding model: {embedding_model_name}")
            
            self.embedding_model = SentenceTransformer(embedding_model_name)
            if torch.cuda.is_available():
                try:
                    self.embedding_model = self.embedding_model.to(self.secondary_device)
                except Exception as e:
                    print(f"⚠️ Could not move embedding model to GPU: {e}")
            print("✅ Embedding model loaded successfully")
        except Exception as e:
            print(f"❌ Error loading embedding model: {e}")
            self.embedding_model = None
            
    def load_ocr_model(self):
        """Load OCR model with error handling"""
        try:
            self.ocr_reader = easyocr.Reader(['en', 'vi'])
            print("✅ OCR model loaded successfully with Vietnamese support")
        except Exception as e:
            print(f"❌ Error loading OCR model: {e}")
            self.ocr_reader = None
    
    def safe_generate_response(self, prompt, user_language='vi'):
        """Safely generate response with comprehensive error handling"""
        try:
            # Check if text generator is available
            if not self.text_generator or not self.tokenizer:
                self.logger.warning("Text generator not available")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['model_error']
                else:
                    return "I apologize, but the AI text generation system is not available right now."
            
            # Validate prompt
            if not prompt or not isinstance(prompt, str) or not prompt.strip():
                self.logger.warning("Invalid prompt provided")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need more information to provide a helpful response."
            
            # Get performance configuration
            performance_config = self.config.get('performance', {})
            
            # Generate with error handling
            try:
                with torch.no_grad():
                    generated = self.text_generator(
                        prompt,
                        max_new_tokens=performance_config.get('max_response_length', 150),
                        min_new_tokens=10,
                        temperature=performance_config.get('temperature', 0.7),
                        top_p=performance_config.get('top_p', 0.9),
                        repetition_penalty=performance_config.get('repetition_penalty', 1.1),
                        do_sample=performance_config.get('do_sample', True),
                        pad_token_id=self.tokenizer.eos_token_id,
                        num_return_sequences=performance_config.get('num_return_sequences', 1),
                        return_full_text=False
                    )
                
                # Safely extract response
                if not generated:
                    self.logger.warning("Empty generation result")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['generation_error']
                    else:
                        return "I'm having trouble generating a response. Please try again."
                
                # Handle different response formats
                if isinstance(generated, list) and len(generated) > 0:
                    first_result = generated[0]
                    if isinstance(first_result, dict) and 'generated_text' in first_result:
                        response = first_result['generated_text']
                    else:
                        response = str(first_result)
                else:
                    response = str(generated)
                
                # Validate response
                if not response or not response.strip():
                    self.logger.warning("Empty response generated")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question but need more details to respond properly."
                
                # Clean and return response
                cleaned_response = self.clean_generated_response(response)
                
                # Final validation
                if not cleaned_response or len(cleaned_response.strip()) < 5:
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question. Could you provide more details?"
                
                return cleaned_response
                
            except torch.cuda.OutOfMemoryError:
                self.logger.error("CUDA out of memory")
                if user_language == 'vi':
                    return "Hệ thống đang quá tải. Vui lòng thử lại sau."
                else:
                    return "System is overloaded. Please try again later."
                    
            except Exception as generation_error:
                self.logger.error(f"Text generation error: {generation_error}")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['generation_error']
                else:
                    return "I encountered an error while generating a response. Please try again."
            
        except Exception as e:
            self.logger.error(f"Safe generation error: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having technical difficulties. Please try again."
    
    def generate_natural_response(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Generate natural language response with Vietnamese support and error handling"""
        try:
            # Validate inputs
            if not user_input or not isinstance(user_input, str):
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need a question or message to respond to."
            
            user_input = user_input.strip()
            if not user_input:
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "Please provide a question or message."
            
            # Check for greeting patterns from config
            try:
                if user_language == 'vi':
                    greeting_flow = self.config['conversation_flows'].get('greeting_vietnamese', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Xin chào!')
                else:
                    greeting_flow = self.config['conversation_flows'].get('greeting', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Hello!')
            except Exception as e:
                self.logger.warning(f"Greeting pattern matching error: {e}")
            
            # Build context-aware prompt
            try:
                prompt = self.build_sales_prompt(user_input, context_data, data_source, user_language)
            except Exception as e:
                self.logger.error(f"Prompt building error: {e}")
                prompt = f"User: {user_input}\nResponse:"
            
            # Generate response safely
            response = self.safe_generate_response(prompt, user_language)
            
            return response
            
        except Exception as e:
            self.logger.error(f"Error generating natural response: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having trouble generating a response right now."
    
    def build_sales_prompt(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Build an optimized prompt for sales conversations with Vietnamese support"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý AI bán hàng chuyên nghiệp. Hãy trả lời bằng tiếng Việt một cách thân thiện và hữu ích.

"""
            else:
                prompt = """You are a professional AI sales assistant. Respond in English in a friendly and helpful manner.

"""
            
            # Add conversation context safely
            try:
                if self.conversation_context and len(self.conversation_context) > 0:
                    if user_language == 'vi':
                        prompt += "Cuộc trò chuyện trước:\n"
                    else:
                        prompt += "Previous conversation:\n"
                    
                    # Only add last 2 turns to avoid prompt overflow
                    recent_context = self.conversation_context[-2:]
                    for turn in recent_context:
                        if isinstance(turn, dict) and 'user' in turn and 'bot' in turn:
                            user_msg = str(turn['user'])[:100]  # Truncate to prevent overflow
                            bot_msg = str(turn['bot'])[:100]
                            prompt += f"User: {user_msg}\nAssistant: {bot_msg}\n"
                    prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Context building error: {e}")
                
            # Add relevant data context safely
            try:
                if context_data and isinstance(context_data, list) and len(context_data) > 0:
                    if data_source == "database":
                        if user_language == 'vi':
                            prompt += "Sản phẩm liên quan:\n"
                        else:
                            prompt += "Relevant products:\n"
                        
                        # Process up to 2 items to avoid prompt overflow
                        for item in context_data[:2]:
                            if isinstance(item, dict):
                                name = item.get('name_vietnamese', item.get('name', '')) if user_language == 'vi' else item.get('name', '')
                                desc = item.get('description_vietnamese', item.get('description', '')) if user_language == 'vi' else item.get('description', '')
                                
                                if name:
                                    prompt += f"- {str(name)[:80]}"  # Truncate to prevent overflow
                                    if desc:
                                        prompt += f": {str(desc)[:100]}...\n"
                                    else:
                                        prompt += "\n"
                                        
                                    if item.get('price'):
                                        try:
                                            price_label = "Giá" if user_language == 'vi' else "Price"
                                            prompt += f"  {price_label}: ${float(item['price']):.2f}\n"
                                        except (ValueError, TypeError):
                                            pass
                        prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Product context building error: {e}")
                
            # Current user question
            try:
                # Truncate user input to prevent prompt overflow
                truncated_input = str(user_input)[:200]
                if user_language == 'vi':
                    prompt += f"Câu hỏi của khách hàng: {truncated_input}\n\nTrả lời bằng tiếng Việt:"
                else:
                    prompt += f"Customer question: {truncated_input}\n\nResponse:"
            except Exception as e:
                self.logger.error(f"Question formatting error: {e}")
                if user_language == 'vi':
                    prompt += f"Câu hỏi: {user_input}\nTrả lời:"
                else:
                    prompt += f"Question: {user_input}\nResponse:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Prompt building error: {e}")
            # Fallback to simple prompt
            if user_language == 'vi':
                return f"Câu hỏi: {user_input}\nTrả lời bằng tiếng Việt:"
            else:
                return f"Question: {user_input}\nResponse:"
    
    def clean_generated_response(self, response):
        """Clean up generated response with error handling"""
        try:
            if not response:
                return ""
            
            response = str(response)
            
            # Remove HTML/XML tags
            response = re.sub(r'<[^>]+>', '', response)
            
            # Remove markdown-style brackets
            response = re.sub(r'\[.*?\]', '', response)
            
            # Normalize whitespace
            response = re.sub(r'\n+', ' ', response)
            response = re.sub(r'\s+', ' ', response)
            
            # Remove common artifacts
            response = response.replace('<|endoftext|>', '')
            response = response.replace('</s>', '')
            response = response.replace('<s>', '')
            
            # Handle sentences
            try:
                sentences = response.split('.')
                if len(sentences) > 1 and len(sentences[-1].strip()) < 10:
                    response = '.'.join(sentences[:-1]) + '.'
            except Exception:
                pass
            
            # Final cleanup
            response = response.strip()
            
            # Ensure minimum length
            if len(response) < 3:
                return ""
                
            return response
            
        except Exception as e:
            self.logger.error(f"Response cleaning error: {e}")
            return str(response) if response else ""
    
    def search_local_database(self, user_input, similarity_threshold=None):
        """Search local database with Vietnamese support and error handling"""
        try:
            if not self.embedding_model:
                self.logger.warning("Embedding model not available")
                return []
            
            if not user_input or not user_input.strip():
                return []
            
            if similarity_threshold is None:
                similarity_threshold = self.config['search_config']['local_similarity_threshold']
                
            # Generate query embedding safely
            try:
                query_embedding = self.embedding_model.encode([user_input.strip()])
                if query_embedding is None or len(query_embedding) == 0:
                    self.logger.warning("Failed to generate query embedding")
                    return []
            except Exception as e:
                self.logger.error(f"Query embedding error: {e}")
                return []
            
            # Query database safely
            try:
                cursor = self.conn.cursor()
                cursor.execute("""
                    SELECT name, name_vietnamese, description, description_vietnamese, 
                           category, price, embedding 
                    FROM products 
                    WHERE embedding IS NOT NULL
                """)
                products = cursor.fetchall()
            except Exception as e:
                self.logger.error(f"Database query error: {e}")
                return []
            
            if not products:
                self.logger.info("No products with embeddings found")
                return []
                
            best_matches = []
            
            for product in products:
                try:
                    if len(product) < 7:  # Ensure product has all expected fields
                        continue
                        
                    stored_embedding_blob = product[6]  # embedding column
                    
                    if stored_embedding_blob:
                        # Safely decode embedding
                        try:
                            stored_embedding = np.frombuffer(stored_embedding_blob, dtype=np.float32)
                            stored_embedding = stored_embedding.reshape(1, -1)
                        except Exception as e:
                            self.logger.warning(f"Embedding decode error: {e}")
                            continue
                        
                        # Calculate similarity safely
                        try:
                            similarity = cosine_similarity(query_embedding, stored_embedding)[0][0]
                        except Exception as e:
                            self.logger.warning(f"Similarity calculation error: {e}")
                            continue
                        
                        if similarity > similarity_threshold:
                            best_matches.append({
                                'name': product[0] or '',
                                'name_vietnamese': product[1] or '',
                                'description': product[2] or '',
                                'description_vietnamese': product[3] or '',
                                'category': product[4] or '',
                                'price': product[5],
                                'similarity': float(similarity)
                            })
                    
                except Exception as e:
                    self.logger.warning(f"Product processing error: {e}")
                    continue
                        
            # Sort by similarity safely
            try:
                best_matches.sort(key=lambda x: x.get('similarity', 0), reverse=True)
                return best_matches[:3]
            except Exception as e:
                self.logger.error(f"Sorting error: {e}")
                return best_matches[:3] if best_matches else []
            
        except Exception as e:
            self.logger.error(f"Database search error: {e}")
            return []

    def process_user_message(self, user_input, voice_input=False):
        """Main processing pipeline with Vietnamese and voice support and comprehensive error handling"""
        start_time = time.time()
        
        try:
            # Validate input
            if not user_input or not isinstance(user_input, str):
                error_msg = "Vui lòng nhập tin nhắn hợp lệ." if self.current_language == 'vi' else "Please enter a valid message."
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'invalid_input'
                }
            
            user_input = user_input.strip()
            if not user_input:
                error_msg = "Vui lòng nhập tin nhắn." if self.current_language == 'vi' else "Please enter a message."
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'empty_input'
                }
            
            # Detect language safely
            try:
                user_language = self.detect_language(user_input)
            except Exception as e:
                self.logger.warning(f"Language detection error: {e}")
                user_language = self.current_language or 'vi'
            
            # Search local database safely
            local_results = []
            try:
                local_results = self.search_local_database(user_input)
            except Exception as e:
                self.logger.error(f"Database search error: {e}")
                local_results = []
            
            # Generate response safely
            try:
                if local_results:
                    response = self.generate_natural_response(
                        user_input, 
                        context_data=local_results,
                        data_source="database",
                        user_language=user_language
                    )
                    data_source = "local_database"
                else:
                    response = self.generate_natural_response(user_input, user_language=user_language)
                    data_source = "general_knowledge"
            except Exception as e:
                self.logger.error(f"Response generation error: {e}")
                if user_language == 'vi':
                    response = self.config['vietnamese_templates']['generation_error']
                else:
                    response = "I encountered an error while generating a response. Please try again."
                data_source = "error"
                    
            # Update conversation context safely
            try:
                self.conversation_context.append({
                    'user': user_input,
                    'bot': response,
                    'user_language': user_language,
                    'timestamp': datetime.now().isoformat()
                })
                
                if len(self.conversation_context) > 10:
                    self.conversation_context = self.conversation_context[-10:]
            except Exception as e:
                self.logger.warning(f"Context update error: {e}")
                
            # Store conversation in database safely
            try:
                processing_time = time.time() - start_time
                self.store_conversation(user_input, response, data_source, processing_time, 
                                      user_language, voice_input)
            except Exception as e:
                self.logger.warning(f"Conversation storage error: {e}")
            
            return {
                'response': response,
                'data_source': data_source,
                'local_results_count': len(local_results),
                'processing_time': time.time() - start_time,
                'user_language': user_language,
                'response_language': user_language
            }
            
        except Exception as e:
            self.logger.error(f"Message processing error: {e}")
            error_language = getattr(self, 'current_language', 'vi')
            if error_language == 'vi':
                error_msg = self.config['vietnamese_templates']['error']
            else:
                error_msg = "I apologize, but I encountered an error processing your request."
            
            return {
                'response': error_msg,
                'data_source': 'error',
                'local_results_count': 0,
                'processing_time': time.time() - start_time,
                'user_language': error_language,
                'response_language': error_language,
                'error': str(e)
            }
            
    def store_conversation(self, user_input, response, data_source, processing_time, 
                          user_language, voice_input=False, voice_output=False, error_type=None, error_message=None):
        """Store conversation in database with language tracking and error logging"""
        try:
            cursor = self.conn.cursor()
            
            # First try with error columns
            try:
                cursor.execute('''
                    INSERT INTO conversations 
                    (timestamp, user_input, user_language, bot_response, bot_language, 
                     data_source, response_time, voice_input, voice_output, error_type, error_message)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    datetime.now().isoformat(),
                    user_input,
                    user_language,
                    response,
                    user_language,
                    data_source,
                    processing_time,
                    voice_input,
                    voice_output,
                    error_type,
                    error_message
                ))
            except sqlite3.OperationalError as e:
                if "no column named error_type" in str(e):
                    # Fallback to original schema without error columns
                    cursor.execute('''
                        INSERT INTO conversations 
                        (timestamp, user_input, user_language, bot_response, bot_language, 
                         data_source, response_time, voice_input, voice_output)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        datetime.now().isoformat(),
                        user_input,
                        user_language,
                        response,
                        user_language,
                        data_source,
                        processing_time,
                        voice_input,
                        voice_output
                    ))
                else:
                    raise e
            
            self.conn.commit()
        except Exception as e:
            self.logger.error(f"Error storing conversation: {e}")
    
    def process_excel_file(self, file_path):
        """Process Excel files for product data with Vietnamese support"""
        try:
            # Read Excel file
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, encoding='utf-8')
            else:
                df = pd.read_excel(file_path)
                
            self.update_training_status(f"📊 Found {len(df)} rows in Excel file")
                
            # Expected columns: name, description, category, price, features, specifications
            # Also support Vietnamese columns
            column_mapping = {
                'name': ['name', 'tên', 'ten', 'product_name', 'sản phẩm'],
                'name_vietnamese': ['name_vietnamese', 'tên_tiếng_việt', 'ten_tieng_viet'],
                'description': ['description', 'mô tả', 'mo_ta', 'desc'],
                'description_vietnamese': ['description_vietnamese', 'mô_tả_tiếng_việt'],
                'category': ['category', 'danh mục', 'danh_muc', 'loại'],
                'category_vietnamese': ['category_vietnamese', 'danh_mục_tiếng_việt'],
                'price': ['price', 'giá', 'gia', 'cost'],
                'features': ['features', 'tính năng', 'tinh_nang'],
                'features_vietnamese': ['features_vietnamese', 'tính_năng_tiếng_việt'],
                'specifications': ['specifications', 'thông số', 'thong_so', 'specs'],
                'specifications_vietnamese': ['specifications_vietnamese', 'thông_số_tiếng_việt']
            }
            
            # Map columns
            mapped_columns = {}
            for standard_col, possible_cols in column_mapping.items():
                for col in possible_cols:
                    if col in df.columns:
                        mapped_columns[standard_col] = col
                        break
            
            if 'name' not in mapped_columns and 'name_vietnamese' not in mapped_columns:
                raise ValueError("Excel file must contain at least a name column (English or Vietnamese)")
                
            cursor = self.conn.cursor()
            added_count = 0
            
            for _, row in df.iterrows():
                try:
                    # Extract data using mapped columns
                    name = str(row.get(mapped_columns.get('name', ''), '')) if 'name' in mapped_columns else ''
                    name_vietnamese = str(row.get(mapped_columns.get('name_vietnamese', ''), '')) if 'name_vietnamese' in mapped_columns else ''
                    description = str(row.get(mapped_columns.get('description', ''), '')) if 'description' in mapped_columns else ''
                    description_vietnamese = str(row.get(mapped_columns.get('description_vietnamese', ''), '')) if 'description_vietnamese' in mapped_columns else ''
                    category = str(row.get(mapped_columns.get('category', ''), 'General')) if 'category' in mapped_columns else 'General'
                    category_vietnamese = str(row.get(mapped_columns.get('category_vietnamese', ''), '')) if 'category_vietnamese' in mapped_columns else ''
                    features = str(row.get(mapped_columns.get('features', ''), '')) if 'features' in mapped_columns else ''
                    features_vietnamese = str(row.get(mapped_columns.get('features_vietnamese', ''), '')) if 'features_vietnamese' in mapped_columns else ''
                    specifications = str(row.get(mapped_columns.get('specifications', ''), '')) if 'specifications' in mapped_columns else ''
                    specifications_vietnamese = str(row.get(mapped_columns.get('specifications_vietnamese', ''), '')) if 'specifications_vietnamese' in mapped_columns else ''

                    # Handle price
                    price = 0
                    if 'price' in mapped_columns:
                        try:
                            price_val = row.get(mapped_columns['price'], 0)
                            if pd.notna(price_val):
                                price = float(str(price_val).replace(',', '').replace('$', ''))
                        except (ValueError, TypeError):
                            price = 0
                    
                    # Skip if no meaningful data
                    if not name and not name_vietnamese:
                        continue
                        
                    # Generate embedding for the product (combine English and Vietnamese)
                    product_text_en = f"{name} {description} {features} {specifications}"
                    product_text_vi = f"{name_vietnamese} {description_vietnamese} {features_vietnamese} {specifications_vietnamese}"
                    combined_text = f"{product_text_en} {product_text_vi}".strip()
                    
                    embedding_blob = None
                    if self.embedding_model and combined_text:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.update_training_status(f"⚠️ Embedding generation failed for {name or name_vietnamese}: {e}")
                    
                    cursor.execute('''
                        INSERT INTO products 
                        (name, name_vietnamese, description, description_vietnamese, 
                         category, category_vietnamese, price, features, features_vietnamese,
                         specifications, specifications_vietnamese, source_file, embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        name or name_vietnamese,  # Use Vietnamese as fallback
                        name_vietnamese,
                        description or description_vietnamese,
                        description_vietnamese,
                        category,
                        category_vietnamese,
                        price,
                        features,
                        features_vietnamese,
                        specifications,
                        specifications_vietnamese,
                        file_path,
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                    added_count += 1
                    
                except Exception as row_error:
                    self.update_training_status(f"⚠️ Error processing row: {row_error}")
                    continue
                
            self.conn.commit()
            self.update_training_status(f"✅ Added {added_count} products from Excel file with embeddings")
            
        except Exception as e:
            raise Exception(f"Error processing Excel file: {e}")
            
    def process_pdf_file(self, file_path):
        """Process PDF files for knowledge base with Vietnamese support"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    self.update_training_status(f"📄 Processed page {page_num + 1}/{len(reader.pages)}")
                except Exception as e:
                    self.update_training_status(f"⚠️ Error extracting page {page_num + 1}: {e}")
                    continue
                
            if text.strip():
                # Generate embedding for knowledge base entry
                embedding_blob = None
                if self.embedding_model:
                    try:
                        # Use first 5000 chars for embedding
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                # Store in knowledge base
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],  # Limit to 10000 chars for storage
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from PDF: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from PDF")
                
        except Exception as e:
            raise Exception(f"Error processing PDF file: {e}")
            
    def process_word_file(self, file_path):
        """Process Word documents for knowledge base with Vietnamese support"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                if para_num % 10 == 0:
                    self.update_training_status(f"📝 Processed {para_num} paragraphs...")
                    
            # Extract text from tables
            for table_num, table in enumerate(doc.tables):
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                self.update_training_status(f"📊 Processed table {table_num + 1}/{len(doc.tables)}")
                
            if text.strip():
                # Generate embedding for knowledge base entry
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                # Store in knowledge base
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],  # Limit to 10000 chars
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from Word doc: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from Word document")
                
        except Exception as e:
            raise Exception(f"Error processing Word file: {e}")
            
    def process_image_file(self, file_path):
        """Process image files using OCR with Vietnamese support"""
        try:
            if not self.ocr_reader:
                raise Exception("OCR model not available")
                
            self.update_training_status(f"🖼️ Running OCR on image...")
            
            # Extract text using OCR
            results = self.ocr_reader.readtext(file_path)
            
            extracted_texts = []
            for (bbox, text, confidence) in results:
                if confidence > 0.5:  # Only keep high-confidence text
                    extracted_texts.append(text.strip())
                    
            if extracted_texts:
                combined_text = " ".join(extracted_texts)
                
                # Generate embedding
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding = self.embedding_model.encode([combined_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                # Store in knowledge base
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    f"OCR: {os.path.basename(file_path)}",
                    combined_text,
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Extracted {len(extracted_texts)} text segments from image")
                self.update_training_status(f"📝 Text preview: {combined_text[:100]}...")
            else:
                self.update_training_status(f"⚠️ No text detected in image")
                
        except Exception as e:
            raise Exception(f"Error processing image file: {e}")
    
    def setup_gui(self):
        """Setup the GUI interface with Vietnamese font support"""
        self.root = tk.Tk()
        self.root.title("Trợ lý AI Bán hàng - Vietnamese Voice ChatBot")
        self.root.geometry("1400x900")
        
        # Configure Vietnamese font support
        self.setup_vietnamese_fonts()
        
        # Create main interface
        self.setup_main_interface()
        
    def setup_vietnamese_fonts(self):
        """Setup fonts that support Vietnamese characters with enhanced input support"""
        # Prioritize fonts with excellent Vietnamese Unicode support
        vietnamese_fonts = [
            ('Segoe UI', 12),           # Excellent Vietnamese support on Windows
            ('Times New Roman', 12),    # Classic with good Vietnamese support
            ('Arial Unicode MS', 12),   # Comprehensive Unicode support
            ('Calibri', 12),           # Modern with Vietnamese support
            ('Tahoma', 12),            # Good Vietnamese rendering
            ('Microsoft Sans Serif', 12), # Fallback option
            ('DejaVu Sans', 12),       # Good for Linux systems
            ('Liberation Sans', 12),    # Good for cross-platform
            ('Verdana', 11)            # Fallback
        ]
        
        self.vietnamese_font = None
        self.input_font = None
        available_fonts = font.families()
        
        for font_name, size in vietnamese_fonts:
            if font_name in available_fonts:
                try:
                    # Test if the font can handle Vietnamese characters
                    test_font = font.Font(family=font_name, size=size)
                    
                    # Main display font
                    self.vietnamese_font = font.Font(family=font_name, size=size-1)
                    # Input font with slightly larger size and specific Unicode encoding
                    self.input_font = font.Font(
                        family=font_name, 
                        size=size, 
                        weight='normal'
                    )
                    
                    print(f"✅ Using Vietnamese font: {font_name}")
                    break
                except Exception as e:
                    print(f"⚠️ Font {font_name} failed: {e}")
                    continue
        
        if not self.vietnamese_font:
            # Ultimate fallback
            self.vietnamese_font = font.Font(family="TkDefaultFont", size=11)
            self.input_font = font.Font(family="TkDefaultFont", size=12)
            print("⚠️ Using default font (Vietnamese characters may not display correctly)")
        
        # Configure root for Vietnamese support
        self.root.option_add('*Font', self.vietnamese_font)
        
        # Set system encoding to UTF-8 for Vietnamese support
        try:
            self.root.tk.call('encoding', 'system', 'utf-8')
            print("✅ System encoding set to UTF-8 for Vietnamese support")
        except Exception as e:
            print(f"⚠️ Could not set UTF-8 encoding: {e}")
    
    def setup_main_interface(self):
        """Setup the main GUI interface with notebook tabs"""
        # Create notebook for tabbed interface
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Setup different tabs
        self.setup_chat_tab()
        self.setup_database_tab()
        self.setup_training_tab()
        self.setup_analytics_tab()
    
    def setup_chat_tab(self):
        """Setup the main chat interface"""
        chat_frame = ttk.Frame(self.notebook)
        self.notebook.add(chat_frame, text="💬 Chat")
        
        # Top frame for language and voice controls
        top_frame = ttk.Frame(chat_frame)
        top_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Language selection
        lang_frame = ttk.Frame(top_frame)
        lang_frame.pack(side=tk.LEFT)
        
        ttk.Label(lang_frame, text="Ngôn ngữ / Language:", font=self.vietnamese_font).pack(side=tk.LEFT, padx=5)
        
        self.language_var = tk.StringVar(value=self.config['language_config']['default_language'])
        language_combo = ttk.Combobox(lang_frame, textvariable=self.language_var, 
                                     values=['vi', 'en'], state='readonly', width=8)
        language_combo.pack(side=tk.LEFT, padx=5)
        language_combo.bind('<<ComboboxSelected>>', self.on_language_change)
        
        # Voice controls
        voice_frame = ttk.Frame(top_frame)
        voice_frame.pack(side=tk.RIGHT)
        
        # Voice status indicator
        voice_status = "🎤✅" if self.voice_enabled else "🎤❌"
        tts_status = "🔊✅" if self.tts_enabled else "🔊❌"
        status_text = f"Voice: {voice_status} TTS: {tts_status}"
        
        ttk.Label(voice_frame, text=status_text, font=self.vietnamese_font).pack(side=tk.LEFT, padx=5)
        
        if self.voice_enabled:
            self.voice_input_btn = ttk.Button(voice_frame, text="🎤 Nói / Speak", 
                                            command=self.start_voice_input)
            self.voice_input_btn.pack(side=tk.LEFT, padx=2)
        
        if self.tts_enabled:
            self.voice_output_var = tk.BooleanVar(value=True)
            voice_output_check = ttk.Checkbutton(voice_frame, text="🔊 Đọc to / Read aloud", 
                                               variable=self.voice_output_var)
            voice_output_check.pack(side=tk.LEFT, padx=2)
        
        # Chat display with styling
        self.chat_display = scrolledtext.ScrolledText(
            chat_frame, 
            height=30, 
            state=tk.DISABLED,
            font=self.vietnamese_font,
            wrap=tk.WORD
        )
        self.chat_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Configure text tags for styling
        self.chat_display.tag_configure("user", foreground="blue", font=(self.vietnamese_font['family'], 10, 'bold'))
        self.chat_display.tag_configure("bot", foreground="green", font=(self.vietnamese_font['family'], 10))
        self.chat_display.tag_configure("system", foreground="gray", font=(self.vietnamese_font['family'], 9, 'italic'))
        self.chat_display.tag_configure("error", foreground="red", font=(self.vietnamese_font['family'], 9, 'italic'))
        
        # Input frame
        input_frame = ttk.Frame(chat_frame)
        input_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # User input
        self.user_input = tk.Text(
            input_frame, 
            height=3, 
            font=self.input_font,
            wrap=tk.WORD
        )
        self.user_input.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        self.user_input.bind('<Control-Return>', self.send_message)
        self.user_input.bind('<Return>', self.on_enter_key)
        
        # Buttons
        button_frame = ttk.Frame(input_frame)
        button_frame.pack(side=tk.RIGHT, fill=tk.Y)
        
        send_btn = ttk.Button(button_frame, text="Gửi / Send", command=self.send_message)
        send_btn.pack(fill=tk.X, pady=(0, 5))
        
        clear_btn = ttk.Button(button_frame, text="Xóa / Clear", command=self.clear_chat)
        clear_btn.pack(fill=tk.X)
        
        # Status frame
        status_frame = ttk.Frame(chat_frame)
        status_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.status_label = ttk.Label(status_frame, text="Sẵn sàng / Ready", font=self.vietnamese_font)
        self.status_label.pack(side=tk.LEFT)
        
        self.processing_label = ttk.Label(status_frame, text="", font=self.vietnamese_font)
        self.processing_label.pack(side=tk.RIGHT)
    
    def setup_database_tab(self):
        """Setup database management interface"""
        db_frame = ttk.Frame(self.notebook)
        self.notebook.add(db_frame, text="🗄️ Database")
        
        # Product management
        product_frame = ttk.LabelFrame(db_frame, text="Product Database Management")
        product_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Buttons
        btn_frame = ttk.Frame(product_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(btn_frame, text="Add Sample Products", 
                  command=self.add_sample_products).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="View Products", 
                  command=self.view_products).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Search Test", 
                  command=self.test_search).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Regenerate Embeddings", 
                  command=self.regenerate_all_embeddings).pack(side=tk.LEFT, padx=5)
        
        # Product display
        self.product_display = scrolledtext.ScrolledText(
            product_frame, height=15, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.product_display.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
    
    def setup_training_tab(self):
        """Setup training interface with file upload capabilities"""
        training_frame = ttk.Frame(self.notebook)
        self.notebook.add(training_frame, text="📚 Training Data")
        
        # File upload section
        upload_frame = ttk.LabelFrame(training_frame, text="Upload Training Data")
        upload_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # File upload buttons
        upload_btn_frame = ttk.Frame(upload_frame)
        upload_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(upload_btn_frame, text="📊 Upload Excel Files", 
                  command=lambda: self.process_files('excel')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📄 Upload PDF Files", 
                  command=lambda: self.process_files('pdf')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📝 Upload Word Files", 
                  command=lambda: self.process_files('word')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="🖼️ Upload Images (OCR)", 
                  command=lambda: self.process_files('image')).pack(side=tk.LEFT, padx=5, pady=5)
        
        # Training options
        options_frame = ttk.LabelFrame(training_frame, text="Training Options")
        options_frame.pack(fill=tk.X, padx=10, pady=5)
        
        options_btn_frame = ttk.Frame(options_frame)
        options_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(options_btn_frame, text="🔄 Regenerate All Embeddings", 
                  command=self.regenerate_all_embeddings).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="🧹 Clear Training Data", 
                  command=self.clear_training_data).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="📈 Training Statistics", 
                  command=self.show_training_stats).pack(side=tk.LEFT, padx=5)
        
        # Training status and log
        self.training_status = scrolledtext.ScrolledText(
            training_frame, height=20, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.training_status.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def setup_analytics_tab(self):
        """Setup analytics interface"""
        analytics_frame = ttk.Frame(self.notebook)
        self.notebook.add(analytics_frame, text="📊 Analytics")
        
        # Metrics frame
        metrics_frame = ttk.LabelFrame(analytics_frame, text="Performance Metrics")
        metrics_frame.pack(fill=tk.X, padx=10, pady=5)
        
        metrics_btn_frame = ttk.Frame(metrics_frame)
        metrics_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(metrics_btn_frame, text="🔄 Refresh Analytics", 
                  command=self.refresh_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📤 Export Data", 
                  command=self.export_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📋 Conversation History", 
                  command=self.view_conversation_history).pack(side=tk.LEFT, padx=5, pady=5)
        
        # Analytics display
        self.analytics_display = scrolledtext.ScrolledText(
            analytics_frame, height=25, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.analytics_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def on_enter_key(self, event):
        """Handle Enter key press"""
        if event.state & 0x4:  # Control key is pressed
            self.send_message()
        else:
            # Insert newline
            return None
    
    def on_language_change(self, event=None):
        """Handle language change"""
        try:
            self.current_language = self.language_var.get()
            print(f"Language changed to: {self.current_language}")
        except Exception as e:
            print(f"Language change error: {e}")
    
    def start_voice_input(self):
        """Start voice input with error handling"""
        if not self.voice_enabled:
            error_msg = self.config['vietnamese_templates']['voice_unavailable']
            self.display_message("Hệ thống / System", error_msg, "error")
            return
        
        try:
            self.voice_input_btn.config(text="🎤 Đang nghe... / Listening...", state='disabled')
            threading.Thread(target=self.voice_input_worker, daemon=True).start()
        except Exception as e:
            print(f"Voice input start error: {e}")
            self.voice_input_error()
    
    def voice_input_worker(self):
        """Worker thread for voice input with robust error handling"""
        try:
            text, language = self.listen_for_speech()
            
            if text:
                self.root.after(0, lambda: self.process_voice_input(text, language))
            else:
                self.root.after(0, lambda: self.voice_input_error())
                
        except Exception as e:
            print(f"Voice input worker error: {e}")
            self.root.after(0, lambda: self.voice_input_error())
    
    def process_voice_input(self, text, language):
        """Process voice input"""
        try:
            # Insert text into input field
            self.user_input.delete(1.0, tk.END)
            self.user_input.insert(1.0, text)
            
            # Reset button
            if hasattr(self, 'voice_input_btn'):
                self.voice_input_btn.config(text="🎤 Nói / Speak", state='normal')
            
            # Display recognition result
            lang_name = "Tiếng Việt" if language == 'vi' else "English"
            self.display_message("Hệ thống / System", f"Đã nhận diện ({lang_name}): {text}", "system")
            
            # Auto-send
            self.send_message(voice_input=True)
        except Exception as e:
            print(f"Voice input processing error: {e}")
            self.voice_input_error()
    
    def voice_input_error(self):
        """Handle voice input error"""
        try:
            if hasattr(self, 'voice_input_btn'):
                self.voice_input_btn.config(text="🎤 Nói / Speak", state='normal')
            
            error_msg = self.config['vietnamese_templates']['voice_error']
            self.display_message("Hệ thống / System", error_msg, "error")
        except Exception as e:
            print(f"Voice error handling error: {e}")
    
    def send_message(self, event=None, voice_input=False):
        """Handle sending user message with error handling"""
        try:
            user_text = self.user_input.get(1.0, tk.END).strip()
            if not user_text:
                return
                
            # Clear input
            self.user_input.delete(1.0, tk.END)
            
            # Display user message
            self.display_message("Bạn / You", user_text, "user")
            
            # Update status
            if hasattr(self, 'processing_label'):
                status_text = "Đang xử lý..." if self.current_language == 'vi' else "Processing..."
                self.processing_label.config(text=status_text)
                self.root.update()
            
            # Process in separate thread
            threading.Thread(target=self.process_message_thread, args=(user_text, voice_input), daemon=True).start()
        except Exception as e:
            print(f"Send message error: {e}")
            error_msg = f"Lỗi gửi tin nhắn / Send error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
    
    def process_message_thread(self, user_input, voice_input=False):
        """Process message in separate thread with comprehensive error handling"""
        try:
            result = self.process_user_message(user_input, voice_input)
            
            # Display response
            response_tag = "error" if result['data_source'] == 'error' else "bot"
            self.display_message("Trợ lý AI / AI Assistant", result['response'], response_tag)
            
            # Text-to-speech output
            if (hasattr(self, 'voice_output_var') and self.voice_output_var.get() and 
                self.tts_enabled and result['data_source'] != 'error'):
                try:
                    threading.Thread(
                        target=lambda: self.speak_text(result['response'], result['response_language']),
                        daemon=True
                    ).start()
                except Exception as e:
                    print(f"TTS thread error: {e}")
            
            # Update status
            if result['user_language'] == 'vi':
                status_text = f"Hoàn thành - {result['data_source']} ({result['processing_time']:.1f}s)"
            else:
                status_text = f"Complete - {result['data_source']} ({result['processing_time']:.1f}s)"
            
            # Safely update processing label
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text=status_text))
            
        except Exception as e:
            self.logger.error(f"Message thread error: {e}")
            error_msg = f"Lỗi xử lý / Processing error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text="Lỗi / Error"))
    
    def display_message(self, sender, message, tag):
        """Display message in chat window with error handling"""
        def update_display():
            try:
                self.chat_display.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime("%H:%M:%S")
                
                self.chat_display.insert(tk.END, f"[{timestamp}] {sender}:\n", tag)
                self.chat_display.insert(tk.END, f"{message}\n\n")
                
                self.chat_display.config(state=tk.DISABLED)
                self.chat_display.see(tk.END)
            except Exception as e:
                print(f"Display update error: {e}")
            
        try:
            if threading.current_thread() != threading.main_thread():
                self.root.after(0, update_display)
            else:
                update_display()
        except Exception as e:
            print(f"Display message error: {e}")
    
    def clear_chat(self):
        """Clear chat display with error handling"""
        try:
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)
            self.chat_display.config(state=tk.DISABLED)
            self.conversation_context.clear()
        except Exception as e:
            print(f"Clear chat error: {e}")
    
    def update_training_status(self, message):
        """Update training status display with Vietnamese support"""
        def update():
            if hasattr(self, 'training_status'):
                self.training_status.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime('%H:%M:%S')
                self.training_status.insert(tk.END, f"[{timestamp}] {message}\n")
                self.training_status.config(state=tk.DISABLED)
                self.training_status.see(tk.END)
            
        if threading.current_thread() != threading.main_thread():
            self.root.after(0, update)
        else:
            update()
    
    def process_files(self, file_type):
        """Process different types of files for training data"""
        file_types = {
            'excel': [('Excel files', '*.xlsx *.xls *.csv')],
            'pdf': [('PDF files', '*.pdf')],
            'word': [('Word files', '*.docx *.doc')],
            'image': [('Image files', '*.png *.jpg *.jpeg *.bmp *.tiff')]
        }
        
        files = filedialog.askopenfilenames(
            title=f"Select {file_type} files",
            filetypes=file_types[file_type]
        )
        
        if files:
            self.update_training_status(f"📂 Starting to process {len(files)} {file_type} file(s)...")
            threading.Thread(
                target=self.process_files_worker,
                args=(files, file_type),
                daemon=True
            ).start()
            
    def process_files_worker(self, files, file_type):
        """Worker thread for processing files with Vietnamese support"""
        for file_path in files:
            try:
                self.update_training_status(f"🔄 Processing: {os.path.basename(file_path)}")
                
                if file_type == 'excel':
                    self.process_excel_file(file_path)
                elif file_type == 'pdf':
                    self.process_pdf_file(file_path)
                elif file_type == 'word':
                    self.process_word_file(file_path)
                elif file_type == 'image':
                    self.process_image_file(file_path)
                    
                self.update_training_status(f"✅ Completed: {os.path.basename(file_path)}")
                
            except Exception as e:
                self.update_training_status(f"❌ Error processing {os.path.basename(file_path)}: {e}")
    
    def add_sample_products(self):
        """Add Vietnamese sample products with error handling"""
        sample_products = [
            {
                'name': 'Gaming Laptop Pro X1',
                'name_vietnamese': 'Laptop Gaming Pro X1',
                'description': 'High-performance gaming laptop with RTX 4070, 16GB RAM, 1TB SSD, perfect for gaming and professional work',
                'description_vietnamese': 'Laptop gaming hiệu suất cao với RTX 4070, RAM 16GB, SSD 1TB, hoàn hảo cho gaming và công việc chuyên nghiệp',
                'category': 'Laptops',
                'category_vietnamese': 'Laptop',
                'price': 1899.99,
                'features': 'RTX 4070, Intel i7, 16GB RAM, 1TB SSD, 15.6" 144Hz display',
                'features_vietnamese': 'RTX 4070, Intel i7, RAM 16GB, SSD 1TB, màn hình 15.6" 144Hz'
            },
            {
                'name': 'Wireless Gaming Mouse RGB Pro',
                'name_vietnamese': 'Chuột Gaming Không dây RGB Pro',
                'description': 'High-precision wireless gaming mouse with RGB lighting, 12000 DPI, programmable buttons',
                'description_vietnamese': 'Chuột gaming không dây độ chính xác cao với đèn RGB, 12000 DPI, nút bấm có thể lập trình',
                'category': 'Gaming Accessories',
                'category_vietnamese': 'Phụ kiện Gaming',
                'price': 79.99,
                'features': '12000 DPI, RGB lighting, 7 programmable buttons, 70-hour battery',
                'features_vietnamese': '12000 DPI, đèn RGB, 7 nút lập trình, pin 70 giờ'
            },
            {
                'name': 'Mechanical Gaming Keyboard',
                'name_vietnamese': 'Bàn phím Gaming Cơ',
                'description': 'Premium mechanical gaming keyboard with RGB backlighting and blue switches',
                'description_vietnamese': 'Bàn phím gaming cơ cao cấp với đèn nền RGB và switch xanh',
                'category': 'Gaming Accessories',
                'category_vietnamese': 'Phụ kiện Gaming',
                'price': 129.99,
                'features': 'Blue mechanical switches, RGB per-key lighting, aluminum frame',
                'features_vietnamese': 'Switch cơ xanh, đèn RGB từng phím, khung nhôm'
            },
            {
                'name': 'Gaming Headset 7.1 Surround',
                'name_vietnamese': 'Tai nghe Gaming 7.1 Surround',
                'description': '7.1 surround sound gaming headset with noise-canceling microphone',
                'description_vietnamese': 'Tai nghe gaming âm thanh vòm 7.1 với microphone chống ồn',
                'category': 'Audio',
                'category_vietnamese': 'Âm thanh',
                'price': 89.99,
                'features': '7.1 surround sound, noise-canceling mic, comfortable padding',
                'features_vietnamese': 'Âm thanh vòm 7.1, mic chống ồn, đệm êm ái'
            },
            {
                'name': 'Business Laptop UltraSlim',
                'name_vietnamese': 'Laptop Văn phòng UltraSlim',
                'description': 'Lightweight business laptop with long battery life, perfect for professionals',
                'description_vietnamese': 'Laptop văn phòng siêu mỏng với thời lượng pin dài, hoàn hảo cho dân văn phòng',
                'category': 'Laptops',
                'category_vietnamese': 'Laptop',
                'price': 899.99,
                'features': 'Intel i5, 8GB RAM, 512GB SSD, 14" display, 12-hour battery',
                'features_vietnamese': 'Intel i5, RAM 8GB, SSD 512GB, màn hình 14", pin 12 giờ'
            }
        ]
        
        try:
            cursor = self.conn.cursor()
            for product in sample_products:
                try:
                    # Generate embedding safely
                    english_text = f"{product['name']} {product['description']} {product.get('features', '')}"
                    vietnamese_text = f"{product['name_vietnamese']} {product['description_vietnamese']} {product.get('features_vietnamese', '')}"
                    combined_text = f"{english_text} {vietnamese_text}"
                    
                    embedding_blob = None
                    if self.embedding_model:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.logger.warning(f"Embedding generation error for {product['name']}: {e}")
                    
                    cursor.execute('''
                        INSERT OR IGNORE INTO products 
                        (name, name_vietnamese, description, description_vietnamese, 
                         category, category_vietnamese, price, features, features_vietnamese,
                         embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        product['name'], product['name_vietnamese'],
                        product['description'], product['description_vietnamese'],
                        product['category'], product['category_vietnamese'],
                        product['price'], product.get('features', ''), product.get('features_vietnamese', ''),
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                except Exception as e:
                    self.logger.error(f"Error inserting product {product.get('name', 'unknown')}: {e}")
                    continue
            
            self.conn.commit()
            print("✅ Sample products added successfully")
            
        except Exception as e:
            self.logger.error(f"Error adding sample products: {e}")
    
    def view_products(self):
        """View products in database with Vietnamese support"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT name, name_vietnamese, description, description_vietnamese, 
                       category, price, source_file, embedding 
                FROM products 
                ORDER BY created_at DESC
            """)
            products = cursor.fetchall()
            
            self.product_display.config(state=tk.NORMAL)
            self.product_display.delete(1.0, tk.END)
            
            if products:
                self.product_display.insert(tk.END, f"📦 Found {len(products)} products in database:\n\n")
                
                for i, product in enumerate(products, 1):
                    name = product[0] or product[1] or "Unnamed Product"
                    name_vi = product[1] or ""
                    desc = product[2] or product[3] or "No description"
                    category = product[4] or "General"
                    price = product[5] or 0
                    source = os.path.basename(product[6]) if product[6] else "Manual"
                    embedding_status = "✅" if product[7] else "❌"
                    
                    self.product_display.insert(tk.END, f"{i}. {name} {embedding_status}\n")
                    if name_vi:
                        self.product_display.insert(tk.END, f"   Vietnamese: {name_vi}\n")
                    self.product_display.insert(tk.END, f"   Description: {desc[:100]}...\n")
                    self.product_display.insert(tk.END, f"   Category: {category} | Price: ${price} | Source: {source}\n\n")
            else:
                self.product_display.insert(tk.END, "📦 No products found in database.\n\n")
                self.product_display.insert(tk.END, "💡 To add products:\n")
                self.product_display.insert(tk.END, "• Click 'Add Sample Products' for demo data\n")
                self.product_display.insert(tk.END, "• Use the Training tab to upload Excel/CSV files\n")
                self.product_display.insert(tk.END, "• Use PDFs/Word docs for knowledge base\n")
                
            self.product_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing products: {e}")
    
    def test_search(self):
        """Test search functionality with Vietnamese support"""
        test_queries = [
            "gaming laptop", "laptop gaming",
            "business computer", "máy tính văn phòng", 
            "wireless mouse", "chuột không dây",
            "cheap laptop under $1500", "laptop rẻ dưới 1500 đô",
            "RGB accessories", "phụ kiện RGB",
            "4K monitor", "màn hình 4K",
            "mechanical keyboard", "bàn phím cơ",
            "laptop for programming", "laptop lập trình"
        ]
        
        self.product_display.config(state=tk.NORMAL)
        self.product_display.delete(1.0, tk.END)
        self.product_display.insert(tk.END, "🔍 Testing search functionality...\n")
        self.product_display.insert(tk.END, f"Threshold: {self.config['search_config']['local_similarity_threshold']}\n\n")
        
        for query in test_queries:
            self.product_display.insert(tk.END, f"Query: '{query}'\n")
            results = self.search_local_database(query)
            
            if results:
                for result in results[:3]:  # Show top 3
                    self.product_display.insert(tk.END, f"  ✓ {result['name']} (similarity: {result['similarity']:.3f})\n")
            else:
                self.product_display.insert(tk.END, "  ✗ No results found\n")
            
            self.product_display.insert(tk.END, "\n")
            
        self.product_display.config(state=tk.DISABLED)
    
    def regenerate_all_embeddings(self):
        """Regenerate all embeddings for products and knowledge base"""
        result = messagebox.askyesno(
            "Regenerate Embeddings", 
            "This will regenerate all embeddings for products and knowledge base entries.\n\n"
            "This may take several minutes depending on the amount of data.\n\n"
            "Continue?"
        )
        
        if result:
            threading.Thread(target=self.regenerate_embeddings_worker, daemon=True).start()
            
    def regenerate_embeddings_worker(self):
        """Worker thread for regenerating embeddings"""
        try:
            if not self.embedding_model:
                self.update_training_status("❌ Embedding model not available")
                return
                
            cursor = self.conn.cursor()
            
            # Regenerate product embeddings
            self.update_training_status("🔄 Regenerating product embeddings...")
            cursor.execute("""
                SELECT id, name, name_vietnamese, description, description_vietnamese, 
                       features, features_vietnamese, specifications, specifications_vietnamese 
                FROM products
            """)
            products = cursor.fetchall()
            
            for i, (product_id, name, name_vi, desc, desc_vi, feat, feat_vi, spec, spec_vi) in enumerate(products):
                try:
                    # Combine all text
                    text_parts = [name or '', name_vi or '', desc or '', desc_vi or '', 
                                feat or '', feat_vi or '', spec or '', spec_vi or '']
                    combined_text = ' '.join(part for part in text_parts if part.strip())
                    
                    if combined_text.strip():
                        embedding = self.embedding_model.encode([combined_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                        
                        cursor.execute("""
                            UPDATE products 
                            SET embedding = ?, updated_at = ?
                            WHERE id = ?
                        """, (embedding_blob, datetime.now().isoformat(), product_id))
                        
                        if (i + 1) % 10 == 0:
                            self.update_training_status(f"📦 Processed {i + 1}/{len(products)} products...")
                            
                except Exception as e:
                    self.update_training_status(f"⚠️ Error processing product {name}: {e}")
                    continue
                    
            # Regenerate knowledge base embeddings
            self.update_training_status("🔄 Regenerating knowledge base embeddings...")
            cursor.execute("SELECT id, content FROM knowledge_base")
            knowledge_entries = cursor.fetchall()
            
            for i, (kb_id, content) in enumerate(knowledge_entries):
                try:
                    if content and content.strip():
                        embedding = self.embedding_model.encode([content[:5000]])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                        
                        cursor.execute("""
                            UPDATE knowledge_base 
                            SET embedding = ?
                            WHERE id = ?
                        """, (embedding_blob, kb_id))
                        
                        if (i + 1) % 5 == 0:
                            self.update_training_status(f"📚 Processed {i + 1}/{len(knowledge_entries)} knowledge entries...")
                            
                except Exception as e:
                    self.update_training_status(f"⚠️ Error processing knowledge entry {kb_id}: {e}")
                    continue
                    
            self.conn.commit()
            self.update_training_status(f"✅ Successfully regenerated embeddings for {len(products)} products and {len(knowledge_entries)} knowledge entries!")
            
        except Exception as e:
            self.update_training_status(f"❌ Error regenerating embeddings: {e}")
    
    def clear_training_data(self):
        """Clear all training data with confirmation"""
        result = messagebox.askyesno(
            "Clear Training Data", 
            "Are you sure you want to clear all training data?\n\n"
            "This will delete:\n"
            "• All products\n"
            "• All knowledge base entries\n"
            "• All embeddings\n\n"
            "This action cannot be undone!"
        )
        
        if result:
            try:
                cursor = self.conn.cursor()
                cursor.execute("DELETE FROM products")
                cursor.execute("DELETE FROM knowledge_base")
                self.conn.commit()
                
                self.update_training_status("🧹 All training data cleared successfully")
                messagebox.showinfo("Success", "All training data has been cleared.")
                
            except Exception as e:
                self.update_training_status(f"❌ Error clearing data: {e}")
                messagebox.showerror("Error", f"Error clearing training data: {e}")
    
    def show_training_stats(self):
        """Show training data statistics"""
        try:
            cursor = self.conn.cursor()
            
            # Get product statistics
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(DISTINCT category) FROM products")
            category_count = cursor.fetchone()[0]
            
            # Get knowledge base statistics
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base WHERE embedding IS NOT NULL")
            knowledge_with_embeddings = cursor.fetchone()[0]
            
            # Display statistics
            self.training_status.config(state=tk.NORMAL)
            self.training_status.delete(1.0, tk.END)
            
            self.training_status.insert(tk.END, "=== TRAINING DATA STATISTICS ===\n\n")
            self.training_status.insert(tk.END, f"📦 Products: {product_count}\n")
            self.training_status.insert(tk.END, f"🔗 Products with embeddings: {products_with_embeddings}/{product_count}\n")
            self.training_status.insert(tk.END, f"📂 Categories: {category_count}\n\n")
            self.training_status.insert(tk.END, f"📚 Knowledge base entries: {knowledge_count}\n")
            self.training_status.insert(tk.END, f"🔗 Knowledge with embeddings: {knowledge_with_embeddings}/{knowledge_count}\n\n")
            
            # Show recent additions
            cursor.execute("""
                SELECT name, created_at, source_file 
                FROM products 
                ORDER BY created_at DESC 
                LIMIT 5
            """)
            recent_products = cursor.fetchall()
            
            if recent_products:
                self.training_status.insert(tk.END, "📈 Recent Products:\n")
                for name, created_at, source_file in recent_products:
                    source = os.path.basename(source_file) if source_file else "Manual"
                    self.training_status.insert(tk.END, f"  • {name} ({source})\n")
                    
            self.training_status.config(state=tk.DISABLED)
            
        except Exception as e:
            self.update_training_status(f"❌ Error showing statistics: {e}")
    
    def refresh_analytics(self):
        """Refresh analytics display with Vietnamese language support"""
        try:
            cursor = self.conn.cursor()
            
            # Get conversation statistics
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_conversations,
                    AVG(response_time) as avg_response_time,
                    data_source,
                    COUNT(*) as count_by_source
                FROM conversations 
                GROUP BY data_source
            """)
            
            source_stats = cursor.fetchall()
            
            # Get recent conversations
            cursor.execute("""
                SELECT timestamp, user_input, user_language, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 10
            """)
            
            recent_convs = cursor.fetchall()
            
            # Get product and knowledge statistics
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            # Language usage statistics
            cursor.execute("""
                SELECT user_language, COUNT(*) as count
                FROM conversations 
                WHERE user_language IS NOT NULL
                GROUP BY user_language
            """)
            language_stats = cursor.fetchall()
            
            # Display analytics
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            # Summary
            self.analytics_display.insert(tk.END, "=== VIETNAMESE AI CHATBOT ANALYTICS ===\n\n")
            self.analytics_display.insert(tk.END, f"📦 Products in Database: {product_count}\n")
            self.analytics_display.insert(tk.END, f"🔗 Products with Embeddings: {products_with_embeddings}/{product_count}\n")
            self.analytics_display.insert(tk.END, f"📚 Knowledge Base Entries: {knowledge_count}\n")
            
            # Language usage
            self.analytics_display.insert(tk.END, f"\n=== LANGUAGE USAGE ===\n")
            for lang, count in language_stats:
                lang_name = "Vietnamese" if lang == 'vi' else "English" if lang == 'en' else lang
                self.analytics_display.insert(tk.END, f"{lang_name}: {count} conversations\n")
            
            # Data source statistics
            self.analytics_display.insert(tk.END, "\n=== RESPONSE SOURCES ===\n")
            for stat in source_stats:
                if len(stat) >= 4:
                    self.analytics_display.insert(tk.END, f"{stat[2]}: {stat[3]} responses\n")
                
            # Performance metrics
            if source_stats:
                total_conversations = sum(stat[3] for stat in source_stats if len(stat) >= 4)
                if total_conversations > 0:
                    avg_response_time = sum(stat[1] * stat[3] for stat in source_stats if stat[1] and len(stat) >= 4) / total_conversations
                    self.analytics_display.insert(tk.END, f"\nTotal Conversations: {total_conversations}\n")
                    self.analytics_display.insert(tk.END, f"Average Response Time: {avg_response_time:.2f} seconds\n")
            
            # Recent conversations
            self.analytics_display.insert(tk.END, "\n=== RECENT CONVERSATIONS ===\n")
            for conv in recent_convs:
                if len(conv) >= 5:
                    timestamp, user_input, user_lang, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    self.analytics_display.insert(tk.END, f"{timestamp} {lang_flag} - {data_source} ({response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input[:100]}...\n\n")
                
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error refreshing analytics: {e}")
    
    def export_analytics(self):
        """Export analytics data to CSV with Vietnamese support"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT * FROM conversations")
            
            # Get column names dynamically
            column_names = [description[0] for description in cursor.description]
            
            df = pd.DataFrame(cursor.fetchall(), columns=column_names)
            
            export_path = f"vietnamese_chatbot_analytics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            df.to_csv(export_path, index=False, encoding='utf-8')
            
            messagebox.showinfo("Success", f"Analytics exported to: {export_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error exporting analytics: {e}")
    
    def view_conversation_history(self):
        """View detailed conversation history"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT timestamp, user_input, user_language, bot_response, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 20
            """)
            conversations = cursor.fetchall()
            
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            self.analytics_display.insert(tk.END, "=== CONVERSATION HISTORY (Last 20) ===\n\n")
            
            for i, conv in enumerate(conversations, 1):
                if len(conv) >= 6:
                    timestamp, user_input, user_lang, bot_response, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    
                    self.analytics_display.insert(tk.END, f"{i}. [{timestamp}] {lang_flag} ({data_source}, {response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input}\n")
                    self.analytics_display.insert(tk.END, f"Bot: {bot_response[:200]}...\n\n")
            
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing conversation history: {e}")
    
    def run(self):
        """Start the application with comprehensive error handling"""
        try:
            # Add sample products
            print("📦 Adding sample products...")
            self.add_sample_products()
            
            # Display welcome message
            self.display_welcome_message()
            
            # Start the GUI
            print("🚀 Starting GUI...")
            self.root.mainloop()
            
        except Exception as e:
            print(f"❌ Error starting application: {e}")
            self.logger.error(f"Application start error: {e}")
    
    def display_welcome_message(self):
        """Display welcome message with system status"""
        try:
            if self.current_language == 'vi':
                welcome_msg = f"""🤖 Chào mừng đến với Trợ lý AI Bán hàng!

Trạng thái hệ thống:
✅ Hỗ trợ tiếng Việt và tiếng Anh
{'✅' if self.text_generator else '❌'} Mô hình AI: {'Hoạt động' if self.text_generator else 'Không khả dụng'}
{'✅' if self.embedding_model else '❌'} Tìm kiếm thông minh: {'Hoạt động' if self.embedding_model else 'Không khả dụng'}
{'✅' if self.voice_enabled else '❌'} Nhận diện giọng nói: {'Có' if self.voice_enabled else 'Không có'}
{'✅' if self.tts_enabled else '❌'} Chuyển văn bản thành giọng nói: {'Có' if self.tts_enabled else 'Không có'}

Tính năng có sẵn:
📦 Tìm kiếm sản phẩm thông minh
💬 Ngữ cảnh cuộc trò chuyện
🌐 Hỗ trợ đa ngôn ngữ
📚 Xử lý file Excel, PDF, Word
🖼️ OCR từ hình ảnh
{f'🎤 Nhập liệu bằng giọng nói' if self.voice_enabled else ''}
{f'🔊 Đọc phản hồi' if self.tts_enabled else ''}

Hãy thử đặt câu hỏi như:
• "Tôi muốn mua laptop gaming"
• "Bạn có chuột không dây nào không?"
• "Giá của laptop là bao nhiêu?"
• "So sánh các sản phẩm gaming"

{f'🎤 Nhấp nút "Nói" để sử dụng giọng nói' if self.voice_enabled else '📝 Sử dụng bàn phím để nhập tin nhắn'}
{f'🔧 Một số tính năng AI có thể bị hạn chế do lỗi tải mô hình' if not self.text_generator else ''}"""
            else:
                welcome_msg = f"""🤖 Welcome to the AI Sales Assistant!

System Status:
✅ Vietnamese and English support
{'✅' if self.text_generator else '❌'} AI Model: {'Available' if self.text_generator else 'Not available'}
{'✅' if self.embedding_model else '❌'} Smart Search: {'Available' if self.embedding_model else 'Not available'}
{'✅' if self.voice_enabled else '❌'} Voice recognition: {'Available' if self.voice_enabled else 'Not available'}
{'✅' if self.tts_enabled else '❌'} Text-to-speech: {'Available' if self.tts_enabled else 'Not available'}

Available features:
📦 Smart product search
💬 Conversation context
🌐 Multi-language support
📚 Excel, PDF, Word processing
🖼️ Image OCR
{f'🎤 Voice input' if self.voice_enabled else ''}
{f'🔊 Voice output' if self.tts_enabled else ''}

Try asking questions like:
• "I want to buy a gaming laptop"
• "What wireless mice do you have?"
• "How much does the laptop cost?"
• "Compare gaming products"

{f'🎤 Click "Speak" button to use voice input' if self.voice_enabled else '📝 Use keyboard to type messages'}
{f'🔧 Some AI features may be limited due to model loading errors' if not self.text_generator else ''}"""

            self.display_message("Hệ thống / System", welcome_msg, "system")
        except Exception as e:
            print(f"Welcome message error: {e}")
        
    def __del__(self):
        """Cleanup with error handling"""
        try:
            if hasattr(self, 'conn'):
                self.conn.close()
        except:
            pass
        
        try:
            if hasattr(self, 'tts_engine') and self.tts_engine:
                self.tts_engine.stop()
        except:
            pass


def create_config_file():
    """Create a sample configuration file"""
    config_content = """# Vietnamese AI Sales ChatBot Configuration
version: '2.0'
environment: 'production'

# AI Models Configuration
ai_models:
  primary_llm: 'microsoft/DialoGPT-medium'
  fallback_llm: 'microsoft/DialoGPT-small'
  embedding_model: 'sentence-transformers/all-MiniLM-L6-v2'

# Language Support
language_config:
  default_language: 'vi'
  supported_languages: ['vi', 'en']
  auto_detect_language: true
  translate_responses: true
  voice_language: 'vi-VN'

# Voice Configuration
voice_config:
  enable_voice_input: true
  enable_voice_output: true
  voice_rate: 150
  voice_volume: 0.8
  graceful_fallback: true

# GPU Configuration
gpu_config:
  primary_device: 'cuda:0'
  use_quantization: true
  mixed_precision: true
  max_memory_per_gpu: 0.85
  batch_size: 2

# Search Configuration
search_config:
  local_similarity_threshold: 0.1
  enable_google_search: true
  max_google_results: 3
  search_timeout: 10

# Performance Settings
performance:
  max_response_length: 150
  temperature: 0.7
  top_p: 0.9
  repetition_penalty: 1.1
  do_sample: true
  num_return_sequences: 1

# Analytics
analytics:
  track_conversations: true
  track_language_usage: true
  track_voice_usage: true
  track_errors: true
"""
    
    try:
        with open('chatbot_config.yaml', 'w', encoding='utf-8') as f:
            f.write(config_content)
        print("✅ Configuration file created: chatbot_config.yaml")
    except Exception as e:
        print(f"❌ Error creating config file: {e}")


def main():
    """Main function with enhanced error handling and setup"""
    print("""
    ===============================================
    🚀 Vietnamese Voice-Enabled AI Sales ChatBot
    ===============================================
    
    🔧 Features:
    ✅ Vietnamese + English language support
    ✅ Voice recognition (if microphone available)
    ✅ Text-to-speech output (if TTS available)
    ✅ Smart product search with embeddings
    ✅ Excel, PDF, Word document processing
    ✅ Image OCR with Vietnamese support
    ✅ Comprehensive error handling
    ✅ Graceful fallback when components unavailable
    
    📋 Requirements:
    Required: transformers, sentence-transformers, torch, tkinter
    Optional (for voice): SpeechRecognition, pyttsx3, pyaudio
    Optional (for translation): googletrans, langdetect
    
    🔧 Installation:
    pip install torch transformers sentence-transformers
    pip install SpeechRecognition pyttsx3 pyaudio  # For voice
    pip install googletrans==4.0.0rc1 langdetect   # For translation
    
    Loading chatbot...
    """)
    
    try:
        # Check for required directories
        os.makedirs('data', exist_ok=True)
        os.makedirs('logs', exist_ok=True)
        
        # Create config file if it doesn't exist
        if not os.path.exists('chatbot_config.yaml'):
            print("📝 Creating configuration file...")
            create_config_file()
        
        # Create and run the chatbot
        print("🤖 Initializing chatbot...")
        chatbot = VietnameseVoiceChatBot()
        
        print("🎬 Starting application...")
        chatbot.run()
        
    except KeyboardInterrupt:
        print("\n👋 ChatBot stopped by user")
    except Exception as e:
        print(f"❌ Error starting chatbot: {e}")
        print("\n🔧 Troubleshooting tips:")
        print("1. Ensure all required packages are installed")
        print("2. Check if CUDA is available for GPU acceleration")
        print("3. Verify microphone/audio device connections")
        print("4. Check internet connection for model downloads")
        print("5. Try running without voice features if audio fails")
        
        import traceback
        traceback.print_exc()
        
        input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()