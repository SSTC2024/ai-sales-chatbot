import os
import json
import sqlite3
import pandas as pd
import numpy as np
from datetime import datetime
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox, font
import threading
import queue
import time
import logging
from pathlib import Path
import re
import yaml
# NEW IMPORTS FOR AI GENERATING SALES PERSON IN WEB CHATBOT
from flask import Flask, render_template, request, jsonify
from flask_cors import CORS

# AI/ML imports
import torch
from transformers import (
    AutoTokenizer, AutoModelForCausalLM, 
    BitsAndBytesConfig, pipeline
)
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Web search imports
import requests
from bs4 import BeautifulSoup
from googlesearch import search
import urllib.parse

# Document processing
import docx
from PyPDF2 import PdfReader
from openpyxl import load_workbook
import easyocr
import socket

# Language detection and translation
try:
    from googletrans import Translator, LANGUAGES
    import langdetect
    TRANSLATION_AVAILABLE = True
    print("✅ Translation libraries loaded successfully")
except ImportError:
    TRANSLATION_AVAILABLE = False
    print("⚠️ Translation libraries not available. Install with: pip install googletrans==4.0.0rc1 langdetect")
    
# Web interface imports
from flask import Flask, render_template, request, jsonify
import argparse
import socket

class VietnameseAISalesBot:
    """
    Enhanced AI Sales ChatBot with Vietnamese language support (Text-only version)
    Features robust error handling for response generation and document processing
    """
    
    def __init__(self, start_gui=False):  # ADD PARAMETER
        self.load_config()
        self.setup_logging()
        self.setup_language_support()
        self.setup_database()
        self.initialize_ai_models()
        self.conversation_context = []
        self.conversation_summary = ""
        if start_gui:  # ADD THIS CONDITION
            self.setup_gui()
    
    def load_config(self):
        """Load configuration with Vietnamese language settings"""
        try:
            with open('chatbot_config.yaml', 'r', encoding='utf-8') as file:
                self.config = yaml.safe_load(file)
            print("✅ Configuration loaded from chatbot_config.yaml")
            
            self.config = self.update_config_for_vietnamese()
            
        except FileNotFoundError:
            print("⚠️ chatbot_config.yaml not found, using default Vietnamese settings")
            self.config = self.get_default_vietnamese_config()
        except Exception as e:
            print(f"❌ Error loading config: {e}")
            self.config = self.get_default_vietnamese_config()
    
    def update_config_for_vietnamese(self):
        """Update existing config to support Vietnamese features"""
        config = self.config.copy()
        
        if 'language_config' not in config:
            config['language_config'] = {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True,
                'translate_responses': True
            }
        
        if 'conversation_flows' not in config:
            config['conversation_flows'] = {}
            
        config['conversation_flows'].update({
            'greeting_vietnamese': {
                'triggers': ['xin chào', 'chào', 'chào bạn', 'chào anh', 'chào chị'],
                'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng của bạn. Tôi có thể giúp bạn tìm kiếm sản phẩm phù hợp với nhu cầu. Bạn đang tìm kiếm gì hôm nay?',
                'next_stage': 'needs_assessment'
            }
        })
        
        config['vietnamese_templates'] = {
            'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp với tiêu chí của bạn.',
            'out_of_stock': 'Rất tiếc, {product_name} hiện đang hết hàng.',
            'processing': 'Đang xử lý yêu cầu của bạn...',
            'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn. Vui lòng thử lại.',
            'model_error': 'Mô hình AI không khả dụng. Đang sử dụng phản hồi dự phòng.',
            'generation_error': 'Có lỗi khi tạo phản hồi. Vui lòng thử lại.',
            'empty_response': 'Tôi hiểu câu hỏi của bạn nhưng cần thêm thông tin để trả lời tốt hơn.',
            'connection_error': 'Lỗi kết nối. Vui lòng kiểm tra mạng và thử lại.'
        }
        
        return config
    
    def get_default_vietnamese_config(self):
        """Get default configuration with Vietnamese support and error handling"""
        return {
            'version': '2.0',
            'environment': 'production',
            'ai_models': {
                'primary_llm': 'microsoft/DialoGPT-medium',
                'fallback_llm': 'microsoft/DialoGPT-small',
                'embedding_model': 'sentence-transformers/all-MiniLM-L6-v2'
            },
            'language_config': {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True,
                'translate_responses': True
            },
            'gpu_config': {
                'primary_device': 'cuda:0' if torch.cuda.is_available() else 'cpu',
                'use_quantization': True,
                'mixed_precision': True,
                'max_memory_per_gpu': 0.85,
                'batch_size': 2
            },
            'search_config': {
                'local_similarity_threshold': 0.1,
                'enable_google_search': True,
                'max_google_results': 3,
                'search_timeout': 10
            },
            'performance': {
                'max_response_length': 150,
                'temperature': 0.7,
                'top_p': 0.9,
                'repetition_penalty': 1.1,
                'do_sample': True,
                'num_return_sequences': 1
            },
            'conversation_flows': {
                'greeting_vietnamese': {
                    'triggers': ['xin chào', 'chào', 'chào bạn'],
                    'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng của bạn.'
                },
                'greeting': {
                    'triggers': ['hello', 'hi', 'hey'],
                    'response_template': 'Hello! I am your AI sales assistant.'
                }
            },
            'vietnamese_templates': {
                'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp.',
                'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn.',
                'model_error': 'Mô hình AI không khả dụng.',
                'generation_error': 'Có lỗi khi tạo phản hồi.',
                'empty_response': 'Tôi cần thêm thông tin để trả lời tốt hơn.'
            },
            'analytics': {
                'track_conversations': True,
                'track_language_usage': True,
                'track_errors': True
            }
        }
    
    def setup_language_support(self):
        """Initialize language detection and translation"""
        self.current_language = self.config['language_config']['default_language']
        
        if TRANSLATION_AVAILABLE:
            try:
                self.translator = Translator()
                print("✅ Translation service initialized")
            except Exception as e:
                print(f"⚠️ Translation service initialization failed: {e}")
                self.translator = None
        else:
            self.translator = None
    
    def detect_language(self, text):
        """Detect language of input text with error handling"""
        if not TRANSLATION_AVAILABLE or not text or not text.strip():
            return 'vi'
        
        try:
            detected = langdetect.detect(text)
            if detected in self.config['language_config']['supported_languages']:
                return detected
            return 'vi'
        except Exception as e:
            print(f"Language detection error: {e}")
            return 'vi'
    
    def translate_text(self, text, target_language):
        """Translate text to target language with error handling"""
        if not self.translator or not text or not text.strip():
            return text
        
        try:
            translated = self.translator.translate(text, dest=target_language)
            return translated.text if translated and translated.text else text
        except Exception as e:
            print(f"Translation error: {e}")
            return text
        
    def setup_logging(self):
        """Setup logging configuration with Vietnamese support"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('chatbot.log', encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
        
    def setup_database(self):
        """Initialize SQLite database with Vietnamese support"""
        self.conn = sqlite3.connect('chatbot_data.db', check_same_thread=False)
        self.conn.execute("PRAGMA encoding = 'UTF-8'")
        
        cursor = self.conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS products (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                name_vietnamese TEXT,
                description TEXT,
                description_vietnamese TEXT,
                category TEXT,
                category_vietnamese TEXT,
                price REAL,
                features TEXT,
                features_vietnamese TEXT,
                specifications TEXT,
                specifications_vietnamese TEXT,
                availability TEXT,
                availability_vietnamese TEXT,
                source_file TEXT,
                embedding BLOB,
                created_at TEXT,
                updated_at TEXT
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT,
                user_input TEXT,
                user_language TEXT,
                bot_response TEXT,
                bot_language TEXT,
                intent TEXT,
                confidence REAL,
                data_source TEXT,
                response_time REAL
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS knowledge_base (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                topic TEXT,
                content TEXT,
                source TEXT,
                embedding BLOB,
                created_at TEXT
            )
        ''')
        
        try:
            cursor.execute('ALTER TABLE conversations ADD COLUMN error_type TEXT')
            print("✅ Added error_type column to conversations table")
        except sqlite3.OperationalError:
            pass
        
        try:
            cursor.execute('ALTER TABLE conversations ADD COLUMN error_message TEXT')
            print("✅ Added error_message column to conversations table")
        except sqlite3.OperationalError:
            pass
        
        self.conn.commit()
        print("✅ Database schema updated successfully")
        
    def initialize_ai_models(self):
        """Initialize AI models with enhanced error handling"""
        self.logger.info("Initializing AI models from configuration...")
        
        try:
            self.setup_devices()
            self.load_language_model()
            self.load_embedding_model()
            self.load_ocr_model()
            
            self.logger.info("✅ All AI models loaded successfully!")
            
        except Exception as e:
            self.logger.error(f"❌ Error initializing AI models: {e}")
            self.text_generator = None
            self.embedding_model = None
            self.ocr_reader = None
    
    def setup_devices(self):
        """Device setup using configuration with fallback"""
        try:
            if torch.cuda.is_available():
                gpu_count = torch.cuda.device_count()
                print(f"Found {gpu_count} GPU(s)")
                
                best_gpu = 0
                max_memory = 0
                
                for i in range(gpu_count):
                    props = torch.cuda.get_device_properties(i)
                    memory_gb = props.total_memory / (1024**3)
                    print(f"GPU {i}: {props.name} - {memory_gb:.1f}GB")
                    
                    if memory_gb > max_memory:
                        max_memory = memory_gb
                        best_gpu = i
                
                primary_device = self.config['gpu_config'].get('primary_device', f'cuda:{best_gpu}')
                
                self.primary_device = primary_device
                self.secondary_device = f'cuda:{(best_gpu + 1) % gpu_count}' if gpu_count > 1 else self.primary_device
                self.primary_memory_gb = max_memory
                
            else:
                self.primary_device = 'cpu'
                self.secondary_device = 'cpu'
                self.primary_memory_gb = 8
                print("CUDA not available. Using CPU mode.")
        except Exception as e:
            print(f"Device setup error: {e}")
            self.primary_device = 'cpu'
            self.secondary_device = 'cpu'
            self.primary_memory_gb = 8
    
    def load_language_model(self):
        """Load language model with comprehensive error handling"""
        try:
            primary_llm = self.config['ai_models']['primary_llm']
            fallback_llm = self.config['ai_models']['fallback_llm']
            use_quantization_config = self.config['gpu_config']['use_quantization']
            
            primary_memory = getattr(self, 'primary_memory_gb', 8)
            
            if primary_memory >= 12:
                use_quantization = use_quantization_config
                model_name = primary_llm
            else:
                use_quantization = True
                model_name = fallback_llm
            
            print(f"Loading language model: {model_name}")
            
            try:
                self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                if self.tokenizer.pad_token is None:
                    self.tokenizer.pad_token = self.tokenizer.eos_token
                print("✅ Tokenizer loaded successfully")
            except Exception as e:
                print(f"❌ Tokenizer loading failed: {e}")
                self.tokenizer = None
                self.text_generator = None
                return
            
            quantization_config = None
            if use_quantization and torch.cuda.is_available():
                try:
                    quantization_config = BitsAndBytesConfig(
                        load_in_8bit=True,
                        llm_int8_threshold=6.0,
                    )
                except Exception as e:
                    print(f"❌ Quantization config failed: {e}")
                    quantization_config = None
            
            try:
                self.llm_model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    quantization_config=quantization_config,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                    trust_remote_code=True,
                    low_cpu_mem_usage=True
                )
                print("✅ Language model loaded successfully")
            except Exception as e:
                print(f"❌ Model loading failed: {e}")
                self.llm_model = None
                self.text_generator = None
                return
            
            try:
                try:
                    self.text_generator = pipeline(
                        "text-generation",
                        model=self.llm_model,
                        tokenizer=self.tokenizer,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
                    )
                    print("✅ Text generation pipeline created successfully (accelerate mode)")
                except Exception as accelerate_error:
                    print(f"⚠️ Accelerate mode failed, trying with device specification: {accelerate_error}")
                    self.text_generator = pipeline(
                        "text-generation",
                        model=self.llm_model,
                        tokenizer=self.tokenizer,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                        device=0 if torch.cuda.is_available() else -1
                    )
                    print("✅ Text generation pipeline created successfully (device mode)")
            except Exception as e:
                print(f"❌ Pipeline creation failed: {e}")
                self.text_generator = None
            
        except Exception as e:
            print(f"❌ Error loading language model: {e}")
            self.text_generator = None
            
    def load_embedding_model(self):
        """Load embedding model with error handling"""
        try:
            embedding_model_name = self.config['ai_models']['embedding_model']
            print(f"Loading embedding model: {embedding_model_name}")
            
            self.embedding_model = SentenceTransformer(embedding_model_name)
            if torch.cuda.is_available():
                try:
                    self.embedding_model = self.embedding_model.to(self.secondary_device)
                except Exception as e:
                    print(f"⚠️ Could not move embedding model to GPU: {e}")
            print("✅ Embedding model loaded successfully")
        except Exception as e:
            print(f"❌ Error loading embedding model: {e}")
            self.embedding_model = None
            
    def load_ocr_model(self):
        """Load OCR model with error handling"""
        try:
            self.ocr_reader = easyocr.Reader(['en', 'vi'])
            print("✅ OCR model loaded successfully with Vietnamese support")
        except Exception as e:
            print(f"❌ Error loading OCR model: {e}")
            self.ocr_reader = None
    
    def safe_generate_response(self, prompt, user_language='vi'):
        """Safely generate response with comprehensive error handling"""
        try:
            if not self.text_generator or not self.tokenizer:
                self.logger.warning("Text generator not available")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['model_error']
                else:
                    return "I apologize, but the AI text generation system is not available right now."
            
            if not prompt or not isinstance(prompt, str) or not prompt.strip():
                self.logger.warning("Invalid prompt provided")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need more information to provide a helpful response."
            
            performance_config = self.config.get('performance', {})
            
            try:
                with torch.no_grad():
                    if self.tokenizer.pad_token_id is None:
                        self.tokenizer.pad_token_id = self.tokenizer.eos_token_id
                    
                    generated = self.text_generator(
                        prompt,
                        max_new_tokens=performance_config.get('max_response_length', 150),
                        min_new_tokens=10,
                        temperature=performance_config.get('temperature', 0.7),
                        top_p=performance_config.get('top_p', 0.9),
                        repetition_penalty=performance_config.get('repetition_penalty', 1.1),
                        do_sample=performance_config.get('do_sample', True),
                        pad_token_id=self.tokenizer.pad_token_id,
                        num_return_sequences=performance_config.get('num_return_sequences', 1),
                        return_full_text=False,
                        batch_size=1
                    )
                
                if not generated:
                    self.logger.warning("Empty generation result")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['generation_error']
                    else:
                        return "I'm having trouble generating a response. Please try again."
                
                if isinstance(generated, list) and len(generated) > 0:
                    first_result = generated[0]
                    if isinstance(first_result, dict) and 'generated_text' in first_result:
                        response = first_result['generated_text']
                    else:
                        response = str(first_result)
                else:
                    response = str(generated)
                
                if not response or not response.strip():
                    self.logger.warning("Empty response generated")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question but need more details to respond properly."
                
                cleaned_response = self.clean_generated_response(response)
                
                if not cleaned_response or len(cleaned_response.strip()) < 5:
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question. Could you provide more details?"
                
                return cleaned_response
                
            except torch.cuda.OutOfMemoryError:
                self.logger.error("CUDA out of memory")
                if user_language == 'vi':
                    return "Hệ thống đang quá tải. Vui lòng thử lại sau."
                else:
                    return "System is overloaded. Please try again later."
                    
            except Exception as generation_error:
                self.logger.error(f"Text generation error: {generation_error}")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['generation_error']
                else:
                    return "I encountered an error while generating a response. Please try again."
            
        except Exception as e:
            self.logger.error(f"Safe generation error: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having technical difficulties. Please try again."
    
    def clean_generated_response(self, response):
        """Clean up generated response with error handling"""
        try:
            if not response:
                return ""
            
            response = str(response)
            
            response = re.sub(r'<[^>]+>', '', response)
            response = re.sub(r'\[.*?\]', '', response)
            response = re.sub(r'\n+', ' ', response)
            response = re.sub(r'\s+', ' ', response)
            
            response = response.replace('<|endoftext|>', '')
            response = response.replace('</s>', '')
            response = response.replace('<s>', '')
            
            try:
                sentences = response.split('.')
                if len(sentences) > 1 and len(sentences[-1].strip()) < 10:
                    response = '.'.join(sentences[:-1]) + '.'
            except Exception:
                pass
            
            response = response.strip()
            
            if len(response) < 3:
                return ""
                
            return response
            
        except Exception as e:
            self.logger.error(f"Response cleaning error: {e}")
            return str(response) if response else ""
    
    def process_message(self, user_input):
        """Main processing pipeline with Vietnamese support and comprehensive error handling"""
        start_time = time.time()
        
        try:
            if not user_input or not isinstance(user_input, str):
                error_msg = "Vui lòng nhập tin nhắn hợp lệ." if self.current_language == 'vi' else "Please enter a valid message."
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'invalid_input'
                }
            
            user_input = user_input.strip()
            if not user_input:
                error_msg = "Vui lòng nhập tin nhắn." if self.current_language == 'vi' else "Please enter a message."
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'empty_input'
                }
            
            try:
                user_language = self.detect_language(user_input)
            except Exception as e:
                self.logger.warning(f"Language detection error: {e}")
                user_language = self.current_language or 'vi'
            
            context_references = self.check_context_references(user_input, user_language)
            
            local_results = []
            knowledge_results = []
            database_searched = False
            
            try:
                if context_references:
                    search_query = self.build_contextual_search_query(user_input, context_references)
                    self.logger.info(f"Contextual search for: {search_query}")
                else:
                    search_query = user_input
                    
                self.logger.info(f"Searching database for: {search_query}")
                local_results = self.search_local_database(search_query)
                database_searched = True
                
                if not local_results:
                    self.logger.info("No products found, searching knowledge base...")
                    knowledge_results = self.search_knowledge_base(search_query)
                    
            except Exception as e:
                self.logger.error(f"Database search error: {e}")
                local_results = []
                knowledge_results = []
            
            try:
                response = ""
                data_source = "unknown"
                
                if local_results:
                    self.logger.info(f"Found {len(local_results)} products, generating AI response with product context")
                    response = self.generate_natural_response(
                        user_input, 
                        context_data=local_results,
                        data_source="database",
                        user_language=user_language
                    )
                    data_source = "local_database"
                    
                elif knowledge_results:
                    self.logger.info(f"Found {len(knowledge_results)} knowledge entries, generating AI response with knowledge context")
                    response = self.generate_natural_response(
                        user_input,
                        context_data=knowledge_results,
                        data_source="knowledge_base",
                        user_language=user_language
                    )
                    data_source = "knowledge_base"
                    
                else:
                    self.logger.info("No database results found, using AI general knowledge")
                    
                    if database_searched:
                        if user_language == 'vi':
                            no_results_msg = "Tôi đã tìm kiếm trong cơ sở dữ liệu nhưng không tìm thấy sản phẩm phù hợp. Để tôi trả lời dựa trên kiến thức chung:\n\n"
                        else:
                            no_results_msg = "I searched our database but couldn't find matching products. Let me answer based on general knowledge:\n\n"
                    else:
                        no_results_msg = ""
                    
                    ai_response = self.generate_natural_response(
                        user_input, 
                        context_data=None,
                        data_source="general",
                        user_language=user_language
                    )
                    
                    response = no_results_msg + ai_response
                    data_source = "ai_knowledge"
                    
            except Exception as e:
                self.logger.error(f"Response generation error: {e}")
                if user_language == 'vi':
                    response = self.config['vietnamese_templates']['generation_error']
                else:
                    response = "I encountered an error while generating a response. Please try again."
                data_source = "error"
            
            conversation_entry = {
                'user': user_input,
                'bot': response,
                'user_language': user_language,
                'timestamp': datetime.now().isoformat(),
                'data_source': data_source,
                'results_found': len(local_results) + len(knowledge_results),
                'products_mentioned': [p.get('name', '') for p in local_results[:3]] if local_results else [],
                'context_references': context_references
            }
            
            try:
                self.conversation_context.append(conversation_entry)
                
                if len(self.conversation_context) > 20:
                    self.conversation_context = self.conversation_context[-20:]
                    
                if len(self.conversation_context) % 5 == 0:
                    self.update_conversation_summary()
                    
            except Exception as e:
                self.logger.warning(f"Context update error: {e}")
                
            try:
                processing_time = time.time() - start_time
                self.store_conversation(user_input, response, data_source, processing_time, user_language)
            except Exception as e:
                self.logger.warning(f"Conversation storage error: {e}")
            
            self.logger.info(f"Response generated using: {data_source}")
            
            return {
                'response': response,
                'data_source': data_source,
                'local_results_count': len(local_results) + len(knowledge_results),
                'processing_time': time.time() - start_time,
                'user_language': user_language,
                'response_language': user_language,
                'database_searched': database_searched,
                'context_used': bool(context_references)
            }
            
        except Exception as e:
            self.logger.error(f"Message processing error: {e}")
            error_language = getattr(self, 'current_language', 'vi')
            if error_language == 'vi':
                error_msg = self.config['vietnamese_templates']['error']
            else:
                error_msg = "I apologize, but I encountered an error processing your request."
            
            return {
                'response': error_msg,
                'data_source': 'error',
                'local_results_count': 0,
                'processing_time': time.time() - start_time,
                'user_language': error_language,
                'response_language': error_language,
                'error': str(e)
            }
    
    def generate_natural_response(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Generate natural language response with Vietnamese support and error handling"""
        try:
            if not user_input or not isinstance(user_input, str):
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need a question or message to respond to."
            
            user_input = user_input.strip()
            if not user_input:
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "Please provide a question or message."
            
            try:
                if user_language == 'vi':
                    greeting_flow = self.config['conversation_flows'].get('greeting_vietnamese', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Xin chào!')
                else:
                    greeting_flow = self.config['conversation_flows'].get('greeting', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Hello!')
            except Exception as e:
                self.logger.warning(f"Greeting pattern matching error: {e}")
            
            if context_data and len(context_data) > 0:
                if data_source == "database":
                    prompt = self.build_product_aware_prompt(user_input, context_data, user_language)
                    ai_response = self.safe_generate_response(prompt, user_language)
                    
                    if not ai_response or ai_response == self.config['vietnamese_templates']['model_error']:
                        return self.format_product_response(context_data, user_input, user_language)
                    
                    return ai_response
                    
                elif data_source == "knowledge_base":
                    prompt = self.build_knowledge_aware_prompt(user_input, context_data, user_language)
                    ai_response = self.safe_generate_response(prompt, user_language)
                    
                    if not ai_response or ai_response == self.config['vietnamese_templates']['model_error']:
                        return self.format_knowledge_response(context_data, user_input, user_language)
                    
                    return ai_response
            
            prompt = self.build_sales_prompt(user_input, None, "general", user_language)
            response = self.safe_generate_response(prompt, user_language)
            
            if not response or response in [self.config['vietnamese_templates']['model_error'], 
                                           self.config['vietnamese_templates']['generation_error']]:
                if user_language == 'vi':
                    return "Xin lỗi, tôi đang gặp khó khăn trong việc tạo phản hồi. Bạn có thể mô tả chi tiết hơn về sản phẩm bạn đang tìm kiếm không?"
                else:
                    return "I apologize, I'm having trouble generating a response. Could you describe in more detail what product you're looking for?"
            
            return response
            
        except Exception as e:
            self.logger.error(f"Error generating natural response: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having trouble generating a response right now."
    
    def build_product_aware_prompt(self, user_input, products, user_language='vi'):
        """Build a prompt that incorporates product data for natural AI response"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý bán hàng AI thân thiện. Dựa trên câu hỏi của khách hàng và thông tin sản phẩm có sẵn, hãy tạo một phản hồi tự nhiên, hữu ích.

"""
            else:
                prompt = """You are a friendly AI sales assistant. Based on the customer's question and available product information, create a natural, helpful response.

"""
            
            summary_prompt = self.get_conversation_context_prompt(user_language)
            if summary_prompt:
                prompt += summary_prompt
            
            if self.conversation_context and len(self.conversation_context) > 0:
                if user_language == 'vi':
                    prompt += "Cuộc trò chuyện gần đây:\n"
                else:
                    prompt += "Recent conversation:\n"
                
                for turn in self.conversation_context[-3:]:
                    user_msg = turn.get('user', '')[:150]
                    bot_msg = turn.get('bot', '')[:150]
                    prompt += f"User: {user_msg}\nAssistant: {bot_msg}...\n"
                prompt += "\n"
            
            if user_language == 'vi':
                prompt += "Sản phẩm có sẵn:\n"
            else:
                prompt += "Available products:\n"
                
            for product in products[:3]:
                name = product.get('name_vietnamese', product.get('name', '')) if user_language == 'vi' else product.get('name', '')
                desc = product.get('description_vietnamese', product.get('description', '')) if user_language == 'vi' else product.get('description', '')
                price = product.get('price', 0)
                features = product.get('features_vietnamese', product.get('features', '')) if user_language == 'vi' else product.get('features', '')
                
                prompt += f"\n- {name}"
                if desc:
                    prompt += f": {desc[:100]}"
                if price > 0:
                    if user_language == 'vi':
                        prompt += f" (Giá: ${price:,.0f})"
                    else:
                        prompt += f" (Price: ${price:,.2f})"
                if features:
                    prompt += f" - {features[:50]}"
            
            if user_language == 'vi':
                prompt += f"\n\nCâu hỏi khách hàng: {user_input}\n\n"
                prompt += "Hãy trả lời một cách tự nhiên, nhớ ngữ cảnh cuộc trò chuyện, giới thiệu sản phẩm phù hợp và giải thích tại sao chúng đáp ứng nhu cầu của khách hàng. "
                prompt += "Nếu khách hàng đề cập đến sản phẩm đã nói trước đó, hãy nhắc lại thông tin đó:"
            else:
                prompt += f"\n\nCustomer question: {user_input}\n\n"
                prompt += "Provide a natural response that remembers the conversation context, introduces suitable products and explains why they meet the customer's needs. "
                prompt += "If the customer refers to previously mentioned products, recall that information:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Error building product prompt: {e}")
            return self.build_sales_prompt(user_input, products, "database", user_language)
    
    def build_knowledge_aware_prompt(self, user_input, knowledge_entries, user_language='vi'):
        """Build a prompt that incorporates knowledge base data for natural AI response"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý AI thông minh. Dựa trên thông tin từ cơ sở tri thức và câu hỏi của người dùng, hãy tạo một câu trả lời tự nhiên và hữu ích.

"""
            else:
                prompt = """You are an intelligent AI assistant. Based on information from the knowledge base and the user's question, create a natural and helpful response.

"""
            
            summary_prompt = self.get_conversation_context_prompt(user_language)
            if summary_prompt:
                prompt += summary_prompt
            
            if self.conversation_context and len(self.conversation_context) > 0:
                if user_language == 'vi':
                    prompt += "Ngữ cảnh cuộc trò chuyện:\n"
                else:
                    prompt += "Conversation context:\n"
                
                for turn in self.conversation_context[-2:]:
                    user_msg = turn.get('user', '')[:100]
                    prompt += f"User: {user_msg}\n"
                prompt += "\n"
            
            if user_language == 'vi':
                prompt += "Thông tin liên quan:\n"
            else:
                prompt += "Related information:\n"
                
            for entry in knowledge_entries[:2]:
                topic = entry.get('topic', '')
                content = entry.get('content', '')[:200]
                
                if topic:
                    prompt += f"\n{topic}: "
                if content:
                    prompt += f"{content}..."
            
            if user_language == 'vi':
                prompt += f"\n\nCâu hỏi: {user_input}\n\n"
                prompt += "Trả lời dựa trên ngữ cảnh cuộc trò chuyện và thông tin có sẵn:"
            else:
                prompt += f"\n\nQuestion: {user_input}\n\n"
                prompt += "Response based on conversation context and available information:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Error building knowledge prompt: {e}")
            return self.build_sales_prompt(user_input, knowledge_entries, "knowledge_base", user_language)
    
    def build_sales_prompt(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Build an optimized prompt for sales conversations with Vietnamese support"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý AI bán hàng chuyên nghiệp. Hãy trả lời bằng tiếng Việt một cách thân thiện và hữu ích.

"""
            else:
                prompt = """You are a professional AI sales assistant. Respond in English in a friendly and helpful manner.

"""
            
            try:
                if self.conversation_context and len(self.conversation_context) > 0:
                    if user_language == 'vi':
                        prompt += "Lịch sử cuộc trò chuyện:\n"
                    else:
                        prompt += "Conversation history:\n"
                    
                    recent_context = self.conversation_context[-5:]
                    for i, turn in enumerate(recent_context):
                        if isinstance(turn, dict) and 'user' in turn and 'bot' in turn:
                            user_msg = str(turn['user'])[:200]
                            bot_msg = str(turn['bot'])[:200]
                            
                            if user_language == 'vi':
                                prompt += f"\nLượt {i+1}:\n"
                            else:
                                prompt += f"\nTurn {i+1}:\n"
                            
                            prompt += f"User: {user_msg}\nAssistant: {bot_msg}\n"
                    
                    if len(self.conversation_context) > 5:
                        if user_language == 'vi':
                            prompt += f"\n(Đã có {len(self.conversation_context)} lượt trò chuyện trước đó)\n"
                        else:
                            prompt += f"\n(There were {len(self.conversation_context)} previous conversation turns)\n"
                    
                    prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Context building error: {e}")
                
            try:
                if context_data and isinstance(context_data, list) and len(context_data) > 0:
                    if data_source == "database":
                        if user_language == 'vi':
                            prompt += "Sản phẩm liên quan:\n"
                        else:
                            prompt += "Relevant products:\n"
                        
                        for item in context_data[:3]:
                            if isinstance(item, dict):
                                name = item.get('name_vietnamese', item.get('name', '')) if user_language == 'vi' else item.get('name', '')
                                desc = item.get('description_vietnamese', item.get('description', '')) if user_language == 'vi' else item.get('description', '')
                                
                                if name:
                                    prompt += f"- {str(name)[:100]}"
                                    if desc:
                                        prompt += f": {str(desc)[:150]}...\n"
                                    else:
                                        prompt += "\n"
                                        
                                    if item.get('price'):
                                        try:
                                            price_label = "Giá" if user_language == 'vi' else "Price"
                                            prompt += f"  {price_label}: ${float(item['price']):.2f}\n"
                                        except (ValueError, TypeError):
                                            pass
                        prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Product context building error: {e}")
                
            try:
                truncated_input = str(user_input)[:300]
                if user_language == 'vi':
                    prompt += f"Câu hỏi hiện tại của khách hàng: {truncated_input}\n\n"
                    prompt += "Hãy trả lời dựa trên lịch sử cuộc trò chuyện và duy trì ngữ cảnh. "
                    prompt += "Nếu khách hàng đề cập đến điều gì đó đã nói trước đó, hãy nhớ và tham chiếu đến nó.\n"
                    prompt += "Trả lời bằng tiếng Việt:"
                else:
                    prompt += f"Current customer question: {truncated_input}\n\n"
                    prompt += "Please respond based on the conversation history and maintain context. "
                    prompt += "If the customer refers to something mentioned earlier, remember and reference it.\n"
                    prompt += "Response:"
            except Exception as e:
                self.logger.error(f"Question formatting error: {e}")
                if user_language == 'vi':
                    prompt += f"Câu hỏi: {user_input}\nTrả lời:"
                else:
                    prompt += f"Question: {user_input}\nResponse:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Prompt building error: {e}")
            if user_language == 'vi':
                return f"Câu hỏi: {user_input}\nTrả lời bằng tiếng Việt:"
            else:
                return f"Question: {user_input}\nResponse:"
    
    def format_product_response(self, products, user_input, user_language='vi'):
        """Format product search results into a natural response"""
        try:
            if user_language == 'vi':
                response = "Tôi tìm thấy các sản phẩm sau phù hợp với yêu cầu của bạn:\n\n"
                
                for i, product in enumerate(products[:3], 1):
                    name = product.get('name_vietnamese') or product.get('name', '')
                    desc = product.get('description_vietnamese') or product.get('description', '')
                    price = product.get('price', 0)
                    
                    response += f"{i}. **{name}**\n"
                    if desc:
                        response += f"   {desc[:150]}...\n"
                    if price > 0:
                        response += f"   💰 Giá: ${price:,.2f}\n"
                    response += "\n"
                
                response += "Bạn muốn biết thêm thông tin chi tiết về sản phẩm nào?"
            else:
                response = "I found the following products that match your request:\n\n"
                
                for i, product in enumerate(products[:3], 1):
                    name = product.get('name', '')
                    desc = product.get('description', '')
                    price = product.get('price', 0)
                    
                    response += f"{i}. **{name}**\n"
                    if desc:
                        response += f"   {desc[:150]}...\n"
                    if price > 0:
                        response += f"   💰 Price: ${price:,.2f}\n"
                    response += "\n"
                
                response += "Which product would you like to know more about?"
                
            return response
        except Exception as e:
            self.logger.error(f"Error formatting product response: {e}")
            if user_language == 'vi':
                return "Tôi tìm thấy một số sản phẩm phù hợp. Bạn có thể cho tôi biết thêm về nhu cầu của bạn không?"
            else:
                return "I found some matching products. Could you tell me more about your needs?"
    
    def format_knowledge_response(self, knowledge_entries, user_input, user_language='vi'):
        """Format knowledge base search results into a natural response"""
        try:
            if user_language == 'vi':
                response = "Dựa trên thông tin trong cơ sở dữ liệu, tôi có thể chia sẻ:\n\n"
                
                for entry in knowledge_entries[:2]:
                    topic = entry.get('topic', '')
                    content = entry.get('content', '')
                    
                    if topic:
                        response += f"📚 **{topic}**\n"
                    if content:
                        response += f"{content[:300]}...\n\n"
                
                response += "Bạn cần thông tin chi tiết hơn về vấn đề nào?"
            else:
                response = "Based on the information in our knowledge base:\n\n"
                
                for entry in knowledge_entries[:2]:
                    topic = entry.get('topic', '')
                    content = entry.get('content', '')
                    
                    if topic:
                        response += f"📚 **{topic}**\n"
                    if content:
                        response += f"{content[:300]}...\n\n"
                
                response += "Would you like more detailed information on any topic?"
                
            return response
        except Exception as e:
            self.logger.error(f"Error formatting knowledge response: {e}")
            if user_language == 'vi':
                return "Tôi tìm thấy một số thông tin liên quan. Bạn có thể hỏi cụ thể hơn không?"
            else:
                return "I found some related information. Could you be more specific?"
    
    def search_local_database(self, user_input, similarity_threshold=None):
        """Search local database with Vietnamese support and error handling"""
        try:
            if not user_input or not user_input.strip():
                return []
            
            if similarity_threshold is None:
                similarity_threshold = self.config['search_config']['local_similarity_threshold']
            
            if self.embedding_model:
                try:
                    query_embedding = self.embedding_model.encode([user_input.strip()])
                    if query_embedding is None or len(query_embedding) == 0:
                        self.logger.warning("Failed to generate query embedding")
                        return self.keyword_search_fallback(user_input)
                except Exception as e:
                    self.logger.error(f"Query embedding error: {e}")
                    return self.keyword_search_fallback(user_input)
                
                try:
                    cursor = self.conn.cursor()
                    cursor.execute("""
                        SELECT name, name_vietnamese, description, description_vietnamese, 
                               category, price, embedding 
                        FROM products 
                        WHERE embedding IS NOT NULL
                    """)
                    products = cursor.fetchall()
                except Exception as e:
                    self.logger.error(f"Database query error: {e}")
                    return []
                
                if not products:
                    self.logger.info("No products with embeddings found, trying keyword search")
                    return self.keyword_search_fallback(user_input)
                    
                best_matches = []
                
                for product in products:
                    try:
                        if len(product) < 7:
                            continue
                            
                        stored_embedding_blob = product[6]
                        
                        if stored_embedding_blob:
                            try:
                                stored_embedding = np.frombuffer(stored_embedding_blob, dtype=np.float32)
                                stored_embedding = stored_embedding.reshape(1, -1)
                            except Exception as e:
                                self.logger.warning(f"Embedding decode error: {e}")
                                continue
                            
                            try:
                                similarity = cosine_similarity(query_embedding, stored_embedding)[0][0]
                            except Exception as e:
                                self.logger.warning(f"Similarity calculation error: {e}")
                                continue
                            
                            if similarity > similarity_threshold:
                                best_matches.append({
                                    'name': product[0] or '',
                                    'name_vietnamese': product[1] or '',
                                    'description': product[2] or '',
                                    'description_vietnamese': product[3] or '',
                                    'category': product[4] or '',
                                    'price': product[5],
                                    'similarity': float(similarity)
                                })
                        
                    except Exception as e:
                        self.logger.warning(f"Product processing error: {e}")
                        continue
                            
                try:
                    best_matches.sort(key=lambda x: x.get('similarity', 0), reverse=True)
                    
                    if not best_matches:
                        self.logger.info("No embedding matches found, trying keyword search")
                        return self.keyword_search_fallback(user_input)
                        
                    return best_matches[:3]
                except Exception as e:
                    self.logger.error(f"Sorting error: {e}")
                    return best_matches[:3] if best_matches else []
            else:
                self.logger.info("Embedding model not available, using keyword search")
                return self.keyword_search_fallback(user_input)
                
        except Exception as e:
            self.logger.error(f"Database search error: {e}")
            return []
    
    def keyword_search_fallback(self, user_input):
        """Fallback keyword-based search when embedding search fails or is unavailable"""
        try:
            if not user_input or not user_input.strip():
                return []
            
            keywords = user_input.lower().split()
            
            cursor = self.conn.cursor()
            
            query = """
                SELECT name, name_vietnamese, description, description_vietnamese, 
                       category, price, NULL as embedding
                FROM products 
                WHERE 1=0
            """
            
            conditions = []
            params = []
            
            for keyword in keywords:
                if len(keyword) > 2:
                    keyword_pattern = f'%{keyword}%'
                    conditions.append("""
                        (LOWER(name) LIKE ? OR 
                         LOWER(name_vietnamese) LIKE ? OR 
                         LOWER(description) LIKE ? OR 
                         LOWER(description_vietnamese) LIKE ? OR
                         LOWER(category) LIKE ? OR
                         LOWER(category_vietnamese) LIKE ?)
                    """)
                    params.extend([keyword_pattern] * 6)
            
            if conditions:
                query = f"""
                    SELECT name, name_vietnamese, description, description_vietnamese, 
                           category, price, NULL as embedding
                    FROM products 
                    WHERE {' OR '.join(conditions)}
                    LIMIT 10
                """
                
                cursor.execute(query, params)
                products = cursor.fetchall()
                
                results = []
                for product in products:
                    if len(product) >= 6:
                        results.append({
                            'name': product[0] or '',
                            'name_vietnamese': product[1] or '',
                            'description': product[2] or '',
                            'description_vietnamese': product[3] or '',
                            'category': product[4] or '',
                            'price': product[5],
                            'similarity': 0.5
                        })
                
                self.logger.info(f"Keyword search found {len(results)} results")
                return results[:3]
            
            return []
            
        except Exception as e:
            self.logger.error(f"Keyword search error: {e}")
            return []

    def search_knowledge_base(self, user_input, similarity_threshold=None):
        """Search knowledge base for relevant information"""
        try:
            if not self.embedding_model:
                self.logger.warning("Embedding model not available for knowledge base search")
                return []
            
            if not user_input or not user_input.strip():
                return []
            
            if similarity_threshold is None:
                similarity_threshold = self.config['search_config']['local_similarity_threshold']
            
            try:
                query_embedding = self.embedding_model.encode([user_input.strip()])
                if query_embedding is None or len(query_embedding) == 0:
                    return []
            except Exception as e:
                self.logger.error(f"Knowledge base query embedding error: {e}")
                return []
            
            try:
                cursor = self.conn.cursor()
                cursor.execute("""
                    SELECT topic, content, source, embedding
                    FROM knowledge_base
                    WHERE embedding IS NOT NULL
                """)
                knowledge_entries = cursor.fetchall()
            except Exception as e:
                self.logger.error(f"Knowledge base query error: {e}")
                return []
            
            if not knowledge_entries:
                self.logger.info("No knowledge base entries with embeddings found")
                return []
            
            best_matches = []
            
            for entry in knowledge_entries:
                try:
                    if len(entry) < 4:
                        continue
                    
                    topic, content, source, embedding_blob = entry
                    
                    if embedding_blob:
                        try:
                            stored_embedding = np.frombuffer(embedding_blob, dtype=np.float32)
                            stored_embedding = stored_embedding.reshape(1, -1)
                        except Exception as e:
                            self.logger.warning(f"Knowledge embedding decode error: {e}")
                            continue
                        
                        try:
                            similarity = cosine_similarity(query_embedding, stored_embedding)[0][0]
                        except Exception as e:
                            self.logger.warning(f"Knowledge similarity calculation error: {e}")
                            continue
                        
                        if similarity > similarity_threshold:
                            best_matches.append({
                                'topic': topic or '',
                                'content': content[:500] if content else '',
                                'source': source or '',
                                'similarity': float(similarity),
                                'type': 'knowledge'
                            })
                
                except Exception as e:
                    self.logger.warning(f"Knowledge entry processing error: {e}")
                    continue
            
            try:
                best_matches.sort(key=lambda x: x.get('similarity', 0), reverse=True)
                return best_matches[:3]
            except Exception as e:
                self.logger.error(f"Knowledge sorting error: {e}")
                return best_matches[:3] if best_matches else []
                
        except Exception as e:
            self.logger.error(f"Knowledge base search error: {e}")
            return []
    
    def check_context_references(self, user_input, user_language='vi'):
        """Check if user is referring to previous conversation items"""
        try:
            if user_language == 'vi':
                context_keywords = [
                    'cái đó', 'cái này', 'sản phẩm đó', 'sản phẩm này',
                    'như trên', 'đã nói', 'vừa nói', 'trước đó',
                    'cái thứ', 'cái đầu', 'cái cuối', 'nó',
                    'chúng', 'những cái', 'mấy cái'
                ]
            else:
                context_keywords = [
                    'that', 'this', 'those', 'these',
                    'the one', 'previous', 'mentioned',
                    'above', 'it', 'them', 'the first',
                    'the last', 'the second'
                ]
            
            user_input_lower = user_input.lower()
            
            references = []
            for keyword in context_keywords:
                if keyword in user_input_lower:
                    references.append(keyword)
            
            if self.conversation_context:
                for turn in self.conversation_context[-3:]:
                    if 'products_mentioned' in turn and turn['products_mentioned']:
                        for product in turn['products_mentioned']:
                            if product.lower() in user_input_lower:
                                references.append(f"product:{product}")
            
            return references
            
        except Exception as e:
            self.logger.error(f"Context reference check error: {e}")
            return []
    
    def build_contextual_search_query(self, user_input, context_references):
        """Build search query that includes context from previous conversations"""
        try:
            search_terms = [user_input]
            
            for turn in self.conversation_context[-3:]:
                if 'products_mentioned' in turn and turn['products_mentioned']:
                    for product in turn['products_mentioned']:
                        if product and product not in search_terms:
                            search_terms.append(product)
            
            enhanced_query = ' '.join(search_terms[:3])
            
            self.logger.info(f"Enhanced search query: {enhanced_query}")
            return enhanced_query
            
        except Exception as e:
            self.logger.error(f"Contextual search query error: {e}")
            return user_input
    
    def update_conversation_summary(self):
        """Update conversation summary for long-term context"""
        try:
            if not self.conversation_context:
                return
            
            topics = []
            products_discussed = set()
            
            for turn in self.conversation_context:
                if 'products_mentioned' in turn and turn['products_mentioned']:
                    products_discussed.update(turn['products_mentioned'])
                
                user_msg = turn.get('user', '').lower()
                if 'laptop' in user_msg or 'máy tính' in user_msg:
                    topics.append('laptops')
                if 'gaming' in user_msg or 'game' in user_msg:
                    topics.append('gaming')
                if 'mouse' in user_msg or 'chuột' in user_msg:
                    topics.append('mouse')
                if 'keyboard' in user_msg or 'bàn phím' in user_msg:
                    topics.append('keyboard')
            
            self.conversation_summary = {
                'topics': list(set(topics)),
                'products_discussed': list(products_discussed),
                'turn_count': len(self.conversation_context),
                'last_update': datetime.now().isoformat()
            }
            
            self.logger.info(f"Updated conversation summary: {self.conversation_summary}")
            
        except Exception as e:
            self.logger.error(f"Conversation summary update error: {e}")
    
    def get_conversation_context_prompt(self, user_language='vi'):
        """Get enhanced conversation context for prompts"""
        try:
            context_prompt = ""
            
            if hasattr(self, 'conversation_summary') and self.conversation_summary:
                if user_language == 'vi':
                    context_prompt += "Tóm tắt cuộc trò chuyện:\n"
                    if self.conversation_summary.get('topics'):
                        context_prompt += f"- Chủ đề đã thảo luận: {', '.join(self.conversation_summary['topics'])}\n"
                    if self.conversation_summary.get('products_discussed'):
                        context_prompt += f"- Sản phẩm đã xem: {', '.join(self.conversation_summary['products_discussed'][:5])}\n"
                else:
                    context_prompt += "Conversation summary:\n"
                    if self.conversation_summary.get('topics'):
                        context_prompt += f"- Topics discussed: {', '.join(self.conversation_summary['topics'])}\n"
                    if self.conversation_summary.get('products_discussed'):
                        context_prompt += f"- Products viewed: {', '.join(self.conversation_summary['products_discussed'][:5])}\n"
                context_prompt += "\n"
            
            return context_prompt
            
        except Exception as e:
            self.logger.error(f"Context prompt error: {e}")
            return ""
    
    def store_conversation(self, user_input, response, data_source, processing_time, 
                          user_language, error_type=None, error_message=None):
        """Store conversation in database with language tracking and error logging"""
        try:
            cursor = self.conn.cursor()
            
            try:
                cursor.execute('''
                    INSERT INTO conversations 
                    (timestamp, user_input, user_language, bot_response, bot_language, 
                     data_source, response_time, error_type, error_message)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    datetime.now().isoformat(),
                    user_input,
                    user_language,
                    response,
                    user_language,
                    data_source,
                    processing_time,
                    error_type,
                    error_message
                ))
            except sqlite3.OperationalError as e:
                if "no column named error_type" in str(e):
                    cursor.execute('''
                        INSERT INTO conversations 
                        (timestamp, user_input, user_language, bot_response, bot_language, 
                         data_source, response_time)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        datetime.now().isoformat(),
                        user_input,
                        user_language,
                        response,
                        user_language,
                        data_source,
                        processing_time
                    ))
                else:
                    raise e
            
            self.conn.commit()
        except Exception as e:
            self.logger.error(f"Error storing conversation: {e}")
    
    def process_excel_file(self, file_path):
        """Process Excel files for product data with Vietnamese support"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, encoding='utf-8')
            else:
                df = pd.read_excel(file_path)
                
            self.update_training_status(f"📊 Found {len(df)} rows in Excel file")
                
            column_mapping = {
                'name': ['name', 'tên', 'ten', 'product_name', 'sản phẩm'],
                'name_vietnamese': ['name_vietnamese', 'tên_tiếng_việt', 'ten_tieng_viet'],
                'description': ['description', 'mô tả', 'mo_ta', 'desc'],
                'description_vietnamese': ['description_vietnamese', 'mô_tả_tiếng_việt'],
                'category': ['category', 'danh mục', 'danh_muc', 'loại'],
                'category_vietnamese': ['category_vietnamese', 'danh_mục_tiếng_việt'],
                'price': ['price', 'giá', 'gia', 'cost'],
                'features': ['features', 'tính năng', 'tinh_nang'],
                'features_vietnamese': ['features_vietnamese', 'tính_năng_tiếng_việt'],
                'specifications': ['specifications', 'thông số', 'thong_so', 'specs'],
                'specifications_vietnamese': ['specifications_vietnamese', 'thông_số_tiếng_việt']
            }
            
            mapped_columns = {}
            for standard_col, possible_cols in column_mapping.items():
                for col in possible_cols:
                    if col in df.columns:
                        mapped_columns[standard_col] = col
                        break
            
            if 'name' not in mapped_columns and 'name_vietnamese' not in mapped_columns:
                raise ValueError("Excel file must contain at least a name column (English or Vietnamese)")
                
            cursor = self.conn.cursor()
            added_count = 0
            
            for _, row in df.iterrows():
                try:
                    name = str(row.get(mapped_columns.get('name', ''), '')) if 'name' in mapped_columns else ''
                    name_vietnamese = str(row.get(mapped_columns.get('name_vietnamese', ''), '')) if 'name_vietnamese' in mapped_columns else ''
                    description = str(row.get(mapped_columns.get('description', ''), '')) if 'description' in mapped_columns else ''
                    description_vietnamese = str(row.get(mapped_columns.get('description_vietnamese', ''), '')) if 'description_vietnamese' in mapped_columns else ''
                    category = str(row.get(mapped_columns.get('category', ''), 'General')) if 'category' in mapped_columns else 'General'
                    category_vietnamese = str(row.get(mapped_columns.get('category_vietnamese', ''), '')) if 'category_vietnamese' in mapped_columns else ''
                    features = str(row.get(mapped_columns.get('features', ''), '')) if 'features' in mapped_columns else ''
                    features_vietnamese = str(row.get(mapped_columns.get('features_vietnamese', ''), '')) if 'features_vietnamese' in mapped_columns else ''
                    specifications = str(row.get(mapped_columns.get('specifications', ''), '')) if 'specifications' in mapped_columns else ''
                    specifications_vietnamese = str(row.get(mapped_columns.get('specifications_vietnamese', ''), '')) if 'specifications_vietnamese' in mapped_columns else ''

                    price = 0
                    if 'price' in mapped_columns:
                        try:
                            price_val = row.get(mapped_columns['price'], 0)
                            if pd.notna(price_val):
                                price = float(str(price_val).replace(',', '').replace('$', ''))
                        except (ValueError, TypeError):
                            price = 0
                    
                    if not name and not name_vietnamese:
                        continue
                        
                    product_text_en = f"{name} {description} {features} {specifications}"
                    product_text_vi = f"{name_vietnamese} {description_vietnamese} {features_vietnamese} {specifications_vietnamese}"
                    combined_text = f"{product_text_en} {product_text_vi}".strip()
                    
                    embedding_blob = None
                    if self.embedding_model and combined_text:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.update_training_status(f"⚠️ Embedding generation failed for {name or name_vietnamese}: {e}")
                    
                    cursor.execute('''
                        INSERT INTO products 
                        (name, name_vietnamese, description, description_vietnamese, 
                         category, category_vietnamese, price, features, features_vietnamese,
                         specifications, specifications_vietnamese, source_file, embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        name or name_vietnamese,
                        name_vietnamese,
                        description or description_vietnamese,
                        description_vietnamese,
                        category,
                        category_vietnamese,
                        price,
                        features,
                        features_vietnamese,
                        specifications,
                        specifications_vietnamese,
                        file_path,
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                    added_count += 1
                    
                except Exception as row_error:
                    self.update_training_status(f"⚠️ Error processing row: {row_error}")
                    continue
                
            self.conn.commit()
            self.update_training_status(f"✅ Added {added_count} products from Excel file with embeddings")
            
        except Exception as e:
            raise Exception(f"Error processing Excel file: {e}")
            
    def process_pdf_file(self, file_path):
        """Process PDF files for knowledge base with Vietnamese support"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    self.update_training_status(f"📄 Processed page {page_num + 1}/{len(reader.pages)}")
                except Exception as e:
                    self.update_training_status(f"⚠️ Error extracting page {page_num + 1}: {e}")
                    continue
                
            if text.strip():
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from PDF: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from PDF")
                
        except Exception as e:
            raise Exception(f"Error processing PDF file: {e}")
            
    def process_word_file(self, file_path):
        """Process Word documents for knowledge base with Vietnamese support"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                if para_num % 10 == 0:
                    self.update_training_status(f"📝 Processed {para_num} paragraphs...")
                    
            for table_num, table in enumerate(doc.tables):
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                self.update_training_status(f"📊 Processed table {table_num + 1}/{len(doc.tables)}")
                
            if text.strip():
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from Word doc: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from Word document")
                
        except Exception as e:
            raise Exception(f"Error processing Word file: {e}")
            
    def process_image_file(self, file_path):
        """Process image files using OCR with Vietnamese support"""
        try:
            if not self.ocr_reader:
                raise Exception("OCR model not available")
                
            self.update_training_status(f"🖼️ Running OCR on image...")
            
            results = self.ocr_reader.readtext(file_path)
            
            extracted_texts = []
            for (bbox, text, confidence) in results:
                if confidence > 0.5:
                    extracted_texts.append(text.strip())
                    
            if extracted_texts:
                combined_text = " ".join(extracted_texts)
                
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding = self.embedding_model.encode([combined_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    f"OCR: {os.path.basename(file_path)}",
                    combined_text,
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Extracted {len(extracted_texts)} text segments from image")
                self.update_training_status(f"📝 Text preview: {combined_text[:100]}...")
            else:
                self.update_training_status(f"⚠️ No text detected in image")
                
        except Exception as e:
            raise Exception(f"Error processing image file: {e}")
    
    def setup_gui(self):
        """Setup the GUI interface with Vietnamese font support"""
        self.root = tk.Tk()
        self.root.title("Trợ lý AI Bán hàng - Vietnamese AI Sales ChatBot")
        self.root.geometry("1400x900")
        
        self.setup_vietnamese_fonts()
        
        self.setup_main_interface()
        
    def setup_vietnamese_fonts(self):
        """Setup fonts that support Vietnamese characters with enhanced input support"""
        vietnamese_fonts = [
            ('Segoe UI', 12),
            ('Times New Roman', 12),
            ('Arial Unicode MS', 12),
            ('Calibri', 12),
            ('Tahoma', 12),
            ('Microsoft Sans Serif', 12),
            ('DejaVu Sans', 12),
            ('Liberation Sans', 12),
            ('Verdana', 11)
        ]
        
        self.vietnamese_font = None
        self.input_font = None
        available_fonts = font.families()
        
        for font_name, size in vietnamese_fonts:
            if font_name in available_fonts:
                try:
                    test_font = font.Font(family=font_name, size=size)
                    
                    self.vietnamese_font = font.Font(family=font_name, size=size-1)
                    self.input_font = font.Font(
                        family=font_name, 
                        size=size, 
                        weight='normal'
                    )
                    
                    print(f"✅ Using Vietnamese font: {font_name}")
                    break
                except Exception as e:
                    print(f"⚠️ Font {font_name} failed: {e}")
                    continue
        
        if not self.vietnamese_font:
            self.vietnamese_font = font.Font(family="TkDefaultFont", size=11)
            self.input_font = font.Font(family="TkDefaultFont", size=12)
            print("⚠️ Using default font (Vietnamese characters may not display correctly)")
        
        self.root.option_add('*Font', self.vietnamese_font)
        
        try:
            self.root.tk.call('encoding', 'system', 'utf-8')
            print("✅ System encoding set to UTF-8 for Vietnamese support")
        except Exception as e:
            print(f"⚠️ Could not set UTF-8 encoding: {e}")
    
    def setup_main_interface(self):
        """Setup the main GUI interface with notebook tabs"""
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.setup_chat_tab()
        self.setup_database_tab()
        self.setup_training_tab()
        self.setup_analytics_tab()
    
    def setup_chat_tab(self):
        """Setup the main chat interface"""
        chat_frame = ttk.Frame(self.notebook)
        self.notebook.add(chat_frame, text="💬 Chat")
        
        top_frame = ttk.Frame(chat_frame)
        top_frame.pack(fill=tk.X, padx=10, pady=5)
        
        lang_frame = ttk.Frame(top_frame)
        lang_frame.pack(side=tk.LEFT)
        
        ttk.Label(lang_frame, text="Ngôn ngữ / Language:", font=self.vietnamese_font).pack(side=tk.LEFT, padx=5)
        
        self.language_var = tk.StringVar(value=self.config['language_config']['default_language'])
        language_combo = ttk.Combobox(lang_frame, textvariable=self.language_var, 
                                     values=['vi', 'en'], state='readonly', width=8)
        language_combo.pack(side=tk.LEFT, padx=5)
        language_combo.bind('<<ComboboxSelected>>', self.on_language_change)
        
        self.chat_display = scrolledtext.ScrolledText(
            chat_frame, 
            height=30, 
            state=tk.DISABLED,
            font=self.vietnamese_font,
            wrap=tk.WORD
        )
        self.chat_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.chat_display.tag_configure("user", foreground="blue", font=(self.vietnamese_font['family'], 10, 'bold'))
        self.chat_display.tag_configure("bot", foreground="green", font=(self.vietnamese_font['family'], 10))
        self.chat_display.tag_configure("system", foreground="gray", font=(self.vietnamese_font['family'], 9, 'italic'))
        self.chat_display.tag_configure("error", foreground="red", font=(self.vietnamese_font['family'], 9, 'italic'))
        
        input_frame = ttk.Frame(chat_frame)
        input_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.user_input = tk.Text(
            input_frame, 
            height=3, 
            font=self.input_font,
            wrap=tk.WORD
        )
        self.user_input.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        self.user_input.bind('<Control-Return>', self.send_message)
        self.user_input.bind('<Return>', self.on_enter_key)
        
        button_frame = ttk.Frame(input_frame)
        button_frame.pack(side=tk.RIGHT, fill=tk.Y)
        
        send_btn = ttk.Button(button_frame, text="Gửi / Send", command=self.send_message)
        send_btn.pack(fill=tk.X, pady=(0, 5))
        
        clear_btn = ttk.Button(button_frame, text="Xóa / Clear", command=self.clear_chat)
        clear_btn.pack(fill=tk.X)
        
        status_frame = ttk.Frame(chat_frame)
        status_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.status_label = ttk.Label(status_frame, text="Sẵn sàng / Ready", font=self.vietnamese_font)
        self.status_label.pack(side=tk.LEFT)
        
        self.processing_label = ttk.Label(status_frame, text="", font=self.vietnamese_font)
        self.processing_label.pack(side=tk.RIGHT)
    
    def setup_database_tab(self):
        """Setup database management interface"""
        db_frame = ttk.Frame(self.notebook)
        self.notebook.add(db_frame, text="🗄️ Database")
        
        product_frame = ttk.LabelFrame(db_frame, text="Product Database Management")
        product_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        btn_frame = ttk.Frame(product_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(btn_frame, text="Add Sample Products", 
                  command=self.add_sample_products).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="View Products", 
                  command=self.view_products).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Search Test", 
                  command=self.test_search).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Regenerate Embeddings", 
                  command=self.regenerate_all_embeddings).pack(side=tk.LEFT, padx=5)
        
        self.product_display = scrolledtext.ScrolledText(
            product_frame, height=15, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.product_display.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
    
    def setup_training_tab(self):
        """Setup training interface with file upload capabilities"""
        training_frame = ttk.Frame(self.notebook)
        self.notebook.add(training_frame, text="📚 Training Data")
        
        upload_frame = ttk.LabelFrame(training_frame, text="Upload Training Data")
        upload_frame.pack(fill=tk.X, padx=10, pady=5)
        
        upload_btn_frame = ttk.Frame(upload_frame)
        upload_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(upload_btn_frame, text="📊 Upload Excel Files", 
                  command=lambda: self.process_files('excel')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📄 Upload PDF Files", 
                  command=lambda: self.process_files('pdf')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📝 Upload Word Files", 
                  command=lambda: self.process_files('word')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="🖼️ Upload Images (OCR)", 
                  command=lambda: self.process_files('image')).pack(side=tk.LEFT, padx=5, pady=5)
        
        options_frame = ttk.LabelFrame(training_frame, text="Training Options")
        options_frame.pack(fill=tk.X, padx=10, pady=5)
        
        options_btn_frame = ttk.Frame(options_frame)
        options_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(options_btn_frame, text="🔄 Regenerate All Embeddings", 
                  command=self.regenerate_all_embeddings).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="🧹 Clear Training Data", 
                  command=self.clear_training_data).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="📈 Training Statistics", 
                  command=self.show_training_stats).pack(side=tk.LEFT, padx=5)
        
        self.training_status = scrolledtext.ScrolledText(
            training_frame, height=20, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.training_status.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def setup_analytics_tab(self):
        """Setup analytics interface"""
        analytics_frame = ttk.Frame(self.notebook)
        self.notebook.add(analytics_frame, text="📊 Analytics")
        
        metrics_frame = ttk.LabelFrame(analytics_frame, text="Performance Metrics")
        metrics_frame.pack(fill=tk.X, padx=10, pady=5)
        
        metrics_btn_frame = ttk.Frame(metrics_frame)
        metrics_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(metrics_btn_frame, text="🔄 Refresh Analytics", 
                  command=self.refresh_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📤 Export Data", 
                  command=self.export_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📋 Conversation History", 
                  command=self.view_conversation_history).pack(side=tk.LEFT, padx=5, pady=5)
        
        self.analytics_display = scrolledtext.ScrolledText(
            analytics_frame, height=25, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.analytics_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def on_enter_key(self, event):
        """Handle Enter key press"""
        if event.state & 0x4:
            self.send_message()
        else:
            return None
    
    def on_language_change(self, event=None):
        """Handle language change"""
        try:
            self.current_language = self.language_var.get()
            print(f"Language changed to: {self.current_language}")
        except Exception as e:
            print(f"Language change error: {e}")
    
    def send_message(self, event=None):
        """Handle sending user message with error handling"""
        try:
            user_text = self.user_input.get(1.0, tk.END).strip()
            if not user_text:
                return
                
            self.user_input.delete(1.0, tk.END)
            
            self.display_message("Bạn / You", user_text, "user")
            
            if hasattr(self, 'processing_label'):
                status_text = "Đang xử lý..." if self.current_language == 'vi' else "Processing..."
                self.processing_label.config(text=status_text)
                self.root.update()
            
            threading.Thread(target=self.process_message_thread, args=(user_text,), daemon=True).start()
        except Exception as e:
            print(f"Send message error: {e}")
            error_msg = f"Lỗi gửi tin nhắn / Send error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
    
    def process_message_thread(self, user_input):
        """Process message in separate thread with comprehensive error handling"""
        try:
            result = self.process_message(user_input)
            
            response_tag = "error" if result.get('data_source') == 'error' else "bot"
            self.display_message("Trợ lý AI / AI Assistant", result.get('response', ''), response_tag)
            
            user_language = result.get('user_language', 'vi')
            if user_language == 'vi':
                if result.get('data_source') == 'local_database':
                    status_text = f"✅ Tìm thấy {result.get('local_results_count', 0)} sản phẩm ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'knowledge_base':
                    status_text = f"📚 Tìm thấy trong cơ sở tri thức ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'error':
                    status_text = f"❌ Lỗi ({result.get('processing_time', 0):.1f}s)"
                else:
                    status_text = f"💬 Hoàn thành ({result.get('processing_time', 0):.1f}s)"
            else:
                if result.get('data_source') == 'local_database':
                    status_text = f"✅ Found {result.get('local_results_count', 0)} products ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'knowledge_base':
                    status_text = f"📚 Found in knowledge base ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'error':
                    status_text = f"❌ Error ({result.get('processing_time', 0):.1f}s)"
                else:
                    status_text = f"💬 Complete ({result.get('processing_time', 0):.1f}s)"
            
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text=status_text))
            
        except AttributeError as ae:
            self.logger.error(f"AttributeError in message thread: {ae}")
            error_msg = f"Method not found error: {str(ae)}"
            self.display_message("Hệ thống / System", error_msg, "error")
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text="Method Error"))
        except Exception as e:
            self.logger.error(f"Message thread error: {e}")
            error_msg = f"Lỗi xử lý / Processing error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text="Lỗi / Error"))
    
    def display_message(self, sender, message, tag):
        """Display message in chat window with error handling"""
        def update_display():
            try:
                self.chat_display.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime("%H:%M:%S")
                
                self.chat_display.insert(tk.END, f"[{timestamp}] {sender}:\n", tag)
                self.chat_display.insert(tk.END, f"{message}\n\n")
                
                self.chat_display.config(state=tk.DISABLED)
                self.chat_display.see(tk.END)
            except Exception as e:
                print(f"Display update error: {e}")
            
        try:
            if threading.current_thread() != threading.main_thread():
                self.root.after(0, update_display)
            else:
                update_display()
        except Exception as e:
            print(f"Display message error: {e}")
    
    def clear_chat(self):
        """Clear chat display with error handling"""
        try:
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)
            self.chat_display.config(state=tk.DISABLED)
            
            self.conversation_context.clear()
            self.conversation_summary = ""
            
            self.logger.info("Chat history and context cleared")
            
        except Exception as e:
            print(f"Clear chat error: {e}")
    
    def update_training_status(self, message):
        """Update training status display with Vietnamese support"""
        def update():
            if hasattr(self, 'training_status'):
                self.training_status.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime('%H:%M:%S')
                self.training_status.insert(tk.END, f"[{timestamp}] {message}\n")
                self.training_status.config(state=tk.DISABLED)
                self.training_status.see(tk.END)
            
        if threading.current_thread() != threading.main_thread():
            self.root.after(0, update)
        else:
            update()
    
    def process_files(self, file_type):
        """Process different types of files for training data"""
        file_types = {
            'excel': [('Excel files', '*.xlsx *.xls *.csv')],
            'pdf': [('PDF files', '*.pdf')],
            'word': [('Word files', '*.docx *.doc')],
            'image': [('Image files', '*.png *.jpg *.jpeg *.bmp *.tiff')]
        }
        
        files = filedialog.askopenfilenames(
            title=f"Select {file_type} files",
            filetypes=file_types[file_type]
        )
        
        if files:
            self.update_training_status(f"📂 Starting to process {len(files)} {file_type} file(s)...")
            threading.Thread(
                target=self.process_files_worker,
                args=(files, file_type),
                daemon=True
            ).start()
            
    def process_files_worker(self, files, file_type):
        """Worker thread for processing files with Vietnamese support"""
        for file_path in files:
            try:
                self.update_training_status(f"🔄 Processing: {os.path.basename(file_path)}")
                
                if file_type == 'excel':
                    self.process_excel_file(file_path)
                elif file_type == 'pdf':
                    self.process_pdf_file(file_path)
                elif file_type == 'word':
                    self.process_word_file(file_path)
                elif file_type == 'image':
                    self.process_image_file(file_path)
                    
                self.update_training_status(f"✅ Completed: {os.path.basename(file_path)}")
                
            except Exception as e:
                self.update_training_status(f"❌ Error processing {os.path.basename(file_path)}: {e}")
    
    def add_sample_products(self):
        """Add Vietnamese sample products with error handling"""
        sample_products = [
            {
                'name': 'Gaming Laptop Pro X1',
                'name_vietnamese': 'Laptop Gaming Pro X1',
                'description': 'High-performance gaming laptop with RTX 4070, 16GB RAM, 1TB SSD, perfect for gaming and professional work',
                'description_vietnamese': 'Laptop gaming hiệu suất cao với RTX 4070, RAM 16GB, SSD 1TB, hoàn hảo cho gaming và công việc chuyên nghiệp',
                'category': 'Laptops',
                'category_vietnamese': 'Laptop',
                'price': 1899.99,
                'features': 'RTX 4070, Intel i7, 16GB RAM, 1TB SSD, 15.6" 144Hz display',
                'features_vietnamese': 'RTX 4070, Intel i7, RAM 16GB, SSD 1TB, màn hình 15.6" 144Hz'
            },
            {
                'name': 'Wireless Gaming Mouse RGB Pro',
                'name_vietnamese': 'Chuột Gaming Không dây RGB Pro',
                'description': 'High-precision wireless gaming mouse with RGB lighting, 12000 DPI, programmable buttons',
                'description_vietnamese': 'Chuột gaming không dây độ chính xác cao với đèn RGB, 12000 DPI, nút bấm có thể lập trình',
                'category': 'Gaming Accessories',
                'category_vietnamese': 'Phụ kiện Gaming',
                'price': 79.99,
                'features': '12000 DPI, RGB lighting, 7 programmable buttons, 70-hour battery',
                'features_vietnamese': '12000 DPI, đèn RGB, 7 nút lập trình, pin 70 giờ'
            },
            {
                'name': 'Mechanical Gaming Keyboard',
                'name_vietnamese': 'Bàn phím Gaming Cơ',
                'description': 'Premium mechanical gaming keyboard with RGB backlighting and blue switches',
                'description_vietnamese': 'Bàn phím gaming cơ cao cấp với đèn nền RGB và switch xanh',
                'category': 'Gaming Accessories',
                'category_vietnamese': 'Phụ kiện Gaming',
                'price': 129.99,
                'features': 'Blue mechanical switches, RGB per-key lighting, aluminum frame',
                'features_vietnamese': 'Switch cơ xanh, đèn RGB từng phím, khung nhôm'
            },
            {
                'name': 'Gaming Headset 7.1 Surround',
                'name_vietnamese': 'Tai nghe Gaming 7.1 Surround',
                'description': '7.1 surround sound gaming headset with noise-canceling microphone',
                'description_vietnamese': 'Tai nghe gaming âm thanh vòm 7.1 với microphone chống ồn',
                'category': 'Audio',
                'category_vietnamese': 'Âm thanh',
                'price': 89.99,
                'features': '7.1 surround sound, noise-canceling mic, comfortable padding',
                'features_vietnamese': 'Âm thanh vòm 7.1, mic chống ồn, đệm êm ái'
            },
            {
                'name': 'Business Laptop UltraSlim',
                'name_vietnamese': 'Laptop Văn phòng UltraSlim',
                'description': 'Lightweight business laptop with long battery life, perfect for professionals',
                'description_vietnamese': 'Laptop văn phòng siêu mỏng với thời lượng pin dài, hoàn hảo cho dân văn phòng',
                'category': 'Laptops',
                'category_vietnamese': 'Laptop',
                'price': 899.99,
                'features': 'Intel i5, 8GB RAM, 512GB SSD, 14" display, 12-hour battery',
                'features_vietnamese': 'Intel i5, RAM 8GB, SSD 512GB, màn hình 14", pin 12 giờ'
            }
        ]
        
        try:
            cursor = self.conn.cursor()
            for product in sample_products:
                try:
                    english_text = f"{product['name']} {product['description']} {product.get('features', '')}"
                    vietnamese_text = f"{product['name_vietnamese']} {product['description_vietnamese']} {product.get('features_vietnamese', '')}"
                    combined_text = f"{english_text} {vietnamese_text}"
                    
                    embedding_blob = None
                    if self.embedding_model:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.logger.warning(f"Embedding generation error for {product['name']}: {e}")
                    
                    cursor.execute('''
                        INSERT OR IGNORE INTO products 
                        (name, name_vietnamese, description, description_vietnamese, 
                         category, category_vietnamese, price, features, features_vietnamese,
                         embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        product['name'], product['name_vietnamese'],
                        product['description'], product['description_vietnamese'],
                        product['category'], product['category_vietnamese'],
                        product['price'], product.get('features', ''), product.get('features_vietnamese', ''),
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                except Exception as e:
                    self.logger.error(f"Error inserting product {product.get('name', 'unknown')}: {e}")
                    continue
            
            self.conn.commit()
            print("✅ Sample products added successfully")
            
        except Exception as e:
            self.logger.error(f"Error adding sample products: {e}")
    
    def view_products(self):
        """View products in database with Vietnamese support"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT name, name_vietnamese, description, description_vietnamese, 
                       category, price, source_file, embedding 
                FROM products 
                ORDER BY created_at DESC
            """)
            products = cursor.fetchall()
            
            self.product_display.config(state=tk.NORMAL)
            self.product_display.delete(1.0, tk.END)
            
            if products:
                self.product_display.insert(tk.END, f"📦 Found {len(products)} products in database:\n\n")
                
                for i, product in enumerate(products, 1):
                    name = product[0] or product[1] or "Unnamed Product"
                    name_vi = product[1] or ""
                    desc = product[2] or product[3] or "No description"
                    category = product[4] or "General"
                    price = product[5] or 0
                    source = os.path.basename(product[6]) if product[6] else "Manual"
                    embedding_status = "✅" if product[7] else "❌"
                    
                    self.product_display.insert(tk.END, f"{i}. {name} {embedding_status}\n")
                    if name_vi:
                        self.product_display.insert(tk.END, f"   Vietnamese: {name_vi}\n")
                    self.product_display.insert(tk.END, f"   Description: {desc[:100]}...\n")
                    self.product_display.insert(tk.END, f"   Category: {category} | Price: ${price} | Source: {source}\n\n")
            else:
                self.product_display.insert(tk.END, "📦 No products found in database.\n\n")
                self.product_display.insert(tk.END, "💡 To add products:\n")
                self.product_display.insert(tk.END, "• Click 'Add Sample Products' for demo data\n")
                self.product_display.insert(tk.END, "• Use the Training tab to upload Excel/CSV files\n")
                self.product_display.insert(tk.END, "• Use PDFs/Word docs for knowledge base\n")
                
            self.product_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing products: {e}")
    
    def test_search(self):
        """Test search functionality with Vietnamese support"""
        test_queries = [
            "gaming laptop", "laptop gaming",
            "business computer", "máy tính văn phòng", 
            "wireless mouse", "chuột không dây",
            "cheap laptop under $1500", "laptop rẻ dưới 1500 đô",
            "RGB accessories", "phụ kiện RGB",
            "4K monitor", "màn hình 4K",
            "mechanical keyboard", "bàn phím cơ",
            "laptop for programming", "laptop lập trình"
        ]
        
        self.product_display.config(state=tk.NORMAL)
        self.product_display.delete(1.0, tk.END)
        self.product_display.insert(tk.END, "🔍 Testing search functionality...\n")
        self.product_display.insert(tk.END, f"Threshold: {self.config['search_config']['local_similarity_threshold']}\n\n")
        
        for query in test_queries:
            self.product_display.insert(tk.END, f"Query: '{query}'\n")
            results = self.search_local_database(query)
            
            if results:
                for result in results[:3]:
                    self.product_display.insert(tk.END, f"  ✓ {result['name']} (similarity: {result['similarity']:.3f})\n")
            else:
                self.product_display.insert(tk.END, "  ✗ No results found\n")
            
            self.product_display.insert(tk.END, "\n")
            
        self.product_display.config(state=tk.DISABLED)
    
    def regenerate_all_embeddings(self):
        """Regenerate all embeddings for products and knowledge base"""
        result = messagebox.askyesno(
            "Regenerate Embeddings", 
            "This will regenerate all embeddings for products and knowledge base entries.\n\n"
            "This may take several minutes depending on the amount of data.\n\n"
            "Continue?"
        )
        
        if result:
            threading.Thread(target=self.regenerate_embeddings_worker, daemon=True).start()
            
    def regenerate_embeddings_worker(self):
        """Worker thread for regenerating embeddings"""
        try:
            if not self.embedding_model:
                self.update_training_status("❌ Embedding model not available")
                return
                
            cursor = self.conn.cursor()
            
            self.update_training_status("🔄 Regenerating product embeddings...")
            cursor.execute("""
                SELECT id, name, name_vietnamese, description, description_vietnamese, 
                       features, features_vietnamese, specifications, specifications_vietnamese 
                FROM products
            """)
            products = cursor.fetchall()
            
            for i, (product_id, name, name_vi, desc, desc_vi, feat, feat_vi, spec, spec_vi) in enumerate(products):
                try:
                    text_parts = [name or '', name_vi or '', desc or '', desc_vi or '', 
                                feat or '', feat_vi or '', spec or '', spec_vi or '']
                    combined_text = ' '.join(part for part in text_parts if part.strip())
                    
                    if combined_text.strip():
                        embedding = self.embedding_model.encode([combined_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                        
                        cursor.execute("""
                            UPDATE products 
                            SET embedding = ?, updated_at = ?
                            WHERE id = ?
                        """, (embedding_blob, datetime.now().isoformat(), product_id))
                        
                        if (i + 1) % 10 == 0:
                            self.update_training_status(f"📦 Processed {i + 1}/{len(products)} products...")
                            
                except Exception as e:
                    self.update_training_status(f"⚠️ Error processing product {name}: {e}")
                    continue
                    
            self.update_training_status("🔄 Regenerating knowledge base embeddings...")
            cursor.execute("SELECT id, content FROM knowledge_base")
            knowledge_entries = cursor.fetchall()
            
            for i, (kb_id, content) in enumerate(knowledge_entries):
                try:
                    if content and content.strip():
                        embedding = self.embedding_model.encode([content[:5000]])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                        
                        cursor.execute("""
                            UPDATE knowledge_base 
                            SET embedding = ?
                            WHERE id = ?
                        """, (embedding_blob, kb_id))
                        
                        if (i + 1) % 5 == 0:
                            self.update_training_status(f"📚 Processed {i + 1}/{len(knowledge_entries)} knowledge entries...")
                            
                except Exception as e:
                    self.update_training_status(f"⚠️ Error processing knowledge entry {kb_id}: {e}")
                    continue
                    
            self.conn.commit()
            self.update_training_status(f"✅ Successfully regenerated embeddings for {len(products)} products and {len(knowledge_entries)} knowledge entries!")
            
        except Exception as e:
            self.update_training_status(f"❌ Error regenerating embeddings: {e}")
    
    def clear_training_data(self):
        """Clear all training data with confirmation"""
        result = messagebox.askyesno(
            "Clear Training Data", 
            "Are you sure you want to clear all training data?\n\n"
            "This will delete:\n"
            "• All products\n"
            "• All knowledge base entries\n"
            "• All embeddings\n\n"
            "This action cannot be undone!"
        )
        
        if result:
            try:
                cursor = self.conn.cursor()
                cursor.execute("DELETE FROM products")
                cursor.execute("DELETE FROM knowledge_base")
                self.conn.commit()
                
                self.update_training_status("🧹 All training data cleared successfully")
                messagebox.showinfo("Success", "All training data has been cleared.")
                
            except Exception as e:
                self.update_training_status(f"❌ Error clearing data: {e}")
                messagebox.showerror("Error", f"Error clearing training data: {e}")
    
    def show_training_stats(self):
        """Show training data statistics"""
        try:
            cursor = self.conn.cursor()
            
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(DISTINCT category) FROM products")
            category_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base WHERE embedding IS NOT NULL")
            knowledge_with_embeddings = cursor.fetchone()[0]
            
            self.training_status.config(state=tk.NORMAL)
            self.training_status.delete(1.0, tk.END)
            
            self.training_status.insert(tk.END, "=== TRAINING DATA STATISTICS ===\n\n")
            self.training_status.insert(tk.END, f"📦 Products: {product_count}\n")
            self.training_status.insert(tk.END, f"🔗 Products with embeddings: {products_with_embeddings}/{product_count}\n")
            self.training_status.insert(tk.END, f"📂 Categories: {category_count}\n\n")
            self.training_status.insert(tk.END, f"📚 Knowledge base entries: {knowledge_count}\n")
            self.training_status.insert(tk.END, f"🔗 Knowledge with embeddings: {knowledge_with_embeddings}/{knowledge_count}\n\n")
            
            cursor.execute("""
                SELECT name, created_at, source_file 
                FROM products 
                ORDER BY created_at DESC 
                LIMIT 5
            """)
            recent_products = cursor.fetchall()
            
            if recent_products:
                self.training_status.insert(tk.END, "📈 Recent Products:\n")
                for name, created_at, source_file in recent_products:
                    source = os.path.basename(source_file) if source_file else "Manual"
                    self.training_status.insert(tk.END, f"  • {name} ({source})\n")
                    
            self.training_status.config(state=tk.DISABLED)
            
        except Exception as e:
            self.update_training_status(f"❌ Error showing statistics: {e}")
    
    def refresh_analytics(self):
        """Refresh analytics display with Vietnamese language support"""
        try:
            cursor = self.conn.cursor()
            
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_conversations,
                    AVG(response_time) as avg_response_time,
                    data_source,
                    COUNT(*) as count_by_source
                FROM conversations 
                GROUP BY data_source
            """)
            
            source_stats = cursor.fetchall()
            
            cursor.execute("""
                SELECT timestamp, user_input, user_language, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 10
            """)
            
            recent_convs = cursor.fetchall()
            
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            cursor.execute("""
                SELECT user_language, COUNT(*) as count
                FROM conversations 
                WHERE user_language IS NOT NULL
                GROUP BY user_language
            """)
            language_stats = cursor.fetchall()
            
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            self.analytics_display.insert(tk.END, "=== VIETNAMESE AI CHATBOT ANALYTICS ===\n\n")
            self.analytics_display.insert(tk.END, f"📦 Products in Database: {product_count}\n")
            self.analytics_display.insert(tk.END, f"🔗 Products with Embeddings: {products_with_embeddings}/{product_count}\n")
            self.analytics_display.insert(tk.END, f"📚 Knowledge Base Entries: {knowledge_count}\n")
            
            self.analytics_display.insert(tk.END, f"\n=== LANGUAGE USAGE ===\n")
            for lang, count in language_stats:
                lang_name = "Vietnamese" if lang == 'vi' else "English" if lang == 'en' else lang
                self.analytics_display.insert(tk.END, f"{lang_name}: {count} conversations\n")
            
            self.analytics_display.insert(tk.END, "\n=== RESPONSE SOURCES ===\n")
            for stat in source_stats:
                if len(stat) >= 4:
                    self.analytics_display.insert(tk.END, f"{stat[2]}: {stat[3]} responses\n")
                
            if source_stats:
                total_conversations = sum(stat[3] for stat in source_stats if len(stat) >= 4)
                if total_conversations > 0:
                    avg_response_time = sum(stat[1] * stat[3] for stat in source_stats if stat[1] and len(stat) >= 4) / total_conversations
                    self.analytics_display.insert(tk.END, f"\nTotal Conversations: {total_conversations}\n")
                    self.analytics_display.insert(tk.END, f"Average Response Time: {avg_response_time:.2f} seconds\n")
            
            self.analytics_display.insert(tk.END, "\n=== RECENT CONVERSATIONS ===\n")
            for conv in recent_convs:
                if len(conv) >= 5:
                    timestamp, user_input, user_lang, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    self.analytics_display.insert(tk.END, f"{timestamp} {lang_flag} - {data_source} ({response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input[:100]}...\n\n")
                
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error refreshing analytics: {e}")
    
    def export_analytics(self):
        """Export analytics data to CSV with Vietnamese support"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT * FROM conversations")
            
            column_names = [description[0] for description in cursor.description]
            
            df = pd.DataFrame(cursor.fetchall(), columns=column_names)
            
            export_path = f"vietnamese_chatbot_analytics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            df.to_csv(export_path, index=False, encoding='utf-8')
            
            messagebox.showinfo("Success", f"Analytics exported to: {export_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error exporting analytics: {e}")
    
    def view_conversation_history(self):
        """View detailed conversation history"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT timestamp, user_input, user_language, bot_response, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 20
            """)
            conversations = cursor.fetchall()
            
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            self.analytics_display.insert(tk.END, "=== CONVERSATION HISTORY (Last 20) ===\n\n")
            
            for i, conv in enumerate(conversations, 1):
                if len(conv) >= 6:
                    timestamp, user_input, user_lang, bot_response, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    
                    self.analytics_display.insert(tk.END, f"{i}. [{timestamp}] {lang_flag} ({data_source}, {response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input}\n")
                    self.analytics_display.insert(tk.END, f"Bot: {bot_response[:200]}...\n\n")
            
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing conversation history: {e}")
    
    def run(self):
        """Start the application with comprehensive error handling"""
        try:
            print("📦 Adding sample products...")
            self.add_sample_products()
            
            self.display_welcome_message()
            
            print("🚀 Starting GUI...")
            self.root.mainloop()
            
        except Exception as e:
            print(f"❌ Error starting application: {e}")
            self.logger.error(f"Application start error: {e}")
    
    def display_welcome_message(self):
        """Display welcome message with system status"""
        try:
            if self.current_language == 'vi':
                welcome_msg = f"""🤖 Chào mừng đến với Trợ lý AI Bán hàng!

Trạng thái hệ thống:
✅ Hỗ trợ tiếng Việt và tiếng Anh
{'✅' if self.text_generator else '❌'} Mô hình AI: {'Hoạt động' if self.text_generator else 'Không khả dụng'}
{'✅' if self.embedding_model else '❌'} Tìm kiếm thông minh: {'Hoạt động' if self.embedding_model else 'Không khả dụng'}

Tính năng có sẵn:
📦 Tìm kiếm sản phẩm thông minh
💬 Ngữ cảnh cuộc trò chuyện
🌐 Hỗ trợ đa ngôn ngữ
📚 Xử lý file Excel, PDF, Word
🖼️ OCR từ hình ảnh

Hãy thử đặt câu hỏi như:
• "Tôi muốn mua laptop gaming"
• "Bạn có chuột không dây nào không?"
• "Giá của laptop là bao nhiêu?"
• "So sánh các sản phẩm gaming"

📝 Sử dụng bàn phím để nhập tin nhắn
{f'🔧 Một số tính năng AI có thể bị hạn chế do lỗi tải mô hình' if not self.text_generator else ''}"""
            else:
                welcome_msg = f"""🤖 Welcome to the AI Sales Assistant!

System Status:
✅ Vietnamese and English support
{'✅' if self.text_generator else '❌'} AI Model: {'Available' if self.text_generator else 'Not available'}
{'✅' if self.embedding_model else '❌'} Smart Search: {'Available' if self.embedding_model else 'Not available'}

Available features:
📦 Smart product search
💬 Conversation context
🌐 Multi-language support
📚 Excel, PDF, Word processing
🖼️ Image OCR

Try asking questions like:
• "I want to buy a gaming laptop"
• "What wireless mice do you have?"
• "How much does the laptop cost?"
• "Compare gaming products"

📝 Use keyboard to type messages
{f'🔧 Some AI features may be limited due to model loading errors' if not self.text_generator else ''}"""

            self.display_message("Hệ thống / System", welcome_msg, "system")
        except Exception as e:
            print(f"Welcome message error: {e}")
        
    def __del__(self):
        """Cleanup with error handling"""
        try:
            if hasattr(self, 'conn'):
                self.conn.close()
        except:
            pass
# NEW WEB CHATBOT FUNCTION FOR AI CREATING SALES PERSON
# ADD THIS NEW FUNCTION HERE (after your VietnameseAISalesBot class ends)
def create_web_app(chatbot_instance):
    """Create Flask web application that shares state with desktop GUI"""
    app = Flask(__name__)
    CORS(app)
    
    @app.route('/')
    def index():
        """Serve the enhanced web interface with full-body photos"""
        return '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI Sales Assistant - Trợ lý AI Bán hàng</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
        }
        .chat-container {
            width: 90%;
            max-width: 1400px;
            height: 90vh;
            background: rgba(255, 255, 255, 0.95);
            border-radius: 20px;
            display: flex;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.1);
        }
        .chat-area {
            flex: 1;
            display: flex;
            flex-direction: column;
            background: rgba(240, 242, 247, 0.8);
        }
        .header {
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            padding: 20px;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }
        .header-title {
            display: flex;
            align-items: center;
            gap: 10px;
        }
        .ai-icon {
            width: 40px;
            height: 40px;
            background: rgba(255, 255, 255, 0.2);
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
        }
        .language-selector {
            background: rgba(255, 255, 255, 0.1);
            border: 1px solid rgba(255, 255, 255, 0.3);
            border-radius: 10px;
            padding: 8px 12px;
            color: white;
            font-size: 14px;
        }
        .messages {
            flex: 1;
            overflow-y: auto;
            padding: 20px;
            display: flex;
            flex-direction: column;
            gap: 15px;
        }
        .message {
            max-width: 80%;
            padding: 15px 20px;
            border-radius: 20px;
            word-wrap: break-word;
            animation: slideIn 0.3s ease-out;
        }
        .user {
            align-self: flex-end;
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            border-bottom-right-radius: 5px;
        }
        .bot {
            align-self: flex-start;
            background: white;
            color: #333;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
            border-bottom-left-radius: 5px;
            border: 1px solid #e0e0e0;
        }
        .system {
            align-self: center;
            background: linear-gradient(135deg, #00b894, #00cec9);
            color: white;
            font-size: 14px;
            max-width: 90%;
        }
        .typing-indicator {
            display: none;
            align-items: center;
            gap: 5px;
            padding: 15px 20px;
            background: white;
            border-radius: 20px;
            border-bottom-left-radius: 5px;
            max-width: 80px;
            margin-bottom: 10px;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
        }
        .typing-dot {
            width: 8px;
            height: 8px;
            background: #6c5ce7;
            border-radius: 50%;
            animation: typing 1.5s infinite;
        }
        .typing-dot:nth-child(2) { animation-delay: 0.2s; }
        .typing-dot:nth-child(3) { animation-delay: 0.4s; }
        .input-area {
            padding: 20px;
            background: white;
            display: flex;
            gap: 15px;
            border-top: 1px solid #e0e0e0;
        }
        .input {
            flex: 1;
            padding: 15px 20px;
            border: 2px solid #e0e0e0;
            border-radius: 25px;
            font-size: 16px;
            outline: none;
            resize: none;
            min-height: 50px;
            max-height: 120px;
        }
        .input:focus { border-color: #6c5ce7; }
        .send-btn {
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            border: none;
            border-radius: 50%;
            width: 50px;
            height: 50px;
            cursor: pointer;
            font-size: 20px;
            transition: transform 0.2s ease;
        }
        .send-btn:hover { transform: scale(1.05); }
        .send-btn:disabled { opacity: 0.6; cursor: not-allowed; }
        .avatar-panel {
            width: 480px;
            background: white;
            padding: 20px;
            display: flex;
            flex-direction: column;
            align-items: center;
            border-left: 1px solid #e0e0e0;
        }
        .avatar-container {
            width: 100%;
            height: 550px;
            border-radius: 15px;
            background: linear-gradient(135deg, #f8f9fa, #e9ecef);
            display: flex;
            align-items: center;
            justify-content: center;
            color: #666;
            margin-bottom: 20px;
            overflow: hidden;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
            border: 3px solid #6c5ce7;
            position: relative;
        }
        .avatar-container img {
            width: 100%;
            height: 100%;
            object-fit: cover;
            border-radius: 12px;
        }
        .status-indicator {
            position: absolute;
            top: 15px;
            right: 15px;
            width: 20px;
            height: 20px;
            background: #00b894;
            border-radius: 50%;
            border: 3px solid white;
            animation: pulse-indicator 1.5s infinite;
        }
        .controls {
            display: flex;
            gap: 10px;
            margin-bottom: 20px;
        }
        .btn {
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            border: none;
            border-radius: 20px;
            padding: 10px 20px;
            cursor: pointer;
            font-size: 14px;
            transition: transform 0.2s ease;
        }
        .btn:hover { transform: translateY(-1px); }
        .info-panel {
            text-align: center;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 10px;
            margin-top: 10px;
        }
        .assistant-name {
            font-size: 20px;
            font-weight: bold;
            color: #2d3436;
            margin-bottom: 5px;
        }
        .assistant-subtitle {
            font-size: 14px;
            color: #636e72;
            margin-bottom: 10px;
        }
        @keyframes slideIn {
            from { opacity: 0; transform: translateY(20px); }
            to { opacity: 1; transform: translateY(0); }
        }
        @keyframes typing {
            0%, 60%, 100% { transform: translateY(0); }
            30% { transform: translateY(-10px); }
        }
        @keyframes pulse-indicator {
            0%, 100% { opacity: 1; }
            50% { opacity: 0.5; }
        }
        @media (max-width: 768px) {
            .chat-container { width: 95%; height: 95vh; flex-direction: column; }
            .avatar-panel { width: 100%; height: 300px; }
            .avatar-container { height: 250px; }
        }
    </style>
</head>
<body>
    <div class="chat-container">
        <div class="chat-area">
            <div class="header">
                <div class="header-title">
                    <div class="ai-icon">🤖</div>
                    <div>
                        <h2>AI Sales Assistant</h2>
                        <p style="font-size: 14px; opacity: 0.9;">Trợ lý AI Bán hàng</p>
                    </div>
                </div>
                <select class="language-selector" id="languageSelector">
                    <option value="vi">🇻🇳 Tiếng Việt</option>
                    <option value="en">🇺🇸 English</option>
                </select>
            </div>
            <div class="messages" id="messages">
                <div class="message system">
                    <strong>🤖 Chào mừng đến với Trợ lý AI Bán hàng!</strong><br><br>
                    Tôi là trợ lý AI xinh đẹp của cửa hàng với kiến thức chuyên sâu về công nghệ. Tôi sử dụng AI models mạnh mẽ để hỗ trợ bạn tìm kiếm sản phẩm phù hợp nhất.<br><br>
                    <strong>Hello! I'm the beautiful AI sales assistant!</strong><br>
                    I use powerful AI models to help you find the perfect products with deep technology knowledge.
                </div>
            </div>
            <div class="typing-indicator" id="typingIndicator">
                <div class="typing-dot"></div>
                <div class="typing-dot"></div>
                <div class="typing-dot"></div>
            </div>
            <div class="input-area">
                <textarea class="input" id="messageInput" placeholder="Nhập tin nhắn của bạn / Type your message..." rows="1"></textarea>
                <button class="send-btn" id="sendButton" onclick="sendMessage()">➤</button>
            </div>
        </div>
        <div class="avatar-panel">
            <div class="avatar-container" id="avatarContainer">
                <div class="status-indicator"></div>
                🎨 Generating Full Body Photo of Beautiful Sales Assistant...
            </div>
            <div class="controls">
                <button class="btn" onclick="generateFullBodyPhoto()">🔄 New Photo</button>
                <button class="btn" onclick="generateDifferentStyle()">✨ New Style</button>
            </div>
            <div class="info-panel">
                <div class="assistant-name">🌟 Beautiful AI Assistant</div>
                <div class="assistant-subtitle">Full Body Professional Photography</div>
                <small style="color: #666;">Powered by Advanced AI Models</small>
            </div>
        </div>
    </div>

    <script>
        let currentLanguage = 'vi';
        let isTyping = false;

        const fullBodyPrompts = [
            'https://image.pollinations.ai/prompt/full%20body%20portrait%20beautiful%20professional%20businesswoman%20sales%20consultant%20standing%20pose%20elegant%20business%20suit%20high%20heels%20confident%20smile%20modern%20office%20background%20photorealistic%20high%20quality%20professional%20photography?width=450&height=700&seed=',
            'https://image.pollinations.ai/prompt/complete%20full%20body%20shot%20stunning%20female%20sales%20representative%20wearing%20professional%20blazer%20skirt%20standing%20confidently%20hands%20on%20hips%20beautiful%20face%20long%20hair%20office%20environment%20realistic%20photo?width=450&height=700&seed=',
            'https://image.pollinations.ai/prompt/full%20length%20body%20gorgeous%20businesswoman%20sales%20agent%20professional%20attire%20standing%20upright%20welcoming%20gesture%20bright%20smile%20modern%20workplace%20head%20to%20toe%20photography%20ultra%20realistic?width=450&height=700&seed=',
            'https://image.pollinations.ai/prompt/entire%20body%20beautiful%20female%20sales%20professional%20business%20dress%20standing%20pose%20confident%20posture%20professional%20makeup%20styled%20hair%20corporate%20office%20setting%20full%20figure%20photorealistic?width=450&height=700&seed=',
            'https://image.pollinations.ai/prompt/complete%20full%20body%20elegant%20sales%20consultant%20woman%20wearing%20business%20outfit%20standing%20professionally%20beautiful%20appearance%20friendly%20expression%20office%20backdrop%20from%20head%20to%20feet%20realistic%20photograph?width=450&height=700&seed='
        ];

        const styleVariations = [
            'https://image.pollinations.ai/prompt/full%20body%20attractive%20asian%20businesswoman%20sales%20representative%20professional%20suit%20standing%20gracefully%20confident%20stance%20office%20setting%20complete%20figure%20photorealistic?width=450&height=700&seed=',
            'https://image.pollinations.ai/prompt/entire%20length%20beautiful%20african%20american%20saleswoman%20business%20professional%20attire%20standing%20pose%20elegant%20posture%20modern%20corporate%20environment%20full%20body%20view%20realistic?width=450&height=700&seed=',
            'https://image.pollinations.ai/prompt/complete%20body%20shot%20gorgeous%20european%20sales%20executive%20wearing%20formal%20business%20wear%20standing%20confidently%20professional%20demeanor%20beautiful%20features%20office%20background%20photographic%20quality?width=450&height=700&seed='
        ];

        function generateFullBodyPhoto() {
            const container = document.getElementById('avatarContainer');
            container.innerHTML = '<div class="status-indicator"></div>🎨 Generating Full Body Photo...';
            
            const randomPrompt = fullBodyPrompts[Math.floor(Math.random() * fullBodyPrompts.length)];
            const randomSeed = Math.floor(Math.random() * 50000);
            const imageUrl = randomPrompt + randomSeed;
            
            loadImage(imageUrl, container);
        }

        function generateDifferentStyle() {
            const container = document.getElementById('avatarContainer');
            container.innerHTML = '<div class="status-indicator"></div>✨ Generating Different Style...';
            
            const randomStyle = styleVariations[Math.floor(Math.random() * styleVariations.length)];
            const randomSeed = Math.floor(Math.random() * 50000);
            const imageUrl = randomStyle + randomSeed;
            
            loadImage(imageUrl, container);
        }

        function loadImage(imageUrl, container) {
            const img = document.createElement('img');
            
            img.onload = function() {
                container.innerHTML = '<div class="status-indicator"></div>';
                container.appendChild(img);
                console.log('✅ Full body photo generated successfully');
            };
            
            img.onerror = function() {
                console.log('❌ Image failed, trying backup...');
                const backupUrl = 'https://image.pollinations.ai/prompt/beautiful%20professional%20businesswoman%20full%20body%20standing%20office%20realistic?width=450&height=700&seed=' + Math.floor(Math.random() * 10000);
                
                const backupImg = document.createElement('img');
                backupImg.onload = function() {
                    container.innerHTML = '<div class="status-indicator"></div>';
                    container.appendChild(backupImg);
                };
                backupImg.onerror = function() {
                    container.innerHTML = '<div class="status-indicator"></div>👩‍💼 Beautiful Sales Assistant<br><small>Click "New Photo" to try again</small>';
                };
                backupImg.src = backupUrl;
            };
            
            img.src = imageUrl;
        }

        function addMessage(text, sender) {
            const messages = document.getElementById('messages');
            const div = document.createElement('div');
            div.className = 'message ' + sender;
            div.innerHTML = text;
            messages.appendChild(div);
            messages.scrollTop = messages.scrollHeight;
        }

        function showTyping() {
            if (isTyping) return;
            isTyping = true;
            document.getElementById('typingIndicator').style.display = 'flex';
            document.getElementById('sendButton').disabled = true;
        }

        function hideTyping() {
            isTyping = false;
            document.getElementById('typingIndicator').style.display = 'none';
            document.getElementById('sendButton').disabled = false;
        }

        async function sendMessage() {
            const input = document.getElementById('messageInput');
            const message = input.value.trim();
            if (!message || isTyping) return;

            addMessage(message, 'user');
            input.value = '';
            autoResize(input);

            showTyping();

            try {
                const response = await fetch('/api/chat', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ message: message, language: currentLanguage })
                });
                
                const data = await response.json();
                hideTyping();
                
                if (data.error) {
                    addMessage(`Error: ${data.response}`, 'system');
                } else {
                    addMessage(data.response, 'bot');
                    console.log(`🧠 AI Response from: ${data.data_source} (${data.processing_time?.toFixed(2)}s)`);
                }
            } catch (error) {
                hideTyping();
                const errorMsg = currentLanguage === 'vi' 
                    ? 'Lỗi kết nối với máy chủ AI. Vui lòng thử lại.' 
                    : 'AI server connection error. Please try again.';
                addMessage(errorMsg, 'system');
                console.error('Error:', error);
            }
        }

        function autoResize(textarea) {
            textarea.style.height = 'auto';
            textarea.style.height = Math.min(textarea.scrollHeight, 120) + 'px';
        }

        function changeLanguage() {
            currentLanguage = document.getElementById('languageSelector').value;
            const input = document.getElementById('messageInput');
            input.placeholder = currentLanguage === 'vi' 
                ? "Nhập tin nhắn của bạn..." 
                : "Type your message...";
        }

        // Event listeners
        document.getElementById('languageSelector').addEventListener('change', changeLanguage);
        document.getElementById('messageInput').addEventListener('keypress', function(e) {
            if (e.key === 'Enter' && !e.shiftKey) {
                e.preventDefault();
                sendMessage();
            }
        });
        document.getElementById('messageInput').addEventListener('input', function(e) {
            autoResize(e.target);
        });

        // Initialize
        generateFullBodyPhoto();
        
        setTimeout(() => {
            addMessage(currentLanguage === 'vi' 
                ? 'Xin chào! Tôi đã sẵn sàng hỗ trợ bạn với AI models mạnh mẽ. Bạn cần tìm sản phẩm gì? 😊'
                : 'Hello! I\'m ready to assist you with powerful AI models. What products are you looking for? 😊', 'bot');
        }, 2000);
    </script>
</body>
</html>'''
    
    @app.route('/api/chat', methods=['POST'])
    def chat():
        """Handle chat messages from web interface - uses full AI models"""
        try:
            data = request.get_json()
            user_message = data.get('message', '')
            language = data.get('language', 'vi')
            
            if not user_message.strip():
                return jsonify({
                    'error': 'Empty message',
                    'response': 'Please enter a message.' if language == 'en' else 'Vui lòng nhập tin nhắn.'
                })
            
            print(f"🌐 Web user ({language}): {user_message}")
            
            # Use the FULL chatbot processing with all AI models
            result = chatbot_instance.process_message(user_message)
            
            print(f"🧠 AI Response generated from: {result.get('data_source')} in {result.get('processing_time', 0):.2f}s")
            
            return jsonify({
                'response': result.get('response', ''),
                'data_source': result.get('data_source', ''),
                'processing_time': result.get('processing_time', 0),
                'user_language': result.get('user_language', language),
                'local_results_count': result.get('local_results_count', 0)
            })
            
        except Exception as e:
            error_msg = f"Server error: {str(e)}"
            print(f"❌ Web interface error: {error_msg}")
            return jsonify({
                'error': error_msg,
                'response': 'Sorry, I encountered an error.' if language == 'en' else 'Xin lỗi, tôi gặp lỗi.'
            }), 500
    
    @app.route('/api/status')
    def status():
        """Check chatbot status"""
        return jsonify({
            'status': 'online',
            'model_available': chatbot_instance.text_generator is not None,
            'embedding_available': chatbot_instance.embedding_model is not None,
            'desktop_gui_running': hasattr(chatbot_instance, 'root') and chatbot_instance.root is not None,
            'conversation_count': len(chatbot_instance.conversation_context),
            'gpu_available': torch.cuda.is_available(),
            'device': str(getattr(chatbot_instance, 'primary_device', 'cpu'))
        })
    
    return app


def main():
    """Main function that starts both desktop and web interfaces with full AI models"""
    print("""
    ===============================================
    🚀 Vietnamese AI Sales ChatBot - Full Power
    ===============================================
    
    🔧 Features:
    ✅ Vietnamese + English language support
    ✅ Full AI models with GPU acceleration
    ✅ Smart product search with embeddings
    ✅ Desktop GUI + Web Interface with full-body photos
    ✅ Advanced natural language processing
    
    🚀 Starting BOTH interfaces with full AI power...
    """)
    
    try:
        os.makedirs('data', exist_ok=True)
        os.makedirs('logs', exist_ok=True)
        os.makedirs('templates', exist_ok=True)
        
        if not os.path.exists('chatbot_config.yaml'):
            print("📝 Creating configuration file...")
            create_config_file()
        
        print("🤖 Initializing full AI chatbot with GPU acceleration...")
        
        # Initialize with FULL model loading (no bypassing)
        chatbot = VietnameseAISalesBot(start_gui=False)
        print("✅ Full AI models loaded successfully!")
        
        # Add sample products
        print("📦 Adding sample products...")
        chatbot.add_sample_products()
        
        # Start web interface in background thread
        def start_web_interface():
            try:
                print("🌐 Starting web interface with full AI integration...")
                app = create_web_app(chatbot)
                print("🔗 Web interface available at: http://localhost:5000")
                print("🎨 Full-body photos + Full AI models included!")
                app.run(debug=False, host='0.0.0.0', port=5000, use_reloader=False, threaded=True)
            except Exception as e:
                print(f"❌ Web interface error: {e}")
                import traceback
                traceback.print_exc()
        
        # Start web interface in background
        web_thread = threading.Thread(target=start_web_interface, daemon=True)
        web_thread.start()
        
        # Give web interface time to start
        time.sleep(3)
        
        # Setup and start desktop GUI
        print("🖥️ Starting desktop GUI with full AI models...")
        chatbot.setup_gui()
        chatbot.display_welcome_message()
        
        print("\n" + "="*60)
        print("🎉 BOTH INTERFACES RUNNING WITH FULL AI POWER!")
        print("🖥️  Desktop GUI: Full AI models + Traditional interface")
        print("🌐 Web Interface: Full AI models + Beautiful full-body photos")
        print("🧠 Shared AI Brain: Same powerful models for both interfaces")
        print("📱 Use either interface - full capabilities on both!")
        print("🔄 Close this window to stop both interfaces")
        print("="*60 + "\n")
        
        # Start desktop GUI (blocking)
        chatbot.root.mainloop()
        
    except KeyboardInterrupt:
        print("\n👋 Both interfaces stopped by user")
    except Exception as e:
        print(f"❌ Error starting full AI interfaces: {e}")
        print("\n🔧 Troubleshooting tips:")
        print("1. Check GPU memory availability")
        print("2. Verify CUDA installation")
        print("3. Ensure model files are accessible")
        
        import traceback
        traceback.print_exc()
        
        input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()