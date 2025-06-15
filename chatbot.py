import os
import json
import sqlite3
import pandas as pd
import numpy as np
from datetime import datetime
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox, font
import threading
import queue
import time
import logging
from pathlib import Path
import re
import yaml
# Excel import functionality
from openpyxl import load_workbook
from typing import Dict, List, Any
import contextlib
# NEW IMPORTS FOR AI GENERATING SALES PERSON IN WEB CHATBOT
from flask import Flask, render_template, request, jsonify
from flask_cors import CORS

# AI/ML imports
import torch
from transformers import (
    AutoTokenizer, AutoModelForCausalLM, 
    BitsAndBytesConfig, pipeline
)
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity

# Web search imports
import requests
from bs4 import BeautifulSoup
from googlesearch import search
import urllib.parse

# Document processing
import docx
from PyPDF2 import PdfReader
from openpyxl import load_workbook
import easyocr
import socket

# Language detection and translation
try:
    from googletrans import Translator, LANGUAGES
    import langdetect
    TRANSLATION_AVAILABLE = True
    print("✅ Translation libraries loaded successfully")
except ImportError:
    TRANSLATION_AVAILABLE = False
    print("⚠️ Translation libraries not available. Install with: pip install googletrans==4.0.0rc1 langdetect")
    
# Web interface imports
from flask import Flask, render_template, request, jsonify
import argparse
import socket

def create_config_file():
    """Create optimized configuration file for RTX 4090"""
    config = {
        'version': '3.0_rtx4090_optimized',
        'environment': 'production',
        
        # RTX 4090 Optimized AI Models Configuration
        'ai_models': {
            'primary_llm': 'microsoft/DialoGPT-large',  # Use large model for RTX 4090
            'fallback_llm': 'microsoft/DialoGPT-medium',
            'embedding_model': 'sentence-transformers/all-MiniLM-L6-v2',
            'model_cache_dir': './models_cache',
            'use_model_caching': True
        },
        
        # RTX 4090 Specific GPU Configuration
        'gpu_config': {
            'primary_device': 'cuda:0',
            'use_quantization': False,  # RTX 4090 has enough VRAM
            'mixed_precision': True,    # Use for speed boost
            'max_memory_per_gpu': 0.90, # Use 90% of 24GB VRAM
            'batch_size': 8,            # Larger batch for RTX 4090
            'gradient_accumulation_steps': 2,
            'torch_compile': True,      # PyTorch 2.0 optimization
            'flash_attention': True,    # For faster attention computation
            'memory_efficient_attention': True
        },
        
        # Enhanced Search Configuration
        'search_config': {
            'local_similarity_threshold': 0.1,
            'enable_google_search': True,
            'max_google_results': 3,
            'search_timeout': 10,
            'use_batch_search': True,   # Process multiple queries together
            'cache_embeddings': True    # Cache embeddings for faster lookup
        },
        
        # RTX 4090 Optimized Response Generation
        'performance': {
            'max_response_length': 200,  # Longer responses for better quality
            'temperature': 0.7,
            'top_p': 0.9,
            'top_k': 50,
            'repetition_penalty': 1.1,
            'do_sample': True,
            'num_return_sequences': 1,
            'use_cache': True,           # Enable KV cache
            'pad_token_id': None,        # Will be set during model loading
            'eos_token_id': None
        },
        
        # Language and Conversation Configuration
        'language_config': {
            'default_language': 'vi',
            'supported_languages': ['vi', 'en'],
            'auto_detect_language': True,
            'translate_responses': True
        },
        
        'conversation_flows': {
            'greeting_vietnamese': {
                'triggers': ['xin chào', 'chào', 'chào bạn', 'chào anh', 'chào chị'],
                'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng với sức mạnh RTX 4090. Tôi có thể giúp bạn tìm kiếm sản phẩm phù hợp với nhu cầu. Bạn đang tìm kiếm gì hôm nay?',
                'next_stage': 'needs_assessment'
            },
            'greeting': {
                'triggers': ['hello', 'hi', 'hey'],
                'response_template': 'Hello! I am your RTX 4090-powered AI sales assistant. How can I help you find the perfect product today?'
            }
        },
        
        'vietnamese_templates': {
            'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp với tiêu chí của bạn.',
            'out_of_stock': 'Rất tiếc, {product_name} hiện đang hết hàng.',
            'processing': 'Đang xử lý với sức mạnh RTX 4090...',
            'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn. Vui lòng thử lại.',
            'model_error': 'Mô hình AI không khả dụng. Đang sử dụng phản hồi dự phòng.',
            'generation_error': 'Có lỗi khi tạo phản hồi. Vui lòng thử lại.',
            'empty_response': 'Tôi hiểu câu hỏi của bạn nhưng cần thêm thông tin để trả lời tốt hơn.',
            'connection_error': 'Lỗi kết nối. Vui lòng kiểm tra mạng và thử lại.'
        },
        
        'analytics': {
            'track_conversations': True,
            'track_language_usage': True,
            'track_errors': True,
            'performance_monitoring': True
        }
    }
    
    with open('chatbot_config.yaml', 'w', encoding='utf-8') as file:
        yaml.dump(config, file, default_flow_style=False, allow_unicode=True)
    
    print("✅ RTX 4090 optimized configuration created: chatbot_config.yaml")
class SmartExcelReader:
    def __init__(self):
        self.supported_formats = {'.xlsx', '.xls', '.xlsm', '.csv'}
    
    def read_file_optimized(self, filepath, sheet_name=0):
        """Automatically select best reading method based on file size and format"""
        file_path = Path(filepath)
        file_size = file_path.stat().st_size
        
        # For files < 50MB, use pandas directly
        if file_size < 50 * 1024 * 1024:
            if file_path.suffix.lower() == '.csv':
                return pd.read_csv(filepath, encoding='utf-8-sig')
            return pd.read_excel(filepath, sheet_name=sheet_name, engine='openpyxl')
        
        # For larger files, use openpyxl read-only mode
        wb = load_workbook(filepath, read_only=True, data_only=True)
        ws = wb[sheet_name] if isinstance(sheet_name, str) else wb.worksheets[sheet_name]
        
        data = list(ws.values)
        return pd.DataFrame(data[1:], columns=data[0]) if data else pd.DataFrame()

class DataProcessor:
    def __init__(self, category: str):
        self.category = category
        self.column_mappings = {
            'SSD': {
                'product_name': ['Sản phẩm', 'Product Name', 'Name', 'Item', 'Product'],
                'model': ['Model', 'Model Name'],
                'interface': ['Giao thức', 'Interface', 'Connection', 'Type'],
                'specifications': ['Thông số', 'Specifications', 'Specs'],
                'stock_status': ['Kho', 'Stock', 'Availability', 'Status'],
                'price': ['Giá bán lẻ', 'Price', 'Cost', 'Amount']
            },
            'Memory': {
                'product_name': ['Sản phẩm', 'Product Name', 'Name', 'Item'],
                'model': ['Model', 'Model Name'],
                'interface': ['Giao thức', 'Interface', 'Type'],
                'specifications': ['Thông số', 'Specifications', 'Specs'],
                'stock_status': ['Kho', 'Stock', 'Availability'],
                'price': ['Giá bán lẻ', 'Price', 'Cost', 'Amount']
            },
            'Motherboard': {
                'product_name': ['Sản phẩm', 'Product Name', 'Name', 'Item'],
                'model': ['Model', 'Model Name'],
                'chipset': ['Chipset', 'Chip Set'],
                'specifications': ['Thông số', 'Specifications', 'Specs'],
                'stock_status': ['Kho', 'Stock', 'Availability'],
                'price': ['Giá bán lẻ', 'Price', 'Cost', 'Amount']
            }
        }
    
    def process_excel_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """Process and validate Excel data for database insertion"""
        # Map columns
        mapped_df = self._map_columns(df)
        
        # Validate data types
        validated_df = self._validate_data_types(mapped_df)
        
        # Clean and transform data
        cleaned_df = self._clean_data(validated_df)
        
        # Add metadata
        cleaned_df['category'] = self.category
        cleaned_df['source_file'] = 'Excel Import'
        cleaned_df['created_at'] = datetime.now().isoformat()
        
        return cleaned_df
    
    def _map_columns(self, df: pd.DataFrame) -> pd.DataFrame:
        """Intelligently map Excel columns to database fields"""
        mapping = self.column_mappings[self.category]
        result_df = pd.DataFrame()
        
        for db_field, possible_names in mapping.items():
            # Find matching column
            matched_column = None
            for excel_col in df.columns:
                if any(name.lower() in excel_col.lower() for name in possible_names):
                    matched_column = excel_col
                    break
            
            if matched_column:
                result_df[db_field] = df[matched_column]
            else:
                # Handle missing columns with defaults
                if db_field in ['specifications', 'stock_status']:
                    result_df[db_field] = 'N/A'
                else:
                    result_df[db_field] = ''
        
        return result_df
    
    def _validate_data_types(self, df: pd.DataFrame) -> pd.DataFrame:
        """Validate and convert data types"""
        # Price validation and conversion
        if 'price' in df.columns:
            df['price'] = pd.to_numeric(
                df['price'].astype(str).str.replace(r'[^\d.]', '', regex=True),
                errors='coerce'
            )
            # Fill missing prices with 0
            df['price'] = df['price'].fillna(0)
        
        return df
    
    def _clean_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """Clean and prepare data for database insertion"""
        # Remove completely empty rows
        df = df.dropna(how='all')
        
        # Clean string columns
        string_columns = ['product_name', 'model', 'interface', 'specifications', 'stock_status']
        for col in string_columns:
            if col in df.columns:
                df[col] = df[col].astype(str).str.strip()
                df[col] = df[col].replace(['nan', 'None', ''], 'N/A')
        
        # Handle missing required fields
        if 'product_name' in df.columns:
            df = df[df['product_name'] != 'N/A']
        
        return df
class VietnameseAISalesBot:
    """
    RTX 4090 Optimized AI Sales ChatBot with Vietnamese language support
    Enhanced for maximum performance with high-end GPU
    """
    
    def __init__(self, start_gui=False):
        self.load_config()
        self.setup_logging()
        self.setup_language_support()
        self.setup_database()
        self.initialize_ai_models_rtx4090()  # RTX 4090 optimized initialization
        self.conversation_context = []
        self.conversation_summary = ""
        # Initialize product databases
        self.ssd_database = []
        self.memory_database = []
        self.motherboard_database = []
        self.response_cache = {}  # Add response caching
        self.customer_profiles = {}  # Store customer profiles by session ID
        if start_gui:
            self.setup_gui()
        # Debug database location
        # Debug database location
        import os
        script_dir = os.path.dirname(os.path.abspath(__file__))
        self.db_path = os.path.join(script_dir, 'chatbot_data.db')

        print(f"📁 Script directory: {script_dir}")
        print(f"📁 Database path: {self.db_path}")
        print(f"📊 Database exists: {os.path.exists(self.db_path)}")

        # Always use the database in the script directory
        self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        print(f"✅ Connected to database at: {self.db_path}")

        # Continue with setup_database
        self.setup_database()
    def load_config(self):
        """Load RTX 4090 optimized configuration"""
        try:
            with open('chatbot_config.yaml', 'r', encoding='utf-8') as file:
                self.config = yaml.safe_load(file)
            print(f"✅ Configuration loaded: {self.config.get('version', 'unknown')}")
            
            self.config = self.update_config_for_vietnamese()
            
        except FileNotFoundError:
            print("⚠️ chatbot_config.yaml not found, creating RTX 4090 optimized config")
            create_config_file()
            with open('chatbot_config.yaml', 'r', encoding='utf-8') as file:
                self.config = yaml.safe_load(file)
        except Exception as e:
            print(f"❌ Error loading config: {e}")
            self.config = self.get_default_rtx4090_config()
    
    def get_default_rtx4090_config(self):
        """Get RTX 4090 optimized default configuration"""
        return {
            'version': '3.0_rtx4090_optimized',
            'environment': 'production',
            'ai_models': {
                'primary_llm': 'microsoft/DialoGPT-large',
                'fallback_llm': 'microsoft/DialoGPT-medium',
                'embedding_model': 'sentence-transformers/all-MiniLM-L6-v2'
            },
            'gpu_config': {
                'primary_device': 'cuda:0',
                'use_quantization': False,
                'mixed_precision': True,
                'max_memory_per_gpu': 0.90,
                'batch_size': 8,
                'torch_compile': True
            },
            'performance': {
                'max_response_length': 200,
                'temperature': 0.7,
                'top_p': 0.9,
                'repetition_penalty': 1.1,
                'do_sample': True,
                'use_cache': True
            },
            'language_config': {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True
            }
        }
    
    def update_config_for_vietnamese(self):
        """Update existing config to support Vietnamese features"""
        config = self.config.copy()
        
        if 'language_config' not in config:
            config['language_config'] = {
                'default_language': 'vi',
                'supported_languages': ['vi', 'en'],
                'auto_detect_language': True,
                'translate_responses': True
            }
        
        if 'conversation_flows' not in config:
            config['conversation_flows'] = {}
            
        config['conversation_flows'].update({
            'greeting_vietnamese': {
                'triggers': ['xin chào', 'chào', 'chào bạn', 'chào anh', 'chào chị'],
                'response_template': 'Xin chào! Tôi là trợ lý AI bán hàng với sức mạnh RTX 4090. Tôi có thể giúp bạn tìm kiếm sản phẩm phù hợp với nhu cầu. Bạn đang tìm kiếm gì hôm nay?',
                'next_stage': 'needs_assessment'
            }
        })
        
        if 'vietnamese_templates' not in config:
            config['vietnamese_templates'] = {
                'no_products_found': 'Tôi không thể tìm thấy sản phẩm nào phù hợp.',
                'error': 'Xin lỗi, tôi gặp lỗi khi xử lý yêu cầu của bạn.',
                'model_error': 'Mô hình AI không khả dụng.',
                'generation_error': 'Có lỗi khi tạo phản hồi.',
                'empty_response': 'Tôi cần thêm thông tin để trả lời tốt hơn.'
            }
        
        return config
    
    def setup_logging(self):
        """Setup logging configuration with Vietnamese support"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('chatbot.log', encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
        
    def setup_language_support(self):
        """Initialize language detection and translation"""
        self.current_language = self.config['language_config']['default_language']
        
        if TRANSLATION_AVAILABLE:
            try:
                self.translator = Translator()
                print("✅ Translation service initialized")
            except Exception as e:
                print(f"⚠️ Translation service initialization failed: {e}")
                self.translator = None
        else:
            self.translator = None
    
    def detect_language(self, text):
        """Detect language of input text with error handling"""
        if not TRANSLATION_AVAILABLE or not text or not text.strip():
            return 'vi'
        
        try:
            detected = langdetect.detect(text)
            if detected in self.config['language_config']['supported_languages']:
                return detected
            return 'vi'
        except Exception as e:
            print(f"Language detection error: {e}")
            return 'vi'
    
    def translate_text(self, text, target_language):
        """Translate text to target language with error handling"""
        if not self.translator or not text or not text.strip():
            return text
        
        try:
            translated = self.translator.translate(text, dest=target_language)
            return translated.text if translated and translated.text else text
        except Exception as e:
            print(f"Translation error: {e}")
            return text
    
    def setup_database(self):
        """Initialize SQLite database with correct path"""
        import os
        
        # Define the exact path where you want the database
        # Option 1: Same folder as the script
        db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'chatbot_data.db')
        
        # Option 2: In the parent folder (E:\SSTCCloud\AI CHATBOT\)
        # db_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'chatbot_data.db')
        
        print(f"📁 Using database at: {db_path}")
        
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.conn.execute("PRAGMA encoding = 'UTF-8'")
        """Initialize SQLite database with RTX 4090 optimizations"""
        self.conn = sqlite3.connect('chatbot_data.db', check_same_thread=False)
        self.conn.execute("PRAGMA encoding = 'UTF-8'")
        
        # RTX 4090 Database Optimizations
        cursor = self.conn.cursor()
        
        # Enable WAL mode for better performance
        cursor.execute("PRAGMA journal_mode=WAL")
        cursor.execute("PRAGMA cache_size=20000")  # Larger cache for RTX 4090 system
        cursor.execute("PRAGMA temp_store=MEMORY")
        cursor.execute("PRAGMA mmap_size=268435456")  # 256MB memory mapping
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS products (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                name_vietnamese TEXT,
                description TEXT,
                description_vietnamese TEXT,
                category TEXT,
                category_vietnamese TEXT,
                price REAL,
                features TEXT,
                features_vietnamese TEXT,
                specifications TEXT,
                specifications_vietnamese TEXT,
                availability TEXT,
                availability_vietnamese TEXT,
                source_file TEXT,
                embedding BLOB,
                created_at TEXT,
                updated_at TEXT
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT,
                user_input TEXT,
                user_language TEXT,
                bot_response TEXT,
                bot_language TEXT,
                intent TEXT,
                confidence REAL,
                data_source TEXT,
                response_time REAL,
                error_type TEXT,
                error_message TEXT
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS knowledge_base (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                topic TEXT,
                content TEXT,
                source TEXT,
                embedding BLOB,
                created_at TEXT
            )
        ''')
        
        # Create optimized indexes for RTX 4090 performance
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_products_name ON products(name)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_products_category ON products(category)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_products_price ON products(price)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_conversations_timestamp ON conversations(timestamp)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_conversations_language ON conversations(user_language)")
        
        self.conn.commit()
        print("✅ RTX 4090 optimized database schema ready")
    
    def initialize_ai_models_rtx4090(self):
        """RTX 4090 optimized AI model initialization"""
        self.logger.info("🚀 Initializing RTX 4090 optimized AI models...")
        
        try:
            self.setup_rtx4090_devices()
            self.load_rtx4090_language_model()
            self.load_rtx4090_embedding_model()
            self.load_ocr_model()
            
            self.logger.info("✅ All RTX 4090 optimized AI models loaded successfully!")
            
        except Exception as e:
            self.logger.error(f"❌ Error initializing RTX 4090 AI models: {e}")
            self.text_generator = None
            self.embedding_model = None
            self.ocr_reader = None
    
    def setup_rtx4090_devices(self):
        """RTX 4090 specific device setup"""
        try:
            if torch.cuda.is_available():
                gpu_count = torch.cuda.device_count()
                print(f"🎮 Found {gpu_count} GPU(s)")
                
                # Check for RTX 4090 specifically
                for i in range(gpu_count):
                    props = torch.cuda.get_device_properties(i)
                    memory_gb = props.total_memory / (1024**3)
                    print(f"GPU {i}: {props.name} - {memory_gb:.1f}GB")
                    
                    # RTX 4090 has ~24GB VRAM
                    if memory_gb > 20:
                        print(f"🚀 RTX 4090 detected on GPU {i}! Enabling maximum performance mode.")
                        self.primary_device = f'cuda:{i}'
                        self.primary_memory_gb = memory_gb
                        self.is_rtx4090 = True
                        break
                else:
                    # Fallback for other GPUs
                    self.primary_device = 'cuda:0'
                    self.primary_memory_gb = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                    self.is_rtx4090 = False
                
                self.secondary_device = f'cuda:{1 % gpu_count}' if gpu_count > 1 else self.primary_device
                
                # Set RTX 4090 specific optimizations
                if hasattr(self, 'is_rtx4090') and self.is_rtx4090:
                    torch.backends.cuda.matmul.allow_tf32 = True
                    torch.backends.cudnn.allow_tf32 = True
                    torch.backends.cudnn.benchmark = True
                    print("✅ RTX 4090 optimizations enabled: TF32, cuDNN benchmark")
                
            else:
                self.primary_device = 'cpu'
                self.secondary_device = 'cpu'
                self.primary_memory_gb = 8
                self.is_rtx4090 = False
                print("⚠️ CUDA not available. Using CPU mode.")
                
        except Exception as e:
            print(f"Device setup error: {e}")
            self.primary_device = 'cpu'
            self.secondary_device = 'cpu'
            self.primary_memory_gb = 8
            self.is_rtx4090 = False
    
    def load_rtx4090_language_model(self):
        """Load language model optimized for RTX 4090"""
        try:
            # RTX 4090 can handle larger models without quantization
            if hasattr(self, 'is_rtx4090') and self.is_rtx4090:
                model_name = self.config['ai_models']['primary_llm']  # Use DialoGPT-large
                use_quantization = False
                print("🚀 RTX 4090 detected: Loading large model without quantization")
            else:
                model_name = self.config['ai_models']['fallback_llm']
                use_quantization = self.config['gpu_config']['use_quantization']
                print(f"📱 Using fallback model with quantization: {use_quantization}")
            
            print(f"🔄 Loading language model: {model_name}")
            
            # Load tokenizer
            try:
                self.tokenizer = AutoTokenizer.from_pretrained(
                    model_name,
                    cache_dir=self.config['ai_models'].get('model_cache_dir', './models_cache')
                )
                if self.tokenizer.pad_token is None:
                    self.tokenizer.pad_token = self.tokenizer.eos_token
                print("✅ Tokenizer loaded successfully")
            except Exception as e:
                print(f"❌ Tokenizer loading failed: {e}")
                self.tokenizer = None
                self.text_generator = None
                return
            
            # Setup quantization config if needed
            quantization_config = None
            if use_quantization and torch.cuda.is_available():
                try:
                    quantization_config = BitsAndBytesConfig(
                        load_in_8bit=True,
                        llm_int8_threshold=6.0,
                    )
                    print("✅ Quantization config created")
                except Exception as e:
                    print(f"⚠️ Quantization config failed: {e}")
                    quantization_config = None
            
            # Load model with RTX 4090 optimizations
            try:
                model_kwargs = {
                    'quantization_config': quantization_config,
                    'torch_dtype': torch.float16 if torch.cuda.is_available() else torch.float32,
                    'trust_remote_code': True,
                    'low_cpu_mem_usage': True,
                    'cache_dir': self.config['ai_models'].get('model_cache_dir', './models_cache')
                }
                
                # RTX 4090 specific optimizations
                if hasattr(self, 'is_rtx4090') and self.is_rtx4090:
                    model_kwargs['device_map'] = {'': self.primary_device}
                    model_kwargs['max_memory'] = {0: "22GB"}  # Leave 2GB for other operations
                
                self.llm_model = AutoModelForCausalLM.from_pretrained(model_name, **model_kwargs)
                
                # Enable PyTorch 2.0 compilation for RTX 4090
                if hasattr(self, 'is_rtx4090') and self.is_rtx4090 and self.config['gpu_config'].get('torch_compile', False):
                    try:
                        self.llm_model = torch.compile(self.llm_model, mode="reduce-overhead")
                        print("🚀 PyTorch 2.0 compilation enabled for RTX 4090")
                    except Exception as e:
                        print(f"⚠️ PyTorch compilation failed: {e}")
                
                print("✅ Language model loaded successfully")
            except Exception as e:
                print(f"❌ Model loading failed: {e}")
                self.llm_model = None
                self.text_generator = None
                return
            
            # Create optimized text generation pipeline - FIXED for accelerate
            try:
                pipeline_kwargs = {
                    "model": self.llm_model,
                    "tokenizer": self.tokenizer,
                    "torch_dtype": torch.float16 if torch.cuda.is_available() else torch.float32,
                    # REMOVED "device" parameter - accelerate handles device placement
                }
                
                # RTX 4090 specific pipeline optimizations
                if hasattr(self, 'is_rtx4090') and self.is_rtx4090:
                    pipeline_kwargs["batch_size"] = self.config['gpu_config'].get('batch_size', 8)
                    print("🚀 RTX 4090: Creating pipeline without device specification")
                
                self.text_generator = pipeline("text-generation", **pipeline_kwargs)
                print("✅ RTX 4090 optimized text generation pipeline created successfully")
                
            except Exception as e:
                print(f"❌ Pipeline creation failed: {e}")
                print("🔄 Trying fallback pipeline creation...")
                
                # Fallback: Create pipeline with minimal parameters
                try:
                    self.text_generator = pipeline(
                        "text-generation",
                        model=self.llm_model,
                        tokenizer=self.tokenizer
                        # No device parameter at all
                    )
                    print("✅ Fallback pipeline created successfully")
                except Exception as fallback_error:
                    print(f"❌ Fallback pipeline creation also failed: {fallback_error}")
                    self.text_generator = None
            
        except Exception as e:
            print(f"❌ Error loading RTX 4090 language model: {e}")
            self.text_generator = None
    
    def load_rtx4090_embedding_model(self):
        """Load embedding model optimized for RTX 4090"""
        try:
            embedding_model_name = self.config['ai_models']['embedding_model']
            print(f"🔄 Loading RTX 4090 optimized embedding model: {embedding_model_name}")
            
            self.embedding_model = SentenceTransformer(embedding_model_name)
            
            if torch.cuda.is_available():
                try:
                    self.embedding_model = self.embedding_model.to(self.secondary_device)
                    
                    # RTX 4090 optimizations for embeddings
                    if hasattr(self, 'is_rtx4090') and self.is_rtx4090:
                        # Enable optimized attention for embeddings
                        if hasattr(self.embedding_model, '_modules'):
                            for module in self.embedding_model._modules.values():
                                if hasattr(module, 'enable_amp'):
                                    module.enable_amp = True
                        print("🚀 RTX 4090 embedding optimizations enabled")
                    
                except Exception as e:
                    print(f"⚠️ Could not optimize embedding model for RTX 4090: {e}")
                    
            print("✅ RTX 4090 optimized embedding model loaded")
            
        except Exception as e:
            print(f"❌ Error loading RTX 4090 embedding model: {e}")
            self.embedding_model = None
    
    def load_ocr_model(self):
        """Load OCR model with error handling"""
        try:
            self.ocr_reader = easyocr.Reader(['en', 'vi'])
            print("✅ OCR model loaded successfully with Vietnamese support")
        except Exception as e:
            print(f"❌ Error loading OCR model: {e}")
            self.ocr_reader = None
    
    def safe_generate_response_rtx4090(self, prompt, user_language='vi'):
        """RTX 4090 optimized response generation with comprehensive error handling"""
        try:
            if not self.text_generator or not self.tokenizer:
                self.logger.warning("Text generator not available")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['model_error']
                else:
                    return "I apologize, but the AI text generation system is not available right now."
            
            if not prompt or not isinstance(prompt, str) or not prompt.strip():
                self.logger.warning("Invalid prompt provided")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need more information to provide a helpful response."
            
            # Check response cache for RTX 4090 optimization
            cache_key = f"{prompt[:100]}_{user_language}"
            if cache_key in self.response_cache:
                return self.response_cache[cache_key]
            
            performance_config = self.config.get('performance', {})
            
            try:
                with torch.no_grad():
                    # RTX 4090 optimized generation parameters
                    generation_kwargs = {
                        'max_new_tokens': performance_config.get('max_response_length', 200),
                        'min_new_tokens': 15,
                        'temperature': performance_config.get('temperature', 0.7),
                        'top_p': performance_config.get('top_p', 0.9),
                        'top_k': performance_config.get('top_k', 50),
                        'repetition_penalty': performance_config.get('repetition_penalty', 1.1),
                        'do_sample': performance_config.get('do_sample', True),
                        'pad_token_id': self.tokenizer.pad_token_id or self.tokenizer.eos_token_id,
                        'num_return_sequences': performance_config.get('num_return_sequences', 1),
                        'return_full_text': False,
                        'use_cache': performance_config.get('use_cache', True)
                    }
                    
                    # RTX 4090 specific optimizations
                    if hasattr(self, 'is_rtx4090') and self.is_rtx4090:
                        generation_kwargs['batch_size'] = min(4, self.config['gpu_config'].get('batch_size', 8))
                    else:
                        generation_kwargs['batch_size'] = 1
                    
                    generated = self.text_generator(prompt, **generation_kwargs)
                
                if not generated:
                    self.logger.warning("Empty generation result")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['generation_error']
                    else:
                        return "I'm having trouble generating a response. Please try again."
                
                if isinstance(generated, list) and len(generated) > 0:
                    first_result = generated[0]
                    if isinstance(first_result, dict) and 'generated_text' in first_result:
                        response = first_result['generated_text']
                    else:
                        response = str(first_result)
                else:
                    response = str(generated)
                
                if not response or not response.strip():
                    self.logger.warning("Empty response generated")
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question but need more details to respond properly."
                
                cleaned_response = self.clean_generated_response(response)
                
                if not cleaned_response or len(cleaned_response.strip()) < 5:
                    if user_language == 'vi':
                        return self.config['vietnamese_templates']['empty_response']
                    else:
                        return "I understand your question. Could you provide more details?"
                
                # Cache successful responses for RTX 4090 optimization
                if len(self.response_cache) < 100:  # Limit cache size
                    self.response_cache[cache_key] = cleaned_response
                
                return cleaned_response
                
            except torch.cuda.OutOfMemoryError:
                self.logger.error("CUDA out of memory")
                # Clear cache and try again with smaller parameters
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                if user_language == 'vi':
                    return "Hệ thống đang quá tải. Vui lòng thử lại sau."
                else:
                    return "System is overloaded. Please try again later."
                    
            except Exception as generation_error:
                self.logger.error(f"Text generation error: {generation_error}")
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['generation_error']
                else:
                    return "I encountered an error while generating a response. Please try again."
            
        except Exception as e:
            self.logger.error(f"Safe generation error: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having technical difficulties. Please try again."
    
    def clean_generated_response(self, response):
        """Clean up generated response with error handling"""
        try:
            if not response:
                return ""
            
            response = str(response)
            
            # Remove HTML tags and special tokens
            response = re.sub(r'<[^>]+>', '', response)
            response = re.sub(r'\[.*?\]', '', response)
            response = re.sub(r'\n+', ' ', response)
            response = re.sub(r'\s+', ' ', response)
            
            # Remove special model tokens
            response = response.replace('<|endoftext|>', '')
            response = response.replace('</s>', '')
            response = response.replace('<s>', '')
            response = response.replace('<pad>', '')
            
            # Clean up incomplete sentences for better quality
            try:
                sentences = response.split('.')
                if len(sentences) > 1 and len(sentences[-1].strip()) < 10:
                    response = '.'.join(sentences[:-1]) + '.'
            except Exception:
                pass
            
            response = response.strip()
            
            if len(response) < 3:
                return ""
                
            return response
            
        except Exception as e:
            self.logger.error(f"Response cleaning error: {e}")
            return str(response) if response else ""
    
    def process_message(self, user_input, session_id=None):
        """Main processing pipeline with RTX 4090 optimizations and comprehensive error handling"""
        start_time = time.time()
        
        try:
            # CREATE SESSION ID IF NOT PROVIDED
            if not session_id:
                session_id = f"web_{int(time.time())}"
            
            # CHECK IF THIS IS A PROFILE SETUP MESSAGE
            if self.is_profile_setup_message(user_input):
                return self.handle_profile_setup(user_input, session_id)
            
            # GET CUSTOMER PROFILE FOR THIS SESSION
            customer_profile = self.customer_profiles.get(session_id, {})
            
            # IF NO PROFILE EXISTS, REQUEST PROFILE SETUP
            if not customer_profile:
                return {
                    'response': 'Kính chào quý khách! Để tiện xưng hô, quý khách vui lòng cho biết tuổi và giới tính (ví dụ: "25 tuổi nam" hoặc "30 tuổi nữ") ạ.',
                    'data_source': 'profile_setup',
                    'processing_time': time.time() - start_time,
                    'user_language': 'vi',
                    'response_language': 'vi',
                    'requires_profile': True
                }
            
            # CONTINUE WITH EXISTING VALIDATION
            if not user_input or not isinstance(user_input, str):
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'invalid_input'
                }
            
            user_input = user_input.strip()
            if not user_input:
                error_msg = "Vui lòng nhập tin nhắn." if self.current_language == 'vi' else "Please enter a message."
                return {
                    'response': error_msg,
                    'data_source': 'error',
                    'local_results_count': 0,
                    'processing_time': time.time() - start_time,
                    'user_language': self.current_language,
                    'response_language': self.current_language,
                    'error': 'empty_input'
                }
            
            # Language detection
            try:
                user_language = self.detect_language(user_input)
            except Exception as e:
                self.logger.warning(f"Language detection error: {e}")
                user_language = self.current_language or 'vi'
            
            # Context processing
            context_references = self.check_context_references(user_input, user_language)
            
            # Database search with RTX 4090 optimizations
            local_results = []
            knowledge_results = []
            database_searched = False
            
            try:
                if context_references:
                    search_query = self.build_contextual_search_query(user_input, context_references)
                    self.logger.info(f"Contextual search for: {search_query}")
                else:
                    search_query = user_input
                    
                self.logger.info(f"RTX 4090 optimized search for: {search_query}")
                local_results = self.search_local_database_rtx4090(search_query)
                database_searched = True
                
                if not local_results:
                    self.logger.info("No products found, searching knowledge base...")
                    knowledge_results = self.search_knowledge_base(search_query)
                    
            except Exception as e:
                self.logger.error(f"Database search error: {e}")
                local_results = []
                knowledge_results = []
            
            # Response generation with RTX 4090 optimization
            try:
                response = ""
                data_source = "unknown"
                
                if local_results:
                    self.logger.info(f"Found {len(local_results)} products, generating RTX 4090 optimized AI response")
                    response = self.generate_natural_response_rtx4090(
                        user_input, 
                        context_data=local_results,
                        data_source="database",
                        user_language=user_language
                    )
                    data_source = "local_database"
                    
                elif knowledge_results:
                    self.logger.info(f"Found {len(knowledge_results)} knowledge entries, generating RTX 4090 optimized response")
                    response = self.generate_natural_response_rtx4090(
                        user_input,
                        context_data=knowledge_results,
                        data_source="knowledge_base",
                        user_language=user_language
                    )
                    data_source = "knowledge_base"
                    
                else:
                    self.logger.info("No database results found, using RTX 4090 AI general knowledge")
                    
                    if database_searched:
                        if user_language == 'vi':
                            no_results_msg = "Tôi đã tìm kiếm trong cơ sở dữ liệu nhưng không tìm thấy sản phẩm phù hợp. Để tôi trả lời dựa trên kiến thức chung với sức mạnh RTX 4090:\n\n"
                        else:
                            no_results_msg = "I searched our database but couldn't find matching products. Let me answer based on general knowledge with RTX 4090 power:\n\n"
                    else:
                        no_results_msg = ""
                    
                    ai_response = self.generate_natural_response_rtx4090(
                        user_input, 
                        context_data=None,
                        data_source="general",
                        user_language=user_language
                    )
                    
                    response = no_results_msg + ai_response
                    data_source = "ai_knowledge"
                    
                # CUSTOMIZE RESPONSE BASED ON VIETNAMESE ADDRESSING
                if user_language == 'vi' and customer_profile:
                    response = self.customize_response_for_addressing(response, customer_profile)
                    
            except Exception as e:
                self.logger.error(f"Response generation error: {e}")
                if user_language == 'vi':
                    response = self.config['vietnamese_templates']['generation_error']
                else:
                    response = "I encountered an error while generating a response. Please try again."
                data_source = "error"
            
            # Context management
            conversation_entry = {
                'user': user_input,
                'bot': response,
                'user_language': user_language,
                'timestamp': datetime.now().isoformat(),
                'data_source': data_source,
                'results_found': len(local_results) + len(knowledge_results),
                'products_mentioned': [p.get('name', '') for p in local_results[:3]] if local_results else [],
                'context_references': context_references
            }
            
            try:
                self.conversation_context.append(conversation_entry)
                
                if len(self.conversation_context) > 20:
                    self.conversation_context = self.conversation_context[-20:]
                    
                if len(self.conversation_context) % 5 == 0:
                    self.update_conversation_summary()
                    
            except Exception as e:
                self.logger.warning(f"Context update error: {e}")
                
            # Store conversation
            try:
                processing_time = time.time() - start_time
                self.store_conversation(user_input, response, data_source, processing_time, user_language)
            except Exception as e:
                self.logger.warning(f"Conversation storage error: {e}")
            
            self.logger.info(f"RTX 4090 optimized response generated using: {data_source}")
            
            return {
                'response': response,
                'data_source': data_source,
                'local_results_count': len(local_results) + len(knowledge_results),
                'processing_time': time.time() - start_time,
                'user_language': user_language,
                'response_language': user_language,
                'database_searched': database_searched,
                'context_used': bool(context_references),
                'rtx4090_optimized': hasattr(self, 'is_rtx4090') and self.is_rtx4090,
                'customer_profile': customer_profile
            }
            
        except Exception as e:
            self.logger.error(f"Message processing error: {e}")
            error_language = getattr(self, 'current_language', 'vi')
            if error_language == 'vi':
                error_msg = self.config['vietnamese_templates']['error']
            else:
                error_msg = "I apologize, but I encountered an error processing your request."
            
            return {
                'response': error_msg,
                'data_source': 'error',
                'local_results_count': 0,
                'processing_time': time.time() - start_time,
                'user_language': error_language,
                'response_language': error_language,
                'error': str(e)
            }
    
    def search_local_database_rtx4090(self, user_input, similarity_threshold=None):
        """RTX 4090 optimized database search with batch processing"""
        try:
            if not user_input or not user_input.strip():
                return []
            
            if similarity_threshold is None:
                similarity_threshold = self.config['search_config']['local_similarity_threshold']
            
            if self.embedding_model:
                try:
                    # RTX 4090 optimized batch embedding generation
                    query_embedding = self.embedding_model.encode([user_input.strip()])
                    if query_embedding is None or len(query_embedding) == 0:
                        self.logger.warning("Failed to generate query embedding")
                        return self.keyword_search_fallback(user_input)
                except Exception as e:
                    self.logger.error(f"Query embedding error: {e}")
                    return self.keyword_search_fallback(user_input)
                
                try:
                    cursor = self.conn.cursor()
                    # Optimized query for RTX 4090 system
                    cursor.execute("""
                        SELECT name, name_vietnamese, description, description_vietnamese, 
                               category, price, embedding, id 
                        FROM products 
                        WHERE embedding IS NOT NULL
                    """)
                    products = cursor.fetchall()
                except Exception as e:
                    self.logger.error(f"Database query error: {e}")
                    return []
                
                if not products:
                    self.logger.info("No products with embeddings found, trying keyword search")
                    return self.keyword_search_fallback(user_input)
                    
                best_matches = []
                
                # RTX 4090 optimized batch similarity computation
                if hasattr(self, 'is_rtx4090') and self.is_rtx4090 and len(products) > 10:
                    # Process in batches for RTX 4090
                    batch_size = 32
                    for i in range(0, len(products), batch_size):
                        batch = products[i:i+batch_size]
                        batch_matches = self.process_product_batch_rtx4090(batch, query_embedding, similarity_threshold)
                        best_matches.extend(batch_matches)
                else:
                    # Standard processing for smaller datasets or non-RTX 4090
                    for product in products:
                        try:
                            if len(product) < 7:
                                continue
                                
                            stored_embedding_blob = product[6]
                            
                            if stored_embedding_blob:
                                try:
                                    stored_embedding = np.frombuffer(stored_embedding_blob, dtype=np.float32)
                                    stored_embedding = stored_embedding.reshape(1, -1)
                                except Exception as e:
                                    self.logger.warning(f"Embedding decode error: {e}")
                                    continue
                                
                                try:
                                    similarity = cosine_similarity(query_embedding, stored_embedding)[0][0]
                                except Exception as e:
                                    self.logger.warning(f"Similarity calculation error: {e}")
                                    continue
                                
                                if similarity > similarity_threshold:
                                    best_matches.append({
                                        'name': product[0] or '',
                                        'name_vietnamese': product[1] or '',
                                        'description': product[2] or '',
                                        'description_vietnamese': product[3] or '',
                                        'category': product[4] or '',
                                        'price': product[5],
                                        'similarity': float(similarity),
                                        'id': product[7] if len(product) > 7 else None
                                    })
                            
                        except Exception as e:
                            self.logger.warning(f"Product processing error: {e}")
                            continue
                            
                try:
                    best_matches.sort(key=lambda x: x.get('similarity', 0), reverse=True)
                    
                    if not best_matches:
                        self.logger.info("No embedding matches found, trying keyword search")
                        return self.keyword_search_fallback(user_input)
                        
                    return best_matches[:5]  # Return more results for RTX 4090
                except Exception as e:
                    self.logger.error(f"Sorting error: {e}")
                    return best_matches[:5] if best_matches else []
            else:
                self.logger.info("Embedding model not available, using keyword search")
                return self.keyword_search_fallback(user_input)
                
        except Exception as e:
            self.logger.error(f"RTX 4090 database search error: {e}")
            return []
    
    def process_product_batch_rtx4090(self, product_batch, query_embedding, similarity_threshold):
        """RTX 4090 optimized batch processing for product similarity"""
        try:
            batch_matches = []
            embeddings_batch = []
            valid_products = []
            
            # Prepare batch embeddings
            for product in product_batch:
                if len(product) >= 7 and product[6]:
                    try:
                        stored_embedding = np.frombuffer(product[6], dtype=np.float32)
                        embeddings_batch.append(stored_embedding)
                        valid_products.append(product)
                    except Exception:
                        continue
            
            if not embeddings_batch:
                return batch_matches
            
            # Batch similarity computation
            try:
                embeddings_matrix = np.vstack(embeddings_batch)
                similarities = cosine_similarity(query_embedding, embeddings_matrix)[0]
                
                for i, similarity in enumerate(similarities):
                    if similarity > similarity_threshold:
                        product = valid_products[i]
                        batch_matches.append({
                            'name': product[0] or '',
                            'name_vietnamese': product[1] or '',
                            'description': product[2] or '',
                            'description_vietnamese': product[3] or '',
                            'category': product[4] or '',
                            'price': product[5],
                            'similarity': float(similarity),
                            'id': product[7] if len(product) > 7 else None
                        })
                        
            except Exception as e:
                self.logger.warning(f"Batch similarity computation error: {e}")
                
            return batch_matches
            
        except Exception as e:
            self.logger.error(f"RTX 4090 batch processing error: {e}")
            return []
    
    def generate_natural_response_rtx4090(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Generate natural language response with RTX 4090 optimizations"""
        try:
            if not user_input or not isinstance(user_input, str):
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "I need a question or message to respond to."
            
            user_input = user_input.strip()
            if not user_input:
                if user_language == 'vi':
                    return self.config['vietnamese_templates']['empty_response']
                else:
                    return "Please provide a question or message."
            
            # Check for greeting patterns
            try:
                if user_language == 'vi':
                    greeting_flow = self.config['conversation_flows'].get('greeting_vietnamese', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Xin chào! Tôi là trợ lý AI với sức mạnh RTX 4090!')
                else:
                    greeting_flow = self.config['conversation_flows'].get('greeting', {})
                    greeting_triggers = greeting_flow.get('triggers', [])
                    if any(trigger.lower() in user_input.lower() for trigger in greeting_triggers):
                        return greeting_flow.get('response_template', 'Hello! I am your RTX 4090-powered Sales Consultant!')
            except Exception as e:
                self.logger.warning(f"Greeting pattern matching error: {e}")
            
            # Generate response based on context
            if context_data and len(context_data) > 0:
                if data_source == "database":
                    prompt = self.build_product_aware_prompt_rtx4090(user_input, context_data, user_language)
                    ai_response = self.safe_generate_response_rtx4090(prompt, user_language)
                    
                    if not ai_response or ai_response == self.config['vietnamese_templates']['model_error']:
                        return self.format_product_response(context_data, user_input, user_language)
                    
                    return ai_response
                    
                elif data_source == "knowledge_base":
                    prompt = self.build_knowledge_aware_prompt(user_input, context_data, user_language)
                    ai_response = self.safe_generate_response_rtx4090(prompt, user_language)
                    
                    if not ai_response or ai_response == self.config['vietnamese_templates']['model_error']:
                        return self.format_knowledge_response(context_data, user_input, user_language)
                    
                    return ai_response
            
            # General AI response with RTX 4090 optimization
            prompt = self.build_sales_prompt_rtx4090(user_input, None, "general", user_language)
            response = self.safe_generate_response_rtx4090(prompt, user_language)
            
            if not response or response in [self.config['vietnamese_templates']['model_error'], 
                                           self.config['vietnamese_templates']['generation_error']]:
                if user_language == 'vi':
                    return "Xin lỗi, tôi đang gặp khó khăn trong việc tạo phản hồi với RTX 4090. Bạn có thể mô tả chi tiết hơn về sản phẩm bạn đang tìm kiếm không?"
                else:
                    return "I apologize, I'm having trouble generating a response with RTX 4090. Could you describe in more detail what product you're looking for?"
            
            return response
            
        except Exception as e:
            self.logger.error(f"Error generating RTX 4090 natural response: {e}")
            if user_language == 'vi':
                return self.config['vietnamese_templates']['error']
            else:
                return "I apologize, but I'm having trouble generating a response right now."
    
    def build_product_aware_prompt_rtx4090(self, user_input, products, user_language='vi'):
        """Build RTX 4090 optimized prompt for product responses"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý bán hàng AI được tăng cường bởi RTX 4090, thân thiện và chuyên nghiệp. Dựa trên câu hỏi của khách hàng và thông tin sản phẩm có sẵn, hãy tạo một phản hồi tự nhiên, hữu ích và chi tiết.

"""
            else:
                prompt = """You are a friendly AI sales assistant powered by RTX 4090. Based on the customer's question and available product information, create a natural, helpful and detailed response.

"""
            
            # Add conversation context
            summary_prompt = self.get_conversation_context_prompt(user_language)
            if summary_prompt:
                prompt += summary_prompt
            
            # Recent conversation history
            if self.conversation_context and len(self.conversation_context) > 0:
                if user_language == 'vi':
                    prompt += "Cuộc trò chuyện gần đây:\n"
                else:
                    prompt += "Recent conversation:\n"
                
                for turn in self.conversation_context[-3:]:
                    user_msg = turn.get('user', '')[:150]
                    bot_msg = turn.get('bot', '')[:150]
                    prompt += f"User: {user_msg}\nAssistant: {bot_msg}...\n"
                prompt += "\n"
            
            # Product information
            if user_language == 'vi':
                prompt += "Sản phẩm có sẵn:\n"
            else:
                prompt += "Available products:\n"
                
            for i, product in enumerate(products[:4], 1):  # Show more products for RTX 4090
                name = product.get('name_vietnamese', product.get('name', '')) if user_language == 'vi' else product.get('name', '')
                desc = product.get('description_vietnamese', product.get('description', '')) if user_language == 'vi' else product.get('description', '')
                price = product.get('price', 0)
                features = product.get('features_vietnamese', product.get('features', '')) if user_language == 'vi' else product.get('features', '')
                similarity = product.get('similarity', 0)
                
                prompt += f"\n{i}. {name}"
                if desc:
                    prompt += f": {desc[:120]}"
                if price > 0:
                    if user_language == 'vi':
                        prompt += f" (Giá: ${price:,.0f})"
                    else:
                        prompt += f" (Price: ${price:,.2f})"
                if features:
                    prompt += f" - {features[:60]}"
                if similarity > 0:
                    prompt += f" [Độ phù hợp: {similarity:.2f}]" if user_language == 'vi' else f" [Match: {similarity:.2f}]"
            
            # Customer question and instructions
            if user_language == 'vi':
                prompt += f"\n\nCâu hỏi khách hàng: {user_input}\n\n"
                prompt += "Với sức mạnh RTX 4090, hãy trả lời một cách tự nhiên và chi tiết, nhớ ngữ cảnh cuộc trò chuyện, giới thiệu sản phẩm phù hợp và giải thích tại sao chúng đáp ứng nhu cầu của khách hàng. "
                prompt += "Nếu khách hàng đề cập đến sản phẩm đã nói trước đó, hãy nhắc lại thông tin đó. Trả lời bằng tiếng Việt:"
            else:
                prompt += f"\n\nCustomer question: {user_input}\n\n"
                prompt += "With RTX 4090 power, provide a natural and detailed response that remembers the conversation context, introduces suitable products and explains why they meet the customer's needs. "
                prompt += "If the customer refers to previously mentioned products, recall that information. Respond in English:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Error building RTX 4090 product prompt: {e}")
            return self.build_sales_prompt_rtx4090(user_input, products, "database", user_language)
    
    def build_sales_prompt_rtx4090(self, user_input, context_data=None, data_source="unknown", user_language='vi'):
        """Build RTX 4090 optimized prompt for sales conversations"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý AI bán hàng chuyên nghiệp được tăng cường bởi RTX 4090. Hãy trả lời bằng tiếng Việt một cách thân thiện, chi tiết và hữu ích.

"""
            else:
                prompt = """You are a professional AI sales assistant powered by RTX 4090. Respond in English in a friendly, detailed and helpful manner.

"""
            
            # Add conversation context
            try:
                if self.conversation_context and len(self.conversation_context) > 0:
                    if user_language == 'vi':
                        prompt += "Lịch sử cuộc trò chuyện:\n"
                    else:
                        prompt += "Conversation history:\n"
                    
                    recent_context = self.conversation_context[-5:]
                    for i, turn in enumerate(recent_context):
                        if isinstance(turn, dict) and 'user' in turn and 'bot' in turn:
                            user_msg = str(turn['user'])[:200]
                            bot_msg = str(turn['bot'])[:200]
                            
                            if user_language == 'vi':
                                prompt += f"\nLượt {i+1}:\n"
                            else:
                                prompt += f"\nTurn {i+1}:\n"
                            
                            prompt += f"User: {user_msg}\nAssistant: {bot_msg}\n"
                    
                    if len(self.conversation_context) > 5:
                        if user_language == 'vi':
                            prompt += f"\n(Đã có {len(self.conversation_context)} lượt trò chuyện trước đó)\n"
                        else:
                            prompt += f"\n(There were {len(self.conversation_context)} previous conversation turns)\n"
                    
                    prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Context building error: {e}")
            
            # Add context data if available
            try:
                if context_data and isinstance(context_data, list) and len(context_data) > 0:
                    if data_source == "database":
                        if user_language == 'vi':
                            prompt += "Sản phẩm liên quan:\n"
                        else:
                            prompt += "Relevant products:\n"
                        
                        for item in context_data[:3]:
                            if isinstance(item, dict):
                                name = item.get('name_vietnamese', item.get('name', '')) if user_language == 'vi' else item.get('name', '')
                                desc = item.get('description_vietnamese', item.get('description', '')) if user_language == 'vi' else item.get('description', '')
                                
                                if name:
                                    prompt += f"- {str(name)[:100]}"
                                    if desc:
                                        prompt += f": {str(desc)[:150]}...\n"
                                    else:
                                        prompt += "\n"
                                        
                                    if item.get('price'):
                                        try:
                                            price_label = "Giá" if user_language == 'vi' else "Price"
                                            prompt += f"  {price_label}: ${float(item['price']):.2f}\n"
                                        except (ValueError, TypeError):
                                            pass
                        prompt += "\n"
            except Exception as e:
                self.logger.warning(f"Product context building error: {e}")
            
            # Add current question
            try:
                truncated_input = str(user_input)[:300]
                if user_language == 'vi':
                    prompt += f"Câu hỏi hiện tại của khách hàng: {truncated_input}\n\n"
                    prompt += "Với sức mạnh RTX 4090, hãy trả lời dựa trên lịch sử cuộc trò chuyện và duy trì ngữ cảnh. "
                    prompt += "Nếu khách hàng đề cập đến điều gì đó đã nói trước đó, hãy nhớ và tham chiếu đến nó.\n"
                    prompt += "Trả lời chi tiết bằng tiếng Việt:"
                else:
                    prompt += f"Current customer question: {truncated_input}\n\n"
                    prompt += "With RTX 4090 power, please respond based on the conversation history and maintain context. "
                    prompt += "If the customer refers to something mentioned earlier, remember and reference it.\n"
                    prompt += "Detailed response:"
            except Exception as e:
                self.logger.error(f"Question formatting error: {e}")
                if user_language == 'vi':
                    prompt += f"Câu hỏi: {user_input}\nTrả lời chi tiết:"
                else:
                    prompt += f"Question: {user_input}\nDetailed response:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"RTX 4090 prompt building error: {e}")
            if user_language == 'vi':
                return f"Câu hỏi: {user_input}\nTrả lời bằng tiếng Việt:"
            else:
                return f"Question: {user_input}\nResponse:"
    
    # Keep all your existing methods (keyword_search_fallback, search_knowledge_base, etc.)
    # They don't need RTX 4090 specific changes, just inherit the original implementations
    
    def keyword_search_fallback(self, user_input):
        """Fallback keyword-based search when embedding search fails or is unavailable"""
        try:
            if not user_input or not user_input.strip():
                return []
            
            keywords = user_input.lower().split()
            
            cursor = self.conn.cursor()
            
            query = """
                SELECT name, name_vietnamese, description, description_vietnamese, 
                       category, price, NULL as embedding
                FROM products 
                WHERE 1=0
            """
            
            conditions = []
            params = []
            
            for keyword in keywords:
                if len(keyword) > 2:
                    keyword_pattern = f'%{keyword}%'
                    conditions.append("""
                        (LOWER(name) LIKE ? OR 
                         LOWER(name_vietnamese) LIKE ? OR 
                         LOWER(description) LIKE ? OR 
                         LOWER(description_vietnamese) LIKE ? OR
                         LOWER(category) LIKE ? OR
                         LOWER(category_vietnamese) LIKE ?)
                    """)
                    params.extend([keyword_pattern] * 6)
            
            if conditions:
                query = f"""
                    SELECT name, name_vietnamese, description, description_vietnamese, 
                           category, price, NULL as embedding
                    FROM products 
                    WHERE {' OR '.join(conditions)}
                    LIMIT 10
                """
                
                cursor.execute(query, params)
                products = cursor.fetchall()
                
                results = []
                for product in products:
                    if len(product) >= 6:
                        results.append({
                            'name': product[0] or '',
                            'name_vietnamese': product[1] or '',
                            'description': product[2] or '',
                            'description_vietnamese': product[3] or '',
                            'category': product[4] or '',
                            'price': product[5],
                            'similarity': 0.5
                        })
                
                self.logger.info(f"Keyword search found {len(results)} results")
                return results[:3]
            
            return []
            
        except Exception as e:
            self.logger.error(f"Keyword search error: {e}")
            return []

    def search_knowledge_base(self, user_input, similarity_threshold=None):
        """Search knowledge base for relevant information"""
        try:
            if not self.embedding_model:
                self.logger.warning("Embedding model not available for knowledge base search")
                return []
            
            if not user_input or not user_input.strip():
                return []
            
            if similarity_threshold is None:
                similarity_threshold = self.config['search_config']['local_similarity_threshold']
            
            try:
                query_embedding = self.embedding_model.encode([user_input.strip()])
                if query_embedding is None or len(query_embedding) == 0:
                    return []
            except Exception as e:
                self.logger.error(f"Knowledge base query embedding error: {e}")
                return []
            
            try:
                cursor = self.conn.cursor()
                cursor.execute("""
                    SELECT topic, content, source, embedding
                    FROM knowledge_base
                    WHERE embedding IS NOT NULL
                """)
                knowledge_entries = cursor.fetchall()
            except Exception as e:
                self.logger.error(f"Knowledge base query error: {e}")
                return []
            
            if not knowledge_entries:
                self.logger.info("No knowledge base entries with embeddings found")
                return []
            
            best_matches = []
            
            for entry in knowledge_entries:
                try:
                    if len(entry) < 4:
                        continue
                    
                    topic, content, source, embedding_blob = entry
                    
                    if embedding_blob:
                        try:
                            stored_embedding = np.frombuffer(embedding_blob, dtype=np.float32)
                            stored_embedding = stored_embedding.reshape(1, -1)
                        except Exception as e:
                            self.logger.warning(f"Knowledge embedding decode error: {e}")
                            continue
                        
                        try:
                            similarity = cosine_similarity(query_embedding, stored_embedding)[0][0]
                        except Exception as e:
                            self.logger.warning(f"Knowledge similarity calculation error: {e}")
                            continue
                        
                        if similarity > similarity_threshold:
                            best_matches.append({
                                'topic': topic or '',
                                'content': content[:500] if content else '',
                                'source': source or '',
                                'similarity': float(similarity),
                                'type': 'knowledge'
                            })
                
                except Exception as e:
                    self.logger.warning(f"Knowledge entry processing error: {e}")
                    continue
            
            try:
                best_matches.sort(key=lambda x: x.get('similarity', 0), reverse=True)
                return best_matches[:3]
            except Exception as e:
                self.logger.error(f"Knowledge sorting error: {e}")
                return best_matches[:3] if best_matches else []
                
        except Exception as e:
            self.logger.error(f"Knowledge base search error: {e}")
            return []
    
    def check_context_references(self, user_input, user_language='vi'):
        """Check if user is referring to previous conversation items"""
        try:
            if user_language == 'vi':
                context_keywords = [
                    'cái đó', 'cái này', 'sản phẩm đó', 'sản phẩm này',
                    'như trên', 'đã nói', 'vừa nói', 'trước đó',
                    'cái thứ', 'cái đầu', 'cái cuối', 'nó',
                    'chúng', 'những cái', 'mấy cái'
                ]
            else:
                context_keywords = [
                    'that', 'this', 'those', 'these',
                    'the one', 'previous', 'mentioned',
                    'above', 'it', 'them', 'the first',
                    'the last', 'the second'
                ]
            
            user_input_lower = user_input.lower()
            
            references = []
            for keyword in context_keywords:
                if keyword in user_input_lower:
                    references.append(keyword)
            
            if self.conversation_context:
                for turn in self.conversation_context[-3:]:
                    if 'products_mentioned' in turn and turn['products_mentioned']:
                        for product in turn['products_mentioned']:
                            if product.lower() in user_input_lower:
                                references.append(f"product:{product}")
            
            return references
            
        except Exception as e:
            self.logger.error(f"Context reference check error: {e}")
            return []
    
    def build_contextual_search_query(self, user_input, context_references):
        """Build search query that includes context from previous conversations"""
        try:
            search_terms = [user_input]
            
            for turn in self.conversation_context[-3:]:
                if 'products_mentioned' in turn and turn['products_mentioned']:
                    for product in turn['products_mentioned']:
                        if product and product not in search_terms:
                            search_terms.append(product)
            
            enhanced_query = ' '.join(search_terms[:3])
            
            self.logger.info(f"Enhanced search query: {enhanced_query}")
            return enhanced_query
            
        except Exception as e:
            self.logger.error(f"Contextual search query error: {e}")
            return user_input
    
    def update_conversation_summary(self):
        """Update conversation summary for long-term context"""
        try:
            if not self.conversation_context:
                return
            
            topics = []
            products_discussed = set()
            
            for turn in self.conversation_context:
                if 'products_mentioned' in turn and turn['products_mentioned']:
                    products_discussed.update(turn['products_mentioned'])
                
                user_msg = turn.get('user', '').lower()
                if 'laptop' in user_msg or 'máy tính' in user_msg:
                    topics.append('laptops')
                if 'gaming' in user_msg or 'game' in user_msg:
                    topics.append('gaming')
                if 'mouse' in user_msg or 'chuột' in user_msg:
                    topics.append('mouse')
                if 'keyboard' in user_msg or 'bàn phím' in user_msg:
                    topics.append('keyboard')
            
            self.conversation_summary = {
                'topics': list(set(topics)),
                'products_discussed': list(products_discussed),
                'turn_count': len(self.conversation_context),
                'last_update': datetime.now().isoformat()
            }
            
            self.logger.info(f"Updated conversation summary: {self.conversation_summary}")
            
        except Exception as e:
            self.logger.error(f"Conversation summary update error: {e}")
    
    def get_conversation_context_prompt(self, user_language='vi'):
        """Get enhanced conversation context for prompts"""
        try:
            context_prompt = ""
            
            if hasattr(self, 'conversation_summary') and self.conversation_summary:
                if user_language == 'vi':
                    context_prompt += "Tóm tắt cuộc trò chuyện:\n"
                    if self.conversation_summary.get('topics'):
                        context_prompt += f"- Chủ đề đã thảo luận: {', '.join(self.conversation_summary['topics'])}\n"
                    if self.conversation_summary.get('products_discussed'):
                        context_prompt += f"- Sản phẩm đã xem: {', '.join(self.conversation_summary['products_discussed'][:5])}\n"
                else:
                    context_prompt += "Conversation summary:\n"
                    if self.conversation_summary.get('topics'):
                        context_prompt += f"- Topics discussed: {', '.join(self.conversation_summary['topics'])}\n"
                    if self.conversation_summary.get('products_discussed'):
                        context_prompt += f"- Products viewed: {', '.join(self.conversation_summary['products_discussed'][:5])}\n"
                context_prompt += "\n"
            
            return context_prompt
            
        except Exception as e:
            self.logger.error(f"Context prompt error: {e}")
            return ""
    
    def store_conversation(self, user_input, response, data_source, processing_time, 
                          user_language, error_type=None, error_message=None):
        """Store conversation in database with language tracking and error logging"""
        try:
            cursor = self.conn.cursor()
            
            cursor.execute('''
                INSERT INTO conversations 
                (timestamp, user_input, user_language, bot_response, bot_language, 
                 data_source, response_time, error_type, error_message)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                datetime.now().isoformat(),
                user_input,
                user_language,
                response,
                user_language,
                data_source,
                processing_time,
                error_type,
                error_message
            ))
            
            self.conn.commit()
        except Exception as e:
            self.logger.error(f"Error storing conversation: {e}")
    
    def build_knowledge_aware_prompt(self, user_input, knowledge_entries, user_language='vi'):
        """Build a prompt that incorporates knowledge base data for natural AI response"""
        try:
            if user_language == 'vi':
                prompt = """Bạn là trợ lý AI thông minh được tăng cường bởi RTX 4090. Dựa trên thông tin từ cơ sở tri thức và câu hỏi của người dùng, hãy tạo một câu trả lời tự nhiên và hữu ích.

"""
            else:
                prompt = """You are an intelligent AI assistant powered by RTX 4090. Based on information from the knowledge base and the user's question, create a natural and helpful response.

"""
            
            summary_prompt = self.get_conversation_context_prompt(user_language)
            if summary_prompt:
                prompt += summary_prompt
            
            if self.conversation_context and len(self.conversation_context) > 0:
                if user_language == 'vi':
                    prompt += "Ngữ cảnh cuộc trò chuyện:\n"
                else:
                    prompt += "Conversation context:\n"
                
                for turn in self.conversation_context[-2:]:
                    user_msg = turn.get('user', '')[:100]
                    prompt += f"User: {user_msg}\n"
                prompt += "\n"
            
            if user_language == 'vi':
                prompt += "Thông tin liên quan:\n"
            else:
                prompt += "Related information:\n"
                
            for entry in knowledge_entries[:2]:
                topic = entry.get('topic', '')
                content = entry.get('content', '')[:200]
                
                if topic:
                    prompt += f"\n{topic}: "
                if content:
                    prompt += f"{content}..."
            
            if user_language == 'vi':
                prompt += f"\n\nCâu hỏi: {user_input}\n\n"
                prompt += "Với sức mạnh RTX 4090, trả lời dựa trên ngữ cảnh cuộc trò chuyện và thông tin có sẵn:"
            else:
                prompt += f"\n\nQuestion: {user_input}\n\n"
                prompt += "With RTX 4090 power, response based on conversation context and available information:"
            
            return prompt
            
        except Exception as e:
            self.logger.error(f"Error building knowledge prompt: {e}")
            return self.build_sales_prompt_rtx4090(user_input, knowledge_entries, "knowledge_base", user_language)
    
    def format_product_response(self, products, user_input, user_language='vi'):
        """Format product search results into a natural response"""
        try:
            if user_language == 'vi':
                response = "Với sức mạnh RTX 4090, tôi tìm thấy các sản phẩm sau phù hợp với yêu cầu của bạn:\n\n"
                
                for i, product in enumerate(products[:4], 1):  # Show more products
                    name = product.get('name_vietnamese') or product.get('name', '')
                    desc = product.get('description_vietnamese') or product.get('description', '')
                    price = product.get('price', 0)
                    similarity = product.get('similarity', 0)
                    
                    response += f"{i}. **{name}**"
                    if similarity > 0:
                        response += f" (Độ phù hợp: {similarity:.1%})"
                    response += "\n"
                    if desc:
                        response += f"   {desc[:180]}...\n"
                    if price > 0:
                        response += f"   💰 Giá: ${price:,.2f}\n"
                    response += "\n"
                
                response += "Bạn muốn biết thêm thông tin chi tiết về sản phẩm nào?"
            else:
                response = "With RTX 4090 power, I found the following products that match your request:\n\n"
                
                for i, product in enumerate(products[:4], 1):
                    name = product.get('name', '')
                    desc = product.get('description', '')
                    price = product.get('price', 0)
                    similarity = product.get('similarity', 0)
                    
                    response += f"{i}. **{name}**"
                    if similarity > 0:
                        response += f" (Match: {similarity:.1%})"
                    response += "\n"
                    if desc:
                        response += f"   {desc[:180]}...\n"
                    if price > 0:
                        response += f"   💰 Price: ${price:,.2f}\n"
                    response += "\n"
                
                response += "Which product would you like to know more about?"
                
            return response
        except Exception as e:
            self.logger.error(f"Error formatting product response: {e}")
            if user_language == 'vi':
                return "Tôi tìm thấy một số sản phẩm phù hợp với RTX 4090. Bạn có thể cho tôi biết thêm về nhu cầu của bạn không?"
            else:
                return "I found some matching products with RTX 4090. Could you tell me more about your needs?"
    
    def format_knowledge_response(self, knowledge_entries, user_input, user_language='vi'):
        """Format knowledge base search results into a natural response"""
        try:
            if user_language == 'vi':
                response = "Dựa trên thông tin trong cơ sở dữ liệu được xử lý bởi RTX 4090, tôi có thể chia sẻ:\n\n"
                
                for entry in knowledge_entries[:2]:
                    topic = entry.get('topic', '')
                    content = entry.get('content', '')
                    
                    if topic:
                        response += f"📚 **{topic}**\n"
                    if content:
                        response += f"{content[:350]}...\n\n"
                
                response += "Bạn cần thông tin chi tiết hơn về vấn đề nào?"
            else:
                response = "Based on the information in our RTX 4090-processed knowledge base:\n\n"
                
                for entry in knowledge_entries[:2]:
                    topic = entry.get('topic', '')
                    content = entry.get('content', '')
                    
                    if topic:
                        response += f"📚 **{topic}**\n"
                    if content:
                        response += f"{content[:350]}...\n\n"
                
                response += "Would you like more detailed information on any topic?"
                
            return response
        except Exception as e:
            self.logger.error(f"Error formatting knowledge response: {e}")
            if user_language == 'vi':
                return "Tôi tìm thấy một số thông tin liên quan với RTX 4090. Bạn có thể hỏi cụ thể hơn không?"
            else:
                return "I found some related information with RTX 4090. Could you be more specific?"

    # Include all your existing file processing methods
    def process_excel_file(self, file_path):
        """Process Excel files for product data with Vietnamese support and RTX 4090 optimization"""
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path, encoding='utf-8')
            else:
                df = pd.read_excel(file_path)
                
            self.update_training_status(f"📊 Found {len(df)} rows in Excel file")
                
            column_mapping = {
                'name': ['name', 'tên', 'ten', 'product_name', 'sản phẩm'],
                'name_vietnamese': ['name_vietnamese', 'tên_tiếng_việt', 'ten_tieng_viet'],
                'description': ['description', 'mô tả', 'mo_ta', 'desc'],
                'description_vietnamese': ['description_vietnamese', 'mô_tả_tiếng_việt'],
                'category': ['category', 'danh mục', 'danh_muc', 'loại'],
                'category_vietnamese': ['category_vietnamese', 'danh_mục_tiếng_việt'],
                'price': ['price', 'giá', 'gia', 'cost'],
                'features': ['features', 'tính năng', 'tinh_nang'],
                'features_vietnamese': ['features_vietnamese', 'tính_năng_tiếng_việt'],
                'specifications': ['specifications', 'thông số', 'thong_so', 'specs'],
                'specifications_vietnamese': ['specifications_vietnamese', 'thông_số_tiếng_việt']
            }
            
            mapped_columns = {}
            for standard_col, possible_cols in column_mapping.items():
                for col in possible_cols:
                    if col in df.columns:
                        mapped_columns[standard_col] = col
                        break
            
            if 'name' not in mapped_columns and 'name_vietnamese' not in mapped_columns:
                raise ValueError("Excel file must contain at least a name column (English or Vietnamese)")
                
            cursor = self.conn.cursor()
            added_count = 0
            
            # RTX 4090 optimization: Process in batches
            batch_size = 50 if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else 25
            
            for batch_start in range(0, len(df), batch_size):
                batch_end = min(batch_start + batch_size, len(df))
                batch_df = df.iloc[batch_start:batch_end]
                
                batch_texts = []
                batch_products = []
                
                for _, row in batch_df.iterrows():
                    try:
                        name = str(row.get(mapped_columns.get('name', ''), '')) if 'name' in mapped_columns else ''
                        name_vietnamese = str(row.get(mapped_columns.get('name_vietnamese', ''), '')) if 'name_vietnamese' in mapped_columns else ''
                        description = str(row.get(mapped_columns.get('description', ''), '')) if 'description' in mapped_columns else ''
                        description_vietnamese = str(row.get(mapped_columns.get('description_vietnamese', ''), '')) if 'description_vietnamese' in mapped_columns else ''
                        category = str(row.get(mapped_columns.get('category', ''), 'General')) if 'category' in mapped_columns else 'General'
                        category_vietnamese = str(row.get(mapped_columns.get('category_vietnamese', ''), '')) if 'category_vietnamese' in mapped_columns else ''
                        features = str(row.get(mapped_columns.get('features', ''), '')) if 'features' in mapped_columns else ''
                        features_vietnamese = str(row.get(mapped_columns.get('features_vietnamese', ''), '')) if 'features_vietnamese' in mapped_columns else ''
                        specifications = str(row.get(mapped_columns.get('specifications', ''), '')) if 'specifications' in mapped_columns else ''
                        specifications_vietnamese = str(row.get(mapped_columns.get('specifications_vietnamese', ''), '')) if 'specifications_vietnamese' in mapped_columns else ''

                        price = 0
                        if 'price' in mapped_columns:
                            try:
                                price_val = row.get(mapped_columns['price'], 0)
                                if pd.notna(price_val):
                                    price = float(str(price_val).replace(',', '').replace('$', ''))
                                    price = float(price_str)
                            except (ValueError, TypeError):
                                price = 0
                        
                        if not name and not name_vietnamese:
                            continue
                            
                        product_text_en = f"{name} {description} {features} {specifications}"
                        product_text_vi = f"{name_vietnamese} {description_vietnamese} {features_vietnamese} {specifications_vietnamese}"
                        combined_text = f"{product_text_en} {product_text_vi}".strip()
                        
                        batch_texts.append(combined_text)
                        batch_products.append({
                            'name': name or name_vietnamese,
                            'name_vietnamese': name_vietnamese,
                            'description': description or description_vietnamese,
                            'description_vietnamese': description_vietnamese,
                            'category': category,
                            'category_vietnamese': category_vietnamese,
                            'price': price,
                            'features': features,
                            'features_vietnamese': features_vietnamese,
                            'specifications': specifications,
                            'specifications_vietnamese': specifications_vietnamese,
                            'source_file': file_path,
                            'combined_text': combined_text
                        })
                        
                    except Exception as row_error:
                        self.update_training_status(f"⚠️ Error processing row: {row_error}")
                        continue
                
                # RTX 4090 optimized batch embedding generation
                if self.embedding_model and batch_texts:
                    try:
                        embeddings = self.embedding_model.encode(batch_texts, batch_size=batch_size)
                        
                        for i, (embedding, product_data) in enumerate(zip(embeddings, batch_products)):
                            try:
                                embedding_blob = embedding.astype(np.float32).tobytes()
                                
                                cursor.execute('''
                                    INSERT INTO products 
                                    (name, name_vietnamese, description, description_vietnamese, 
                                     category, category_vietnamese, price, features, features_vietnamese,
                                     specifications, specifications_vietnamese, source_file, embedding, created_at)
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                ''', (
                                    product_data['name'],
                                    product_data['name_vietnamese'],
                                    product_data['description'],
                                    product_data['description_vietnamese'],
                                    product_data['category'],
                                    product_data['category_vietnamese'],
                                    product_data['price'],
                                    product_data['features'],
                                    product_data['features_vietnamese'],
                                    product_data['specifications'],
                                    product_data['specifications_vietnamese'],
                                    product_data['source_file'],
                                    embedding_blob,
                                    datetime.now().isoformat()
                                ))
                                added_count += 1
                                
                            except Exception as insert_error:
                                self.update_training_status(f"⚠️ Error inserting product: {insert_error}")
                                continue
                                
                    except Exception as batch_error:
                        self.update_training_status(f"⚠️ Batch embedding error: {batch_error}")
                        # Fallback to individual processing
                        for product_data in batch_products:
                            try:
                                if product_data['combined_text']:
                                    embedding = self.embedding_model.encode([product_data['combined_text']])[0]
                                    embedding_blob = embedding.astype(np.float32).tobytes()
                                else:
                                    embedding_blob = None
                                    
                                cursor.execute('''
                                    INSERT INTO products 
                                    (name, name_vietnamese, description, description_vietnamese, 
                                     category, category_vietnamese, price, features, features_vietnamese,
                                     specifications, specifications_vietnamese, source_file, embedding, created_at)
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                ''', (
                                    product_data['name'],
                                    product_data['name_vietnamese'],
                                    product_data['description'],
                                    product_data['description_vietnamese'],
                                    product_data['category'],
                                    product_data['category_vietnamese'],
                                    product_data['price'],
                                    product_data['features'],
                                    product_data['features_vietnamese'],
                                    product_data['specifications'],
                                    product_data['specifications_vietnamese'],
                                    product_data['source_file'],
                                    embedding_blob,
                                    datetime.now().isoformat()
                                ))
                                added_count += 1
                                
                            except Exception as fallback_error:
                                self.update_training_status(f"⚠️ Fallback processing error: {fallback_error}")
                                continue
                
                # Update progress
                self.update_training_status(f"📦 Processed batch {batch_start//batch_size + 1}: {added_count} products added")
                
            self.conn.commit()
            self.update_training_status(f"✅ RTX 4090 optimized processing: Added {added_count} products from Excel file with embeddings")
            
        except Exception as e:
            raise Exception(f"Error processing Excel file with RTX 4090: {e}")

    def process_pdf_file(self, file_path):
        """Process PDF files for knowledge base with Vietnamese support"""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    self.update_training_status(f"📄 Processed page {page_num + 1}/{len(reader.pages)}")
                except Exception as e:
                    self.update_training_status(f"⚠️ Error extracting page {page_num + 1}: {e}")
                    continue
                
            if text.strip():
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from PDF: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from PDF")
                
        except Exception as e:
            raise Exception(f"Error processing PDF file: {e}")
            
    def process_word_file(self, file_path):
        """Process Word documents for knowledge base with Vietnamese support"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                if para_num % 10 == 0:
                    self.update_training_status(f"📝 Processed {para_num} paragraphs...")
                    
            for table_num, table in enumerate(doc.tables):
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                self.update_training_status(f"📊 Processed table {table_num + 1}/{len(doc.tables)}")
                
            if text.strip():
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding_text = text[:5000]
                        embedding = self.embedding_model.encode([embedding_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    os.path.basename(file_path),
                    text[:10000],
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Added knowledge from Word doc: {len(text)} characters with embedding")
            else:
                self.update_training_status(f"⚠️ No text extracted from Word document")
                
        except Exception as e:
            raise Exception(f"Error processing Word file: {e}")
            
    def process_image_file(self, file_path):
        """Process image files using OCR with Vietnamese support"""
        try:
            if not self.ocr_reader:
                raise Exception("OCR model not available")
                
            self.update_training_status(f"🖼️ Running OCR on image...")
            
            results = self.ocr_reader.readtext(file_path)
            
            extracted_texts = []
            for (bbox, text, confidence) in results:
                if confidence > 0.5:
                    extracted_texts.append(text.strip())
                    
            if extracted_texts:
                combined_text = " ".join(extracted_texts)
                
                embedding_blob = None
                if self.embedding_model:
                    try:
                        embedding = self.embedding_model.encode([combined_text])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                    except Exception as e:
                        self.update_training_status(f"⚠️ Embedding generation failed: {e}")
                
                cursor = self.conn.cursor()
                cursor.execute('''
                    INSERT INTO knowledge_base 
                    (topic, content, source, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    f"OCR: {os.path.basename(file_path)}",
                    combined_text,
                    file_path,
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
                self.conn.commit()
                self.update_training_status(f"✅ Extracted {len(extracted_texts)} text segments from image")
                self.update_training_status(f"📝 Text preview: {combined_text[:100]}...")
            else:
                self.update_training_status(f"⚠️ No text detected in image")
                
        except Exception as e:
            raise Exception(f"Error processing image file: {e}")
    
    def setup_gui(self):
        """Setup the GUI interface with Vietnamese font support"""
        self.root = tk.Tk()
        self.root.title("SSTC SUPER SALES")
        self.root.geometry("1400x900")
        
        self.setup_vietnamese_fonts()
        self.setup_main_interface()
        
    def setup_vietnamese_fonts(self):
        """Setup fonts that support Vietnamese characters with enhanced input support"""
        vietnamese_fonts = [
            ('Segoe UI', 12),
            ('Times New Roman', 12),
            ('Arial Unicode MS', 12),
            ('Calibri', 12),
            ('Tahoma', 12),
            ('Microsoft Sans Serif', 12),
            ('DejaVu Sans', 12),
            ('Liberation Sans', 12),
            ('Verdana', 11)
        ]
        
        self.vietnamese_font = None
        self.input_font = None
        available_fonts = font.families()
        
        for font_name, size in vietnamese_fonts:
            if font_name in available_fonts:
                try:
                    test_font = font.Font(family=font_name, size=size)
                    
                    self.vietnamese_font = font.Font(family=font_name, size=size-1)
                    self.input_font = font.Font(
                        family=font_name, 
                        size=size, 
                        weight='normal'
                    )
                    
                    print(f"✅ Using Vietnamese font: {font_name}")
                    break
                except Exception as e:
                    print(f"⚠️ Font {font_name} failed: {e}")
                    continue
        
        if not self.vietnamese_font:
            self.vietnamese_font = font.Font(family="TkDefaultFont", size=11)
            self.input_font = font.Font(family="TkDefaultFont", size=12)
            print("⚠️ Using default font (Vietnamese characters may not display correctly)")
        
        self.root.option_add('*Font', self.vietnamese_font)
        
        try:
            self.root.tk.call('encoding', 'system', 'utf-8')
            print("✅ System encoding set to UTF-8 for Vietnamese support")
        except Exception as e:
            print(f"⚠️ Could not set UTF-8 encoding: {e}")
    
    def setup_main_interface(self):
        """Setup the main GUI interface with notebook tabs"""
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.setup_chat_tab()
        self.setup_database_tab()
        self.setup_training_tab()
        self.setup_analytics_tab()
    
    def setup_chat_tab(self):
        """Setup the main chat interface"""
        chat_frame = ttk.Frame(self.notebook)
        self.notebook.add(chat_frame, text="💬Chat")
        
        top_frame = ttk.Frame(chat_frame)
        top_frame.pack(fill=tk.X, padx=10, pady=5)
        
        lang_frame = ttk.Frame(top_frame)
        lang_frame.pack(side=tk.LEFT)
        
        ttk.Label(lang_frame, text="Ngôn ngữ / Language:", font=self.vietnamese_font).pack(side=tk.LEFT, padx=5)
        
        self.language_var = tk.StringVar(value=self.config['language_config']['default_language'])
        language_combo = ttk.Combobox(lang_frame, textvariable=self.language_var, 
                                     values=['vi', 'en'], state='readonly', width=8)
        language_combo.pack(side=tk.LEFT, padx=5)
        language_combo.bind('<<ComboboxSelected>>', self.on_language_change)
        
        # RTX 4090 status indicator
        status_frame = ttk.Frame(top_frame)
        status_frame.pack(side=tk.RIGHT)
        
        rtx_status = "🚀 RTX 4090" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else "📱 Standard GPU"
        ttk.Label(status_frame, text=rtx_status, font=self.vietnamese_font).pack(side=tk.RIGHT, padx=5)
        
        self.chat_display = scrolledtext.ScrolledText(
            chat_frame, 
            height=30, 
            state=tk.DISABLED,
            font=self.vietnamese_font,
            wrap=tk.WORD
        )
        self.chat_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.chat_display.tag_configure("user", foreground="blue", font=(self.vietnamese_font['family'], 10, 'bold'))
        self.chat_display.tag_configure("bot", foreground="green", font=(self.vietnamese_font['family'], 10))
        self.chat_display.tag_configure("system", foreground="gray", font=(self.vietnamese_font['family'], 9, 'italic'))
        self.chat_display.tag_configure("error", foreground="red", font=(self.vietnamese_font['family'], 9, 'italic'))
        
        input_frame = ttk.Frame(chat_frame)
        input_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.user_input = tk.Text(
            input_frame, 
            height=3, 
            font=self.input_font,
            wrap=tk.WORD
        )
        self.user_input.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        self.user_input.bind('<Control-Return>', self.send_message)
        self.user_input.bind('<Return>', self.on_enter_key)
        
        button_frame = ttk.Frame(input_frame)
        button_frame.pack(side=tk.RIGHT, fill=tk.Y)
        
        send_btn = ttk.Button(button_frame, text="Gửi / Send", command=self.send_message)
        send_btn.pack(fill=tk.X, pady=(0, 5))
        
        clear_btn = ttk.Button(button_frame, text="Xóa / Clear", command=self.clear_chat)
        clear_btn.pack(fill=tk.X)
        
        status_frame = ttk.Frame(chat_frame)
        status_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.status_label = ttk.Label(status_frame, text="RTX 4090 Ready / Sẵn sàng", font=self.vietnamese_font)
        self.status_label.pack(side=tk.LEFT)
        
        self.processing_label = ttk.Label(status_frame, text="", font=self.vietnamese_font)
        self.processing_label.pack(side=tk.RIGHT)
        
    def on_language_change(self, event=None):
        """Handle language change"""
        try:
            self.current_language = self.language_var.get()
            print(f"Language changed to: {self.current_language}")
        except Exception as e:
            print(f"Language change error: {e}")
    
    def setup_database_tab(self):
        """Setup database management interface with product categories"""
        db_frame = ttk.Frame(self.notebook)
        self.notebook.add(db_frame, text="🗄️ Database")
        
        # Add control panel with import button
        control_frame = ttk.Frame(db_frame)
        control_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Button(control_frame, text="📊 Import Excel File", 
                  command=self.open_import_dialog).pack(side='left', padx=5)
        
        ttk.Button(control_frame, text="🔄 Refresh All", 
                  command=self.refresh_all_tabs).pack(side='left', padx=5)
        # ADD Check Database BUTTON:
        ttk.Button(control_frame, text="🔍 Check Database", 
                  command=self.view_database_contents).pack(side='left', padx=5)
        # ADD Remove Item Button:
        ttk.Button(control_frame, text="🗑️ Remove Item", 
                  command=self.remove_selected_item).pack(side='left', padx=5)
        # Create sub-notebook for product categories
        self.db_sub_notebook = ttk.Notebook(db_frame)
        self.db_sub_notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
            
        # Initialize product databasess
        self.ssd_database = []
        self.memory_database = []
        self.motherboard_database = []
        
        # Create individual product tabs
        self.create_ssd_tab()
        self.create_memory_tab()
        self.create_motherboard_tab()
    def create_ssd_tab(self):
        """Create SSD products tab based on Excel template"""
        ssd_frame = ttk.Frame(self.db_sub_notebook)
        self.db_sub_notebook.add(ssd_frame, text="SSD")
        
        # Create input fields based on Excel template structure
        fields_frame = ttk.LabelFrame(ssd_frame, text="Add SSD Product", padding="10")
        fields_frame.pack(fill='x', padx=10, pady=10)
        
        # Input fields matching Excel template: STT, Sản phẩm, Giao thức, Model, Thông số, Kho, Giá bán lẻ
        ttk.Label(fields_frame, text="Sản phẩm:").grid(row=0, column=0, sticky='w', padx=5, pady=5)
        self.ssd_product = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.ssd_product.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Giao thức:").grid(row=1, column=0, sticky='w', padx=5, pady=5)
        self.ssd_interface = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.ssd_interface.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Model:").grid(row=2, column=0, sticky='w', padx=5, pady=5)
        self.ssd_model = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.ssd_model.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Thông số:").grid(row=3, column=0, sticky='w', padx=5, pady=5)
        self.ssd_specs = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.ssd_specs.grid(row=3, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Kho:").grid(row=4, column=0, sticky='w', padx=5, pady=5)
        self.ssd_stock = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.ssd_stock.grid(row=4, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Giá bán lẻ:").grid(row=5, column=0, sticky='w', padx=5, pady=5)
        self.ssd_price = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.ssd_price.grid(row=5, column=1, padx=5, pady=5)
        
        # Add Product button
        ttk.Button(fields_frame, text="Add Product", 
                   command=lambda: self.add_product('ssd')).grid(row=6, column=0, columnspan=2, pady=10)
        
        # Product display area
        display_frame = ttk.LabelFrame(ssd_frame, text="SSD Products", padding="10")
        display_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview for displaying products
        self.ssd_tree = ttk.Treeview(display_frame, columns=('STT', 'Product', 'Interface', 'Model', 'Specs', 'Stock', 'Price'), 
                                     show='headings', height=10)
        
        # Configure columns
        self.ssd_tree.heading('STT', text='STT')
        self.ssd_tree.heading('Product', text='Sản phẩm')
        self.ssd_tree.heading('Interface', text='Giao thức')
        self.ssd_tree.heading('Model', text='Model')
        self.ssd_tree.heading('Specs', text='Thông số')
        self.ssd_tree.heading('Stock', text='Kho')
        self.ssd_tree.heading('Price', text='Giá bán lẻ')
        
        # Column widths
        self.ssd_tree.column('STT', width=50)
        self.ssd_tree.column('Product', width=200)
        self.ssd_tree.column('Interface', width=100)
        self.ssd_tree.column('Model', width=150)
        self.ssd_tree.column('Specs', width=200)
        self.ssd_tree.column('Stock', width=100)
        self.ssd_tree.column('Price', width=100)
        
        self.ssd_tree.pack(fill=tk.BOTH, expand=True)
        
        # Scrollbar
        scrollbar_ssd = ttk.Scrollbar(display_frame, orient='vertical', command=self.ssd_tree.yview)
        scrollbar_ssd.pack(side='right', fill='y')
        self.ssd_tree.configure(yscrollcommand=scrollbar_ssd.set)

    def create_memory_tab(self):
        """Create Memory products tab based on Excel template"""
        memory_frame = ttk.Frame(self.db_sub_notebook)
        self.db_sub_notebook.add(memory_frame, text="Memory")
        
        # Create input fields based on Excel template structure
        fields_frame = ttk.LabelFrame(memory_frame, text="Add Memory Product", padding="10")
        fields_frame.pack(fill='x', padx=10, pady=10)
        
        # Input fields matching Excel template: STT, Sản phẩm, Giao thức, Model, Thông số, Kho, Giá bán lẻ
        ttk.Label(fields_frame, text="Sản phẩm:").grid(row=0, column=0, sticky='w', padx=5, pady=5)
        self.memory_product = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.memory_product.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Giao thức:").grid(row=1, column=0, sticky='w', padx=5, pady=5)
        self.memory_interface = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.memory_interface.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Model:").grid(row=2, column=0, sticky='w', padx=5, pady=5)
        self.memory_model = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.memory_model.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Thông số:").grid(row=3, column=0, sticky='w', padx=5, pady=5)
        self.memory_specs = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.memory_specs.grid(row=3, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Kho:").grid(row=4, column=0, sticky='w', padx=5, pady=5)
        self.memory_stock = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.memory_stock.grid(row=4, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Giá bán lẻ:").grid(row=5, column=0, sticky='w', padx=5, pady=5)
        self.memory_price = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.memory_price.grid(row=5, column=1, padx=5, pady=5)
        
        # Add Product button
        ttk.Button(fields_frame, text="Add Product", 
                   command=lambda: self.add_product('memory')).grid(row=6, column=0, columnspan=2, pady=10)
        
        # Product display area
        display_frame = ttk.LabelFrame(memory_frame, text="Memory Products", padding="10")
        display_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview for displaying products
        self.memory_tree = ttk.Treeview(display_frame, columns=('STT', 'Product', 'Interface', 'Model', 'Specs', 'Stock', 'Price'), 
                                       show='headings', height=10)
        
        # Configure columns
        self.memory_tree.heading('STT', text='STT')
        self.memory_tree.heading('Product', text='Sản phẩm')
        self.memory_tree.heading('Interface', text='Giao thức')
        self.memory_tree.heading('Model', text='Model')
        self.memory_tree.heading('Specs', text='Thông số')
        self.memory_tree.heading('Stock', text='Kho')
        self.memory_tree.heading('Price', text='Giá bán lẻ')
        
        # Column widths
        self.memory_tree.column('STT', width=50)
        self.memory_tree.column('Product', width=200)
        self.memory_tree.column('Interface', width=120)
        self.memory_tree.column('Model', width=150)
        self.memory_tree.column('Specs', width=200)
        self.memory_tree.column('Stock', width=100)
        self.memory_tree.column('Price', width=100)
        
        self.memory_tree.pack(fill=tk.BOTH, expand=True)
        
        # Scrollbar
        scrollbar_memory = ttk.Scrollbar(display_frame, orient='vertical', command=self.memory_tree.yview)
        scrollbar_memory.pack(side='right', fill='y')
        self.memory_tree.configure(yscrollcommand=scrollbar_memory.set)

    def create_motherboard_tab(self):
        """Create Motherboard products tab based on Excel template"""
        motherboard_frame = ttk.Frame(self.db_sub_notebook)
        self.db_sub_notebook.add(motherboard_frame, text="Motherboard")
        
        # Create input fields based on Excel template structure
        fields_frame = ttk.LabelFrame(motherboard_frame, text="Add Motherboard Product", padding="10")
        fields_frame.pack(fill='x', padx=10, pady=10)
        
        # Input fields matching Excel template: STT, Sản phẩm, Chipset, Model, Thông số, Kho, Giá bán lẻ
        ttk.Label(fields_frame, text="Sản phẩm:").grid(row=0, column=0, sticky='w', padx=5, pady=5)
        self.mb_product = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.mb_product.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Chipset:").grid(row=1, column=0, sticky='w', padx=5, pady=5)
        self.mb_chipset = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.mb_chipset.grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Model:").grid(row=2, column=0, sticky='w', padx=5, pady=5)
        self.mb_model = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.mb_model.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Thông số:").grid(row=3, column=0, sticky='w', padx=5, pady=5)
        self.mb_specs = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.mb_specs.grid(row=3, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Kho:").grid(row=4, column=0, sticky='w', padx=5, pady=5)
        self.mb_stock = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.mb_stock.grid(row=4, column=1, padx=5, pady=5)
        
        ttk.Label(fields_frame, text="Giá bán lẻ:").grid(row=5, column=0, sticky='w', padx=5, pady=5)
        self.mb_price = ttk.Entry(fields_frame, width=40, font=self.input_font)
        self.mb_price.grid(row=5, column=1, padx=5, pady=5)
        
        # Add Product button
        ttk.Button(fields_frame, text="Add Product", 
                   command=lambda: self.add_product('motherboard')).grid(row=6, column=0, columnspan=2, pady=10)
        
        # Product display area
        display_frame = ttk.LabelFrame(motherboard_frame, text="Motherboard Products", padding="10")
        display_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Treeview for displaying products
        self.mb_tree = ttk.Treeview(display_frame, columns=('STT', 'Product', 'Chipset', 'Model', 'Specs', 'Stock', 'Price'), 
                                   show='headings', height=10)
        
        # Configure columns
        self.mb_tree.heading('STT', text='STT')
        self.mb_tree.heading('Product', text='Sản phẩm')
        self.mb_tree.heading('Chipset', text='Chipset')
        self.mb_tree.heading('Model', text='Model')
        self.mb_tree.heading('Specs', text='Thông số')
        self.mb_tree.heading('Stock', text='Kho')
        self.mb_tree.heading('Price', text='Giá bán lẻ')
        
        # Column widths
        self.mb_tree.column('STT', width=50)
        self.mb_tree.column('Product', width=200)
        self.mb_tree.column('Chipset', width=100)
        self.mb_tree.column('Model', width=150)
        self.mb_tree.column('Specs', width=250)
        self.mb_tree.column('Stock', width=100)
        self.mb_tree.column('Price', width=100)
        
        self.mb_tree.pack(fill=tk.BOTH, expand=True)
        
        # Scrollbar
        scrollbar_mb = ttk.Scrollbar(display_frame, orient='vertical', command=self.mb_tree.yview)
        scrollbar_mb.pack(side='right', fill='y')
        self.mb_tree.configure(yscrollcommand=scrollbar_mb.set)
   
    def add_product(self, product_type):
        """Add a product to the appropriate database"""
        try:
            if product_type == 'ssd':
                # Get values from SSD input fields
                product = {
                    'product': self.ssd_product.get(),
                    'interface': self.ssd_interface.get(),
                    'model': self.ssd_model.get(),
                    'specifications': self.ssd_specs.get(),
                    'stock_status': self.ssd_stock.get(),
                    'price': self.ssd_price.get(),
                    'category': 'SSD'
                }
                
                # Validate inputs
                if all(product.values()):
                    # Generate auto-incrementing number
                    stt = len(self.ssd_database) + 1
                    
                    # Add to database list
                    product['stt'] = stt
                    self.ssd_database.append(product)
                    
                    # Parse price for database storage
                    try:
                        price_num = float(product['price'].replace(',', '').replace(' ', ''))
                    except:
                        price_num = 0
                    
                    # STANDARDIZED DATABASE INSERTION - Same as sample products
                    cursor = self.conn.cursor()
                    
                    # Build text for embedding - SAME AS SAMPLE PRODUCTS
                    combined_text = f"{product.get('product', '')} {product.get('model', '')} {product.get('specifications', '')}"
                    
                    embedding_blob = None
                    if self.embedding_model:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.logger.warning(f"Embedding generation error: {e}")
                    
                    # EXACT SAME INSERT STRUCTURE AS SAMPLE PRODUCTS
                    cursor.execute('''
                        INSERT INTO products 
                        (name, description, category, price, features, specifications, availability,
                         source_file, embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        product['model'],           # model -> name (SAME mapping)
                        product['product'],         # product -> description (SAME mapping)
                        product['category'],
                        price_num,
                        product['interface'],       # interface -> features (SAME mapping)
                        product['specifications'],
                        product['stock_status'],    # stock_status -> availability (SAME mapping)
                        'Manual Entry',
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                    self.conn.commit()
                    
                    # Add to treeview
                    self.ssd_tree.insert('', 'end', values=(
                        stt, product['product'], product['interface'], product['model'],
                        product['specifications'], product['stock_status'], product['price']
                    ))
                    
                    # Clear input fields
                    self.clear_ssd_fields()
                    self.show_message("Success", "SSD product added successfully!")
                else:
                    self.show_message("Error", "Please fill all fields!")
            
            elif product_type == 'memory':
                # Same structure for Memory products
                product = {
                    'product': self.memory_product.get(),
                    'interface': self.memory_interface.get(),
                    'model': self.memory_model.get(),
                    'specifications': self.memory_specs.get(),
                    'stock_status': self.memory_stock.get(),
                    'price': self.memory_price.get(),
                    'category': 'Memory'
                }
                # ... rest of memory code with SAME database insertion pattern
                
            elif product_type == 'motherboard':
                # Note: Motherboard uses 'chipset' in GUI but maps to 'features' in database
                product = {
                    'product': self.mb_product.get(),
                    'chipset': self.mb_chipset.get(),  # GUI field name
                    'model': self.mb_model.get(),
                    'specifications': self.mb_specs.get(),
                    'stock_status': self.mb_stock.get(),
                    'price': self.mb_price.get(),
                    'category': 'Motherboard'
                }
                
                # ... validation code ...
                
                # In database insertion, map chipset to features
                cursor.execute('''
                    INSERT INTO products 
                    (name, description, category, price, features, specifications, availability,
                     source_file, embedding, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    product['model'],
                    product['product'],
                    product['category'],
                    price_num,
                    product['chipset'],     # chipset -> features for motherboard
                    product['specifications'],
                    product['stock_status'],
                    'Manual Entry',
                    embedding_blob,
                    datetime.now().isoformat()
                ))
                
        except Exception as e:
            self.show_message("Error", f"Error adding product: {str(e)}")

    def clear_ssd_fields(self):
        """Clear SSD input fields"""
        self.ssd_product.delete(0, 'end')
        self.ssd_interface.delete(0, 'end')
        self.ssd_model.delete(0, 'end')
        self.ssd_specs.delete(0, 'end')
        self.ssd_stock.delete(0, 'end')
        self.ssd_price.delete(0, 'end')

    def clear_memory_fields(self):
        """Clear Memory input fields"""
        self.memory_product.delete(0, 'end')
        self.memory_interface.delete(0, 'end')
        self.memory_model.delete(0, 'end')
        self.memory_specs.delete(0, 'end')
        self.memory_stock.delete(0, 'end')
        self.memory_price.delete(0, 'end')

    def clear_motherboard_fields(self):
        """Clear Motherboard input fields"""
        self.mb_product.delete(0, 'end')
        self.mb_chipset.delete(0, 'end')
        self.mb_model.delete(0, 'end')
        self.mb_specs.delete(0, 'end')
        self.mb_stock.delete(0, 'end')
        self.mb_price.delete(0, 'end')

    def show_message(self, title, message):
        """Show message dialog"""
        from tkinter import messagebox
        messagebox.showinfo(title, message)
    
    def bulk_insert_products(self, category: str, df: pd.DataFrame) -> Dict[str, int]:
        """Perform bulk insert with comprehensive error handling"""
        result = {
            'total_records': len(df),
            'successful': 0,
            'duplicates': 0,
            'errors': 0,
            'error_details': []
        }
        try:
            cursor = self.conn.cursor()
            
            for index, row in df.iterrows():
                try:
                    # Check for existing product
                    cursor.execute("""
                        SELECT COUNT(*) FROM products 
                        WHERE name = ? AND category = ?
                    """, (row.get('model', ''), category))
                    
                    if cursor.fetchone()[0] > 0:
                        result['duplicates'] += 1
                        continue
                    
                    # Create searchable text
                    combined_text = f"{row.get('product_name', '')} {row.get('model', '')} {row.get('specifications', '')}"
                    
                    # Generate embedding if model available
                    embedding_blob = None
                    if self.embedding_model and combined_text.strip():
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            print(f"Embedding generation failed: {e}")
                    
                    # Insert into database
                    cursor.execute('''
                        INSERT INTO products 
                        (name, description, category, price, features, specifications, availability,
                         source_file, embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        row.get('model', ''),
                        row.get('product_name', ''),
                        category,
                        float(row.get('price', 0)),
                        row.get('specifications', ''),
                        f"{row.get('interface', '')} {row.get('chipset', '')}".strip(),
                        row.get('stock_status', 'Available'),
                        'Excel Import',
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                    
                    result['successful'] += 1
                    
                except Exception as e:
                    result['errors'] += 1
                    result['error_details'].append({
                        'product': row.get('product_name', f'Row {index}'),
                        'error': str(e)
                    })
            
            self.conn.commit()
            
            # Refresh the appropriate tab
            self.refresh_category_view_after_import(category)
            
        except Exception as e:
            self.conn.rollback()
            raise Exception(f"Database operation failed: {str(e)}")
        return result

    def refresh_category_view_after_import(self, category):
        """Refresh TreeView after import"""
        try:
            if category == 'SSD' and hasattr(self, 'ssd_tree'):
                self.refresh_ssd_tree()
            elif category == 'Memory' and hasattr(self, 'memory_tree'):
                self.refresh_memory_tree()
            elif category == 'Motherboard' and hasattr(self, 'mb_tree'):
                self.refresh_motherboard_tree()
        except Exception as e:
            print(f"Error refreshing tree view: {e}")
            
    def open_import_dialog(self):
        """Open the Excel import dialog"""
        try:
            ExcelImportDialog(self.root, self)
        except Exception as e:
            messagebox.showerror("Error", f"Could not open import dialog: {str(e)}")

    def refresh_all_tabs(self):
        """Refresh all product tabs"""
        try:
            self.refresh_ssd_tree()
            self.refresh_memory_tree()
            self.refresh_motherboard_tree()
            messagebox.showinfo("Success", "All tabs refreshed successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Error refreshing tabs: {str(e)}")                
        except Exception as e:
            print(f"Error refreshing tree view: {e}")

    def refresh_ssd_tree(self):
        """Refresh SSD TreeView"""
        try:
            # Clear existing items
            for item in self.ssd_tree.get_children():
                self.ssd_tree.delete(item)
            
            # Reload from database
            cursor = self.conn.cursor()
            cursor.execute("SELECT * FROM products WHERE category = 'SSD' ORDER BY created_at")
            products = cursor.fetchall()
            
            for i, product in enumerate(products, 1):
                if len(product) >= 6:
                    self.ssd_tree.insert('', 'end', values=(
                        i, product[2] or 'SSD Product', product[6] or 'Interface', 
                        product[1] or '', product[6] or '', product[8] or 'Available', 
                        f"{product[5] or 0:,.0f}"
                    ))
        except Exception as e:
            print(f"Error refreshing SSD tree: {e}")

    def refresh_memory_tree(self):
        """Refresh Memory TreeView"""
        try:
            for item in self.memory_tree.get_children():
                self.memory_tree.delete(item)
            
            cursor = self.conn.cursor()
            cursor.execute("SELECT * FROM products WHERE category = 'Memory' ORDER BY created_at")
            products = cursor.fetchall()
            
            for i, product in enumerate(products, 1):
                if len(product) >= 6:
                    self.memory_tree.insert('', 'end', values=(
                        i, product[2] or 'Memory Product', product[6] or 'Interface', 
                        product[1] or '', product[6] or '', product[8] or 'Available', 
                        f"{product[5] or 0:,.0f}"
                    ))
        except Exception as e:
            print(f"Error refreshing Memory tree: {e}")

    def refresh_motherboard_tree(self):
        """Refresh Motherboard TreeView"""
        try:
            for item in self.mb_tree.get_children():
                self.mb_tree.delete(item)
            
            cursor = self.conn.cursor()
            cursor.execute("SELECT * FROM products WHERE category = 'Motherboard' ORDER BY created_at")
            products = cursor.fetchall()
            
            for i, product in enumerate(products, 1):
                if len(product) >= 6:
                    self.mb_tree.insert('', 'end', values=(
                        i, product[2] or 'Motherboard Product', product[7] or 'Chipset', 
                        product[1] or '', product[6] or '', product[8] or 'Available', 
                        f"{product[5] or 0:,.0f}"
                    ))
        except Exception as e:
            print(f"Error refreshing Motherboard tree: {e}")
            
    def load_existing_products_to_tabs(self):
        """Load existing products from database into respective tabs"""
        try:
            cursor = self.conn.cursor()
            
            # Load SSD products
            cursor.execute("SELECT * FROM products WHERE category = 'SSD' ORDER BY created_at")
            ssd_products = cursor.fetchall()
            for i, product in enumerate(ssd_products, 1):
                if len(product) >= 6:  # Ensure we have enough columns
                    # Map database fields to display format
                    product_name = product[1] if product[1] else 'SSD Product'
                    model = product[1] if product[1] else ''
                    price = product[5] if product[5] else 0
                    features = product[6] if len(product) > 6 and product[6] else ''
                    specs = product[7] if len(product) > 7 and product[7] else ''
                    availability = product[8] if len(product) > 8 and product[8] else 'Còn hàng'
                    
                    self.ssd_tree.insert('', 'end', values=(
                        i, product_name, 'M2 NVMe', model, features, availability, f"{price:,.0f}"
                    ))
            
            # Load Memory products  
            cursor.execute("SELECT * FROM products WHERE category = 'Memory' ORDER BY created_at")
            memory_products = cursor.fetchall()
            for i, product in enumerate(memory_products, 1):
                if len(product) >= 6:
                    product_name = product[1] if product[1] else 'Memory Product'
                    model = product[1] if product[1] else ''
                    price = product[5] if product[5] else 0
                    features = product[6] if len(product) > 6 and product[6] else ''
                    specs = product[7] if len(product) > 7 and product[7] else ''
                    availability = product[8] if len(product) > 8 and product[8] else 'Còn hàng'
                    
                    self.memory_tree.insert('', 'end', values=(
                        i, product_name, 'DDR4/DDR5', model, features, availability, f"{price:,.0f}"
                    ))
            
            # Load Motherboard products
            cursor.execute("SELECT * FROM products WHERE category = 'Motherboard' ORDER BY created_at")
            mb_products = cursor.fetchall()
            for i, product in enumerate(mb_products, 1):
                if len(product) >= 6:
                    product_name = product[1] if product[1] else 'Motherboard Product'
                    model = product[1] if product[1] else ''
                    price = product[5] if product[5] else 0
                    features = product[6] if len(product) > 6 and product[6] else ''
                    specs = product[7] if len(product) > 7 and product[7] else ''
                    availability = product[8] if len(product) > 8 and product[8] else 'Còn hàng'
                    
                    self.mb_tree.insert('', 'end', values=(
                        i, product_name, 'Intel/AMD', model, features, availability, f"{price:,.0f}"
                    ))
                    
            print(f"✅ Loaded {len(ssd_products)} SSD, {len(memory_products)} Memory, {len(mb_products)} Motherboard products to tabs")
                    
        except Exception as e:
            print(f"Error loading existing products: {e}")

        
    def setup_training_tab(self):
        """Setup training interface with file upload capabilities"""
        training_frame = ttk.Frame(self.notebook)
        self.notebook.add(training_frame, text="📚 Data Training")
        
        upload_frame = ttk.LabelFrame(training_frame, text="RTX 4090 Optimized Training Data Upload")
        upload_frame.pack(fill=tk.X, padx=10, pady=5)
        
        upload_btn_frame = ttk.Frame(upload_frame)
        upload_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(upload_btn_frame, text="📊 Upload Excel Files", 
                  command=lambda: self.process_files('excel')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📄 Upload PDF Files", 
                  command=lambda: self.process_files('pdf')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="📝 Upload Word Files", 
                  command=lambda: self.process_files('word')).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(upload_btn_frame, text="🖼️ Upload Images (OCR)", 
                  command=lambda: self.process_files('image')).pack(side=tk.LEFT, padx=5, pady=5)
        
        options_frame = ttk.LabelFrame(training_frame, text="RTX 4090 Training Options")
        options_frame.pack(fill=tk.X, padx=10, pady=5)
        
        options_btn_frame = ttk.Frame(options_frame)
        options_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(options_btn_frame, text="🔄 Regenerate All Embeddings", 
                  command=self.regenerate_all_embeddings).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="🧹 Clear Training Data", 
                  command=self.clear_training_data).pack(side=tk.LEFT, padx=5)
        ttk.Button(options_btn_frame, text="📈 Training Statistics", 
                  command=self.show_training_stats).pack(side=tk.LEFT, padx=5)
        
        self.training_status = scrolledtext.ScrolledText(
            training_frame, height=20, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.training_status.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def setup_analytics_tab(self):
        """Setup analytics interface"""
        analytics_frame = ttk.Frame(self.notebook)
        self.notebook.add(analytics_frame, text="📊 Data Analytics")
        
        metrics_frame = ttk.LabelFrame(analytics_frame, text="RTX 4090 Performance Metrics")
        metrics_frame.pack(fill=tk.X, padx=10, pady=5)
        
        metrics_btn_frame = ttk.Frame(metrics_frame)
        metrics_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(metrics_btn_frame, text="🔄 Refresh Analytics", 
                  command=self.refresh_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📤 Export Data", 
                  command=self.export_analytics).pack(side=tk.LEFT, padx=5, pady=5)
        ttk.Button(metrics_btn_frame, text="📋 Conversation History", 
                  command=self.view_conversation_history).pack(side=tk.LEFT, padx=5, pady=5)
        
        self.analytics_display = scrolledtext.ScrolledText(
            analytics_frame, height=25, state=tk.DISABLED, font=self.vietnamese_font
        )
        self.analytics_display.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
    
    def on_enter_key(self, event):
        """Handle Enter key press"""
        if event.state & 0x4:
            self.send_message()
        else:
            return None
    
    def on_language_change(self, event=None):
        """Handle language change"""
        try:
            self.current_language = self.language_var.get()
            print(f"Language changed to: {self.current_language}")
        except Exception as e:
            print(f"Language change error: {e}")
    
    def send_message(self, event=None):
        """Handle sending user message with error handling"""
        try:
            user_text = self.user_input.get(1.0, tk.END).strip()
            if not user_text:
                return
                
            self.user_input.delete(1.0, tk.END)
            
            self.display_message("Bạn / You", user_text, "user")
            
            if hasattr(self, 'processing_label'):
                status_text = "Đang xử lý với RTX 4090..." if self.current_language == 'vi' else "Processing with RTX 4090..."
                self.processing_label.config(text=status_text)
                self.root.update()
            
            threading.Thread(target=self.process_message_thread, args=(user_text,), daemon=True).start()
        except Exception as e:
            print(f"Send message error: {e}")
            error_msg = f"Lỗi gửi tin nhắn / Send error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
    
    def process_message_thread(self, user_input):
        """Process message in separate thread with comprehensive error handling"""
        try:
            result = self.process_message(user_input)
            
            response_tag = "error" if result.get('data_source') == 'error' else "bot"
            self.display_message("RTX 4090 AI Assistant", result.get('response', ''), response_tag)
            
            user_language = result.get('user_language', 'vi')
            rtx_indicator = "🚀" if result.get('rtx4090_optimized', False) else "📱"
            
            if user_language == 'vi':
                if result.get('data_source') == 'local_database':
                    status_text = f"{rtx_indicator} Tìm thấy {result.get('local_results_count', 0)} sản phẩm ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'knowledge_base':
                    status_text = f"{rtx_indicator} Tìm thấy trong cơ sở tri thức ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'error':
                    status_text = f"❌ Lỗi ({result.get('processing_time', 0):.1f}s)"
                else:
                    status_text = f"{rtx_indicator} Hoàn thành ({result.get('processing_time', 0):.1f}s)"
            else:
                if result.get('data_source') == 'local_database':
                    status_text = f"{rtx_indicator} Found {result.get('local_results_count', 0)} products ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'knowledge_base':
                    status_text = f"{rtx_indicator} Found in knowledge base ({result.get('processing_time', 0):.1f}s)"
                elif result.get('data_source') == 'error':
                    status_text = f"❌ Error ({result.get('processing_time', 0):.1f}s)"
                else:
                    status_text = f"{rtx_indicator} Complete ({result.get('processing_time', 0):.1f}s)"
            
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text=status_text))
            
        except Exception as e:
            self.logger.error(f"Message thread error: {e}")
            error_msg = f"Lỗi xử lý RTX 4090 / RTX 4090 processing error: {str(e)}"
            self.display_message("Hệ thống / System", error_msg, "error")
            if hasattr(self, 'processing_label'):
                self.root.after(0, lambda: self.processing_label.config(text="RTX 4090 Error"))
    
    def display_message(self, sender, message, tag):
        """Display message in chat window with error handling"""
        def update_display():
            try:
                self.chat_display.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime("%H:%M:%S")
                
                self.chat_display.insert(tk.END, f"[{timestamp}] {sender}:\n", tag)
                self.chat_display.insert(tk.END, f"{message}\n\n")
                
                self.chat_display.config(state=tk.DISABLED)
                self.chat_display.see(tk.END)
            except Exception as e:
                print(f"Display update error: {e}")
            
        try:
            if threading.current_thread() != threading.main_thread():
                self.root.after(0, update_display)
            else:
                update_display()
        except Exception as e:
            print(f"Display message error: {e}")
    
    def clear_chat(self):
        """Clear chat display with error handling"""
        try:
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete(1.0, tk.END)
            self.chat_display.config(state=tk.DISABLED)
            
            self.conversation_context.clear()
            self.conversation_summary = ""
            self.response_cache.clear()  # Clear RTX 4090 response cache
            
            # Clear GPU cache if available
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
            
            self.logger.info("Chat history, context, and GPU cache cleared")
            
        except Exception as e:
            print(f"Clear chat error: {e}")
    
    def update_training_status(self, message):
        """Update training status display with Vietnamese support"""
        def update():
            if hasattr(self, 'training_status'):
                self.training_status.config(state=tk.NORMAL)
                timestamp = datetime.now().strftime('%H:%M:%S')
                self.training_status.insert(tk.END, f"[{timestamp}] {message}\n")
                self.training_status.config(state=tk.DISABLED)
                self.training_status.see(tk.END)
            
        if threading.current_thread() != threading.main_thread():
            self.root.after(0, update)
        else:
            update()
    
    def process_files(self, file_type):
        """Process different types of files for training data"""
        file_types = {
            'excel': [('Excel files', '*.xlsx *.xls *.csv')],
            'pdf': [('PDF files', '*.pdf')],
            'word': [('Word files', '*.docx *.doc')],
            'image': [('Image files', '*.png *.jpg *.jpeg *.bmp *.tiff')]
        }
        
        files = filedialog.askopenfilenames(
            title=f"Select {file_type} files for RTX 4090 processing",
            filetypes=file_types[file_type]
        )
        
        if files:
            self.update_training_status(f"🚀 Starting RTX 4090 optimized processing of {len(files)} {file_type} file(s)...")
            threading.Thread(
                target=self.process_files_worker,
                args=(files, file_type),
                daemon=True
            ).start()
            
    def process_files_worker(self, files, file_type):
        """Worker thread for processing files with RTX 4090 optimization"""
        for file_path in files:
            try:
                self.update_training_status(f"🔄 RTX 4090 processing: {os.path.basename(file_path)}")
                
                if file_type == 'excel':
                    self.process_excel_file(file_path)
                elif file_type == 'pdf':
                    self.process_pdf_file(file_path)
                elif file_type == 'word':
                    self.process_word_file(file_path)
                elif file_type == 'image':
                    self.process_image_file(file_path)
                    
                self.update_training_status(f"✅ RTX 4090 completed: {os.path.basename(file_path)}")
                
                # Clear GPU cache between files for RTX 4090 optimization
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                
            except Exception as e:
                self.update_training_status(f"❌ RTX 4090 error processing {os.path.basename(file_path)}: {e}")
                
    def view_products(self):
        """View products in database with RTX 4090 optimization info"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT name, name_vietnamese, description, description_vietnamese, 
                       category, price, source_file, embedding 
                FROM products 
                ORDER BY created_at DESC
            """)
            products = cursor.fetchall()
            
            self.product_display.config(state=tk.NORMAL)
            self.product_display.delete(1.0, tk.END)
            
            if products:
                rtx_status = "🚀 RTX 4090 Optimized" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else "📱 Standard GPU"
                self.product_display.insert(tk.END, f"📦 {rtx_status} Database: Found {len(products)} products:\n\n")
                
                for i, product in enumerate(products, 1):
                    name = product[0] or product[1] or "Unnamed Product"
                    name_vi = product[1] or ""
                    desc = product[2] or product[3] or "No description"
                    category = product[4] or "General"
                    price = product[5] or 0
                    source = os.path.basename(product[6]) if product[6] else "Manual"
                    embedding_status = "✅" if product[7] else "❌"
                    
                    self.product_display.insert(tk.END, f"{i}. {name} {embedding_status}\n")
                    if name_vi:
                        self.product_display.insert(tk.END, f"   Vietnamese: {name_vi}\n")
                    self.product_display.insert(tk.END, f"   Description: {desc[:100]}...\n")
                    self.product_display.insert(tk.END, f"   Category: {category} | Price: ${price} | Source: {source}\n\n")
            else:
                self.product_display.insert(tk.END, "📦 No products found in RTX 4090 optimized database.\n\n")
                self.product_display.insert(tk.END, "💡 To add products:\n")
                self.product_display.insert(tk.END, "• Click 'Add Sample Products' for RTX 4090 demo data\n")
                self.product_display.insert(tk.END, "• Use the Training tab to upload Excel/CSV files\n")
                self.product_display.insert(tk.END, "• Use PDFs/Word docs for knowledge base\n")
                
            self.product_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing RTX 4090 products: {e}")
    
    def test_search(self):
        """Test RTX 4090 optimized search functionality"""
        test_queries = [
            "RTX 4090 gaming", "RTX 4090 gaming",
            "AI development", "phát triển AI", 
            "VR gaming", "gaming VR",
            "content creation", "sáng tạo nội dung",
            "workstation", "máy trạm",
            "4K gaming", "gaming 4K",
            "machine learning", "học máy",
            "deep learning", "học sâu"
        ]
        
        self.product_display.config(state=tk.NORMAL)
        self.product_display.delete(1.0, tk.END)
        
        rtx_status = "🚀 RTX 4090" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else "📱 Standard GPU"
        self.product_display.insert(tk.END, f"🔍 Testing {rtx_status} optimized search functionality...\n")
        self.product_display.insert(tk.END, f"Threshold: {self.config['search_config']['local_similarity_threshold']}\n\n")
        
        for query in test_queries:
            self.product_display.insert(tk.END, f"Query: '{query}'\n")
            start_time = time.time()
            results = self.search_local_database_rtx4090(query)
            search_time = time.time() - start_time
            
            if results:
                for result in results[:3]:
                    self.product_display.insert(tk.END, f"  ✓ {result['name']} (similarity: {result['similarity']:.3f})\n")
                self.product_display.insert(tk.END, f"  ⚡ Search time: {search_time:.3f}s\n")
            else:
                self.product_display.insert(tk.END, "  ✗ No results found\n")
            
            self.product_display.insert(tk.END, "\n")
            
        self.product_display.config(state=tk.DISABLED)
    
    def regenerate_all_embeddings(self):
        """Regenerate all embeddings with RTX 4090 optimization"""
        result = messagebox.askyesno(
            "Regenerate RTX 4090 Embeddings", 
            "This will regenerate all embeddings using RTX 4090 optimization.\n\n"
            "This may take several minutes depending on the amount of data.\n\n"
            "RTX 4090 will process embeddings in larger batches for better performance.\n\n"
            "Continue?"
        )
        
        if result:
            threading.Thread(target=self.regenerate_embeddings_worker_rtx4090, daemon=True).start()
            
    def regenerate_embeddings_worker_rtx4090(self):
        """RTX 4090 optimized worker thread for regenerating embeddings"""
        try:
            if not self.embedding_model:
                self.update_training_status("❌ Embedding model not available")
                return
                
            cursor = self.conn.cursor()
            
            self.update_training_status("🚀 RTX 4090 regenerating product embeddings...")
            cursor.execute("""
                SELECT id, name, name_vietnamese, description, description_vietnamese, 
                       features, features_vietnamese, specifications, specifications_vietnamese 
                FROM products
            """)
            products = cursor.fetchall()
            
            # RTX 4090 optimization: Process in larger batches
            batch_size = 64 if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else 32
            
            for i in range(0, len(products), batch_size):
                batch = products[i:i+batch_size]
                batch_texts = []
                batch_ids = []
                
                for product in batch:
                    product_id = product[0]
                    text_parts = [str(part) for part in product[1:] if part]
                    combined_text = ' '.join(text_parts).strip()
                    
                    if combined_text:
                        batch_texts.append(combined_text)
                        batch_ids.append(product_id)
                
                if batch_texts:
                    try:
                        # RTX 4090 batch embedding generation
                        embeddings = self.embedding_model.encode(batch_texts, batch_size=batch_size)
                        
                        for product_id, embedding in zip(batch_ids, embeddings):
                            embedding_blob = embedding.astype(np.float32).tobytes()
                            cursor.execute("""
                                UPDATE products 
                                SET embedding = ?, updated_at = ?
                                WHERE id = ?
                            """, (embedding_blob, datetime.now().isoformat(), product_id))
                        
                        self.update_training_status(f"🚀 RTX 4090 processed {i + len(batch)}/{len(products)} products...")
                        
                    except Exception as batch_error:
                        self.update_training_status(f"⚠️ RTX 4090 batch error: {batch_error}")
                        # Fallback to individual processing
                        for j, (product_id, text) in enumerate(zip(batch_ids, batch_texts)):
                            try:
                                embedding = self.embedding_model.encode([text])[0]
                                embedding_blob = embedding.astype(np.float32).tobytes()
                                cursor.execute("""
                                    UPDATE products 
                                    SET embedding = ?, updated_at = ?
                                    WHERE id = ?
                                """, (embedding_blob, datetime.now().isoformat(), product_id))
                            except Exception as e:
                                self.update_training_status(f"⚠️ Error processing product {product_id}: {e}")
                                continue
            
            # Knowledge base embeddings
            self.update_training_status("🚀 RTX 4090 regenerating knowledge base embeddings...")
            cursor.execute("SELECT id, content FROM knowledge_base")
            knowledge_entries = cursor.fetchall()
            
            for i, (kb_id, content) in enumerate(knowledge_entries):
                try:
                    if content and content.strip():
                        embedding = self.embedding_model.encode([content[:5000]])[0]
                        embedding_blob = embedding.astype(np.float32).tobytes()
                        
                        cursor.execute("""
                            UPDATE knowledge_base 
                            SET embedding = ?
                            WHERE id = ?
                        """, (embedding_blob, kb_id))
                        
                        if (i + 1) % 10 == 0:
                            self.update_training_status(f"📚 RTX 4090 processed {i + 1}/{len(knowledge_entries)} knowledge entries...")
                            
                except Exception as e:
                    self.update_training_status(f"⚠️ Error processing knowledge entry {kb_id}: {e}")
                    continue
                    
            self.conn.commit()
            
            # Clear GPU cache after processing
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                
            self.update_training_status(f"✅ RTX 4090 successfully regenerated embeddings for {len(products)} products and {len(knowledge_entries)} knowledge entries!")
            
        except Exception as e:
            self.update_training_status(f"❌ RTX 4090 error regenerating embeddings: {e}")
    
    def clear_training_data(self):
        """Clear all training data with confirmation"""
        result = messagebox.askyesno(
            "Clear RTX 4090 Training Data", 
            "Are you sure you want to clear all RTX 4090 optimized training data?\n\n"
            "This will delete:\n"
            "• All products\n"
            "• All knowledge base entries\n"
            "• All embeddings\n"
            "• RTX 4090 response cache\n\n"
            "This action cannot be undone!"
        )
        
        if result:
            try:
                cursor = self.conn.cursor()
                cursor.execute("DELETE FROM products")
                cursor.execute("DELETE FROM knowledge_base")
                self.conn.commit()
                
                # Clear RTX 4090 specific caches
                if hasattr(self, 'response_cache'):
                    self.response_cache.clear()
                
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                
                self.update_training_status("🧹 All RTX 4090 training data and caches cleared successfully")
                messagebox.showinfo("Success", "All RTX 4090 training data has been cleared.")
                
            except Exception as e:
                self.update_training_status(f"❌ Error clearing RTX 4090 data: {e}")
                messagebox.showerror("Error", f"Error clearing RTX 4090 training data: {e}")
    
    def show_training_stats(self):
        """Show RTX 4090 optimized training data statistics"""
        try:
            cursor = self.conn.cursor()
            
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(DISTINCT category) FROM products")
            category_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base WHERE embedding IS NOT NULL")
            knowledge_with_embeddings = cursor.fetchone()[0]
            
            # RTX 4090 specific stats
            rtx_status = "🚀 RTX 4090 Optimized" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else "📱 Standard GPU"
            cache_size = len(getattr(self, 'response_cache', {}))
            
            self.training_status.config(state=tk.NORMAL)
            self.training_status.delete(1.0, tk.END)
            
            self.training_status.insert(tk.END, f"=== {rtx_status} TRAINING DATA STATISTICS ===\n\n")
            self.training_status.insert(tk.END, f"📦 Products: {product_count}\n")
            self.training_status.insert(tk.END, f"🔗 Products with embeddings: {products_with_embeddings}/{product_count}\n")
            self.training_status.insert(tk.END, f"📂 Categories: {category_count}\n\n")
            self.training_status.insert(tk.END, f"📚 Knowledge base entries: {knowledge_count}\n")
            self.training_status.insert(tk.END, f"🔗 Knowledge with embeddings: {knowledge_with_embeddings}/{knowledge_count}\n\n")
            self.training_status.insert(tk.END, f"🚀 RTX 4090 Response cache: {cache_size} entries\n")
            
            if torch.cuda.is_available():
                memory_allocated = torch.cuda.memory_allocated(0) / (1024**3)
                memory_reserved = torch.cuda.memory_reserved(0) / (1024**3)
                self.training_status.insert(tk.END, f"💾 GPU Memory: {memory_allocated:.1f}GB allocated, {memory_reserved:.1f}GB reserved\n\n")
            
            cursor.execute("""
                SELECT name, created_at, source_file 
                FROM products 
                ORDER BY created_at DESC 
                LIMIT 5
            """)
            recent_products = cursor.fetchall()
            
            if recent_products:
                self.training_status.insert(tk.END, "📈 Recent Products:\n")
                for name, created_at, source_file in recent_products:
                    source = os.path.basename(source_file) if source_file else "Manual"
                    self.training_status.insert(tk.END, f"  • {name} ({source})\n")
                    
            self.training_status.config(state=tk.DISABLED)
            
        except Exception as e:
            self.update_training_status(f"❌ Error showing RTX 4090 statistics: {e}")
    
    def refresh_analytics(self):
        """Refresh analytics display with RTX 4090 performance metrics"""
        try:
            cursor = self.conn.cursor()
            
            cursor.execute("""
                SELECT 
                    COUNT(*) as total_conversations,
                    AVG(response_time) as avg_response_time,
                    data_source,
                    COUNT(*) as count_by_source
                FROM conversations 
                GROUP BY data_source
            """)
            
            source_stats = cursor.fetchall()
            
            cursor.execute("""
                SELECT timestamp, user_input, user_language, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 10
            """)
            
            recent_convs = cursor.fetchall()
            
            cursor.execute("SELECT COUNT(*) FROM products")
            product_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM products WHERE embedding IS NOT NULL")
            products_with_embeddings = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM knowledge_base")
            knowledge_count = cursor.fetchone()[0]
            
            cursor.execute("""
                SELECT user_language, COUNT(*) as count
                FROM conversations 
                WHERE user_language IS NOT NULL
                GROUP BY user_language
            """)
            language_stats = cursor.fetchall()
            
            # RTX 4090 specific metrics
            rtx_status = "🚀 RTX 4090 Optimized" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else "📱 Standard GPU"
            cache_size = len(getattr(self, 'response_cache', {}))
            
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            self.analytics_display.insert(tk.END, f"=== {rtx_status} AI CHATBOT ANALYTICS ===\n\n")
            self.analytics_display.insert(tk.END, f"📦 Products in Database: {product_count}\n")
            self.analytics_display.insert(tk.END, f"🔗 Products with Embeddings: {products_with_embeddings}/{product_count}\n")
            self.analytics_display.insert(tk.END, f"📚 Knowledge Base Entries: {knowledge_count}\n")
            self.analytics_display.insert(tk.END, f"🚀 Response Cache: {cache_size} entries\n")
            
            if torch.cuda.is_available():
                gpu_name = torch.cuda.get_device_name(0)
                memory_allocated = torch.cuda.memory_allocated(0) / (1024**3)
                memory_total = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                self.analytics_display.insert(tk.END, f"🎮 GPU: {gpu_name}\n")
                self.analytics_display.insert(tk.END, f"💾 GPU Memory: {memory_allocated:.1f}GB / {memory_total:.1f}GB\n")
            
            self.analytics_display.insert(tk.END, f"\n=== LANGUAGE USAGE ===\n")
            for lang, count in language_stats:
                lang_name = "Vietnamese" if lang == 'vi' else "English" if lang == 'en' else lang
                self.analytics_display.insert(tk.END, f"{lang_name}: {count} conversations\n")
            
            self.analytics_display.insert(tk.END, "\n=== RESPONSE SOURCES ===\n")
            for stat in source_stats:
                if len(stat) >= 4:
                    self.analytics_display.insert(tk.END, f"{stat[2]}: {stat[3]} responses\n")
                
            if source_stats:
                total_conversations = sum(stat[3] for stat in source_stats if len(stat) >= 4)
                if total_conversations > 0:
                    avg_response_time = sum(stat[1] * stat[3] for stat in source_stats if stat[1] and len(stat) >= 4) / total_conversations
                    self.analytics_display.insert(tk.END, f"\nTotal Conversations: {total_conversations}\n")
                    self.analytics_display.insert(tk.END, f"Average Response Time: {avg_response_time:.2f} seconds\n")
            
            self.analytics_display.insert(tk.END, "\n=== RECENT CONVERSATIONS ===\n")
            for conv in recent_convs:
                if len(conv) >= 5:
                    timestamp, user_input, user_lang, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    self.analytics_display.insert(tk.END, f"{timestamp} {lang_flag} - {data_source} ({response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input[:100]}...\n\n")
                
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error refreshing RTX 4090 analytics: {e}")
    
    def export_analytics(self):
        """Export analytics data to CSV with RTX 4090 metrics"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT * FROM conversations")
            
            column_names = [description[0] for description in cursor.description]
            
            df = pd.DataFrame(cursor.fetchall(), columns=column_names)
            
            rtx_prefix = "rtx4090_" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else ""
            export_path = f"{rtx_prefix}vietnamese_chatbot_analytics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            df.to_csv(export_path, index=False, encoding='utf-8')
            
            messagebox.showinfo("Success", f"RTX 4090 analytics exported to: {export_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error exporting RTX 4090 analytics: {e}")
    
    def view_conversation_history(self):
        """View detailed conversation history with RTX 4090 performance info"""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                SELECT timestamp, user_input, user_language, bot_response, data_source, response_time
                FROM conversations 
                ORDER BY timestamp DESC 
                LIMIT 20
            """)
            conversations = cursor.fetchall()
            
            rtx_status = "🚀 RTX 4090" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else "📱 Standard GPU"
            
            self.analytics_display.config(state=tk.NORMAL)
            self.analytics_display.delete(1.0, tk.END)
            
            self.analytics_display.insert(tk.END, f"=== {rtx_status} CONVERSATION HISTORY (Last 20) ===\n\n")
            
            for i, conv in enumerate(conversations, 1):
                if len(conv) >= 6:
                    timestamp, user_input, user_lang, bot_response, data_source, response_time = conv
                    lang_flag = "🇻🇳" if user_lang == 'vi' else "🇺🇸" if user_lang == 'en' else "🌐"
                    
                    performance_icon = "🚀" if response_time < 2.0 else "⚡" if response_time < 5.0 else "🐌"
                    
                    self.analytics_display.insert(tk.END, f"{i}. [{timestamp}] {lang_flag} {performance_icon} ({data_source}, {response_time:.2f}s)\n")
                    self.analytics_display.insert(tk.END, f"User: {user_input}\n")
                    self.analytics_display.insert(tk.END, f"Bot: {bot_response[:200]}...\n\n")
            
            self.analytics_display.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error viewing RTX 4090 conversation history: {e}")
    
    def run(self):
        """Start the application with comprehensive error handling"""
        try:
            print("📦 Adding sample products...")
            self.add_sample_products()
            
            print("📋 Loading existing products to tabs...")
            self.load_existing_products_to_tabs()
            
            self.display_welcome_message()
            
            print("🚀 Starting GUI...")
            self.root.mainloop()
            
        except Exception as e:
            print(f"❌ Error starting RTX 4090 application: {e}")
            self.logger.error(f"RTX 4090 application start error: {e}")
    
    def display_welcome_message(self):
        """Display welcome message with RTX 4090 system status"""
        try:
            rtx_status = "🚀 RTX 4090 Optimized" if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else "📱 Standard GPU"
            
            if self.current_language == 'vi':
                welcome_msg = f"""🤖 Chào mừng đến với Trợ lý AI Bán hàng {rtx_status}!

            Trạng thái hệ thống:
            ✅ Hỗ trợ tiếng Việt và tiếng Anh
            {rtx_status} {'RTX 4090 với 24GB VRAM' if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else 'GPU tiêu chuẩn'}
            {'✅' if self.text_generator else '❌'} Mô hình AI: {'RTX 4090 Optimized' if self.text_generator else 'Không khả dụng'}
            {'✅' if self.embedding_model else '❌'} Tìm kiếm thông minh: {'Batch Processing RTX 4090' if self.embedding_model else 'Không khả dụng'}

            Tính năng RTX 4090:
            📦 Tìm kiếm sản phẩm siêu nhanh với batch processing
            💬 Ngữ cảnh cuộc trò chuyện được tối ưu
            🌐 Hỗ trợ đa ngôn ngữ với AI mạnh mẽ
            📚 Xử lý file hàng loạt với RTX 4090
            🖼️ OCR tốc độ cao
            🚀 Response caching cho hiệu suất tối đa

            Hãy thử đặt câu hỏi như:
            • "Tôi muốn RTX 4090 cho AI development"
            • "Bạn có máy trạm RTX 4090 nào không?"
            • "So sánh các hệ thống RTX 4090"
            • "Giá RTX 4090 gaming rig là bao nhiêu?"

            📝 Sử dụng bàn phím để nhập tin nhắn
            {f'🔧 Một số tính năng AI có thể bị hạn chế do lỗi tải mô hình' if not self.text_generator else '✅ Tất cả tính năng RTX 4090 hoạt động hoàn hảo'}"""
            else:
                welcome_msg = f"""🤖 Welcome to the {rtx_status} AI Sales Assistant!

            System Status:
            ✅ Vietnamese and English support
            {rtx_status} {'RTX 4090 with 24GB VRAM' if hasattr(self, 'is_rtx4090') and self.is_rtx4090 else 'Standard GPU'}
            {'✅' if self.text_generator else '❌'} AI Model: {'RTX 4090 Optimized' if self.text_generator else 'Not available'}
            {'✅' if self.embedding_model else '❌'} Smart Search: {'Batch Processing RTX 4090' if self.embedding_model else 'Not available'}

            RTX 4090 Features:
            📦 Ultra-fast product search with batch processing
            💬 Optimized conversation context
            🌐 Multi-language support with powerful AI
            📚 Batch file processing with RTX 4090
            🖼️ High-speed OCR
            🚀 Response caching for maximum performance

            Try asking questions like:
            • "I need RTX 4090 for AI development"
            • "What RTX 4090 workstations do you have?"
            • "Compare RTX 4090 systems"
            • "How much does RTX 4090 gaming rig cost?"

            📝 Use keyboard to type messages
            {f'🔧 Some AI features may be limited due to model loading errors' if not self.text_generator else '✅ All RTX 4090 features working perfectly'}"""

            self.display_message("RTX 4090 System", welcome_msg, "system")
        except Exception as e:
            print(f"Welcome message error: {e}")
            
    def add_sample_products(self):
        """Add sample SSD, Memory, and Motherboard products from Excel template"""
        sample_products = [
            # SSD Products - matching Database tab fields exactly
            {
                'product': 'SSD (SOLID STATE DRIVE)',     # Sản phẩm
                'interface': 'M2 NVMe',                   # Giao thức
                'model': 'E130 256GB',                    # Model
                'specifications': '3200M/s Read 2700MB Write',  # Thông số
                'stock_status': 'Còn hàng',               # Kho
                'price': '600000',                        # Giá bán lẻ
                'category': 'SSD'
            },
            {
                'product': 'SSD (SOLID STATE DRIVE)',
                'interface': 'M2 NVME',
                'model': 'E130 512GB',
                'specifications': 'Speed 3500M/s Read 3200MB Write',
                'stock_status': 'Còn hàng',
                'price': '900000',
                'category': 'SSD'
            },
            # Memory Products - matching Database tab fields
            {
                'product': 'Desktop Memory',
                'interface': 'DDR4 UDIMM',
                'model': 'U3200I-C22 16GB',
                'specifications': '3200Mhz 1.2V (Jedec), CAS Latency 22 chỉ tương thích với CPU Intel',
                'stock_status': 'còn hàng',
                'price': '600000',
                'category': 'Memory'
            },
            {
                'product': 'Laptop Memory',
                'interface': 'DDR4 SODIM',
                'model': 'S3200I-C22 8GB',
                'specifications': '3200Mhz 1.2V (Jedec) CAS Latency 22 chỉ tương thích với CPU Intel',
                'stock_status': 'Còn hàng',
                'price': '400000',
                'category': 'Memory'
            },
            {
                'product': 'Desktop Memory',
                'interface': 'DDR5 UDIMM',
                'model': 'U5600-C46 16GB',
                'specifications': '5600Mh 1.1V (Jedec) CAS Latency 46 tương thích với CPU Intel và AMD',
                'stock_status': 'còn hàng',
                'price': '1300000',
                'category': 'Memory'
            },
            # Motherboard Products - note: uses 'chipset' instead of 'interface'
            {
                'product': 'Intel Motherboard',
                'chipset': 'H610',                # Note: Motherboard uses 'chipset' not 'interface'
                'model': 'H610M-HDV',
                'specifications': 'Hỗ trợ CPU Intel thế hệ 12 13 14, sử dụng DDR4 cổng xuất hình: HDMI, Display Port, VGA. 2 USB 3.2, 2 USB 2.0',
                'stock_status': 'Còn hàng',
                'price': '1650000',
                'category': 'Motherboard'
            },
            {
                'product': 'Intel Motherboard',
                'chipset': 'B760',
                'model': 'B760M-HDV',
                'specifications': 'Hỗ trợ CPU Intel thế hệ 12 13 14, sử dụng DDR4 cổng xuất hình: HDMI, Display Port, VGA. 4 USB 3.2, 2 USB 2.0, Có đèn led chẩn đoán lỗi.',
                'stock_status': 'Còn hàng',
                'price': '1890000',
                'category': 'Motherboard'
            }
        ]
        
        try:
            cursor = self.conn.cursor()
            for product in sample_products:
                try:
                    # Build text for embedding - using exact fields
                    combined_text = f"{product.get('product', '')} {product.get('model', '')} {product.get('specifications', '')}"
                    
                    embedding_blob = None
                    if self.embedding_model:
                        try:
                            embedding = self.embedding_model.encode([combined_text])[0]
                            embedding_blob = embedding.astype(np.float32).tobytes()
                        except Exception as e:
                            self.logger.warning(f"Embedding generation error for {product.get('model', '')}: {e}")
                    
                    # Parse price
                    try:
                        price_num = float(str(product.get('price', '0')).replace(',', '').replace(' ', ''))
                    except:
                        price_num = 0
                    
                    # Insert using the same structure as add_product method
                    cursor.execute('''
                        INSERT OR IGNORE INTO products 
                        (name, description, category, price, features, specifications, availability,
                         source_file, embedding, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        product.get('model', ''),      # model -> name
                        product.get('product', ''),    # product -> description
                        product.get('category', ''),
                        price_num,
                        product.get('interface', product.get('chipset', '')),  # interface/chipset -> features
                        product.get('specifications', ''),
                        product.get('stock_status', 'Available'),
                        'Sample Data',
                        embedding_blob,
                        datetime.now().isoformat()
                    ))
                    
                except Exception as e:
                    self.logger.error(f"Error inserting sample product {product.get('model', 'unknown')}: {e}")
                    continue
            
            self.conn.commit()
            print("✅ Sample products added successfully with exact Database tab field structure")
            
        except Exception as e:
            self.logger.error(f"Error adding sample products: {e}")
            print(f"❌ Error adding sample products: {e}")  
    
    # ADD THESE 4 NEW METHODS HERE:
    def get_vietnamese_addressing(self, age, gender):
        """Determine Vietnamese addressing based on age and gender"""
        try:
            age = int(age)
            gender = str(gender).lower()
            
            if age < 25:
                return {
                    'customer_title': 'bạn',
                    'bot_title': 'tôi',
                    'addressing_style': 'casual'
                }
            elif 25 <= age <= 45:
                if gender == 'male':
                    return {
                        'customer_title': 'anh',
                        'bot_title': 'em',
                        'addressing_style': 'respectful_peer'
                    }
                else:
                    return {
                        'customer_title': 'chị',
                        'bot_title': 'em',
                        'addressing_style': 'respectful_peer'
                    }
            else:
                if gender == 'male':
                    return {
                        'customer_title': 'chú',
                        'bot_title': 'cháu',
                        'addressing_style': 'respectful_elder'
                    }
                else:
                    return {
                        'customer_title': 'cô',
                        'bot_title': 'cháu',
                        'addressing_style': 'respectful_elder'
                    }
        except:
            return {
                'customer_title': 'bạn',
                'bot_title': 'tôi',
                'addressing_style': 'casual'
            }

    def customize_response_for_addressing(self, response, customer_profile):
        """Customize response based on Vietnamese addressing conventions"""
        try:
            if not customer_profile:
                return response
            
            addressing = customer_profile.get('addressing', {})
            customer_title = addressing.get('customer_title', 'bạn')
            bot_title = addressing.get('bot_title', 'tôi')
            
            response = response.replace('bạn', customer_title)
            response = response.replace('tôi', bot_title)
            
            addressing_style = addressing.get('addressing_style', 'casual')
            if addressing_style == 'respectful_elder':
                if not response.endswith(' ạ.') and not response.endswith(' ạ!'):
                    response = response.rstrip('.!') + ' ạ.'
            
            return response
        except Exception as e:
            self.logger.error(f"Error customizing response: {e}")
            return response

    def is_profile_setup_message(self, user_input):
        """Check if message is for profile setup"""
        user_input_lower = user_input.lower()
        profile_keywords = ['tuổi', 'tuoi', 'nam', 'nữ', 'nu', 'male', 'female']
        age_pattern = r'\d{1,2}'
        
        has_age = bool(re.search(age_pattern, user_input))
        has_gender = any(keyword in user_input_lower for keyword in profile_keywords)
        
        return has_age and has_gender

    def handle_profile_setup(self, user_input, session_id):
        """Handle customer profile setup"""
        try:
            age_match = re.search(r'(\d{1,2})', user_input)
            age = int(age_match.group(1)) if age_match else 25
            
            user_input_lower = user_input.lower()
            if any(word in user_input_lower for word in ['nam', 'male']):
                gender = 'male'
            elif any(word in user_input_lower for word in ['nữ', 'nu', 'female']):
                gender = 'female'
            else:
                gender = 'male'
            
            addressing = self.get_vietnamese_addressing(age, gender)
            
            self.customer_profiles[session_id] = {
                'age': age,
                'gender': gender,
                'addressing': addressing,
                'created_at': time.time()
            }
            
            customer_title = addressing['customer_title']
            bot_title = addressing['bot_title']
            
            if addressing['addressing_style'] == 'casual':
                greeting = f"Xin chào bạn! Tôi là trợ lý AI bán hàng với sức mạnh RTX 4090. Bạn đang quan tâm đến sản phẩm gì?"
            elif addressing['addressing_style'] == 'respectful_peer':
                greeting = f"Xin chào {customer_title}! {bot_title.capitalize()} là trợ lý AI bán hàng RTX 4090. {customer_title.capitalize()} quan tâm đến sản phẩm nào ạ?"
            else:
                greeting = f"Kính chào {customer_title}! {bot_title.capitalize()} là trợ lý AI RTX 4090. {customer_title.capitalize()} muốn tìm hiểu sản phẩm gì ạ?"
            
            return {
                'response': greeting,
                'data_source': 'profile_created',
                'processing_time': 0.1,
                'user_language': 'vi',
                'response_language': 'vi',
                'customer_profile': self.customer_profiles[session_id],
                'profile_setup_complete': True
            }
            
        except Exception as e:
            self.logger.error(f"Profile setup error: {e}")
            return {
                'response': 'Xin lỗi, vui lòng thử lại với định dạng "tuổi + giới tính" (ví dụ: "25 tuổi nam") ạ.',
                'data_source': 'profile_error',
                'processing_time': 0.1,
                'user_language': 'vi',
                'response_language': 'vi'
            }  
    def check_database_contents(self):
        """Check what's actually in the database"""
        try:
            cursor = self.conn.cursor()
            
            # Check if products table exists
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='products'")
            if not cursor.fetchone():
                print("❌ Products table doesn't exist!")
                return
                
            # Get all products
            cursor.execute("SELECT * FROM products")
            products = cursor.fetchall()
            print(f"\n📊 Total products in database: {len(products)}")
            
            # Show column names
            cursor.execute("PRAGMA table_info(products)")
            columns = cursor.fetchall()
            print("\nDatabase columns:")
            for col in columns:
                print(f"  - {col[1]} ({col[2]})")
            
            # Show products by category
            cursor.execute("SELECT category, COUNT(*) FROM products GROUP BY category")
            categories = cursor.fetchall()
            print("\nProducts by category:")
            for cat, count in categories:
                print(f"  - {cat}: {count} products")
                
            # Show first 5 products
            cursor.execute("SELECT name, category, created_at FROM products LIMIT 5")
            sample = cursor.fetchall()
            print("\nFirst 5 products:")
            for name, cat, created in sample:
                print(f"  - {name} ({cat}) - Created: {created}")
                
        except Exception as e:
            print(f"❌ Error checking database: {e}")
        
    def __del__(self):
        """Cleanup with RTX 4090 optimization"""
        try:
            if hasattr(self, 'conn'):
                self.conn.close()
            if hasattr(self, 'response_cache'):
                self.response_cache.clear()
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
        except:
            pass
            
    def view_database_contents(self, category=None):
        """View products in the SQLite database, optionally filtered by category"""
        try:
            cursor = self.conn.cursor()
            
            # Auto-detect category from current tab if not specified
            if not category:
                try:
                    current_tab_index = self.db_sub_notebook.index('current')
                    current_tab_text = self.db_sub_notebook.tab(current_tab_index, 'text')
                    
                    # Extract category from tab text
                    if 'SSD' in current_tab_text:
                        category = 'SSD'
                    elif 'Memory' in current_tab_text:
                        category = 'Memory'
                    elif 'Motherboard' in current_tab_text:
                        category = 'Motherboard'
                    else:
                        print(f"Unknown tab: {current_tab_text}")
                        return
                except Exception as e:
                    print(f"Error detecting current tab: {e}")
                    return
            
            # Query products for the specific category
            cursor.execute("""
                SELECT id, name, description, category, category_vietnamese, 
                       price, features, specifications, availability, 
                       source_file, created_at
                FROM products 
                WHERE category = ? 
                ORDER BY created_at DESC
            """, (category,))
            products = cursor.fetchall()
            
            print(f"\n=== {category} Products in Database ===")
            print(f"Total {category} products: {len(products)}")
            
            # Update the appropriate TreeView based on category
            if category == 'SSD' and hasattr(self, 'ssd_tree'):
                self.populate_tree_from_database(self.ssd_tree, products, category)
                
            elif category == 'Memory' and hasattr(self, 'memory_tree'):
                self.populate_tree_from_database(self.memory_tree, products, category)
                
            elif category == 'Motherboard' and hasattr(self, 'mb_tree'):
                self.populate_tree_from_database(self.mb_tree, products, category)
            
            # Show success message
            self.show_message("Success", f"Loaded {len(products)} {category} products from database")
            
        except Exception as e:
            print(f"Error viewing database: {e}")
            self.show_message("Error", f"Error viewing database: {str(e)}")
            
    def populate_tree_from_database(self, tree_widget, products, category):
        """Helper method to populate a TreeView with products from database"""
        try:
            # Clear existing items
            for item in tree_widget.get_children():
                tree_widget.delete(item)
            
            # Add products to TreeView
            for i, product in enumerate(products, 1):
                # Extract values based on database structure
                # product = (id, name, description, category, category_vietnamese, 
                #           price, features, specifications, availability, source_file, created_at)
                
                stt = i
                product_name = product[2] or product[1] or 'N/A'  # description or name
                interface_or_chipset = product[6] or 'N/A'        # features field
                model = product[1] or 'N/A'                       # name field
                specifications = product[7] or 'N/A'              # specifications
                stock_status = product[8] or 'Available'          # availability
                price = f"{product[5] or 0:,.0f}" if product[5] else "0"
                
                # Insert into TreeView
                tree_widget.insert('', 'end', values=(
                    stt,
                    product_name,
                    interface_or_chipset,
                    model,
                    specifications,
                    stock_status,
                    price
                ))
            
            print(f"✅ Populated {category} TreeView with {len(products)} products")
            
        except Exception as e:
            print(f"Error populating TreeView: {e}")
            
    def remove_selected_item(self):
        """Remove selected item(s) from current category's TreeView and database"""
        try:
            # Determine current category tab
            current_tab_index = self.db_sub_notebook.index('current')
            current_tab_text = self.db_sub_notebook.tab(current_tab_index, 'text')
            
            # Get the appropriate tree widget
            tree_widget = None
            category = None
            
            if 'SSD' in current_tab_text:
                tree_widget = self.ssd_tree
                category = 'SSD'
            elif 'Memory' in current_tab_text:
                tree_widget = self.memory_tree
                category = 'Memory'
            elif 'Motherboard' in current_tab_text:
                tree_widget = self.mb_tree
                category = 'Motherboard'
            else:
                self.show_message("Error", "Unknown product category")
                return
            
            # Get selected items (can be multiple)
            selected_items = tree_widget.selection()
            if not selected_items:
                self.show_message("Warning", "Please select item(s) to remove")
                return
            
            # Collect items to delete
            items_to_delete = []
            for item in selected_items:
                item_values = tree_widget.item(item, 'values')
                if item_values and len(item_values) >= 4:
                    items_to_delete.append({
                        'tree_item': item,
                        'product': item_values[1],
                        'model': item_values[3]
                    })
            
            if not items_to_delete:
                self.show_message("Error", "No valid items selected")
                return
            
            # Confirm deletion
            item_count = len(items_to_delete)
            message = f"Are you sure you want to delete {item_count} item(s)?\n\n"
            if item_count <= 5:
                for item in items_to_delete:
                    message += f"• {item['model']} - {item['product']}\n"
            message += f"\nCategory: {category}\nThis action cannot be undone!"
            
            result = messagebox.askyesno("Confirm Deletion", message)
            
            if result:
                # Delete from database
                cursor = self.conn.cursor()
                total_deleted = 0
                
                for item in items_to_delete:
                    cursor.execute("""
                        DELETE FROM products 
                        WHERE (name = ? OR description = ?) AND category = ?
                    """, (item['model'], item['product'], category))
                    total_deleted += cursor.rowcount
                    
                    # Remove from TreeView
                    tree_widget.delete(item['tree_item'])
                
                self.conn.commit()
                
                if total_deleted > 0:
                    self.show_message("Success", f"Removed {total_deleted} product(s) successfully")
                    # Refresh to update row numbers
                    self.view_database_contents(category)
                else:
                    self.show_message("Warning", "No products were removed from database")
                    
        except Exception as e:
            print(f"Error removing items: {e}")
            self.show_message("Error", f"Error removing items: {str(e)}")  
            
class ExcelImportDialog:
    def __init__(self, parent, database_manager):
        self.parent = parent
        self.db_manager = database_manager
        self.setup_ui()
        
    def setup_ui(self):
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("Import Excel Data")
        self.dialog.geometry("600x400")
        self.dialog.transient(self.parent)
        self.dialog.grab_set()
        
        # File selection frame
        file_frame = ttk.LabelFrame(self.dialog, text="1. Select Excel File")
        file_frame.pack(fill='x', padx=10, pady=5)
        
        self.file_path_var = tk.StringVar()
        ttk.Entry(file_frame, textvariable=self.file_path_var, 
                 state='readonly').pack(side='left', fill='x', expand=True, padx=5)
        ttk.Button(file_frame, text="Browse", 
                  command=self.browse_file).pack(side='right', padx=5)
        
        # Category selection frame
        category_frame = ttk.LabelFrame(self.dialog, text="2. Select Product Category")
        category_frame.pack(fill='x', padx=10, pady=5)
        
        self.category_var = tk.StringVar(value="SSD")
        for category in ["SSD", "Memory", "Motherboard"]:
            ttk.Radiobutton(category_frame, text=category, 
                           variable=self.category_var, 
                           value=category).pack(side='left', padx=10)
        
        # Progress frame
        progress_frame = ttk.LabelFrame(self.dialog, text="Import Progress")
        progress_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        self.progress_bar = ttk.Progressbar(progress_frame, mode='determinate')
        self.progress_bar.pack(fill='x', padx=10, pady=10)
        
        self.status_label = ttk.Label(progress_frame, text="Ready to import")
        self.status_label.pack(pady=5)
        
        self.details_text = scrolledtext.ScrolledText(progress_frame, height=8, width=60)
        self.details_text.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Action buttons
        button_frame = ttk.Frame(self.dialog)
        button_frame.pack(fill='x', padx=10, pady=10)
        
        self.import_btn = ttk.Button(button_frame, text="Import", 
                                    command=self.start_import)
        self.import_btn.pack(side='left', padx=5)
        
        self.cancel_btn = ttk.Button(button_frame, text="Cancel", 
                                    command=self.cancel_import)
        self.cancel_btn.pack(side='left', padx=5)
        
        ttk.Button(button_frame, text="Close", 
                  command=self.dialog.destroy).pack(side='right', padx=5)
        
        self.cancel_event = threading.Event()
        
    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="Select Excel File",
            filetypes=[
                ("Excel files", "*.xlsx *.xls *.xlsm"),
                ("CSV files", "*.csv"),
                ("All files", "*.*")
            ]
        )
        if filename:
            self.file_path_var.set(filename)
            self.validate_file(filename)
    
    def validate_file(self, filepath):
        """Quick validation of Excel file"""
        try:
            reader = SmartExcelReader()
            if filepath.endswith('.csv'):
                df_sample = pd.read_csv(filepath, nrows=5)
            else:
                df_sample = pd.read_excel(filepath, nrows=5)
            
            self.details_text.delete(1.0, tk.END)
            self.details_text.insert(tk.END, f"File: {Path(filepath).name}\n")
            self.details_text.insert(tk.END, f"Columns found: {', '.join(df_sample.columns)}\n")
            self.details_text.insert(tk.END, f"Sample rows: {len(df_sample)}\n")
            
        except Exception as e:
            messagebox.showerror("File Error", f"Cannot read file: {str(e)}")
    
    def start_import(self):
        if not self.file_path_var.get():
            messagebox.showwarning("No File", "Please select a file first")
            return
        
        self.import_btn.config(state='disabled')
        self.cancel_event.clear()
        
        import_thread = threading.Thread(
            target=self._import_worker,
            args=(self.file_path_var.get(), self.category_var.get())
        )
        import_thread.daemon = True
        import_thread.start()
    
    def _import_worker(self, filepath, category):
        """Background worker for import process"""
        try:
            # Update UI
            self.dialog.after(0, lambda: self.status_label.config(text="Reading Excel file..."))
            self.dialog.after(0, lambda: self.progress_bar.config(value=10))
            
            # Read Excel file
            reader = SmartExcelReader()
            df = reader.read_file_optimized(filepath)
            total_rows = len(df)
            
            if self.cancel_event.is_set():
                return
            
            # Process data
            self.dialog.after(0, lambda: self.status_label.config(text=f"Processing {total_rows} rows..."))
            self.dialog.after(0, lambda: self.progress_bar.config(value=30))
            
            # Validate and map columns
            processor = DataProcessor(category)
            processed_df = processor.process_excel_data(df)
            
            if self.cancel_event.is_set():
                return
            
            # Insert into database
            self.dialog.after(0, lambda: self.status_label.config(text="Inserting into database..."))
            self.dialog.after(0, lambda: self.progress_bar.config(value=60))
            
            result = self.db_manager.bulk_insert_products(category, processed_df)
            
            # Update TreeView
            self.dialog.after(0, lambda: self._update_treeview(category, result))
            
            # Complete
            self.dialog.after(0, lambda: self.progress_bar.config(value=100))
            self.dialog.after(0, lambda: self._import_complete(result))
            
        except Exception as e:
            self.dialog.after(0, lambda: self._import_error(str(e)))
    
    def _update_treeview(self, category, result):
        """Update the main application's TreeView"""
        if hasattr(self.db_manager, 'refresh_category_view'):
            self.db_manager.refresh_category_view(category)
    
    def _import_complete(self, result):
        """Handle successful import completion"""
        self.import_btn.config(state='normal')
        self.status_label.config(text="Import completed successfully!")
        
        summary = f"\n\nIMPORT SUMMARY:\n"
        summary += f"Total records: {result.get('total_records', 0)}\n"
        summary += f"Successfully imported: {result.get('successful', 0)}\n"
        summary += f"Skipped (duplicates): {result.get('duplicates', 0)}\n"
        summary += f"Errors: {result.get('errors', 0)}\n"
        
        self.details_text.insert(tk.END, summary)
        
        messagebox.showinfo("Import Complete", 
                           f"Successfully imported {result.get('successful', 0)} records!")
    
    def _import_error(self, error_message):
        """Handle import errors"""
        self.import_btn.config(state='normal')
        self.status_label.config(text="Import failed!")
        self.details_text.insert(tk.END, f"\nERROR: {error_message}\n")
        messagebox.showerror("Import Error", f"Import failed: {error_message}")
    
    def cancel_import(self):
        self.cancel_event.set()
        self.import_btn.config(state='normal')
        self.status_label.config(text="Import cancelled")
 

# FIXED WEB CHATBOT FUNCTION FOR RTX 4090 OPTIMIZATION
def create_web_app(chatbot_instance):
    """Create Flask web application optimized for RTX 4090 with FIXED message sending"""
    app = Flask(__name__)
    CORS(app, 
     origins="*", 
     methods=["GET", "POST", "OPTIONS"], 
     allow_headers=["Content-Type", "Accept"],
     supports_credentials=False)
    
    @app.route('/')
    def index():
        """Serve the enhanced web interface with RTX 4090 optimization and FIXED chat functionality"""
        rtx_status = "🚀 SSTC Powered" if hasattr(chatbot_instance, 'is_rtx4090') and chatbot_instance.is_rtx4090 else "📱 Standard GPU"
        
        return f'''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{rtx_status} Trợ lý bán hàng siêu thông minh </title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
        }}
        .chat-container {{
            width: 90%;
            max-width: 1400px;
            height: 90vh;
            background: rgba(255, 255, 255, 0.95);
            border-radius: 20px;
            display: flex;
            overflow: hidden;
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.1);
        }}
        .chat-area {{
            flex: 1;
            display: flex;
            flex-direction: column;
            background: rgba(240, 242, 247, 0.8);
        }}
        .header {{
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            padding: 20px;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }}
        .header-title {{
            display: flex;
            align-items: center;
            gap: 10px;
        }}
        .ai-icon {{
            width: 40px;
            height: 40px;
            background: rgba(255, 255, 255, 0.2);
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
        }}
        .language-selector {{
            background: rgba(255, 255, 255, 0.1);
            border: 1px solid rgba(255, 255, 255, 0.3);
            border-radius: 10px;
            padding: 8px 12px;
            color: white;
            font-size: 14px;
        }}
        .messages {{
            flex: 1;
            overflow-y: auto;
            padding: 20px;
            display: flex;
            flex-direction: column;
            gap: 15px;
        }}
        .message {{
            max-width: 80%;
            padding: 15px 20px;
            border-radius: 20px;
            word-wrap: break-word;
            animation: slideIn 0.3s ease-out;
        }}
        .user {{
            align-self: flex-end;
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            border-bottom-right-radius: 5px;
        }}
        .bot {{
            align-self: flex-start;
            background: white;
            color: #333;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
            border-bottom-left-radius: 5px;
            border: 1px solid #e0e0e0;
        }}
        .system {{
            align-self: center;
            background: linear-gradient(135deg, #00b894, #00cec9);
            color: white;
            font-size: 14px;
            max-width: 90%;
        }}
        .typing-indicator {{
            display: none;
            align-items: center;
            gap: 5px;
            padding: 15px 20px;
            background: white;
            border-radius: 20px;
            border-bottom-left-radius: 5px;
            max-width: 80px;
            margin-bottom: 10px;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
        }}
        .typing-dot {{
            width: 8px;
            height: 8px;
            background: #6c5ce7;
            border-radius: 50%;
            animation: typing 1.5s infinite;
        }}
        .typing-dot:nth-child(2) {{ animation-delay: 0.2s; }}
        .typing-dot:nth-child(3) {{ animation-delay: 0.4s; }}
        .input-area {{
            padding: 20px;
            background: white;
            display: flex;
            gap: 15px;
            border-top: 1px solid #e0e0e0;
        }}
        .input {{
            flex: 1;
            padding: 15px 20px;
            border: 2px solid #e0e0e0;
            border-radius: 25px;
            font-size: 16px;
            outline: none;
            resize: none;
            min-height: 50px;
            max-height: 120px;
        }}
        .input:focus {{ border-color: #6c5ce7; }}
        .send-btn {{
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            border: none;
            border-radius: 50%;
            width: 50px;
            height: 50px;
            cursor: pointer;
            font-size: 20px;
            transition: transform 0.2s ease;
        }}
        .send-btn:hover {{ transform: scale(1.05); }}
        .send-btn:disabled {{ opacity: 0.6; cursor: not-allowed; }}
        .avatar-panel {{
            width: 480px;
            background: white;
            padding: 20px;
            display: flex;
            flex-direction: column;
            align-items: center;
            border-left: 1px solid #e0e0e0;
        }}
        .avatar-container {{
            width: 100%;
            height: 550px;
            border-radius: 15px;
            background: linear-gradient(135deg, #f8f9fa, #e9ecef);
            display: flex;
            align-items: center;
            justify-content: center;
            color: #666;
            margin-bottom: 20px;
            overflow: hidden;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
            border: 3px solid #6c5ce7;
            position: relative;
        }}
        .avatar-container img {{
            width: 100%;
            height: 100%;
            object-fit: cover;
            border-radius: 12px;
        }}
        .status-indicator {{
            position: absolute;
            top: 15px;
            right: 15px;
            width: 20px;
            height: 20px;
            background: #00b894;
            border-radius: 50%;
            border: 3px solid white;
            animation: pulse-indicator 1.5s infinite;
        }}
        .controls {{
            display: flex;
            gap: 10px;
            margin-bottom: 20px;
        }}
        .btn {{
            background: linear-gradient(135deg, #6c5ce7, #a29bfe);
            color: white;
            border: none;
            border-radius: 20px;
            padding: 10px 20px;
            cursor: pointer;
            font-size: 14px;
            transition: transform 0.2s ease;
        }}
        .btn:hover {{ transform: translateY(-1px); }}
        .info-panel {{
            text-align: center;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 10px;
            margin-top: 10px;
        }}
        .assistant-name {{
            font-size: 20px;
            font-weight: bold;
            color: #2d3436;
            margin-bottom: 5px;
        }}
        .assistant-subtitle {{
            font-size: 14px;
            color: #636e72;
            margin-bottom: 10px;
        }}
        .rtx-badge {{
            display: inline-block;
            background: linear-gradient(135deg, #00b894, #00cec9);
            color: white;
            padding: 4px 8px;
            border-radius: 12px;
            font-size: 12px;
            font-weight: bold;
            margin-left: 10px;
        }}
        .error-message {{
            background: #ff7675;
            color: white;
            padding: 10px;
            border-radius: 10px;
            margin: 10px;
            text-align: center;
        }}
        @keyframes slideIn {{
            from {{ opacity: 0; transform: translateY(20px); }}
            to {{ opacity: 1; transform: translateY(0); }}
        }}
        @keyframes typing {{
            0%, 60%, 100% {{ transform: translateY(0); }}
            30% {{ transform: translateY(-10px); }}
        }}
        @keyframes pulse-indicator {{
            0%, 100% {{ opacity: 1; }}
            50% {{ opacity: 0.5; }}
        }}
        @media (max-width: 768px) {{
            .chat-container {{ width: 95%; height: 95vh; flex-direction: column; }}
            .avatar-panel {{ width: 100%; height: 300px; }}
            .avatar-container {{ height: 250px; }}
        }}
    </style>
</head>
<body>
    <div class="chat-container">
        <div class="chat-area">
            <div class="header">
                <div class="header-title">
                    <div class="ai-icon">🤖</div>
                    <div>
                        <h2>{rtx_status} SSTC Super Sales Consultant</h2>
                        <p style="font-size: 14px; opacity: 0.9;">Siêu trợ lý SSTC</p>
                    </div>
                </div>
                <select class="language-selector" id="languageSelector">
                    <option value="vi">🇻🇳 Tiếng Việt</option>
                    <option value="en">🇺🇸 English</option>
                </select>
            </div>
            <div class="messages" id="messages">
                <div class="message system">
                    <strong>🎉 Chào mừng bạn đến với SSTC AI STORE! 🎉</strong><br><br>
                    🛍️ Chào mừng bạn đến với cửa hàng công nghệ thông minh của chúng tôi!<br><br>
                    <strong>🎉 Welcome to SSTC AI STORE! 🎉</strong><br>
                    🛍️ Welcome to our smart technology store!
                </div>
            </div>
            <div class="typing-indicator" id="typingIndicator">
                <div class="typing-dot"></div>
                <div class="typing-dot"></div>
                <div class="typing-dot"></div>
            </div>
            <div class="input-area">
                <textarea class="input" id="messageInput" placeholder="Nhập tin nhắn của bạn / Type your message..." rows="1"></textarea>
                <button class="send-btn" id="sendButton" onclick="sendMessage()">➤</button>
            </div>
        </div>
        <div class="avatar-panel">
            <div class="avatar-container" id="avatarContainer">
                <div class="status-indicator"></div>
                🎨 Generating SSTC Super Sales Consultant Photo...
            </div>
            <div class="controls">
                <button class="btn" onclick="generateFullBodyPhoto()">🔄 New Photo</button>
                <button class="btn" onclick="generateDifferentStyle()">✨ New Style</button>
            </div>
            <div class="info-panel">
                <div class="assistant-name">🌟 {rtx_status} Sales Consultant</div>
                <div class="assistant-subtitle">
                    Professional AI Sales Consultant
                    <span class="rtx-badge">SSTC</span>
                </div>
                <small style="color: #666;">Siêu đẹp, siêu nhanh và siêu chuẩn</small>
            </div>
        </div>
    </div>

<script>
        let currentLanguage = 'vi';
        let isTyping = false;

        function generateFullBodyPhoto() {{
            const container = document.getElementById('avatarContainer');
            container.innerHTML = '<div class="status-indicator"></div>🚀 Generating SSTC Super Sales Consultant Photo...';
            
            const randomSeed = Math.floor(Math.random() * 50000);
            const imageUrl = 'https://image.pollinations.ai/prompt/full%20body%20portrait%20professional%20businesswoman%20AI%20consultant?width=450&height=700&seed=' + randomSeed;
            
            loadImage(imageUrl, container);
        }}

        function generateDifferentStyle() {{
            const container = document.getElementById('avatarContainer');
            container.innerHTML = '<div class="status-indicator"></div>✨ Generating Different AI Style...';
            
            const randomSeed = Math.floor(Math.random() * 50000);
            const imageUrl = 'https://image.pollinations.ai/prompt/beautiful%20professional%20AI%20businesswoman%20full%20body?width=450&height=700&seed=' + randomSeed;
            
            loadImage(imageUrl, container);
        }}

        function loadImage(imageUrl, container) {{
            const img = document.createElement('img');
            
            img.onload = function() {{
                container.innerHTML = '<div class="status-indicator"></div>';
                container.appendChild(img);
                console.log('✅ SSTC Super Sales Consultant photo generated successfully');
            }};
            
            img.onerror = function() {{
                console.log('❌ Image failed, trying backup...');
                container.innerHTML = '<div class="status-indicator"></div> SSTC Super Sales Consultant<br><small>Click "New Photo" to try again</small>';
            }};
            
            img.src = imageUrl;
        }}

        function addMessage(text, sender) {{
            const messages = document.getElementById('messages');
            const div = document.createElement('div');
            div.className = 'message ' + sender;
            div.innerHTML = text;
            messages.appendChild(div);
            messages.scrollTop = messages.scrollHeight;
        }}

        function showTyping() {{
            if (isTyping) return;
            isTyping = true;
            document.getElementById('typingIndicator').style.display = 'flex';
            document.getElementById('sendButton').disabled = true;
        }}

        function hideTyping() {{
            isTyping = false;
            document.getElementById('typingIndicator').style.display = 'none';
            document.getElementById('sendButton').disabled = false;
        }}

        async function sendMessage() {{
            const input = document.getElementById('messageInput');
            const message = input.value.trim();
            if (!message || isTyping) return;

            console.log('🚀 Sending message:', message);
            addMessage(message, 'user');
            input.value = '';
            autoResize(input);

            showTyping();

            try {{
                // Get or create session ID
                let sessionId = localStorage.getItem('chatbot_session_id');
                if (!sessionId) {{
                    sessionId = 'web_' + Date.now() + '_' + Math.random().toString(36).substr(2, 9);
                    localStorage.setItem('chatbot_session_id', sessionId);
                }}

                const requestData = {{
                    message: message,
                    language: currentLanguage,
                    session_id: sessionId
                }};

                const response = await fetch('/api/chat', {{
                    method: 'POST',
                    headers: {{
                        'Content-Type': 'application/json',
                        'Accept': 'application/json'
                    }},
                    body: JSON.stringify(requestData)
                }});

                if (!response.ok) {{
                    throw new Error('HTTP ' + response.status);
                }}

                const data = await response.json();
                hideTyping();

                if (data.error) {{
                    addMessage('❌ Error: ' + data.response, 'system');
                }} else {{
                    const rtxIcon = data.rtx4090_optimized ? '🚀' : '📱';
                    addMessage(rtxIcon + ' ' + data.response, 'bot');
                }}
            }} catch (error) {{
                hideTyping();
                console.error('Error:', error);
                addMessage('❌ Connection error. Please try again.', 'system');
            }}
        }}

        function autoResize(textarea) {{
            textarea.style.height = 'auto';
            textarea.style.height = Math.min(textarea.scrollHeight, 120) + 'px';
        }}

        function changeLanguage() {{
            currentLanguage = document.getElementById('languageSelector').value;
            const input = document.getElementById('messageInput');
            input.placeholder = currentLanguage === 'vi' 
                ? "Nhập tin nhắn của bạn..." 
                : "Type your message...";
        }}

        // Event listeners
        document.getElementById('languageSelector').addEventListener('change', changeLanguage);
        document.getElementById('messageInput').addEventListener('keypress', function(e) {{
            if (e.key === 'Enter' && !e.shiftKey) {{
                e.preventDefault();
                sendMessage();
            }}
        }});
        document.getElementById('messageInput').addEventListener('input', function(e) {{
            autoResize(e.target);
        }});

        // Initialize
        console.log('🚀 RTX 4090 Web Interface initializing...');
        generateFullBodyPhoto();
        
        setTimeout(function() {{
            const initialGreeting = '🌟 Kính chào quý khách! Để tiện xưng hô, quý khách vui lòng cho biết tuổi và giới tính (ví dụ: "25 tuổi nam" hoặc "30 tuổi nữ") ạ.';
            addMessage(initialGreeting, 'bot');
        }}, 1000);
        
        console.log('✅ RTX 4090 Web Interface ready!');
</script>
</body>
</html>'''    
    
    @app.route('/api/chat', methods=['POST', 'OPTIONS'])
    @app.route('/api/chat', methods=['POST', 'OPTIONS'])
    def chat():
        """Enhanced chat with Vietnamese cultural addressing"""
        if request.method == 'OPTIONS':
            response = jsonify({})
            response.headers.add('Access-Control-Allow-Origin', '*')
            response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
            response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
            return response, 200
            
        try:
            if not request.is_json:
                return jsonify({
                    'error': 'Invalid content type',
                    'response': 'Content-Type must be application/json.'
                }), 400
            
            data = request.get_json()
            if not data:
                return jsonify({
                    'error': 'No data provided',
                    'response': 'Please provide message data.'
                }), 400
                
            user_message = data.get('message', '').strip()
            language = data.get('language', 'vi')
            session_id = data.get('session_id', f"web_{request.remote_addr}_{int(time.time())}")
            
            print(f"🌐 Web Interface: Message from session {session_id}: '{user_message}'")
            
            if not user_message:
                error_response = 'Please enter a message.' if language == 'en' else 'Vui lòng nhập tin nhắn.'
                return jsonify({
                    'error': 'Empty message',
                    'response': error_response
                }), 400
            
            # Process message with session ID for profile tracking
            result = chatbot_instance.process_message(user_message, session_id)
            
            print(f"✅ Web Interface: Response generated from {result.get('data_source')}")
            
            # Prepare enhanced response
            response_data = {
                'response': result.get('response', ''),
                'data_source': result.get('data_source', ''),
                'processing_time': result.get('processing_time', 0),
                'user_language': result.get('user_language', language),
                'local_results_count': result.get('local_results_count', 0),
                'rtx4090_optimized': hasattr(chatbot_instance, 'is_rtx4090') and chatbot_instance.is_rtx4090,
                'session_id': session_id,
                'customer_profile': result.get('customer_profile', {}),
                'requires_profile': result.get('requires_profile', False),
                'profile_setup_complete': result.get('profile_setup_complete', False)
            }
            
            response = jsonify(response_data)
            response.headers.add('Access-Control-Allow-Origin', '*')
            return response, 200
            
        except Exception as e:
            error_msg = f"Server error: {str(e)}"
            print(f"❌ Web Interface Error: {error_msg}")
            
            response = jsonify({
                'error': error_msg,
                'response': 'Sorry, I encountered an error.' if language == 'en' else 'Xin lỗi, tôi gặp lỗi.'
            })
            response.headers.add('Access-Control-Allow-Origin', '*')
            return response, 500
    
    @app.route('/api/status')
    def status():
        """Check RTX 4090 optimized chatbot status"""
        return jsonify({
            'status': 'online',
            'rtx4090_optimized': hasattr(chatbot_instance, 'is_rtx4090') and chatbot_instance.is_rtx4090,
            'model_available': chatbot_instance.text_generator is not None,
            'embedding_available': chatbot_instance.embedding_model is not None,
            'desktop_gui_running': hasattr(chatbot_instance, 'root') and chatbot_instance.root is not None,
            'conversation_count': len(chatbot_instance.conversation_context),
            'response_cache_size': len(getattr(chatbot_instance, 'response_cache', {})),
            'gpu_available': torch.cuda.is_available(),
            'device': str(getattr(chatbot_instance, 'primary_device', 'cpu')),
            'gpu_memory_gb': getattr(chatbot_instance, 'primary_memory_gb', 0)
        })
    
    return app


def main():
    import os
    print(f"🏃 Running from: {os.getcwd()}")
    print(f"📄 Script location: {os.path.abspath(__file__)}")
    
    # List all .db files in current and parent directories
    for root, dirs, files in os.walk('.'):
        for file in files:
            if file.endswith('.db'):
                full_path = os.path.join(root, file)
                mod_time = datetime.fromtimestamp(os.path.getmtime(full_path))
                size = os.path.getsize(full_path) / 1024  # KB
                print(f"  Found: {full_path} - {size:.1f}KB - Modified: {mod_time}")
        # Add this temporary code to your main() function
    import shutil
    import os

    # Find the active database with 248 products
    active_db = None
    for path in ['chatbot_data.db', '../chatbot_data.db', 'venv/chatbot_data.db']:
        if os.path.exists(path):
            conn = sqlite3.connect(path)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM products")
            count = cursor.fetchone()[0]
            conn.close()
            if count == 248:
                active_db = path
                break

    if active_db and active_db != 'chatbot_data.db':
        print(f"📋 Copying active database from {active_db} to main folder")
        shutil.copy2(active_db, 'chatbot_data.db')
    
    """Main function that starts both RTX 4090 optimized desktop and web interfaces"""
    print("""
    ===============================================
    🚀 SSTC AI SUPER SALES
    ===============================================
    
    🔧 RTX 4090 Features:
    ✅ Vietnamese + English language support
    ✅ RTX 4090 optimized AI models (24GB VRAM)
    ✅ Batch processing for ultra-fast performance
    ✅ Smart product search with RTX 4090 acceleration
    ✅ Desktop GUI + Web Interface with FIXED messaging
    ✅ Advanced natural language processing
    ✅ Response caching for instant replies
    ✅ GPU memory optimization
    
    🚀 Starting BOTH interfaces with RTX 4090 power...
    """)
    
    try:
        os.makedirs('data', exist_ok=True)
        os.makedirs('logs', exist_ok=True)
        os.makedirs('templates', exist_ok=True)
        
        if not os.path.exists('chatbot_config.yaml'):
            print("📝 Creating RTX 4090 optimized configuration file...")
            create_config_file()
        
        print("🚀 Initializing RTX 4090 optimized AI chatbot...")
        
        # Initialize with RTX 4090 optimizations
        chatbot = VietnameseAISalesBot(start_gui=False)
        print("✅ RTX 4090 optimized AI models loaded successfully!")
        # Check database contents right after initialization
        print("\n=== Checking Database Contents ===")
        chatbot.check_database_contents()
        
        # Add sample products
        chatbot.add_sample_products()
        
        # Check again after adding samples
        print("\n=== After Adding Sample Products ===")
        chatbot.check_database_contents()
        
        # Add RTX 4090 sample products
        print("📦 Adding RTX 4090 sample products...")
        chatbot.add_sample_products()
        # DEBUG CODE for add_sample_products()
        cursor = chatbot.conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM products")
        count = cursor.fetchone()[0]
        print(f"📊 Total products in database: {count}")
        
        cursor.execute("SELECT name, category FROM products LIMIT 5")
        products = cursor.fetchall()
        for p in products:
            print(f"  - {p[0]} ({p[1]})")
        
        # Start web interface in background thread
        def start_web_interface():
            try:
                print("🌐 Starting FIXED RTX 4090 web interface...")
                app = create_web_app(chatbot)
                print("🔗 RTX 4090 Web interface available at: http://localhost:5000")
                print("🎨 Full-body photos + RTX 4090 AI models + FIXED messaging!")
                app.run(debug=False, host='0.0.0.0', port=5000, use_reloader=False, threaded=True)
            except Exception as e:
                print(f"❌ RTX 4090 Web interface error: {e}")
                import traceback
                traceback.print_exc()
        
        # Start web interface in background
        web_thread = threading.Thread(target=start_web_interface, daemon=True)
        web_thread.start()
        
        # Give web interface time to start
        time.sleep(3)
        
        # Setup and start desktop GUI
        print("🖥️ Starting RTX 4090 optimized desktop GUI...")
        chatbot.setup_gui()
        chatbot.display_welcome_message()
        
        rtx_status = "🚀 RTX 4090 OPTIMIZED" if hasattr(chatbot, 'is_rtx4090') and chatbot.is_rtx4090 else "📱 STANDARD GPU"
        
        print("\n" + "="*70)
        print(f"🎉 BOTH INTERFACES RUNNING WITH {rtx_status}!")
        print("🖥️  Desktop GUI: RTX 4090 optimized models + Enhanced interface")
        print("🌐 Web Interface: RTX 4090 AI models + Beautiful photos + FIXED chat")
        print("🧠 Shared AI Brain: Same powerful RTX 4090 models for both interfaces")
        print("📱 Use either interface - full RTX 4090 capabilities on both!")
        print("🔄 Close this window to stop both interfaces")
        print("="*70 + "\n")
        
        # Start desktop GUI (blocking)
        chatbot.root.mainloop()
        
    except KeyboardInterrupt:
        print("\n👋 Both RTX 4090 interfaces stopped by user")
    except Exception as e:
        print(f"❌ Error starting RTX 4090 interfaces: {e}")
        print("\n🔧 RTX 4090 Troubleshooting tips:")
        print("1. Ensure RTX 4090 is properly installed")
        print("2. Check CUDA 11.8+ installation")
        print("3. Verify 24GB VRAM availability")
        print("4. Ensure sufficient system RAM (32GB+ recommended)")
        
        import traceback
        traceback.print_exc()
        
        input("\nPress Enter to exit...")


if __name__ == "__main__":
    main()
